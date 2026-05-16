
# --- Flask App and Config (must be before any route decorators) ---
import os
import re
import io
import csv
import json
import math
import zipfile
import random
import logging
import pandas as pd
import mysql.connector
from math import isfinite
from datetime import datetime
from functools import wraps, lru_cache
from itertools import combinations
from io import BytesIO, StringIO

from flask import (
    Flask, render_template, request, redirect, url_for, session,
    flash, make_response, get_flashed_messages, send_file, jsonify, abort
)
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from sqlalchemy import func, or_, text
from sqlalchemy.inspection import inspect
from sqlalchemy.exc import SQLAlchemyError

# --- Initialize Flask App ---
app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-change-this-in-production'
app.config['SQLALCHEMY_DATABASE_URI'] = (
    'mysql+pymysql://DOMASSIGNMENT:rcmouli123@DOMASSIGNMENT.mysql.pythonanywhere-services.com/DOMASSIGNMENT$default'
)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_pre_ping': True,
    'pool_recycle': 300
}

# --- Initialize Database and Migration ---
db = SQLAlchemy(app)
migrate = Migrate(app, db)

# --- Import Models and Utilities (after db is defined) ---
#from models import User, KDMLabQuizQuestion, KDMLabQuizSettings, KDMLabQuizAttempt
#from utils import find_admin_for_student  # make sure utils.py doesn't import app back

# --- Optional DOCX imports ---
try:
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    Inches = None
    WD_PARAGRAPH_ALIGNMENT = None

# --- Upload Folder Setup ---
UPLOAD_FOLDER = os.path.join(os.getcwd(), "tmp_uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)



@app.route("/kdm_lab_reset_individual_quiz", methods=["POST"])
def kdm_lab_reset_individual_quiz():
    """Reset all quiz attempts and marks for a specific student (admin-specific)"""
    if "loggedin" not in session:
        return jsonify({"error": "Not logged in"}), 401

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        return jsonify({"error": "Admin access required"}), 403

    try:
        data = request.get_json()
        student_id = data.get('student_id')
        if not student_id:
            return jsonify({"error": "Missing student_id"}), 400

        # Delete all quiz attempts for this student for this admin
        attempts = KDMLabQuizAttempt.query.filter_by(
            student_id=student_id, admin_id=current_admin.id).all()
        for attempt in attempts:
            db.session.delete(attempt)

        # Optionally, reset individual release (set is_released to False)
        release = KDMLabStudentQuizRelease.query.filter_by(
            admin_id=current_admin.id, student_id=student_id).first()
        if release:
            release.is_released = False

        db.session.commit()
        return jsonify({"success": True})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

# --- Download Database Route ---


# ==========================================================
# ✅ KDM LAB MODELS (Independent from main users table)
# ==========================================================
class AdminDashboardVisibility(db.Model):
    __tablename__ = 'admin_dashboard_visibility'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    dashboard_type = db.Column(db.String(50), nullable=False)
    is_visible = db.Column(db.Boolean, default=True)


class KDMLabSettings(db.Model):
    __tablename__ = 'kdm_lab_settings'

    id = db.Column(db.Integer, primary_key=True)
    num_experiments = db.Column(db.Integer, nullable=False)
    num_criteria = db.Column(db.Integer, nullable=False)
    student_visibility = db.Column(db.Boolean, default=False)
    admin_edit_locked = db.Column(db.Boolean, default=False)


class KDMLabExperiment(db.Model):
    __tablename__ = 'kdm_lab_experiment'
    id = db.Column(db.Integer, primary_key=True)
    experiment_number = db.Column(db.Integer, nullable=False)
    title = db.Column(db.String(200), nullable=False)
    is_released = db.Column(db.Boolean, default=False)
    admin_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True)
    criteria = db.relationship(
        'KDMLabCriteria', backref='experiment', cascade="all, delete-orphan")


class KDMLabCriteria(db.Model):
    __tablename__ = 'kdm_lab_criteria'
    id = db.Column(db.Integer, primary_key=True)
    experiment_id = db.Column(
        db.Integer, db.ForeignKey('kdm_lab_experiment.id'))
    criteria_number = db.Column(db.Integer)
    question_text = db.Column(db.Text)
    option_a = db.Column(db.String(255))
    option_b = db.Column(db.String(255))
    option_c = db.Column(db.String(255))
    option_d = db.Column(db.String(255))
    marks_a = db.Column(db.Float)
    marks_b = db.Column(db.Float)
    marks_c = db.Column(db.Float)
    marks_d = db.Column(db.Float)
    max_marks = db.Column(db.Float, default=4.0)
    withheld = db.Column(db.Boolean, default=False)
    responses = db.relationship(
        'KDMLabResponse', backref='criteria', cascade="all, delete-orphan")


class KDMLabStudent(db.Model):
    __tablename__ = 'kdm_lab_student'

    id = db.Column(db.Integer, primary_key=True)
    rollnumber = db.Column(db.String(50), nullable=False, unique=True)
    password = db.Column(db.String(255), nullable=False)
    name = db.Column(db.String(100))
    email = db.Column(db.String(100))
    phonenumber = db.Column(db.String(20))
    total_lab_score = db.Column(db.Float, default=0)
    completed_experiments = db.Column(db.Integer, default=0)

    # ✅ Add this new line
    created_by = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='SET NULL'))


class KDMLabResponse(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))  # ✅ match DB
    experiment_id = db.Column(
        db.Integer, db.ForeignKey('kdm_lab_experiment.id'))
    criteria_id = db.Column(db.Integer, db.ForeignKey('kdm_lab_criteria.id'))
    selected_option = db.Column(db.String(1))
    obtained_points = db.Column(db.Float)
    marks_earned = db.Column(db.Float)
    attempt_number = db.Column(db.Integer, default=1)  # ✅ Track attempt number


class KDMLabAttempt(db.Model):
    __tablename__ = 'kdm_lab_attempts'

    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    student_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    experiment_id = db.Column(db.Integer, db.ForeignKey(
        'kdm_lab_experiment.id', ondelete='CASCADE'), nullable=False)
    attempt_count = db.Column(db.Integer, default=0)
    last_attempt = db.Column(db.TIMESTAMP, server_default=db.func.current_timestamp(),
                             onupdate=db.func.current_timestamp())


class KDMLabManualMarks(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(
        db.Integer, db.ForeignKey('users.id'))  # ✅ match MySQL
    experiment_id = db.Column(db.Integer)
    criteria_id = db.Column(db.Integer)
    marks_given = db.Column(db.Float)


class KDMlabAdminRelease(db.Model):
    __tablename__ = 'kdm_lab_admin_release'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'))
    experiment_id = db.Column(db.Integer, db.ForeignKey(
        'kdm_lab_experiment.id', ondelete='CASCADE'))
    is_released = db.Column(db.Boolean, default=False)


class KDMLabStudentExperimentRelease(db.Model):
    __tablename__ = 'kdm_lab_student_experiment_releases'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'))
    student_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'))
    experiment_id = db.Column(db.Integer, db.ForeignKey(
        'kdm_lab_experiment.id', ondelete='CASCADE'))
    is_released = db.Column(db.Boolean, default=False)


class KDMLabQuizQuestion(db.Model):
    __tablename__ = 'kdm_lab_quiz_questions'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    slno = db.Column(db.Integer)
    question = db.Column(db.Text, nullable=False)
    option_a = db.Column(db.String(500))
    option_b = db.Column(db.String(500))
    option_c = db.Column(db.String(500))
    option_d = db.Column(db.String(500))
    correct_answer = db.Column(db.String(1))  # A, B, C, or D
    created_at = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())


class KDMLabQuizAttempt(db.Model):
    __tablename__ = 'kdm_lab_quiz_attempts'

    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    quiz_questions = db.Column(db.Text)  # JSON string of question IDs
    student_answers = db.Column(db.Text)  # JSON string of answers
    score = db.Column(db.Integer, default=0)
    total_questions = db.Column(db.Integer, default=10)
    attempt_number = db.Column(db.Integer, default=1)
    start_time = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())
    end_time = db.Column(db.TIMESTAMP)
    is_completed = db.Column(db.Boolean, default=False)


class KDMLabQuizSettings(db.Model):
    __tablename__ = 'kdm_lab_quiz_settings'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    is_quiz_released = db.Column(db.Boolean, default=False)
    quiz_duration_minutes = db.Column(db.Integer, default=10)
    questions_per_quiz = db.Column(db.Integer, default=10)
    created_at = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())
    updated_at = db.Column(db.TIMESTAMP, server_default=db.func.current_timestamp(
    ), onupdate=db.func.current_timestamp())


class KDMLabStudentQuizRelease(db.Model):
    __tablename__ = 'kdm_lab_student_quiz_releases'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'))
    student_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'))
    is_released = db.Column(db.Boolean, default=False)


# ==========================================================
# ✅ PYTHON LAB MODELS (Python Programming Lab Quiz System)
# ==========================================================

class PythonLabQuizSettings(db.Model):
    __tablename__ = 'python_lab_quiz_settings'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, nullable=False)
    module_number = db.Column(
        db.Integer, nullable=False)  # Module number (1-5)
    is_active = db.Column(db.Boolean, default=False)
    is_quiz_released = db.Column(db.Boolean, default=False)
    # Global setting for student question visibility (only for super admin)
    questions_visible_to_students = db.Column(db.Boolean, default=False)
    max_attempts = db.Column(db.Integer, default=1)
    # Quiz duration in minutes
    quiz_duration = db.Column(db.Integer, default=60)
    # Questions per module for the quiz
    questions_per_module = db.Column(db.Integer, default=4)
    start_time = db.Column(db.DateTime, nullable=True)
    end_time = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())
    updated_at = db.Column(db.TIMESTAMP, server_default=db.func.current_timestamp(
    ), onupdate=db.func.current_timestamp())


class PythonLabStudentQuizRelease(db.Model):
    __tablename__ = 'python_lab_student_quiz_releases'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    student_id = db.Column(
        db.Integer, db.ForeignKey('users.id'), nullable=False)
    is_released = db.Column(db.Boolean, default=False)
    created_at = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())
    updated_at = db.Column(db.TIMESTAMP, server_default=db.func.current_timestamp(
    ), onupdate=db.func.current_timestamp())

    # Ensure one release record per admin-student pair
    __table_args__ = (db.UniqueConstraint('admin_id', 'student_id'),)


class PythonLabQuizQuestion(db.Model):
    __tablename__ = 'python_lab_quiz_questions'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    # 1-5 for the 5 modules
    module_number = db.Column(db.Integer, nullable=False)
    module_name = db.Column(db.String(100), nullable=False)
    slno = db.Column(db.Integer)
    question = db.Column(db.Text, nullable=False)
    option_a = db.Column(db.String(500))
    option_b = db.Column(db.String(500))
    option_c = db.Column(db.String(500))
    option_d = db.Column(db.String(500))
    correct_answer = db.Column(db.String(1))  # A, B, C, or D
    points = db.Column(db.Integer, default=1)
    created_at = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())
    updated_at = db.Column(db.TIMESTAMP, server_default=db.func.current_timestamp(
    ), onupdate=db.func.current_timestamp())


class PythonLabQuizAttempt(db.Model):
    __tablename__ = 'python_lab_quiz_attempts'

    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    attempt_number = db.Column(db.Integer, default=1)
    quiz_questions = db.Column(db.Text)  # Comma-separated question IDs
    score = db.Column(db.Integer, default=0)  # Percentage score (0-100)
    # 4 per module × 5 modules
    total_questions = db.Column(db.Integer, default=20)
    started_at = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())
    completed_at = db.Column(db.TIMESTAMP, nullable=True)
    is_completed = db.Column(db.Boolean, default=False)
    created_at = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())
    updated_at = db.Column(db.TIMESTAMP, server_default=db.func.current_timestamp(
    ), onupdate=db.func.current_timestamp())


class PythonLabQuizResponse(db.Model):
    __tablename__ = 'python_lab_quiz_responses'

    id = db.Column(db.Integer, primary_key=True)
    attempt_id = db.Column(db.Integer, db.ForeignKey(
        'python_lab_quiz_attempts.id', ondelete='CASCADE'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey(
        'python_lab_quiz_questions.id', ondelete='CASCADE'), nullable=False)
    student_answer = db.Column(db.String(1))  # A, B, C, D, or NULL
    correct_answer = db.Column(db.String(1), nullable=False)
    is_correct = db.Column(db.Boolean, default=False)
    module_number = db.Column(db.Integer, nullable=False)
    created_at = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())


# ==========================================================
# ✅ PYTHON LAB EXPERIMENTS MODELS (Similar to KDM Lab)
# ==========================================================

class PythonLabSettings(db.Model):
    __tablename__ = 'python_lab_settings'

    id = db.Column(db.Integer, primary_key=True)
    num_experiments = db.Column(db.Integer, nullable=False)
    num_criteria = db.Column(db.Integer, nullable=False)
    student_visibility = db.Column(db.Boolean, default=False)
    admin_edit_locked = db.Column(db.Boolean, default=False)


class PythonLabExperiment(db.Model):
    __tablename__ = 'python_lab_experiment'
    id = db.Column(db.Integer, primary_key=True)
    experiment_number = db.Column(db.Integer, nullable=False)
    title = db.Column(db.String(200), nullable=False)
    is_released = db.Column(db.Boolean, default=False)
    admin_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True)
    criteria = db.relationship(
        'PythonLabCriteria', backref='experiment', cascade="all, delete-orphan")


class PythonLabCriteria(db.Model):
    __tablename__ = 'python_lab_criteria'
    id = db.Column(db.Integer, primary_key=True)
    experiment_id = db.Column(
        db.Integer, db.ForeignKey('python_lab_experiment.id'))
    criteria_number = db.Column(db.Integer)
    question_text = db.Column(db.Text)
    option_a = db.Column(db.String(255))
    option_b = db.Column(db.String(255))
    option_c = db.Column(db.String(255))
    option_d = db.Column(db.String(255))
    marks_a = db.Column(db.Float)
    marks_b = db.Column(db.Float)
    marks_c = db.Column(db.Float)
    marks_d = db.Column(db.Float)
    max_marks = db.Column(db.Float, default=4.0)
    withheld = db.Column(db.Boolean, default=False)
    responses = db.relationship(
        'PythonLabResponse', backref='criteria', cascade="all, delete-orphan")


class PythonLabStudent(db.Model):
    __tablename__ = 'python_lab_student'

    id = db.Column(db.Integer, primary_key=True)
    rollnumber = db.Column(db.String(50), nullable=False, unique=True)
    password = db.Column(db.String(255), nullable=False)
    name = db.Column(db.String(100))
    email = db.Column(db.String(100))
    phonenumber = db.Column(db.String(20))
    total_lab_score = db.Column(db.Float, default=0)
    completed_experiments = db.Column(db.Integer, default=0)
    created_by = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='SET NULL'))


class PythonLabResponse(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    experiment_id = db.Column(
        db.Integer, db.ForeignKey('python_lab_experiment.id'))
    criteria_id = db.Column(
        db.Integer, db.ForeignKey('python_lab_criteria.id'))
    selected_option = db.Column(db.String(1))
    obtained_points = db.Column(db.Float)
    marks_earned = db.Column(db.Float)
    attempt_number = db.Column(db.Integer, default=1)


class PythonLabAttempt(db.Model):
    __tablename__ = 'python_lab_attempts'

    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    student_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    experiment_id = db.Column(db.Integer, db.ForeignKey(
        'python_lab_experiment.id', ondelete='CASCADE'), nullable=False)
    attempt_count = db.Column(db.Integer, default=0)
    last_attempt = db.Column(db.TIMESTAMP, server_default=db.func.current_timestamp(),
                             onupdate=db.func.current_timestamp())


class PythonLabManualMarks(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(
        db.Integer, db.ForeignKey('users.id'))
    experiment_id = db.Column(db.Integer)
    criteria_id = db.Column(db.Integer)
    marks_given = db.Column(db.Float)


class PythonLabAdminRelease(db.Model):
    __tablename__ = 'python_lab_admin_release'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'))
    experiment_id = db.Column(db.Integer, db.ForeignKey(
        'python_lab_experiment.id', ondelete='CASCADE'))
    is_released = db.Column(db.Boolean, default=False)


class PythonLabStudentExperimentRelease(db.Model):
    __tablename__ = 'python_lab_student_experiment_releases'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    student_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'), nullable=False)
    experiment_id = db.Column(db.Integer, db.ForeignKey(
        'python_lab_experiment.id', ondelete='CASCADE'), nullable=False)
    is_released = db.Column(db.Boolean, default=False)
    created_at = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())
    updated_at = db.Column(db.TIMESTAMP, server_default=db.func.current_timestamp(
    ), onupdate=db.func.current_timestamp())

    # Ensure one release record per admin-student-experiment combination
    __table_args__ = (db.UniqueConstraint(
        'admin_id', 'student_id', 'experiment_id'),)


# ------------------------------------------------------------------------------------------------------
# ========================= KDM LAB — ROUTES & HELPERS =========================
# Paste this into app.py (after models/imports). Adjust imports if needed.


# --- Helpers -----------------------------------------------------------------


def _require_login():
    if "loggedin" not in session:
        return False
    return True


def _current_admin_user():
    # Uses your existing User model
    if "rollnumber" not in session:
        return None
    return User.query.filter_by(rollnumber=session["rollnumber"]).first()


def _is_admin(user):
    return bool(user and user.role in ["admin", "super_admin"])


def _is_super_admin(user):
    return bool(user and user.admin_level == "super_admin")


def _get_lab_settings(create_if_missing=True):
    s = KDMLabSettings.query.first()
    if not s and create_if_missing:
        s = KDMLabSettings(num_experiments=12, num_criteria=10,
                           student_visibility=False, admin_edit_locked=False)
        db.session.add(s)
        db.session.commit()
    return s


def _ensure_experiments_exist(num_experiments):
    # Ensure EX-1..EX-n exist and remove any > n
    existing = {e.experiment_number: e for e in KDMLabExperiment.query.all()}
    # Create missing
    created = 0
    for i in range(1, num_experiments + 1):
        if i not in existing:
            ex = KDMLabExperiment(experiment_number=i,
                                  title=f"EX-{i}", is_released=False)
            db.session.add(ex)
            created += 1
    # Delete beyond
    for ex in KDMLabExperiment.query.all():
        if ex.experiment_number > num_experiments:
            # cascades will remove criteria/responses/manual marks
            db.session.delete(ex)
    if created:
        db.session.commit()


def _sync_lab_students_for_admin(admin_user):
    """
    Auto-sync KDM LAB student list from main users assigned to this admin.
    Super admin -> all students; admin -> assigned only.
    """
    main_students = get_students_for_admin(admin_user)  # you already have this
    main_rolls = [s.rollnumber for s in main_students]

    existing = {s.rollnumber: s for s in KDMLabStudent.query.filter(
        KDMLabStudent.rollnumber.in_(main_rolls)).all()}
    to_add = []
    for u in main_students:
        if u.rollnumber not in existing:
            to_add.append(KDMLabStudent(
                rollnumber=u.rollnumber,
                password=u.password,
                name=getattr(u, "name", None),
                email=getattr(u, "email", None),
                phonenumber=getattr(u, "phonenumber", None),
            ))
    if to_add:
        db.session.add_all(to_add)
        db.session.commit()
        print(
            f"🔁 Synced {len(to_add)} KDM LAB students for admin {admin_user.rollnumber}")


def _lab_students_for_admin(admin_user):
    """Return KDMLabStudent list filtered by admin scope."""
    if _is_super_admin(admin_user):
        return KDMLabStudent.query.order_by(KDMLabStudent.rollnumber).all()
    assigned = get_students_for_admin(admin_user)
    assigned_rolls = [s.rollnumber for s in assigned]
    if not assigned_rolls:
        return []
    return KDMLabStudent.query.filter(KDMLabStudent.rollnumber.in_(assigned_rolls)).order_by(KDMLabStudent.rollnumber).all()


def _lab_student_for_logged_in_user():
    """Ensure KDMLabStudent exists for the logged-in user (student)."""
    if "rollnumber" not in session:
        return None
    roll = session["rollnumber"]
    lab_s = KDMLabStudent.query.filter_by(rollnumber=roll).first()
    if not lab_s:
        # Create a minimal record using existing users table (read-only)
        u = User.query.filter_by(rollnumber=roll).first()
        if not u:
            return None
        lab_s = KDMLabStudent(
            rollnumber=u.rollnumber,
            password=u.password,
            name=getattr(u, "name", None),
            email=getattr(u, "email", None),
            phonenumber=getattr(u, "phonenumber", None),
        )
        db.session.add(lab_s)
        db.session.commit()
    return lab_s


def _criteria_for_experiment(experiment_id, include_withheld=False):
    q = KDMLabCriteria.query.filter_by(experiment_id=experiment_id)
    if not include_withheld:
        q = q.filter_by(withheld=False)
    return q.order_by(KDMLabCriteria.criteria_number).all()


def _calc_marks_earned(max_marks, obtained_points):
    try:
        if not isfinite(max_marks) or not isfinite(obtained_points):
            return 0.0
        return float(max_marks) * (float(obtained_points) / 4.0)
    except Exception:
        return 0.0


def _get_best_kdm_lab_response(user_id, experiment_id, criteria_id):
    """
    Get the best response (highest obtained_points) for a user, experiment, and criteria
    across all attempts.
    """
    responses = KDMLabResponse.query.filter_by(
        user_id=user_id,
        experiment_id=experiment_id,
        criteria_id=criteria_id
    ).all()

    if not responses:
        return None

    # Return the response with the highest obtained_points
    best_response = max(responses, key=lambda r: r.obtained_points or 0)
    return best_response


def _get_best_kdm_quiz_score(student_id, admin_id):
    """
    Get the best quiz score for a student from KDM lab internal quiz attempts.
    Returns the score converted to 10 marks scale.
    """
    # Get all completed quiz attempts for this student and admin
    attempts = KDMLabQuizAttempt.query.filter_by(
        student_id=student_id,
        admin_id=admin_id,
        is_completed=True
    ).all()

    # If no attempts found with specific admin, try all admins (in case super admin released quiz)
    if not attempts:
        print(
            f"⚠️ No quiz attempts found for student_id={student_id}, admin_id={admin_id}")
        attempts = KDMLabQuizAttempt.query.filter_by(
            student_id=student_id,
            is_completed=True
        ).all()
        print(
            f"🔍 Found {len(attempts)} attempts for student_id={student_id} across all admins")

    if not attempts:
        print(f"❌ No quiz attempts at all for student_id={student_id}")
        return 0  # No quiz attempts

    # Get the best score from all attempts
    best_score = max(attempt.score for attempt in attempts)
    print(f"✅ Best score for student_id={student_id}: {best_score}")

    # Get quiz settings to determine total questions
    quiz_settings = KDMLabQuizSettings.query.filter_by(
        admin_id=admin_id).first()

    # If no quiz settings for this admin, try to find any quiz settings
    if not quiz_settings:
        quiz_settings = KDMLabQuizSettings.query.first()
        print(
            f"⚠️ No quiz settings for admin_id={admin_id}, using fallback settings")

    total_questions = quiz_settings.questions_per_quiz if quiz_settings else 10

    print(
        f"📋 Quiz settings for admin_id={admin_id}: {quiz_settings is not None}, total_questions={total_questions}")

    # Convert to 10 marks scale: (best_score / total_questions) * 10
    viva_marks = (best_score / total_questions) * \
        10 if total_questions > 0 else 0

    print(
        f"🎯 Final calculation: ({best_score} / {total_questions}) * 10 = {viva_marks}")

    # Round viva marks to 2 decimal places (normal rounding for precision)
    # Note: Final internal marks will be rounded up, but keeping viva precise for intermediate calculation
    return round(viva_marks, 2)


def _recompute_student_experiment_total(student_id, experiment_id):
    # Get all criteria for this experiment
    criteria = KDMLabCriteria.query.filter_by(
        experiment_id=experiment_id).all()
    total = 0.0

    for c in criteria:
        if c.withheld:
            # Manual marks for withheld criteria
            mm = KDMLabManualMarks.query.filter_by(
                student_id=student_id,
                experiment_id=experiment_id,
                criteria_id=c.id,
            ).first()
            obtained_marks = mm.marks_given if mm and mm.marks_given is not None else 0.0
            # Apply weighted formula: (obtained_marks * max_marks) / 4
            weighted_mark = obtained_marks * (c.max_marks or 4.0) / 4.0
            total += weighted_mark
        else:
            # Student response for regular criteria - get BEST attempt
            resp = _get_best_kdm_lab_response(student_id, experiment_id, c.id)
            if resp and resp.obtained_points is not None:
                # Apply weighted formula: (obtained_points * max_marks) / 4
                weighted_mark = resp.obtained_points * \
                    (c.max_marks or 4.0) / 4.0
                total += weighted_mark

    return round(total, 2)


def _recompute_student_lab_totals(student_id):
    # optional: compute cumulative lab totals across all experiments
    ex_ids = [e.id for e in KDMLabExperiment.query.all()]
    grand = 0.0
    completed = 0
    for exid in ex_ids:
        t = _recompute_student_experiment_total(student_id, exid)
        if t > 0:
            completed += 1
        grand += t
    stu = KDMLabStudent.query.get(student_id)
    if stu:
        stu.total_lab_score = round(grand, 2)
        stu.completed_experiments = completed
        db.session.commit()


def get_admin_for_student(student):
    """
    Finds which admin a student belongs to using the same logic as assignment dashboard.
    Returns the User (admin) object or None.
    """
    roll = getattr(student, "rollnumber", None)
    if not roll:
        return None

    admin = None

    try:
        # Get all admins (normal + super admins)
        all_admins = User.query.filter(User.role == "admin").all()
        for a in all_admins:
            students_under_admin = get_students_for_admin(a)
            for s in students_under_admin:
                if s.rollnumber == roll:
                    return a
    except Exception as e:
        print(f"⚠️ Error mapping admin for student {roll}: {e}")

    # If not found, fallback to super admin
    super_admin = User.query.filter_by(admin_level="super_admin").first()
    if super_admin:
        return super_admin

    return admin


# --- ADMIN: Dashboard ---------------------------------------------------------


@app.route("/kdm_lab_admin_dashboard")
def kdm_lab_admin_dashboard():
    if "loggedin" not in session:
        return redirect(url_for("login"))
    from flask import get_flashed_messages
    get_flashed_messages()

    # ✅ Ensure only admin or super_admin can access
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    # ✅ Load lab settings safely
    lab_settings = KDMLabSettings.query.first()

    # ✅ Get experiments
    experiments = KDMLabExperiment.query.order_by(
        KDMLabExperiment.experiment_number).all()

    # ✅ If super admin, show all students
    if current_admin.admin_level == "super_admin":
        students = User.query.filter_by(role="student").all()
    else:
        # ✅ Reuse existing logic — same as in assignment dashboard
        students = get_students_for_admin(current_admin)

    # Sort students by roll number in ascending order
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # ✅ Pre-load all data in bulk for performance (same optimization as consolidated marks)
    student_ids = [s.id for s in students]

    # Bulk load all criteria once
    all_criteria = KDMLabCriteria.query.all()
    criteria_by_experiment = {}  # {experiment_id: [criteria_objects]}
    for c in all_criteria:
        if c.experiment_id not in criteria_by_experiment:
            criteria_by_experiment[c.experiment_id] = []
        criteria_by_experiment[c.experiment_id].append(c)

    # Bulk load all manual marks for all students
    all_manual_marks = KDMLabManualMarks.query.filter(
        KDMLabManualMarks.student_id.in_(student_ids)).all()
    # {(student_id, experiment_id, criteria_id): marks_given}
    manual_marks_lookup = {}
    for mm in all_manual_marks:
        manual_marks_lookup[(mm.student_id, mm.experiment_id,
                             mm.criteria_id)] = mm.marks_given

    # Bulk load all responses for all students
    all_responses = KDMLabResponse.query.filter(
        KDMLabResponse.user_id.in_(student_ids)).all()

    # Group responses by (student_id, experiment_id, criteria_id) and keep best
    # {(student_id, experiment_id, criteria_id): best_response}
    best_responses = {}
    for resp in all_responses:
        key = (resp.user_id, resp.experiment_id, resp.criteria_id)
        if key not in best_responses or (resp.obtained_points or 0) > (best_responses[key].obtained_points or 0):
            best_responses[key] = resp

    # Bulk load all attempts for all students
    all_attempts = KDMLabAttempt.query.filter(
        KDMLabAttempt.student_id.in_(student_ids)).all()
    attempts_lookup = {}  # {(student_id, experiment_id): attempt_count}
    for att in all_attempts:
        attempts_lookup[(att.student_id, att.experiment_id)
                        ] = att.attempt_count

    # ✅ Prepare marks & attempts data using pre-loaded data
    lab_data = {}
    for student in students:
        lab_data[student.id] = {}
        total_marks_sum = 0.0
        completed_count = 0   # ✅ Initialize here before counting

        for ex in experiments:
            # ✅ Calculate total marks using pre-loaded data
            experiment_total = 0.0

            # Get criteria for this experiment from pre-loaded data
            criteria_list = criteria_by_experiment.get(ex.id, [])

            for c in criteria_list:
                if c.withheld:
                    # Look up manual marks from pre-loaded data
                    obtained_marks = manual_marks_lookup.get(
                        (student.id, ex.id, c.id), 0.0)
                    # Apply weighted formula: (obtained_marks * max_marks) / 4
                    weighted_mark = obtained_marks * (c.max_marks or 4.0) / 4.0
                    experiment_total += weighted_mark
                else:
                    # Look up best response from pre-loaded data
                    resp = best_responses.get((student.id, ex.id, c.id))
                    if resp and resp.obtained_points is not None:
                        # Apply weighted formula: (obtained_points * max_marks) / 4
                        weighted_mark = resp.obtained_points * \
                            (c.max_marks or 4.0) / 4.0
                        experiment_total += weighted_mark

            # Look up attempts from pre-loaded data
            attempts = attempts_lookup.get((student.id, ex.id), 0)

            lab_data[student.id][ex.id] = {
                "marks": round(experiment_total, 2),
                "attempts": attempts,
            }

            # ✅ Count experiment as completed if marks > 0 or attempt made
            if experiment_total > 0 or attempts > 0:
                completed_count += 1

            total_marks_sum += experiment_total

        # ✅ Compute average by dividing by ALL experiments (not just ones with marks > 0)
        avg_score = math.ceil(
            total_marks_sum / len(experiments)) if len(experiments) > 0 else 0
        student.total_lab_score = avg_score

        # ✅ Save completed experiments count
        student.completed_experiments = completed_count

    # ✅ Fetch release info per admin
    admin_releases = {}
    global_releases = {}

    if current_admin.admin_level == "super_admin":
        # Super admin: Check if experiments are released for ANY admin
        # If an experiment is released for any admin, show it as released
        all_releases = KDMlabAdminRelease.query.all()
        experiment_release_counts = {}

        # Count how many admins have each experiment released
        for release in all_releases:
            if release.is_released:
                experiment_release_counts[release.experiment_id] = experiment_release_counts.get(
                    release.experiment_id, 0) + 1

        # If any admin has the experiment released, show as released for super admin
        # Use global_releases for super admin (template expects this)
        for exp_id, count in experiment_release_counts.items():
            global_releases[exp_id] = count > 0

    else:
        # Regular admin: Only check their own release records
        admin_releases = {
            r.experiment_id: r.is_released
            for r in KDMlabAdminRelease.query.filter_by(admin_id=current_admin.id).all()
        }

    return render_template(
        "kdm_lab_admin_dashboard.html",
        students=students,
        experiments=experiments,
        lab_data=lab_data,
        lab_settings=lab_settings,
        current_admin=current_admin,
        admin_releases=admin_releases,
        global_releases=global_releases,   # ✅ Added this
    )


# ✏️ EDIT student’s KDM LAB marks & attempts
@app.route("/edit_kdm_lab_student/<int:student_id>", methods=["GET", "POST"])
def edit_kdm_lab_student(student_id):
    if "loggedin" not in session or session.get("role") != "admin":
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    student = User.query.get_or_404(student_id)
    experiments = KDMLabExperiment.query.order_by(
        KDMLabExperiment.experiment_number).all()

    # ✅ On form submit
    if request.method == "POST":
        for ex in experiments:
            marks_field = request.form.get(f"marks_{ex.id}")
            attempt_field = request.form.get(f"attempts_{ex.id}")

            # Update marks - Note: This is handled through the best attempt logic now
            # Individual marks should be updated via quiz attempts or manual marks, not here
            # This edit function is mainly for attempt count adjustments

            # Update attempts
            if attempt_field is not None:
                attempts_value = int(
                    attempt_field) if attempt_field.strip() else 0
                attempt_record = KDMLabAttempt.query.filter_by(
                    student_id=student.id, experiment_id=ex.id
                ).first()

                if not attempt_record:
                    attempt_record = KDMLabAttempt(
                        student_id=student.id,
                        experiment_id=ex.id,
                        attempt_count=attempts_value
                    )
                    db.session.add(attempt_record)
                else:
                    attempt_record.attempt_count = attempts_value

        db.session.commit()
        flash("✅ Student marks and attempts updated successfully!", "success")
        return redirect(url_for("kdm_lab_admin_dashboard"))

    # ✅ For GET request - Use best attempt logic for consistent calculation
    lab_data = {}
    for ex in experiments:
        # Calculate experiment total using same logic as admin dashboard
        experiment_total = 0.0

        # Get all criteria for this experiment
        criteria = KDMLabCriteria.query.filter_by(experiment_id=ex.id).all()

        for c in criteria:
            if c.withheld:
                # Manual marks for withheld criteria
                mm = KDMLabManualMarks.query.filter_by(
                    student_id=student.id,
                    experiment_id=ex.id,
                    criteria_id=c.id,
                ).first()
                obtained_marks = mm.marks_given if mm and mm.marks_given is not None else 0.0
                # Apply weighted formula: (obtained_marks * max_marks) / 4
                weighted_mark = obtained_marks * (c.max_marks or 4.0) / 4.0
                experiment_total += weighted_mark
            else:
                # Student response for regular criteria - get BEST attempt
                resp = _get_best_kdm_lab_response(student.id, ex.id, c.id)
                if resp and resp.obtained_points is not None:
                    # Apply weighted formula: (obtained_points * max_marks) / 4
                    weighted_mark = resp.obtained_points * \
                        (c.max_marks or 4.0) / 4.0
                    experiment_total += weighted_mark

        attempt = KDMLabAttempt.query.filter_by(
            student_id=student.id, experiment_id=ex.id).first()
        attempts = attempt.attempt_count if attempt else 0
        lab_data[ex.id] = {"marks": round(
            experiment_total, 2), "attempts": attempts}

    return render_template("edit_kdm_lab_student.html", student=student, experiments=experiments, lab_data=lab_data)


# 🗑️ DELETE all lab data for one student
@app.route("/delete_kdm_lab_student/<int:student_id>")
def delete_kdm_lab_student(student_id):
    if "loggedin" not in session or session.get("role") != "admin":
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    student = User.query.get_or_404(student_id)

    # Delete all related lab responses, manual marks & attempts
    KDMLabResponse.query.filter_by(user_id=student.id).delete()
    KDMLabManualMarks.query.filter_by(student_id=student.id).delete()
    KDMLabAttempt.query.filter_by(student_id=student.id).delete()

    db.session.commit()
    flash(f"🗑️ Deleted all KDM Lab data for {student.rollnumber}.", "warning")
    return redirect(url_for("kdm_lab_admin_dashboard"))


# 🔄 RESET attempts ,marks
@app.route("/reset_kdm_lab_attempts/<int:student_id>")
def reset_kdm_lab_attempts(student_id):
    if "loggedin" not in session or session.get("role") != "admin":
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    student = User.query.get_or_404(student_id)

    # ✅ Delete all responses (removes quiz answers + marks)
    KDMLabResponse.query.filter_by(user_id=student.id).delete()

    # ✅ Delete all manual marks
    KDMLabManualMarks.query.filter_by(student_id=student.id).delete()

    # ✅ Reset attempts count
    KDMLabAttempt.query.filter_by(student_id=student.id).delete()

    # ✅ Reset total score + completed
    student.total_lab_score = 0
    student.completed_experiments = 0

    db.session.commit()

    flash(
        f"✅ Reset all attempts and marks for {student.rollnumber}.", "success")
    return redirect(url_for("kdm_lab_admin_dashboard"))


# 🔄 RESET Python Lab quiz attempts for a student
@app.route("/reset_python_lab_attempts", methods=["POST"])
def reset_python_lab_attempts():
    if "loggedin" not in session:
        return jsonify({"success": False, "message": "Login required."})

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        return jsonify({"success": False, "message": "Admin access required."})

    try:
        data = request.get_json()
        student_id = data.get("student_id")

        if not student_id:
            return jsonify({"success": False, "message": "Student ID is required."})

        student = User.query.get(student_id)
        if not student:
            return jsonify({"success": False, "message": "Student not found."})

        # First get all attempt IDs for this student
        attempt_ids = [attempt.id for attempt in PythonLabQuizAttempt.query.filter_by(
            student_id=student.id).all()]

        # Delete all Python Lab quiz responses for these attempts
        if attempt_ids:
            PythonLabQuizResponse.query.filter(
                PythonLabQuizResponse.attempt_id.in_(attempt_ids)
            ).delete(synchronize_session=False)

        # Delete all Python Lab quiz attempts for this student
        PythonLabQuizAttempt.query.filter_by(student_id=student.id).delete()

        db.session.commit()

        return jsonify({
            "success": True,
            "message": f"Successfully reset all Python Lab quiz attempts for {student.rollnumber}."
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": f"Error: {str(e)}"})


@app.route("/python_lab_reset_all_quiz_attempts", methods=["POST"])
def python_lab_reset_all_quiz_attempts():
    """Reset all Python Lab internal quiz attempts for all students (Super Admin only)"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.admin_level != "super_admin":
        flash("Super admin access required!", "danger")
        return redirect(url_for("python_lab_admin_dashboard"))

    try:
        # Delete all Python Lab quiz responses
        PythonLabQuizResponse.query.delete()

        # Delete all Python Lab quiz attempts
        PythonLabQuizAttempt.query.delete()

        db.session.commit()
        flash("All Python Lab internal quiz attempts have been reset for all students!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error resetting quiz attempts: {str(e)}", "danger")

    return redirect(url_for("python_lab_admin_dashboard"))


# --- ADMIN: Set Settings (super admin only) -----------------------------------

@app.route("/set_kdm_lab_settings", methods=["GET", "POST"])
def set_kdm_lab_settings():
    if not _require_login():
        return redirect(url_for("login"))

    current_admin = _current_admin_user()
    if not _is_super_admin(current_admin):
        flash("Only super admin can modify lab settings.", "danger")
        return redirect(url_for("kdm_lab_admin_dashboard"))

    settings = _get_lab_settings(create_if_missing=True)

    if request.method == "POST":
        try:
            num_experiments = int(request.form.get(
                "num_experiments", settings.num_experiments))
            num_criteria = int(request.form.get(
                "num_criteria", settings.num_criteria))
            if num_experiments < 1 or num_criteria < 1:
                raise ValueError

            settings.num_experiments = num_experiments
            settings.num_criteria = num_criteria
            db.session.commit()

            _ensure_experiments_exist(num_experiments)
            flash("KDM LAB settings updated.", "success")
            return redirect(url_for("kdm_lab_admin_dashboard"))
        except Exception:
            db.session.rollback()
            flash("Invalid values for experiments/criteria.", "danger")

    current_experiments = KDMLabExperiment.query.order_by(
        KDMLabExperiment.experiment_number).all()

    return render_template(
        "set_kdm_lab_settings.html",
        lab_settings=settings,
        current_experiments=current_experiments
    )

# --- ADMIN: Toggle student visibility (super admin only) ----------------------


@app.route("/toggle_kdm_lab_visibility", methods=["POST"])
def toggle_kdm_lab_visibility():
    if not _require_login():
        return redirect(url_for("login"))
    current_admin = _current_admin_user()
    if not _is_super_admin(current_admin):
        flash("Only super admin can toggle visibility.", "danger")
        return redirect(url_for("kdm_lab_admin_dashboard"))

    s = _get_lab_settings()
    s.student_visibility = not bool(s.student_visibility)
    db.session.commit()
    flash(
        f"Student visibility set to: {'ON' if s.student_visibility else 'OFF'}", "success")
    return redirect(url_for("kdm_lab_admin_dashboard"))

# --- ADMIN: Toggle admin edit lock (super admin only) -------------------------


@app.route("/toggle_kdm_lab_edit_lock", methods=["POST"])
def toggle_kdm_lab_edit_lock():
    if not _require_login():
        return redirect(url_for("login"))
    current_admin = _current_admin_user()
    if not _is_super_admin(current_admin):
        flash("Only super admin can lock/unlock editing.", "danger")
        return redirect(url_for("kdm_lab_admin_dashboard"))

    s = _get_lab_settings()
    s.admin_edit_locked = not bool(s.admin_edit_locked)
    db.session.commit()
    flash(
        f"Admin edit lock is now: {'LOCKED' if s.admin_edit_locked else 'UNLOCKED'}", "success")
    return redirect(url_for("kdm_lab_admin_dashboard"))

# --- ADMIN: Create/Edit Quiz (criteria) ---------------------------------------


@app.route("/create_kdm_lab_quiz", methods=["GET", "POST"])
def create_kdm_lab_quiz():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    lab_settings = KDMLabSettings.query.first()
    if not lab_settings:
        flash("KDM LAB settings not found. Please set them first.", "warning")
        return redirect(url_for("kdm_lab_admin_dashboard"))

    experiments = KDMLabExperiment.query.order_by(
        KDMLabExperiment.experiment_number).all()

    selected_experiment_number = None
    experiment = None
    existing_criteria = []

    # ------------------- POST -------------------
    if request.method == "POST":
        selected_experiment_number = int(request.form.get("experiment_number"))
        title = request.form.get("title").strip() if request.form.get(
            "title") else f"EX-{selected_experiment_number}"
        criteria_count = int(request.form.get(
            "criteria_count") or lab_settings.num_criteria)

        experiment = KDMLabExperiment.query.filter_by(
            experiment_number=selected_experiment_number).first()
        if not experiment:
            experiment = KDMLabExperiment(
                experiment_number=selected_experiment_number,
                title=title,
                is_released=False
            )
            db.session.add(experiment)
        else:
            experiment.title = title

        db.session.commit()

        # --- Save / update criteria ---
        for i in range(1, criteria_count + 1):
            q_text = request.form.get(f"question_{i}")
            if not q_text:
                continue

            existing = KDMLabCriteria.query.filter_by(
                experiment_id=experiment.id, criteria_number=i
            ).first()

            if not existing:
                existing = KDMLabCriteria(
                    experiment_id=experiment.id, criteria_number=i)

            existing.question_text = q_text
            existing.option_a = request.form.get(f"optionA_{i}")
            existing.option_b = request.form.get(f"optionB_{i}")
            existing.option_c = request.form.get(f"optionC_{i}")
            existing.option_d = request.form.get(f"optionD_{i}")
            existing.marks_a = float(request.form.get(f"marksA_{i}") or 0)
            existing.marks_b = float(request.form.get(f"marksB_{i}") or 0)
            existing.marks_c = float(request.form.get(f"marksC_{i}") or 0)
            existing.marks_d = float(request.form.get(f"marksD_{i}") or 0)
            existing.max_marks = float(request.form.get(f"max_marks_{i}") or 4)
            # --- Handle Withhold (Admin Only) checkbox ---
            checkbox_name = f"withhold_{i}"
            if checkbox_name in request.form:
                existing.withheld = True
            else:
                existing.withheld = False
            print(f"🟢 Criterion {i}: withheld={existing.withheld}")

            db.session.add(existing)

        db.session.commit()

        flash(
            f"KDM LAB Quiz saved successfully for EX-{selected_experiment_number}.", "success")

        # ✅ Redirect to same page with ?experiment=<saved_experiment_number>
        return redirect(url_for("create_kdm_lab_quiz", experiment=selected_experiment_number))

    # ------------------- GET -------------------
    selected_experiment_number = request.args.get("experiment")
    if selected_experiment_number:
        experiment = KDMLabExperiment.query.filter_by(
            experiment_number=int(selected_experiment_number)
        ).first()
        if experiment:
            existing_criteria = KDMLabCriteria.query.filter_by(
                experiment_id=experiment.id
            ).order_by(KDMLabCriteria.criteria_number).all()

    return render_template(
        "create_kdm_lab_quiz.html",
        current_admin=current_admin,
        lab_settings=lab_settings,
        experiments=experiments,
        experiment=experiment,
        existing_criteria=existing_criteria
    )


@app.route("/attempt_kdm_lab_experiment/<int:experiment_id>", methods=["GET", "POST"])
def attempt_kdm_lab_experiment(experiment_id):
    if "loggedin" not in session:
        return redirect(url_for("login"))

    # ✅ Get current user
    current_user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not current_user:
        flash("User not found.", "danger")
        return redirect(url_for("login"))

    experiment = KDMLabExperiment.query.get_or_404(experiment_id)

    # If user is admin or super_admin, allow direct access (admin preview)
    if current_user.role in ["admin", "super_admin"]:
        pass  # allow access, no release check
    else:
        # Student: Check if experiment is released for them
        admin = get_admin_for_student(current_user)
        if not admin:
            flash("Admin not found for this student.", "danger")
            return redirect(url_for("kdm_lab_student_dashboard"))

        # Check bulk release
        bulk_released = KDMlabAdminRelease.query.filter(
            KDMlabAdminRelease.experiment_id == experiment_id,
            or_(
                KDMlabAdminRelease.admin_id == admin.id,
                KDMlabAdminRelease.admin_id.is_(None)  # global release
            ),
            KDMlabAdminRelease.is_released == True
        ).first()

        # Check individual release
        individual_released = KDMLabStudentExperimentRelease.query.filter_by(
            admin_id=admin.id,
            student_id=current_user.id,
            experiment_id=experiment_id,
            is_released=True
        ).first()

        # Deny access if neither bulk nor individually released
        if not bulk_released and not individual_released:
            flash("This experiment is not accessible to you.", "warning")
            return redirect(url_for("kdm_lab_student_dashboard"))

    # ✅ Fetch existing attempt record (for display only, not restriction)
    attempt_record = KDMLabAttempt.query.filter_by(
        student_id=current_user.id, experiment_id=experiment.id
    ).first()

    # ✅ Get visible criteria
    criteria = KDMLabCriteria.query.filter_by(
        experiment_id=experiment.id, withheld=False
    ).all()

    import random
    randomized_criteria = []
    for c in criteria:
        options = [
            {"key": "A", "text": c.option_a or "", "marks": c.marks_a or 0},
            {"key": "B", "text": c.option_b or "", "marks": c.marks_b or 0},
            {"key": "C", "text": c.option_c or "", "marks": c.marks_c or 0},
            {"key": "D", "text": c.option_d or "", "marks": c.marks_d or 0},
        ]
        options = [opt for opt in options if opt["text"].strip()]
        random.shuffle(options)
        randomized_criteria.append({"criteria": c, "options": options})

    # ✅ Handle POST (quiz submission)
    if request.method == "POST":
        total_marks = 0.0

        all_criteria = KDMLabCriteria.query.filter_by(
            experiment_id=experiment_id).all()
        total_possible = sum((crit.max_marks or 4) for crit in all_criteria)

        for item in randomized_criteria:
            crit = item["criteria"]
            selected = request.form.get(f"criteria_{crit.id}")
            if not selected:
                continue

            marks_map = {
                "A": crit.marks_a or 0,
                "B": crit.marks_b or 0,
                "C": crit.marks_c or 0,
                "D": crit.marks_d or 0,
            }

            obtained_points = marks_map.get(selected, 0)
            marks_earned = (crit.max_marks or 4) * (obtained_points / 4)
            total_marks += marks_earned

            # ✅ Determine current attempt number for admin
            current_attempt = (attempt_record.attempt_count +
                               1) if attempt_record else 1

            # ✅ Save new response for this attempt (admin gets new attempt each time)
            response = KDMLabResponse(
                user_id=current_user.id,
                experiment_id=experiment_id,
                criteria_id=crit.id,
                selected_option=selected,
                obtained_points=obtained_points,
                marks_earned=marks_earned,
                attempt_number=current_attempt
            )
            db.session.add(response)

        # ✅ Update attempt count (for info only)
        if not attempt_record:
            attempt_record = KDMLabAttempt(
                student_id=current_user.id,
                experiment_id=experiment.id,
                attempt_count=1
            )
            db.session.add(attempt_record)
        else:
            # Admin can attempt unlimited times
            attempt_record.attempt_count = (
                attempt_record.attempt_count or 0) + 1

        db.session.commit()

        # ✅ Include manual marks
        manual_marks = db.session.query(db.func.sum(KDMLabManualMarks.marks_given)) \
            .filter_by(student_id=current_user.id, experiment_id=experiment.id).scalar() or 0.0
        total_marks += manual_marks

        # ✅ Show result
        return render_template(
            "kdm_lab_result.html",
            experiment=experiment,
            total_marks=round(total_marks, 2),
            total_possible=round(total_possible, 2),
            current_user=current_user,
        )

    # ✅ Render quiz page with attempt info
    return render_template(
        "attempt_kdm_lab_experiment.html",
        current_user=current_user,
        experiment=experiment,
        randomized_criteria=randomized_criteria,
        attempt_record=attempt_record
    )


# --- ADMIN: Toggle Experiment Release -----------------------------------------

@app.route("/toggle_experiment_release/<int:experiment_id>", methods=["POST"])
def toggle_experiment_release(experiment_id):
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Admin or Super Admin access required.", "danger")
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin:
        flash("Admin not found.", "danger")
        return redirect(url_for("admin_dashboard"))

    experiment = KDMLabExperiment.query.get_or_404(experiment_id)

    # ✅ Case 1: Super Admin releases for all admins
    if current_admin.admin_level == "super_admin":
        all_admins = User.query.filter_by(admin_level="admin").all()
        for admin in all_admins:
            release_entry = KDMlabAdminRelease.query.filter_by(
                admin_id=admin.id, experiment_id=experiment.id
            ).first()

            if not release_entry:
                db.session.add(KDMlabAdminRelease(
                    admin_id=admin.id,
                    experiment_id=experiment.id,
                    is_released=True
                ))
            else:
                release_entry.is_released = not release_entry.is_released

        db.session.commit()
        flash(
            f"✅ {experiment.title} release status updated for all admins.", "success")

    # ✅ Case 2: Regular admin releases only for their own students
    else:
        release_entry = KDMlabAdminRelease.query.filter_by(
            admin_id=current_admin.id, experiment_id=experiment.id
        ).first()

        if not release_entry:
            release_entry = KDMlabAdminRelease(
                admin_id=current_admin.id,
                experiment_id=experiment.id,
                is_released=True
            )
            db.session.add(release_entry)
            flash(f"✅ {experiment.title} released for your students.", "success")
        else:
            release_entry.is_released = not release_entry.is_released
            status = "released" if release_entry.is_released else "locked"
            flash(f"{experiment.title} has been {status} for your students.", "info")

        db.session.commit()

    return redirect(url_for("kdm_lab_admin_dashboard"))


# --- ADMIN: View Submissions (All Students for an Experiment) -----------------

@app.route("/view_kdm_lab_submissions/<int:experiment_id>")
def view_kdm_lab_submissions(experiment_id):
    if not _require_login():
        return redirect(url_for("login"))

    current_admin = _current_admin_user()
    if not _is_admin(current_admin):
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # ✅ Experiment & criteria
    ex = KDMLabExperiment.query.get_or_404(experiment_id)
    criteria = (
        KDMLabCriteria.query.filter_by(experiment_id=experiment_id)
        .order_by(KDMLabCriteria.criteria_number)
        .all()
    )

    # ✅ Students in this admin’s scope
    students = _lab_students_for_admin(current_admin)

    # Sort students by roll number in ascending order
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    rows = []
    for s in students:
        # ✅ Get corresponding User record for proper foreign key references
        user = User.query.filter_by(rollnumber=s.rollnumber).first()
        if not user:
            # Skip this student if no corresponding user record exists
            continue

        row = {"student": s, "criteria_marks": [], "total": 0.0}
        total = 0.0

        for c in criteria:
            display_mark = 0.0  # What to show in the column
            weighted_mark = 0.0  # What to add to total

            if c.withheld:
                # ✅ Admin-given marks (use user.id for foreign key)
                mm = KDMLabManualMarks.query.filter_by(
                    student_id=user.id,
                    experiment_id=experiment_id,
                    criteria_id=c.id,
                ).first()
                obtained_marks = mm.marks_given if mm and mm.marks_given is not None else 0.0
                display_mark = obtained_marks
                # Apply weighted formula: (obtained_marks * max_marks) / 4
                weighted_mark = obtained_marks * (c.max_marks or 4.0) / 4.0

            else:
                # ✅ Student response - get BEST attempt (use user.id for foreign key)
                resp = _get_best_kdm_lab_response(
                    user.id, experiment_id, c.id)

                if resp and resp.obtained_points is not None:
                    # Show obtained points in column (1, 2, 3, or 4)
                    display_mark = resp.obtained_points
                    # Apply weighted formula: (obtained_points * max_marks) / 4
                    weighted_mark = resp.obtained_points * \
                        (c.max_marks or 4.0) / 4.0
                elif resp and resp.marks_earned is not None:
                    # fallback if marks_earned is already stored weighted
                    display_mark = resp.marks_earned
                    weighted_mark = resp.marks_earned
                else:
                    display_mark = 0.0
                    weighted_mark = 0.0

            row["criteria_marks"].append(round(display_mark, 2))
            total += weighted_mark

        # ✅ Final total (optionally rounded up like admin dashboard)
        row["total"] = math.ceil(total) if total % 1 > 0 else round(total, 2)
        rows.append(row)

    return render_template(
        "view_kdm_lab_submissions.html",
        current_admin=current_admin,
        experiment=ex,
        criteria=criteria,
        rows=rows,
    )


# --- ADMIN: Fill Manual Marks for Withheld Criteria ---------------------------

@app.route("/fill_kdm_lab_manual_marks/<int:experiment_id>", methods=["GET", "POST"])
def fill_kdm_lab_manual_marks(experiment_id):
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin:
        flash("Admin not found.", "danger")
        return redirect(url_for("admin_dashboard"))

    experiment = KDMLabExperiment.query.get_or_404(experiment_id)

    # ✅ Get withheld criteria only
    withheld_criteria = KDMLabCriteria.query.filter_by(
        experiment_id=experiment.id, withheld=True
    ).order_by(KDMLabCriteria.criteria_number).all()

    # ✅ Get students in admin’s scope
    if current_admin.admin_level == "super_admin":
        students = User.query.filter_by(role="student").all()
    else:
        student_rolls = [
            s.rollnumber for s in get_students_for_admin(current_admin)]
        students = User.query.filter(User.rollnumber.in_(student_rolls)).all()

    # Sort students by roll number in ascending order
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # ✅ Load existing manual marks (by student_id)
    existing_marks = {}
    manual_marks = KDMLabManualMarks.query.filter_by(
        experiment_id=experiment.id).all()
    for mark in manual_marks:
        existing_marks[(mark.student_id, mark.criteria_id)] = mark.marks_given

    # ✅ Handle form submission
    if request.method == "POST":
        bulk_criteria_id = request.form.get("bulk_criteria_id")
        bulk_mark = request.form.get("bulk_mark")

        # --- Bulk marks applied to all students ---
        if bulk_criteria_id and bulk_mark:
            for student in students:
                existing = KDMLabManualMarks.query.filter_by(
                    student_id=student.id,
                    experiment_id=experiment.id,
                    criteria_id=bulk_criteria_id
                ).first()
                if existing:
                    existing.marks_given = float(bulk_mark)
                else:
                    db.session.add(KDMLabManualMarks(
                        student_id=student.id,
                        experiment_id=experiment.id,
                        criteria_id=bulk_criteria_id,
                        marks_given=float(bulk_mark)
                    ))
            db.session.commit()
            flash("✅ Bulk marks applied successfully!", "success")
            return redirect(url_for("fill_kdm_lab_manual_marks", experiment_id=experiment.id))

        # --- Individual marks per student ---
        for student in students:
            for c in withheld_criteria:
                field_name = f"mark_{student.id}_{c.id}"
                mark_val = request.form.get(field_name)
                if mark_val is not None and mark_val.strip() != "":
                    existing = KDMLabManualMarks.query.filter_by(
                        student_id=student.id,
                        experiment_id=experiment.id,
                        criteria_id=c.id
                    ).first()
                    if existing:
                        existing.marks_given = float(mark_val)
                    else:
                        db.session.add(KDMLabManualMarks(
                            student_id=student.id,
                            experiment_id=experiment.id,
                            criteria_id=c.id,
                            marks_given=float(mark_val)
                        ))

        db.session.commit()
        flash("✅ Manual marks saved successfully!", "success")
        return redirect(url_for("fill_kdm_lab_manual_marks", experiment_id=experiment.id))

    # ✅ Render page
    return render_template(
        "fill_kdm_lab_manual_marks.html",
        experiment=experiment,
        withheld_criteria=withheld_criteria,
        students=students,
        existing_marks=existing_marks,
    )


# --- STUDENT: KDM LAB Dashboard (hidden until visible) ------------------------

@app.route("/kdm_lab_student_dashboard")
def kdm_lab_student_dashboard():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    # ✅ Get both student record and user record
    student = KDMLabStudent.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    if not student:
        flash("Student record not found in KDM LAB. Please contact admin.", "danger")
        return redirect(url_for("student_dashboard"))

    if not user:
        # Try to create user record from student record
        try:
            user = User(
                rollnumber=student.rollnumber,
                password=student.password,
                email=student.email,
                phonenumber=student.phonenumber,
                role="student"
            )
            db.session.add(user)
            db.session.commit()
            print(
                f"Created user record for KDM lab student: {student.rollnumber}")
        except Exception as e:
            print(f"Error creating user record: {e}")
            flash("Unable to create user record. Please contact admin.", "danger")
            return redirect(url_for("student_dashboard"))

    print(
        f"Debug: Dashboard - Student ID: {student.id}, User ID: {user.id}, Rollnumber: {session['rollnumber']}")

    # ✅ Get admin using helper
    admin = get_admin_for_student(student)
    if not admin:
        flash("Admin not found for this student.", "warning")
        experiments = []
    else:
        # ✅ Get experiments released by this admin OR globally released OR individually released to this student
        try:
            # First get bulk released experiments (existing logic)
            bulk_experiments = (
                db.session.query(KDMLabExperiment)
                .join(KDMlabAdminRelease, KDMlabAdminRelease.experiment_id == KDMLabExperiment.id)
                .filter(
                    KDMlabAdminRelease.is_released == True
                )
                .all()
            )

            # Get individually released experiments for this student by ANY admin
            individual_experiments = (
                db.session.query(KDMLabExperiment)
                .join(KDMLabStudentExperimentRelease, KDMLabStudentExperimentRelease.experiment_id == KDMLabExperiment.id)
                .filter(
                    KDMLabStudentExperimentRelease.student_id == user.id,
                    KDMLabStudentExperimentRelease.is_released == True
                )
                .all()
            )

            # Combine both lists and remove duplicates
            experiment_ids = set()
            experiments = []

            # Add bulk released experiments
            for exp in bulk_experiments:
                if exp.id not in experiment_ids:
                    experiments.append(exp)
                    experiment_ids.add(exp.id)

            # Add individually released experiments
            for exp in individual_experiments:
                if exp.id not in experiment_ids:
                    experiments.append(exp)
                    experiment_ids.add(exp.id)

            # Sort by experiment number
            experiments.sort(key=lambda x: x.experiment_number)

            print(
                f"Debug: Found {len(experiments)} experiments for student {student.rollnumber} (any admin individual release included)")

            # ✅ Add attempt records to each experiment
            for experiment in experiments:
                experiment.attempt_record = KDMLabAttempt.query.filter_by(
                    student_id=user.id,
                    experiment_id=experiment.id
                ).first()

        except Exception as e:
            print(f"Error loading experiments: {e}")
            flash("Error loading experiments.", "danger")
            experiments = []

    # Get sum of all criteria max marks for button text
    all_criteria_numbers = db.session.query(
        KDMLabCriteria.criteria_number).distinct().all()
    total_max_marks = 0
    for criteria_num_tuple in all_criteria_numbers:
        criteria_num = criteria_num_tuple[0]
        sample_criteria = KDMLabCriteria.query.filter_by(
            criteria_number=criteria_num).first()
        if sample_criteria:
            total_max_marks += sample_criteria.max_marks or 4.0

    return render_template(
        "kdm_lab_student_dashboard.html",
        student=student,
        experiments=experiments,
        total_max_marks=int(total_max_marks)
    )


# --- STUDENT: View Own Submission for an Experiment --------------------------

@app.route("/view_my_kdm_lab_submission/<int:experiment_id>")
def view_my_kdm_lab_submission(experiment_id):
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    # ✅ Get both student record and user record
    student = KDMLabStudent.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    if not student or not user:
        flash("Student record not found.", "danger")
        return redirect(url_for("student_dashboard"))

    # ✅ Get experiment and criteria
    experiment = KDMLabExperiment.query.get_or_404(experiment_id)
    criteria = (
        KDMLabCriteria.query.filter_by(experiment_id=experiment_id)
        .order_by(KDMLabCriteria.criteria_number)
        .all()
    )

    # ✅ Calculate student's marks using same logic as admin view
    row = {"student": student, "criteria_marks": [], "total": 0.0}
    total = 0.0

    for c in criteria:
        display_mark = 0.0  # What to show in the column
        weighted_mark = 0.0  # What to add to total

        if c.withheld:
            # ✅ Admin-given marks (use user.id for foreign key)
            mm = KDMLabManualMarks.query.filter_by(
                student_id=user.id,
                experiment_id=experiment_id,
                criteria_id=c.id,
            ).first()
            obtained_marks = mm.marks_given if mm and mm.marks_given is not None else 0.0
            display_mark = obtained_marks
            # Apply weighted formula: (obtained_marks * max_marks) / 4
            weighted_mark = obtained_marks * (c.max_marks or 4.0) / 4.0

        else:
            # ✅ Student response - get BEST attempt (use user.id for foreign key)
            resp = _get_best_kdm_lab_response(
                user.id, experiment_id, c.id)

            if resp and resp.obtained_points is not None:
                # Show obtained points in column (1, 2, 3, or 4)
                display_mark = resp.obtained_points
                # Apply weighted formula: (obtained_points * max_marks) / 4
                weighted_mark = resp.obtained_points * \
                    (c.max_marks or 4.0) / 4.0
            elif resp and resp.marks_earned is not None:
                # fallback if marks_earned is already stored weighted
                display_mark = resp.marks_earned
                weighted_mark = resp.marks_earned
            else:
                display_mark = 0.0
                weighted_mark = 0.0

        row["criteria_marks"].append(round(display_mark, 2))
        total += weighted_mark

    # ✅ Final total (optionally rounded up like admin dashboard)
    row["total"] = math.ceil(total) if total % 1 > 0 else round(total, 2)

    # ✅ Get attempt information
    attempt_record = KDMLabAttempt.query.filter_by(
        student_id=user.id, experiment_id=experiment_id
    ).first()

    return render_template(
        "view_my_kdm_lab_submission.html",
        student=student,
        experiment=experiment,
        criteria=criteria,
        row=row,
        attempt_record=attempt_record
    )


# --- STUDENT: View All Marks (Criteria-wise across all experiments) ---------

@app.route("/kdm_lab_student_all_marks")
def kdm_lab_student_all_marks():
    """Show student their marks criteria-wise across all experiments"""
    if "rollnumber" not in session:
        flash("Please login first.", "danger")
        return redirect(url_for("login"))

    # Get student and user records
    student = KDMLabStudent.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    if not student:
        flash("Student not found in KDM LAB records.", "danger")
        return redirect(url_for("student_dashboard"))

    if not user:
        flash("User record not found.", "danger")
        return redirect(url_for("student_dashboard"))

    # Get all experiments
    experiments = KDMLabExperiment.query.order_by(
        KDMLabExperiment.experiment_number).all()

    # Get all unique criteria numbers across all experiments
    all_criteria_numbers = db.session.query(KDMLabCriteria.criteria_number).distinct(
    ).order_by(KDMLabCriteria.criteria_number).all()
    criteria_numbers = [c[0] for c in all_criteria_numbers]

    # Build criteria max marks mapping: criteria_number -> max_marks
    criteria_max_marks = {}
    for criteria_num in criteria_numbers:
        # Find a criteria with this number to get max_marks (should be consistent across experiments)
        sample_criteria = KDMLabCriteria.query.filter_by(
            criteria_number=criteria_num).first()
        criteria_max_marks[criteria_num] = sample_criteria.max_marks if sample_criteria else 4.0

    # Build marks matrix: criteria_number -> experiment_id -> marks
    marks_matrix = {}
    raw_points_matrix = {}  # Store raw obtained points for total calculation
    experiment_totals = {}

    for criteria_num in criteria_numbers:
        marks_matrix[criteria_num] = {}
        raw_points_matrix[criteria_num] = {}

        for exp in experiments:
            # Find criteria with this number in this experiment
            criteria = KDMLabCriteria.query.filter_by(
                experiment_id=exp.id,
                criteria_number=criteria_num
            ).first()

            if criteria:
                if criteria.withheld:
                    # Manual marks for withheld criteria
                    mm = KDMLabManualMarks.query.filter_by(
                        student_id=user.id,
                        experiment_id=exp.id,
                        criteria_id=criteria.id,
                    ).first()
                    obtained_marks = mm.marks_given if mm and mm.marks_given is not None else 0.0
                    # Store raw points for both display and calculation
                    raw_points_matrix[criteria_num][exp.id] = obtained_marks
                    marks_matrix[criteria_num][exp.id] = round(
                        obtained_marks, 2)
                else:
                    # Student response - get BEST attempt
                    resp = _get_best_kdm_lab_response(
                        user.id, exp.id, criteria.id)
                    if resp and resp.obtained_points is not None:
                        # Store raw points for both display and calculation
                        raw_points_matrix[criteria_num][exp.id] = resp.obtained_points
                        marks_matrix[criteria_num][exp.id] = round(
                            resp.obtained_points, 2)
                    else:
                        raw_points_matrix[criteria_num][exp.id] = 0.0
                        marks_matrix[criteria_num][exp.id] = 0.0
            else:
                # No criteria with this number in this experiment
                raw_points_matrix[criteria_num][exp.id] = 0.0
                marks_matrix[criteria_num][exp.id] = 0.0

    # Calculate experiment totals using raw points and proper weighting
    for exp in experiments:
        total = 0.0
        for criteria_num in criteria_numbers:
            # Find the criteria for this experiment to get max_marks
            criteria = KDMLabCriteria.query.filter_by(
                experiment_id=exp.id,
                criteria_number=criteria_num
            ).first()

            if criteria:
                raw_points = raw_points_matrix[criteria_num].get(exp.id, 0.0)
                # Apply weighting formula: (raw_points * max_marks) / 4
                weighted_contribution = raw_points * \
                    (criteria.max_marks or 4.0) / 4.0
                total += weighted_contribution
        experiment_totals[exp.id] = round(total, 2)

    # Calculate final rubrics mark: sum of all experiment totals / number of experiments
    total_sum = sum(experiment_totals.values())
    final_rubrics_mark = round(
        total_sum / len(experiments), 2) if len(experiments) > 0 else 0.0

    # Calculate total possible marks (sum of all criteria max marks)
    total_possible_marks = sum(criteria_max_marks.values())

    return render_template(
        "kdm_lab_student_all_marks.html",
        student=student,
        experiments=experiments,
        criteria_numbers=criteria_numbers,
        marks_matrix=marks_matrix,
        experiment_totals=experiment_totals,
        criteria_max_marks=criteria_max_marks,
        final_rubrics_mark=final_rubrics_mark,
        total_possible_marks=total_possible_marks
    )


@app.route("/kdm_lab_student_all_marks_csv")
def kdm_lab_student_all_marks_csv():
    """Download CSV of student marks criteria-wise across all experiments"""
    if "rollnumber" not in session:
        flash("Please login first.", "danger")
        return redirect(url_for("login"))

    # Get student and user records
    student = KDMLabStudent.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    if not student or not user:
        flash("Student/User record not found.", "danger")
        return redirect(url_for("student_dashboard"))

    # Get all experiments
    experiments = KDMLabExperiment.query.order_by(
        KDMLabExperiment.experiment_number).all()

    # Get all unique criteria numbers
    all_criteria_numbers = db.session.query(KDMLabCriteria.criteria_number).distinct(
    ).order_by(KDMLabCriteria.criteria_number).all()
    criteria_numbers = [c[0] for c in all_criteria_numbers]

    # Build criteria max marks mapping
    criteria_max_marks = {}
    for criteria_num in criteria_numbers:
        sample_criteria = KDMLabCriteria.query.filter_by(
            criteria_number=criteria_num).first()
        criteria_max_marks[criteria_num] = sample_criteria.max_marks if sample_criteria else 4.0

    # Build marks matrix (same logic as above)
    marks_matrix = {}
    raw_points_matrix = {}
    experiment_totals = {}

    for criteria_num in criteria_numbers:
        marks_matrix[criteria_num] = {}
        raw_points_matrix[criteria_num] = {}

        for exp in experiments:
            criteria = KDMLabCriteria.query.filter_by(
                experiment_id=exp.id,
                criteria_number=criteria_num
            ).first()

            if criteria:
                if criteria.withheld:
                    mm = KDMLabManualMarks.query.filter_by(
                        student_id=user.id,
                        experiment_id=exp.id,
                        criteria_id=criteria.id,
                    ).first()
                    obtained_marks = mm.marks_given if mm and mm.marks_given is not None else 0.0
                    raw_points_matrix[criteria_num][exp.id] = obtained_marks
                    marks_matrix[criteria_num][exp.id] = round(
                        obtained_marks, 2)
                else:
                    resp = _get_best_kdm_lab_response(
                        user.id, exp.id, criteria.id)
                    if resp and resp.obtained_points is not None:
                        raw_points_matrix[criteria_num][exp.id] = resp.obtained_points
                        marks_matrix[criteria_num][exp.id] = round(
                            resp.obtained_points, 2)
                    else:
                        raw_points_matrix[criteria_num][exp.id] = 0.0
                        marks_matrix[criteria_num][exp.id] = 0.0
            else:
                raw_points_matrix[criteria_num][exp.id] = 0.0
                marks_matrix[criteria_num][exp.id] = 0.0

    # Calculate experiment totals using raw points and proper weighting
    for exp in experiments:
        total = 0.0
        for criteria_num in criteria_numbers:
            # Find the criteria for this experiment to get max_marks
            criteria = KDMLabCriteria.query.filter_by(
                experiment_id=exp.id,
                criteria_number=criteria_num
            ).first()

            if criteria:
                raw_points = raw_points_matrix[criteria_num].get(exp.id, 0.0)
                # Apply weighting formula: (raw_points * max_marks) / 4
                weighted_contribution = raw_points * \
                    (criteria.max_marks or 4.0) / 4.0
                total += weighted_contribution
        experiment_totals[exp.id] = round(total, 2)

    # Calculate final rubrics mark: sum of all experiment totals / number of experiments
    total_sum = sum(experiment_totals.values())
    final_rubrics_mark = round(
        total_sum / len(experiments), 2) if len(experiments) > 0 else 0.0

    # Generate CSV content
    import io
    import csv
    from flask import make_response

    output = io.StringIO()
    writer = csv.writer(output)

    # Header row
    header = ['Criteria'] + \
        [f'Experiment {exp.experiment_number}' for exp in experiments]
    writer.writerow(header)

    # Criteria rows
    for criteria_num in criteria_numbers:
        max_marks = criteria_max_marks.get(criteria_num, 4.0)
        row = [f'Criterion {criteria_num} (Max: {max_marks})']
        for exp in experiments:
            row.append(marks_matrix[criteria_num].get(exp.id, 0.0))
        writer.writerow(row)

    # Total row
    total_row = ['TOTAL']
    for exp in experiments:
        total_row.append(experiment_totals[exp.id])
    writer.writerow(total_row)

    # Final rubrics mark row
    final_row = ['FINAL RUBRICS MARK (Average)', final_rubrics_mark]
    # Add empty cells for other experiments to align properly
    for _ in range(len(experiments) - 1):
        final_row.append('')
    writer.writerow(final_row)

    # Create response
    output.seek(0)
    response = make_response(output.getvalue())
    response.headers[
        "Content-Disposition"] = f"attachment; filename=kdm_lab_marks_{student.rollnumber}.csv"
    response.headers["Content-type"] = "text/csv"

    return response


# --- STUDENT: View Individual Consolidated Marks -------------------------

@app.route("/kdm_lab_student_consolidated_marks")
def kdm_lab_student_consolidated_marks():
    """Show student their individual consolidated marks (rubrics + internal + attendance)"""
    if "rollnumber" not in session:
        flash("Please login first.", "danger")
        return redirect(url_for("login"))

    # Get student and user records
    student = KDMLabStudent.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    if not student or not user:
        flash("Student not found in KDM LAB records.", "danger")
        return redirect(url_for("student_dashboard"))

    # Calculate KDM Lab Rubrics Marks (same logic as admin consolidated marks)
    rubrics_marks = 0.0
    try:
        # Get all experiments for this student
        experiments = KDMLabExperiment.query.order_by(
            KDMLabExperiment.experiment_number).all()

        # Get all unique criteria numbers
        all_criteria_numbers = db.session.query(KDMLabCriteria.criteria_number).distinct(
        ).order_by(KDMLabCriteria.criteria_number).all()
        criteria_numbers = [c[0] for c in all_criteria_numbers]

        # Calculate total weighted marks across all experiments
        total_weighted_marks = 0.0
        for exp in experiments:
            for criteria_num in criteria_numbers:
                # Find criteria for this experiment
                criteria = KDMLabCriteria.query.filter_by(
                    experiment_id=exp.id,
                    criteria_number=criteria_num
                ).first()

                if criteria:
                    if criteria.withheld:
                        # Manual marks
                        mm = KDMLabManualMarks.query.filter_by(
                            student_id=user.id,
                            experiment_id=exp.id,
                            criteria_id=criteria.id,
                        ).first()
                        obtained_points = mm.marks_given if mm and mm.marks_given is not None else 0.0
                    else:
                        # Best attempt from quiz responses
                        resp = _get_best_kdm_lab_response(
                            user.id, exp.id, criteria.id)
                        obtained_points = resp.obtained_points if resp and resp.obtained_points is not None else 0.0

                    # Apply weighting formula: (obtained_points * max_marks) / 4
                    weighted_mark = obtained_points * \
                        (criteria.max_marks or 4.0) / 4.0
                    total_weighted_marks += weighted_mark

        # Calculate final rubrics mark: sum of all weighted marks / number of experiments
        if len(experiments) > 0:
            avg_score = total_weighted_marks / len(experiments)
            # Round up (e.g. 12.1 → 13) - same logic as KDM lab admin dashboard
            rubrics_marks = int(avg_score) if avg_score == int(
                avg_score) else int(avg_score) + 1
    except Exception as e:
        print(
            f"Error calculating rubrics marks for student {student.rollnumber}: {e}")
        rubrics_marks = 0.0

    # Try to get script and attendance marks from uploaded admin data
    script_marks = attendance_marks = "N/A"
    try:
        # Check all admin files for this student's marks
        import os
        import pandas as pd

        upload_dir = os.path.join(os.getcwd(), "tmp_uploads")
        if os.path.exists(upload_dir):
            # Look for any admin's uploaded KDM marks file
            for filename in os.listdir(upload_dir):
                if filename.startswith("kdm_internal_marks_") and filename.endswith(".csv"):
                    try:
                        file_path = os.path.join(upload_dir, filename)
                        df = pd.read_csv(file_path)

                        # Find this student's row
                        for _, row in df.iterrows():
                            if str(row.get("Roll Number", "")).strip() == str(student.rollnumber):
                                script_marks = row.get(
                                    "Internal Script Marks (15)", "N/A")
                                attendance_marks = row.get(
                                    "Attendance Marks (5)", "N/A")
                                break

                        if script_marks != "N/A":
                            break  # Found the marks, stop searching
                    except Exception:
                        continue
    except Exception as e:
        print(
            f"Error reading uploaded marks for student {student.rollnumber}: {e}")

    # Calculate Viva Marks (10) from best quiz attempt - using same logic as admin consolidated marks
    viva_marks = 0.0
    if user:
        # Use the same logic as admin consolidated marks for consistency
        responsible_admin = find_admin_for_student(student.rollnumber)
        if responsible_admin:
            viva_marks = _get_best_kdm_quiz_score(
                user.id, responsible_admin.id)
            print(
                f"🔍 Student View: Student {student.rollnumber} -> Responsible Admin {responsible_admin.rollnumber} -> Viva: {viva_marks}")

            # Add same debugging as admin consolidated marks for 0 viva cases
            if viva_marks == 0:
                all_attempts = KDMLabQuizAttempt.query.filter_by(
                    student_id=user.id, is_completed=True).all()
                print(
                    f"   📊 Student View DEBUG: All quiz attempts for {student.rollnumber}: {len(all_attempts)} attempts")
                for att in all_attempts:
                    att_admin = User.query.get(att.admin_id)
                    print(
                        f"      - Admin: {att_admin.rollnumber if att_admin else 'Unknown'}, Score: {att.score}")
        else:
            print(
                f"⚠️ Student View: No responsible admin found for student {student.rollnumber}")
    else:
        print(
            f"⚠️ Student View: User record not found for student {student.rollnumber}")

    # Calculate Total Internal Marks (25) = Script (15) + Viva (10)
    total_internal_marks = "N/A"
    try:
        if (str(script_marks).replace('.', '', 1).replace('-', '', 1).isdigit() and
                isinstance(viva_marks, (int, float))):
            total_internal_marks = round(float(script_marks) + viva_marks, 2)
    except Exception:
        total_internal_marks = "N/A"

    # Calculate Final Internal Marks (50)
    final_internal = "N/A"
    try:
        if (isinstance(rubrics_marks, (int, float)) and
            isinstance(total_internal_marks, (int, float)) and
                str(attendance_marks).replace('.', '', 1).replace('-', '', 1).isdigit()):
            total_marks = rubrics_marks + \
                total_internal_marks + float(attendance_marks)
            # Round up (e.g. 42.1 → 43) - consistent with other final internal marks
            final_internal = math.ceil(total_marks)
    except Exception:
        final_internal = "N/A"

    # Prepare student data
    student_data = {
        "Roll Number": student.rollnumber,
        "Rubrics Marks (20)": rubrics_marks,
        "Internal Script Marks (15)": script_marks,
        "Internal Viva Marks (10)": viva_marks,
        "Total Internal Marks (25)": total_internal_marks,
        "Attendance Marks (5)": attendance_marks,
        "Final Internal Marks (50)": final_internal
    }

    return render_template("kdm_lab_student_consolidated_marks.html",
                           student_data=student_data,
                           student=student)


# --- STUDENT: Take Quiz for an Experiment ------------------------------------


@app.route("/take_kdm_lab_quiz/<int:experiment_id>", methods=["GET", "POST"])
def take_kdm_lab_quiz(experiment_id):
    flash(f"🟢 Entered route for experiment_id={experiment_id}", "info")

    if "rollnumber" not in session:
        flash("No rollnumber in session - please login again", "danger")
        return redirect(url_for("login"))

    # ✅ Get both student record and user record
    student = KDMLabStudent.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    if not student:
        flash("Student not found in KDM LAB records. Please contact admin.", "danger")
        return redirect(url_for("student_dashboard"))

    if not user:
        # Try to create user record from student record
        try:
            user = User(
                rollnumber=student.rollnumber,
                password=student.password,
                email=student.email,
                phonenumber=student.phonenumber,
                role="student"
            )
            db.session.add(user)
            db.session.commit()
            flash(
                f"Created user record for KDM lab student: {student.rollnumber}", "success")
        except Exception as e:
            flash(f"Error creating user record: {str(e)}", "danger")
            return redirect(url_for("student_dashboard"))

    experiment = KDMLabExperiment.query.get_or_404(experiment_id)

    # ✅ Check if experiment is released for this student's admin
    try:
        admin = get_admin_for_student(student)
    except Exception as e:
        admin = None
        flash(f"Error finding admin: {str(e)}", "danger")

    if not admin:
        flash("⚠️ No admin assigned for your batch — contact faculty.", "warning")
        return redirect(url_for("kdm_lab_student_dashboard"))

    # ✅ Verify if experiment is released for that admin (bulk/global) OR individually for this student
    try:
        bulk_released = (
            db.session.query(KDMlabAdminRelease)
            .filter(
                or_(
                    KDMlabAdminRelease.admin_id == admin.id,
                    KDMlabAdminRelease.admin_id.is_(None)  # global release
                ),
                KDMlabAdminRelease.experiment_id == experiment_id,
                KDMlabAdminRelease.is_released == True
            )
            .first()
        )
        individual_released = KDMLabStudentExperimentRelease.query.filter_by(
            admin_id=admin.id,
            student_id=user.id,
            experiment_id=experiment_id,
            is_released=True
        ).first()
    except Exception as e:
        flash(f"Error checking experiment availability: {str(e)}", "danger")
        return redirect(url_for("kdm_lab_student_dashboard"))

    if not bulk_released and not individual_released:
        flash("⚠️ This experiment is not yet released for your batch or individually for you.", "warning")
        return redirect(url_for("kdm_lab_student_dashboard"))

    # ✅ Check if student already attempted twice
    try:
        attempt_record = KDMLabAttempt.query.filter_by(
            student_id=user.id, experiment_id=experiment.id
        ).first()

    except Exception as e:
        flash(f"Error checking attempt history: {str(e)}", "danger")
        return redirect(url_for("kdm_lab_student_dashboard"))

    if attempt_record and attempt_record.attempt_count >= 2:
        flash("⚠️ You have already completed 2 attempts for this experiment.", "warning")
        return redirect(url_for("kdm_lab_student_dashboard"))

    # ✅ Get non-withheld criteria
    criteria = KDMLabCriteria.query.filter_by(
        experiment_id=experiment.id, withheld=False).all()

    # ✅ Handle quiz submission
    if request.method == "POST":
        total_marks = 0.0

        for c in criteria:
            selected_option = request.form.get(f"criteria_{c.id}")
            if not selected_option:
                continue

            marks_map = {
                "A": c.marks_a or 0,
                "B": c.marks_b or 0,
                "C": c.marks_c or 0,
                "D": c.marks_d or 0,
            }

            obtained_points = marks_map.get(selected_option, 0)
            marks_earned = (c.max_marks or 4) * (obtained_points / 4)
            total_marks += marks_earned

            # ✅ Determine current attempt number
            current_attempt = (attempt_record.attempt_count +
                               1) if attempt_record else 1

            # ✅ Save new response for this attempt (don't update existing)
            db.session.add(KDMLabResponse(
                user_id=user.id,
                experiment_id=experiment.id,
                criteria_id=c.id,
                selected_option=selected_option,
                obtained_points=obtained_points,
                marks_earned=marks_earned,
                attempt_number=current_attempt
            ))

        # ✅ Step 1: Calculate current total with same logic as view submissions
        current_total = 0.0
        all_criteria = KDMLabCriteria.query.filter_by(
            experiment_id=experiment.id).all()

        for c in all_criteria:
            if c.withheld:
                # Manual marks for withheld criteria
                mm = KDMLabManualMarks.query.filter_by(
                    student_id=user.id,
                    experiment_id=experiment.id,
                    criteria_id=c.id,
                ).first()
                obtained_marks = mm.marks_given if mm and mm.marks_given is not None else 0.0
                # Apply weighted formula: (obtained_marks * max_marks) / 4
                weighted_mark = obtained_marks * (c.max_marks or 4.0) / 4.0
                current_total += weighted_mark
            else:
                # Student response for regular criteria - get BEST attempt
                resp = _get_best_kdm_lab_response(
                    user.id, experiment.id, c.id)

                if resp and resp.obtained_points is not None:
                    # Apply weighted formula: (obtained_points * max_marks) / 4
                    weighted_mark = resp.obtained_points * \
                        (c.max_marks or 4.0) / 4.0
                    current_total += weighted_mark

        # ✅ Step 2: Use current total as best marks (since this is the most recent attempt)
        best_marks = current_total

        # ✅ Step 3: Update attempt count (increment only on quiz submission)
        if not attempt_record:
            attempt_record = KDMLabAttempt(
                student_id=user.id,
                experiment_id=experiment.id,
                attempt_count=1
            )
            db.session.add(attempt_record)
        else:
            attempt_record.attempt_count += 1

        db.session.commit()

        # ✅ Step 4: Compute average of all experiments (best of each)
        all_experiments = KDMLabExperiment.query.all()
        experiment_scores = []

        for ex in all_experiments:
            best_for_ex = db.session.query(db.func.sum(KDMLabResponse.marks_earned)) \
                .filter_by(user_id=user.id, experiment_id=ex.id).scalar() or 0.0
            experiment_scores.append(best_for_ex)

        if experiment_scores:
            avg_score = sum(experiment_scores) / len(experiment_scores)
            # Round up (e.g. 35.1 → 36)
            student.total_lab_score = int(
                avg_score) if avg_score.is_integer() else int(avg_score) + 1
        else:
            student.total_lab_score = 0

        db.session.commit()

        # ✅ Step 5: Compute total possible marks using weighted formula
        total_possible = 0.0
        for crit in all_criteria:
            # Maximum possible score for each criterion using weighted formula
            # If student gets max points (4) or max manual marks, weighted = (max_points * max_marks) / 4
            if crit.withheld:
                # For manual criteria, assume max possible is the max_marks itself
                max_weighted = crit.max_marks or 4.0
            else:
                # For quiz criteria, max is when student gets 4 points: (4 * max_marks) / 4 = max_marks
                max_weighted = crit.max_marks or 4.0
            total_possible += max_weighted
        session["rollnumber"] = student.rollnumber
        user_role = session.get("role", "student")  # default to student if not set
        # ✅ Step 6: Show result page
        return render_template(
            "kdm_lab_result.html",
            experiment=experiment,
            total_marks=round(best_marks, 2),
            total_possible=round(total_possible, 2),
            current_user=student,
            user_role=user_role,   # ✅ Pass role info to template
        )

    # ✅ For GET — show quiz form
    shuffled_criteria = []
    for c in criteria:
        options = [
            {"label": "A", "text": c.option_a, "marks": c.marks_a},
            {"label": "B", "text": c.option_b, "marks": c.marks_b},
            {"label": "C", "text": c.option_c, "marks": c.marks_c},
            {"label": "D", "text": c.option_d, "marks": c.marks_d},
        ]
        options = [opt for opt in options if opt["text"]]  # remove blanks
        random.shuffle(options)
        shuffled_criteria.append({
            "id": c.id,
            "criteria_number": c.criteria_number,
            "question_text": c.question_text,
            "options": options,
        })

    return render_template(
        "take_kdm_lab_quiz.html",
        experiment=experiment,
        criteria=shuffled_criteria,
        attempt_record=attempt_record,
    )

# --- ADMIN: (Optional) Reset All KDM LAB Users/Marks (super admin only) -------

# --- KDM LAB: Bulk Release/Lock All Experiments ------------------------------


@app.route("/kdm_lab_release_all_experiments", methods=["POST"])
def kdm_lab_release_all_experiments():
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin:
        flash("Admin not found.", "danger")
        return redirect(url_for("login"))

    try:
        # Get all experiments
        experiments = KDMLabExperiment.query.all()

        # If super admin, release for ALL admins; if regular admin, only for themselves
        if current_admin.admin_level == 'super_admin':
            # Get all admins
            all_admins = User.query.filter(
                User.role.in_(['admin', 'super_admin'])).all()
            admin_ids = [admin.id for admin in all_admins]
        else:
            admin_ids = [current_admin.id]

        for experiment in experiments:
            for admin_id in admin_ids:
                # Check if release record exists for this admin and experiment
                existing_release = KDMlabAdminRelease.query.filter_by(
                    admin_id=admin_id,
                    experiment_id=experiment.id
                ).first()

                if existing_release:
                    # Update existing record to released
                    existing_release.is_released = True
                else:
                    # Create new release record
                    new_release = KDMlabAdminRelease(
                        admin_id=admin_id,
                        experiment_id=experiment.id,
                        is_released=True
                    )
                    db.session.add(new_release)

        db.session.commit()
        if current_admin.admin_level == 'super_admin':
            flash(
                f"🚀 Successfully released all {len(experiments)} experiments for ALL admins!", "success")
        else:
            flash(
                f"🚀 Successfully released all {len(experiments)} experiments!", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"❌ Error releasing experiments: {str(e)}", "danger")

    return redirect(url_for("kdm_lab_admin_dashboard"))


@app.route("/kdm_lab_lock_all_experiments", methods=["POST"])
def kdm_lab_lock_all_experiments():
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin:
        flash("Admin not found.", "danger")
        return redirect(url_for("login"))

    try:
        # If super admin, lock for ALL admins; if regular admin, only for themselves
        if current_admin.admin_level == 'super_admin':
            # Update all existing release records for ALL admins to locked
            KDMlabAdminRelease.query.update(
                {KDMlabAdminRelease.is_released: False})
            flash("🔒 Successfully locked all experiments for ALL admins!", "success")
        else:
            # Update all existing release records for this admin to locked
            KDMlabAdminRelease.query.filter_by(admin_id=current_admin.id).update(
                {KDMlabAdminRelease.is_released: False}
            )
            flash("🔒 Successfully locked all experiments!", "success")

        db.session.commit()

    except Exception as e:
        db.session.rollback()
        flash(f"❌ Error locking experiments: {str(e)}", "danger")

    return redirect(url_for("kdm_lab_admin_dashboard"))


@app.route("/reset_kdm_lab_experiment/<int:experiment_id>", methods=["POST"])
def reset_kdm_lab_experiment(experiment_id):
    """Reset all attempts and marks for all students of this admin for a specific experiment"""
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin:
        flash("Admin not found.", "danger")
        return redirect(url_for("login"))

    experiment = KDMLabExperiment.query.get_or_404(experiment_id)

    try:
        # Get all students in this admin's scope
        students = _lab_students_for_admin(current_admin)
        student_count = 0

        for student in students:
            # Get corresponding User record for proper foreign key references
            user = User.query.filter_by(rollnumber=student.rollnumber).first()
            if not user:
                continue

            # Delete all responses for this experiment and student
            KDMLabResponse.query.filter_by(
                user_id=user.id,
                experiment_id=experiment_id
            ).delete()

            # Delete all manual marks for this experiment and student
            KDMLabManualMarks.query.filter_by(
                student_id=user.id,
                experiment_id=experiment_id
            ).delete()

            # Reset attempt count for this experiment and student
            KDMLabAttempt.query.filter_by(
                student_id=user.id,
                experiment_id=experiment_id
            ).delete()

            student_count += 1

        db.session.commit()
        flash(
            f"🔄 Successfully reset Experiment #{experiment.experiment_number} for {student_count} students!", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"❌ Error resetting experiment: {str(e)}", "danger")

    return redirect(url_for("kdm_lab_admin_dashboard"))


@app.route("/fill_kdm_lab_manual_marks_all_experiments", methods=["GET", "POST"])
def fill_kdm_lab_manual_marks_all_experiments():
    """Fill manual marks for all experiments at once"""
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin:
        flash("Admin not found.", "danger")
        return redirect(url_for("admin_dashboard"))

    # Get all experiments
    experiments = KDMLabExperiment.query.order_by(
        KDMLabExperiment.experiment_number).all()

    # Get all withheld criteria across all experiments
    all_withheld_criteria = []
    for exp in experiments:
        criteria = KDMLabCriteria.query.filter_by(
            experiment_id=exp.id, withheld=True
        ).order_by(KDMLabCriteria.criteria_number).all()
        for c in criteria:
            all_withheld_criteria.append({
                'experiment': exp,
                'criteria': c,
                'id': c.id,
                'display_name': f"Exp {exp.experiment_number} - Criterion {c.criteria_number}"
            })

    # Get students in admin's scope - Super admin sees all students
    if current_admin.admin_level == 'super_admin':
        # Super admin sees all students directly
        user_students = User.query.filter(User.role == 'student').all()
    else:
        # Regular admin sees only assigned students
        students = _lab_students_for_admin(current_admin)
        # Convert to User objects for consistency
        user_students = []
        for student in students:
            user = User.query.filter_by(rollnumber=student.rollnumber).first()
            if user:
                user_students.append(user)

    # Sort students by roll number in ascending order
    user_students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # Load existing manual marks
    existing_marks = {}
    manual_marks = KDMLabManualMarks.query.all()
    for mark in manual_marks:
        existing_marks[(mark.student_id, mark.criteria_id)] = mark.marks_given

    # Handle form submission
    if request.method == "POST":
        bulk_criteria_id = request.form.get("bulk_criteria_id")
        bulk_mark = request.form.get("bulk_mark")

        # Bulk marks applied to all students for selected criterion
        if bulk_criteria_id and bulk_mark:
            criteria_obj = None
            experiment_obj = None

            # Find the criteria and experiment
            for item in all_withheld_criteria:
                if str(item['id']) == bulk_criteria_id:
                    criteria_obj = item['criteria']
                    experiment_obj = item['experiment']
                    break

            if criteria_obj and experiment_obj:
                for student in user_students:
                    existing = KDMLabManualMarks.query.filter_by(
                        student_id=student.id,
                        experiment_id=experiment_obj.id,
                        criteria_id=criteria_obj.id
                    ).first()
                    if existing:
                        existing.marks_given = float(bulk_mark)
                    else:
                        db.session.add(KDMLabManualMarks(
                            student_id=student.id,
                            experiment_id=experiment_obj.id,
                            criteria_id=criteria_obj.id,
                            marks_given=float(bulk_mark)
                        ))
                db.session.commit()
                flash(
                    f"✅ Bulk marks applied to all students for {criteria_obj.criteria_number} in Experiment {experiment_obj.experiment_number}!", "success")
                return redirect(url_for("fill_kdm_lab_manual_marks_all_experiments"))

        # Individual marks per student per criterion
        for student in user_students:
            for item in all_withheld_criteria:
                field_name = f"mark_{student.id}_{item['id']}"
                mark_val = request.form.get(field_name)
                if mark_val is not None and mark_val.strip() != "":
                    existing = KDMLabManualMarks.query.filter_by(
                        student_id=student.id,
                        experiment_id=item['experiment'].id,
                        criteria_id=item['id']
                    ).first()
                    if existing:
                        existing.marks_given = float(mark_val)
                    else:
                        db.session.add(KDMLabManualMarks(
                            student_id=student.id,
                            experiment_id=item['experiment'].id,
                            criteria_id=item['id'],
                            marks_given=float(mark_val)
                        ))

        db.session.commit()
        flash("✅ Manual marks updated successfully for all experiments!", "success")
        return redirect(url_for("kdm_lab_admin_dashboard"))

    return render_template(
        "fill_kdm_lab_manual_marks_all_experiments.html",
        experiments=experiments,
        all_withheld_criteria=all_withheld_criteria,
        students=user_students,
        existing_marks=existing_marks,
        current_admin=current_admin
    )


@app.route("/kdm_lab_consolidated_marks")
def kdm_lab_consolidated_marks():
    import os
    import pandas as pd

    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    current_admin = user

    # --- Fetch students based on admin level ---
    if current_admin.admin_level == 'super_admin':
        students = User.query.filter(User.role == 'student').all()
    else:
        students = get_students_for_admin(current_admin)
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # --- Load uploaded CSV (if exists) ---
    uploaded_data = []
    try:
        if current_admin.admin_level == 'super_admin':
            # Super admin: Load ALL CSV files from all admins
            import glob
            upload_dir = os.path.join(os.getcwd(), "tmp_uploads")
            csv_files = glob.glob(os.path.join(
                upload_dir, "kdm_internal_marks_*.csv"))

            for csv_file in csv_files:
                try:
                    df = pd.read_csv(csv_file)
                    file_data = df.to_dict(orient="records")
                    uploaded_data.extend(file_data)
                except Exception as e:
                    print(f"⚠️ Error reading {csv_file}: {e}")
        else:
            # Regular admin: Load only their own CSV file
            admin_id = current_admin.id
            upload_path = os.path.join(
                os.getcwd(), "tmp_uploads", f"kdm_internal_marks_{admin_id}.csv")
            if os.path.exists(upload_path):
                df = pd.read_csv(upload_path)
                uploaded_data = df.to_dict(orient="records")
    except Exception as e:
        print(f"⚠️ Unable to read uploaded KDM internal marks file: {e}")

    # --- Pre-load all data in bulk for performance ---
    # Get all experiments once
    experiments = KDMLabExperiment.query.order_by(
        KDMLabExperiment.experiment_number).all()

    # Get all criteria once
    all_criteria = KDMLabCriteria.query.all()
    criteria_numbers = sorted(
        list(set([c.criteria_number for c in all_criteria])))

    # Create criteria lookup: {(experiment_id, criteria_number): criteria_object}
    criteria_lookup = {}
    for c in all_criteria:
        criteria_lookup[(c.experiment_id, c.criteria_number)] = c

    # Bulk load all manual marks for all students
    student_ids = [s.id for s in students]
    all_manual_marks = KDMLabManualMarks.query.filter(
        KDMLabManualMarks.student_id.in_(student_ids)).all()
    # {(student_id, experiment_id, criteria_id): marks_given}
    manual_marks_lookup = {}
    for mm in all_manual_marks:
        manual_marks_lookup[(mm.student_id, mm.experiment_id,
                             mm.criteria_id)] = mm.marks_given

    # Bulk load all responses for all students
    all_responses = KDMLabResponse.query.filter(
        KDMLabResponse.user_id.in_(student_ids)).all()

    # Group responses by (student_id, experiment_id, criteria_id) and keep best
    # {(student_id, experiment_id, criteria_id): best_response}
    best_responses = {}
    for resp in all_responses:
        key = (resp.user_id, resp.experiment_id, resp.criteria_id)
        if key not in best_responses or (resp.obtained_points or 0) > (best_responses[key].obtained_points or 0):
            best_responses[key] = resp

    # --- Compute all marks ---
    data = []
    for s in students:
        roll = s.rollnumber

        # Calculate KDM Lab Rubrics Marks (out of 20)
        rubrics_marks = 0.0
        try:
            # Calculate total weighted marks across all experiments
            total_weighted_marks = 0.0
            for exp in experiments:
                for criteria_num in criteria_numbers:
                    # Look up criteria from pre-loaded data
                    criteria = criteria_lookup.get((exp.id, criteria_num))

                    if criteria:
                        if criteria.withheld:
                            # Look up manual marks from pre-loaded data
                            obtained_points = manual_marks_lookup.get(
                                (s.id, exp.id, criteria.id), 0.0)
                        else:
                            # Look up best response from pre-loaded data
                            resp = best_responses.get(
                                (s.id, exp.id, criteria.id))
                            obtained_points = resp.obtained_points if resp and resp.obtained_points is not None else 0.0

                        # Apply weighting formula: (obtained_points * max_marks) / 4
                        weighted_mark = obtained_points * \
                            (criteria.max_marks or 4.0) / 4.0
                        total_weighted_marks += weighted_mark

            # Calculate final rubrics mark: sum of all weighted marks / number of experiments
            if len(experiments) > 0:
                avg_score = total_weighted_marks / len(experiments)
                # Round up (e.g. 12.1 → 13) - same logic as KDM lab admin dashboard
                rubrics_marks = int(avg_score) if avg_score == int(
                    avg_score) else int(avg_score) + 1
        except Exception as e:
            print(f"Error calculating rubrics marks for {roll}: {e}")
            rubrics_marks = 0.0

        # Default Script & Attendance marks
        script_marks = attendance_marks = "N/A"

        # If uploaded, read from CSV
        for row in uploaded_data:
            if str(row.get("Roll Number")).strip() == str(roll):
                script_marks = row.get("Internal Script Marks (15)", "N/A")
                attendance_marks = row.get("Attendance Marks (5)", "N/A")

                # Convert numeric 0 values to actual 0 instead of N/A
                if script_marks == 0 or script_marks == "0" or script_marks == 0.0:
                    script_marks = 0
                if attendance_marks == 0 or attendance_marks == "0" or attendance_marks == 0.0:
                    attendance_marks = 0

        # Calculate Viva Marks (10) from best quiz attempt
        student_user = User.query.filter_by(rollnumber=roll).first()
        viva_marks = 0.0
        if student_user:
            if current_admin.admin_level == 'super_admin':
                # Super admin: find the actual admin responsible for this student
                responsible_admin = find_admin_for_student(roll)
                if responsible_admin:
                    viva_marks = _get_best_kdm_quiz_score(
                        student_user.id, responsible_admin.id)
                    print(
                        f"🔍 Super Admin: Student {roll} -> Responsible Admin {responsible_admin.rollnumber} ({responsible_admin.admin_level}) -> Viva: {viva_marks}")

                    # Extra debugging: check all quiz attempts for this student
                    all_attempts = KDMLabQuizAttempt.query.filter_by(
                        student_id=student_user.id, is_completed=True).all()
                    print(
                        f"   📊 All quiz attempts for {roll}: {len(all_attempts)} attempts")
                    for att in all_attempts:
                        att_admin = User.query.get(att.admin_id)
                        print(
                            f"      - Admin: {att_admin.rollnumber if att_admin else 'Unknown'}, Score: {att.score}")
                else:
                    print(
                        f"⚠️ Super Admin: No responsible admin found for student {roll}")
            else:
                # Regular admin: use current admin
                viva_marks = _get_best_kdm_quiz_score(
                    student_user.id, current_admin.id)
                print(
                    f"🔍 Admin Consolidated: Student {roll} -> Admin {current_admin.rollnumber} -> Viva: {viva_marks}")

        # Calculate Total Internal Marks (25) = Script (15) + Viva (10)
        total_internal_marks = "N/A"
        try:
            # Check if script_marks is a number (including 0) or N/A
            script_numeric = None
            if script_marks != "N/A":
                if isinstance(script_marks, (int, float)) or str(script_marks).replace('.', '', 1).replace('-', '', 1).isdigit():
                    script_numeric = float(script_marks)

            if script_numeric is not None and isinstance(viva_marks, (int, float)):
                total_internal_marks = round(script_numeric + viva_marks, 2)
        except Exception:
            total_internal_marks = "N/A"

        # Calculate Final Internal Marks (50)
        final_internal = "N/A"
        try:
            # Check if attendance_marks is a number (including 0) or N/A
            attendance_numeric = None
            if attendance_marks != "N/A":
                if isinstance(attendance_marks, (int, float)) or str(attendance_marks).replace('.', '', 1).replace('-', '', 1).isdigit():
                    attendance_numeric = float(attendance_marks)

            if (isinstance(rubrics_marks, (int, float)) and
                isinstance(total_internal_marks, (int, float)) and
                    attendance_numeric is not None):
                total_marks = rubrics_marks + total_internal_marks + attendance_numeric
                # Round up (e.g. 42.1 → 43) - consistent with other final internal marks
                final_internal = math.ceil(total_marks)
        except Exception:
            final_internal = "N/A"

        data.append({
            "Roll Number": roll,
            "Rubrics Marks (20)": rubrics_marks,
            "Internal Script Marks (15)": script_marks,
            "Internal Viva Marks (10)": viva_marks,
            "Total Internal Marks (25)": total_internal_marks,
            "Attendance Marks (5)": attendance_marks,
            "Final Internal Marks (50)": final_internal
        })

    return render_template("kdm_lab_consolidated_marks.html",
                           data=data,
                           current_admin=current_admin)


@app.route("/download_kdm_template")
def download_kdm_template():
    import pandas as pd
    from io import BytesIO
    from flask import send_file

    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Get students
    if current_admin.admin_level == 'super_admin':
        students = User.query.filter(User.role == 'student').order_by(
            User.rollnumber.asc()).all()
    else:
        students = get_students_for_admin(current_admin)
        students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # Create template DataFrame
    data = []
    for student in students:
        data.append({
            "Roll Number": student.rollnumber,
            "Internal Script Marks (15)": "",
            "Attendance Marks (5)": ""
        })

    df = pd.DataFrame(data)

    # Create Excel file in memory
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='KDM_Internal_Marks', index=False)

    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name=f'kdm_internal_marks_template_{current_admin.id}.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


@app.route("/upload_kdm_marks", methods=["POST"])
def upload_kdm_marks():
    import os
    import pandas as pd

    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    if "file" not in request.files:
        flash("No file selected!", "danger")
        return redirect(url_for("kdm_lab_consolidated_marks"))

    file = request.files["file"]
    if file.filename == "":
        flash("No file selected!", "danger")
        return redirect(url_for("kdm_lab_consolidated_marks"))

    try:
        # Create tmp_uploads directory if it doesn't exist
        upload_dir = os.path.join(os.getcwd(), "tmp_uploads")
        os.makedirs(upload_dir, exist_ok=True)

        # Save uploaded file
        admin_id = current_admin.id
        upload_path = os.path.join(
            upload_dir, f"kdm_internal_marks_{admin_id}.csv")

        # Read and validate the uploaded file
        df = pd.read_excel(file) if file.filename.endswith(
            '.xlsx') else pd.read_csv(file)

        # Validate required columns (updated for new structure)
        required_columns = ["Roll Number",
                            "Internal Script Marks (15)", "Attendance Marks (5)"]
        if not all(col in df.columns for col in required_columns):
            flash(
                f"Invalid file format! Required columns: {', '.join(required_columns)}", "danger")
            return redirect(url_for("kdm_lab_consolidated_marks"))

        # Save as CSV
        df.to_csv(upload_path, index=False)
        flash("✅ KDM Internal marks file uploaded successfully!", "success")

    except Exception as e:
        flash(f"❌ Error uploading file: {str(e)}", "danger")

    return redirect(url_for("kdm_lab_consolidated_marks"))


# ======================= END KDM LAB — ROUTES & HELPERS =======================


@app.route('/download_database_zip')
def download_database_zip():
    if 'rollnumber' not in session:
        flash('Please log in to access this page.', 'error')
        return redirect(url_for('login'))

    user = User.query.filter_by(rollnumber=session['rollnumber']).first()
    if not user or user.admin_level != 'super_admin':
        flash('Access denied. Only super admins can download the database.', 'error')
        return redirect(url_for('admin_dashboard'))

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:

        # Create comprehensive database schema documentation
        schema_info = StringIO()
        schema_info.write("DATABASE SCHEMA INFORMATION\n")
        schema_info.write("=" * 50 + "\n")
        schema_info.write(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        schema_info.write(
            f"Total Tables: {len(db.Model.registry.mappers)}\n\n")

        # Create SQL file with CREATE statements
        sql_output = StringIO()
        sql_output.write("-- DATABASE CREATION SCRIPT\n")
        sql_output.write(
            f"-- Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        sql_output.write(
            "-- Use this file to recreate the database structure\n\n")

        for mapper in db.Model.registry.mappers:
            model = mapper.class_
            try:
                table_name = model.__tablename__

                # === DATA CSV FILE ===
                filename = f"{table_name}.csv"
                output = StringIO()
                writer = csv.writer(output)

                # Get all column information
                columns_info = [(c.key, c) for c in inspect(model).columns]
                column_names = [c[0] for c in columns_info]
                writer.writerow(column_names)

                # Dump all rows
                for row in model.query.all():
                    writer.writerow([getattr(row, c) for c in column_names])

                zip_file.writestr(filename, output.getvalue())

                # === SCHEMA DOCUMENTATION ===
                schema_info.write(f"TABLE: {table_name.upper()}\n")
                schema_info.write("-" * 40 + "\n")

                # Model docstring if available
                if model.__doc__:
                    schema_info.write(
                        f"Description: {model.__doc__.strip()}\n")

                schema_info.write(f"Columns ({len(columns_info)}):\n")
                for col_name, col in columns_info:
                    col_type = str(col.type)
                    nullable = "NULL" if col.nullable else "NOT NULL"
                    primary_key = "PRIMARY KEY" if col.primary_key else ""
                    foreign_key = ""
                    if col.foreign_keys:
                        fk = list(col.foreign_keys)[0]
                        foreign_key = f"FOREIGN KEY -> {fk.column}"

                    default = f"DEFAULT {col.default.arg}" if col.default else ""

                    schema_info.write(
                        f"  - {col_name}: {col_type} {nullable} {primary_key} {foreign_key} {default}\n".strip() + "\n")

                # Count records
                record_count = model.query.count()
                schema_info.write(f"Records: {record_count}\n")
                schema_info.write(f"CSV File: {filename}\n\n")

                # === SQL CREATE STATEMENT ===
                sql_output.write(f"-- Table: {table_name}\n")
                sql_output.write(f"DROP TABLE IF EXISTS `{table_name}`;\n")
                sql_output.write(f"CREATE TABLE `{table_name}` (\n")

                col_definitions = []
                for col_name, col in columns_info:
                    # Convert SQLAlchemy types to MySQL types
                    col_type_str = str(col.type)
                    if 'VARCHAR' in col_type_str.upper():
                        mysql_type = col_type_str
                    elif 'INTEGER' in col_type_str.upper():
                        mysql_type = "INT"
                    elif 'TEXT' in col_type_str.upper():
                        mysql_type = "TEXT"
                    elif 'DATETIME' in col_type_str.upper():
                        mysql_type = "DATETIME"
                    elif 'BOOLEAN' in col_type_str.upper():
                        mysql_type = "TINYINT(1)"
                    elif 'FLOAT' in col_type_str.upper():
                        mysql_type = "FLOAT"
                    else:
                        mysql_type = col_type_str

                    col_def = f"  `{col_name}` {mysql_type}"

                    if not col.nullable:
                        col_def += " NOT NULL"

                    if col.primary_key:
                        if 'INT' in mysql_type.upper():
                            col_def += " AUTO_INCREMENT"
                        col_def += " PRIMARY KEY"

                    if col.default:
                        if isinstance(col.default.arg, str):
                            col_def += f" DEFAULT '{col.default.arg}'"
                        else:
                            col_def += f" DEFAULT {col.default.arg}"

                    col_definitions.append(col_def)

                sql_output.write(",\n".join(col_definitions))
                sql_output.write("\n);\n\n")

            except Exception as e:
                schema_info.write(
                    f"ERROR processing {model.__name__}: {str(e)}\n\n")
                print(f"⚠️ Skipping {model.__name__}: {e}")

        # Add schema files to ZIP
        zip_file.writestr("DATABASE_SCHEMA.txt", schema_info.getvalue())
        zip_file.writestr("CREATE_TABLES.sql", sql_output.getvalue())

        # Add README for deployment instructions
        readme = StringIO()
        readme.write("DATABASE DEPLOYMENT INSTRUCTIONS\n")
        readme.write("=" * 40 + "\n\n")
        readme.write("Files included:\n")
        readme.write("- DATABASE_SCHEMA.txt: Complete table documentation\n")
        readme.write("- CREATE_TABLES.sql: SQL script to create all tables\n")
        readme.write("- *.csv: Data files for each table\n\n")
        readme.write("Deployment Steps:\n")
        readme.write("1. Create new MySQL database\n")
        readme.write("2. Run CREATE_TABLES.sql to create table structure\n")
        readme.write("3. Import CSV files into respective tables\n")
        readme.write("4. Update Flask app database connection settings\n")
        readme.write("5. Run Flask app and verify functionality\n\n")
        readme.write(
            "Note: Adjust data types in CREATE_TABLES.sql as needed for your MySQL version\n")

        zip_file.writestr("README_DEPLOYMENT.txt", readme.getvalue())

    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        mimetype="application/zip",
        as_attachment=True,
        download_name="database_complete_backup.zip"
    )


@app.route('/download_database_csv')
def download_database_csv():
    # Only super_admins can download
    if 'rollnumber' not in session:
        flash('Please log in to access this page.', 'error')
        return redirect(url_for('login'))
    user = User.query.filter_by(rollnumber=session['rollnumber']).first()
    if not user or user.admin_level != 'super_admin':
        flash('Access denied. Only super admins can download the database.', 'error')
        return redirect(url_for('admin_dashboard'))

    output = StringIO()
    writer = csv.writer(output)

    # Export Users
    writer.writerow(['USERS'])
    writer.writerow([c.name for c in User.__table__.columns])
    for u in User.query.all():
        writer.writerow([getattr(u, c.name) for c in User.__table__.columns])
    writer.writerow([])

    # Export Quiz Questions
    writer.writerow(['QUIZ_QUESTIONS'])
    writer.writerow([c.name for c in QuizQuestion.__table__.columns])
    for q in QuizQuestion.query.all():
        writer.writerow([getattr(q, c.name)
                        for c in QuizQuestion.__table__.columns])
    writer.writerow([])

    # Export Quiz Attempts
    writer.writerow(['QUIZ_ATTEMPTS'])
    writer.writerow([c.name for c in QuizAttempt.__table__.columns])
    for a in QuizAttempt.query.all():
        writer.writerow([getattr(a, c.name)
                        for c in QuizAttempt.__table__.columns])
    writer.writerow([])

    # Export Quiz Responses
    writer.writerow(['QUIZ_RESPONSES'])
    writer.writerow([c.name for c in QuizResponse.__table__.columns])
    for r in QuizResponse.query.all():
        writer.writerow([getattr(r, c.name)
                        for c in QuizResponse.__table__.columns])
    writer.writerow([])

    # Export CO Mappings
    writer.writerow(['CO_MAPPING'])
    writer.writerow([c.name for c in COMapping.__table__.columns])
    for m in COMapping.query.all():
        writer.writerow([getattr(m, c.name)
                        for c in COMapping.__table__.columns])
    writer.writerow([])

    # Export CO Marks
    writer.writerow(['CO_MARKS'])
    writer.writerow([c.name for c in COMarks.__table__.columns])
    for m in COMarks.query.all():
        writer.writerow([getattr(m, c.name)
                        for c in COMarks.__table__.columns])
    writer.writerow([])

    # Export Admin Roll Assignments
    writer.writerow(['ADMIN_ROLL_ASSIGNMENT'])
    writer.writerow([c.name for c in AdminRollAssignment.__table__.columns])
    for a in AdminRollAssignment.query.all():
        writer.writerow([getattr(a, c.name)
                        for c in AdminRollAssignment.__table__.columns])
    writer.writerow([])

    output.seek(0)
    mem = BytesIO()
    mem.write(output.getvalue().encode('utf-8'))
    mem.seek(0)
    return send_file(
        mem,
        mimetype='text/csv',
        as_attachment=True,
        download_name='database_export.csv'
    )


@app.route("/admin_assignment_view")
def admin_assignment_view():
    """Allow admin to access assignment problems like a student"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    # Ensure only admin or super_admin can access
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    total_max_marks = calculate_total_max_marks()

    # For admin users, show all problems as available (bypass release restrictions)
    problem_visibility = {}
    for i in range(1, 12):  # Problems 1-11
        problem_visibility[i] = True  # Admin can see all problems as released

    return render_template("admin_assignment_view.html", user=user, total_max_marks=total_max_marks, problem_visibility=problem_visibility)

# --- Config Model for global settings ---


def get_config(key, default=None):
    c = Config.query.filter_by(key=key).first()
    return c.value if c else default


def set_config(key, value):
    c = Config.query.filter_by(key=key).first()
    if c:
        c.value = value
    else:
        c = Config(key=key, value=value)
        db.session.add(c)
    db.session.commit()


# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('debug.log'),
        logging.StreamHandler()  # This will also print to console
    ]
)

# --- Config Model for global settings ---


class Config(db.Model):
    __tablename__ = 'config'
    id = db.Column(db.Integer, primary_key=True)
    key = db.Column(db.String(100), unique=True, nullable=False)
    value = db.Column(db.Text, nullable=False)


# ---------------- Index → Login ----------------
ALLOWED_ROLLS_FILE = "/home/DOMASSIGNMENT/allowed_rolls.csv"

# Fallback path for local development
# ALLOWED_ROLLS_FILE = os.path.join(
#     os.path.dirname(__file__), "allowed_rolls.csv")


def load_allowed_rolls():
    if not os.path.exists(ALLOWED_ROLLS_FILE):
        # 🚀 Return empty list instead of crashing
        return []
    import csv
    with open(ALLOWED_ROLLS_FILE, newline='') as csvfile:
        # Read CSV without headers - each row contains just the roll number
        reader = csv.reader(csvfile)
        return [row[0].strip() for row in reader if row]  # Skip empty rows


# Load once at startup
ALLOWED_ROLLS = load_allowed_rolls()


def user_model_has_admin_id():
    """Return True only if the mapped User model exposes an admin_id attribute."""
    return hasattr(User, "admin_id")


# ---------------- Login Required Decorator ----------------
def login_required(f):
    """Decorator to require login for routes"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'loggedin' not in session:
            flash('Please log in to access this page.', 'warning')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


def is_admin():
    """Check if current user is an admin"""
    return session.get('role') == 'admin'


@app.route("/")
def index():
    return redirect(url_for("login"))

# ---------------- Register ----------------

# --- User Model ---


class User(db.Model):
    __tablename__ = 'users'   # ✅ make sure this matches your actual table name
    __table_args__ = {'extend_existing': True}  # ✅ Allow table redefinition

    id = db.Column(db.Integer, primary_key=True)
    rollnumber = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    email = db.Column(db.String(120))
    phonenumber = db.Column(db.String(20))
    marks = db.Column(db.Integer, default=0)
    completed = db.Column(db.Integer, default=0)
    role = db.Column(db.String(20), default="student")

    # ✅ Attempts and scores for each problem
    p1_attempts = db.Column(db.Integer, default=0)
    p1_score = db.Column(db.Integer, default=0)
    p2_attempts = db.Column(db.Integer, default=0)
    p2_score = db.Column(db.Integer, default=0)
    p3_attempts = db.Column(db.Integer, default=0)
    p3_score = db.Column(db.Integer, default=0)
    p4_attempts = db.Column(db.Integer, default=0)
    p4_score = db.Column(db.Integer, default=0)
    p5_attempts = db.Column(db.Integer, default=0)
    p5_score = db.Column(db.Integer, default=0)
    p6_attempts = db.Column(db.Integer, default=0)
    p6_score = db.Column(db.Integer, default=0)
    p7_attempts = db.Column(db.Integer, default=0)
    p7_score = db.Column(db.Integer, default=0)
    p8_attempts = db.Column(db.Integer, default=0)
    p8_score = db.Column(db.Integer, default=0)
    p9_attempts = db.Column(db.Integer, default=0)
    p9_score = db.Column(db.Integer, default=0)
    p10_attempts = db.Column(db.Integer, default=0)
    p10_score = db.Column(db.Integer, default=0)
    p11_attempts = db.Column(db.Integer, default=0)
    p11_score = db.Column(db.Integer, default=0)

    # ✅ Admin management fields
    admin_level = db.Column(db.Enum(
        'super_admin', 'admin', 'student', name='admin_level_enum'), default='student')
    created_by = db.Column(
        db.Integer, db.ForeignKey('users.id'), nullable=True)
    reg_date = db.Column(db.DateTime, default=db.func.current_timestamp())

    # Relationship to track who created this user (for admins)
    creator = db.relationship(
        'User', remote_side=[id], backref='created_admins')


# --- Admin Roll Assignment Model ---
class AdminRollAssignment(db.Model):
    __tablename__ = 'admin_roll_assignment'
    __table_args__ = {'extend_existing': True}

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    assignment_type = db.Column(
        db.Enum('range', 'individual', name='assignment_type_enum'), nullable=False)
    roll_start = db.Column(db.String(20), nullable=True)  # For range type
    roll_end = db.Column(db.String(20), nullable=True)    # For range type
    # For individual type (comma-separated)
    roll_numbers = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=db.func.current_timestamp())

    # Relationship to admin user
    admin = db.relationship('User', backref='roll_assignments')


# --- CO Mapping Model ---
class COMapping(db.Model):
    __tablename__ = 'co_mapping'
    __table_args__ = {'extend_existing': True}

    id = db.Column(db.Integer, primary_key=True)
    question_number = db.Column(
        db.Integer, nullable=False, unique=True)  # 1-11 for problems 1-11
    # CO1, CO2, CO3, CO4, CO5
    co_number = db.Column(db.String(10), nullable=False)
    # Maximum marks for this question
    max_marks = db.Column(db.Float, nullable=False, default=10.0)
    created_at = db.Column(db.DateTime, default=db.func.current_timestamp())
    updated_at = db.Column(db.DateTime, default=db.func.current_timestamp(
    ), onupdate=db.func.current_timestamp())


# --- CO Marks Model ---
class COMarks(db.Model):
    __tablename__ = 'co_marks'
    __table_args__ = {'extend_existing': True}

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    # CO1, CO2, CO3, CO4, CO5
    co_number = db.Column(db.String(10), nullable=False)
    marks_obtained = db.Column(db.Float, nullable=False, default=0.0)
    max_marks = db.Column(db.Float, nullable=False, default=0.0)
    percentage = db.Column(db.Float, nullable=False, default=0.0)
    last_updated = db.Column(db.DateTime, default=db.func.current_timestamp(
    ), onupdate=db.func.current_timestamp())

    # Relationship to user
    user = db.relationship('User', backref=db.backref('co_marks', lazy=True))

# ---------------- MID Exam Question Bank Models ----------------
# models.py (or wherever your models are)


class GlobalSettings(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    allow_student_mid_exam_questions = db.Column(db.Boolean, default=False)
    allow_student_python_lab = db.Column(db.Boolean, default=True)


class MidExamQuestion(db.Model):
    __tablename__ = 'mid_exam_question'
    id = db.Column(db.Integer, primary_key=True)
    question_text = db.Column(db.Text, nullable=False)
    bloom_level = db.Column(db.String(50), nullable=False)
    pi_level = db.Column(db.String(50), nullable=False)
    total_marks = db.Column(db.Integer, nullable=False)
    co_mappings = db.relationship(
        'MidExamQuestionCO', backref='question', lazy=True)


class MidExamQuestionCO(db.Model):
    __tablename__ = 'mid_exam_question_co'
    id = db.Column(db.Integer, primary_key=True)
    question_id = db.Column(db.Integer, db.ForeignKey(
        'mid_exam_question.id'), nullable=False)
    co_number = db.Column(db.Integer, nullable=False)
    co_marks = db.Column(db.Integer, nullable=False)


class MidExamCO(db.Model):
    __tablename__ = 'mid_exam_co'
    co_number = db.Column(db.Integer, primary_key=True)
    co_description = db.Column(db.Text)

# --- MID Exam Question Bank Routes ---


@app.route('/mid_exam_questions')
def mid_exam_questions():
    if 'loggedin' not in session:
        flash('Please log in to continue.', 'danger')
        return redirect(url_for('login'))

    role = session.get('role')
    current_admin = None
    problem_visibility = {}

    # Fetch all questions & COs
    questions = MidExamQuestion.query.all()
    cos = MidExamCO.query.all()

    # If admin/super_admin → include problem visibility logic
    if role in ['admin', 'super_admin']:
        current_admin = User.query.filter_by(
            rollnumber=session['rollnumber']).first()

        for i in range(1, 12):
            if current_admin.admin_level == 'super_admin':
                visibility = ProblemVisibility.query.filter_by(
                    problem_number=i, admin_id=None).first()
            else:
                global_visibility = ProblemVisibility.query.filter_by(
                    problem_number=i, admin_id=None).first()
                own_visibility = ProblemVisibility.query.filter_by(
                    problem_number=i, admin_id=current_admin.id).first()

                is_globally_released = global_visibility.is_released if global_visibility else False
                is_own_released = own_visibility.is_released if own_visibility else False

                class VisibilityStatus:
                    def __init__(self, is_released):
                        self.is_released = is_released
                visibility = VisibilityStatus(
                    is_globally_released or is_own_released)

            problem_visibility[i] = visibility.is_released if visibility else False

    # If student → no editing allowed, just view
    elif role == 'student':
        # Students don't need problem_visibility (optional)
        # show all released
        problem_visibility = {i: True for i in range(1, 12)}
    else:
        flash("Access denied.", "danger")
        return redirect(url_for("login"))

    return render_template(
        'mid_exam_questions.html',
        questions=questions,
        cos=cos,
        current_admin=current_admin,
        problem_visibility=problem_visibility,
        role=role,
        show_admin_panel=(role in ['admin', 'super_admin'])
    )


@app.route('/mid_exam_question/add', methods=['GET', 'POST'])
def add_mid_exam_question():
    if 'loggedin' not in session or session.get('role') not in ['admin', 'super_admin']:
        flash('Access denied.', 'danger')
        return redirect(url_for('login'))
    current_admin = User.query.filter_by(
        rollnumber=session['rollnumber']).first()
    cos = MidExamCO.query.all()
    if request.method == 'POST':
        text = request.form['question_text']
        bloom = request.form['bloom_level']
        pi = request.form['pi_level']
        marks = int(request.form['total_marks'])
        co_num = int(request.form['co_number'])  # Only one CO now
        try:
            q = MidExamQuestion(question_text=text,
                                bloom_level=bloom, pi_level=pi, total_marks=marks)
            db.session.add(q)
            db.session.flush()  # Get q.id before commit
            db.session.add(MidExamQuestionCO(question_id=q.id,
                           co_number=co_num, co_marks=marks))
            db.session.commit()
            flash('Question added!', 'success')
            return redirect(url_for('mid_exam_questions'))
        except SQLAlchemyError:
            db.session.rollback()
            flash('Error adding question.', 'danger')
    return render_template('add_mid_exam_question.html', cos=cos, current_admin=current_admin)


@app.route('/mid_exam_question/edit/<int:qid>', methods=['GET', 'POST'])
def edit_mid_exam_question(qid):
    if 'loggedin' not in session or session.get('role') not in ['admin', 'super_admin']:
        flash('Access denied.', 'danger')
        return redirect(url_for('login'))
    current_admin = User.query.filter_by(
        rollnumber=session['rollnumber']).first()
    question = MidExamQuestion.query.get_or_404(qid)
    cos = MidExamCO.query.all()
    if request.method == 'POST':
        question.question_text = request.form['question_text']
        question.bloom_level = request.form['bloom_level']
        question.pi_level = request.form['pi_level']
        question.total_marks = int(request.form['total_marks'])
        db.session.query(MidExamQuestionCO).filter_by(
            question_id=question.id).delete()
        co_num = int(request.form['co_number'])
        db.session.add(MidExamQuestionCO(question_id=question.id,
                       co_number=co_num, co_marks=question.total_marks))
        db.session.commit()
        flash('Question updated!', 'success')
        return redirect(url_for('mid_exam_questions'))
    co_mappings = MidExamQuestionCO.query.filter_by(
        question_id=question.id).all()
    return render_template('edit_mid_exam_question.html', question=question, co_mappings=co_mappings, cos=cos, current_admin=current_admin)


@app.route('/mid_exam_question/delete/<int:qid>', methods=['POST'])
def delete_mid_exam_question(qid):
    if 'loggedin' not in session or session.get('role') not in ['admin', 'super_admin']:
        flash('Access denied.', 'danger')
        return redirect(url_for('login'))
    question = MidExamQuestion.query.get_or_404(qid)
    db.session.query(MidExamQuestionCO).filter_by(
        question_id=question.id).delete()
    db.session.delete(question)
    db.session.commit()
    flash('Question deleted!', 'success')
    return redirect(url_for('mid_exam_questions'))


@app.route('/toggle_student_mid_exam_questions')
def toggle_student_mid_exam_questions():
    if 'loggedin' not in session or session.get('role') != 'super_admin':
        flash("Access denied.", "danger")
        return redirect(url_for("login"))

    setting = GlobalSettings.query.first()
    if not setting:
        setting = GlobalSettings(allow_student_mid_exam_questions=False)
        db.session.add(setting)
        db.session.commit()

    # Toggle the flag
    setting.allow_student_mid_exam_questions = not setting.allow_student_mid_exam_questions
    db.session.commit()

    flash("Student MID exam question view updated.", "success")
    return redirect(url_for("admin_dashboard"))


@app.route('/toggle_student_python_lab')
def toggle_student_python_lab():
    if 'loggedin' not in session or session.get('role') != 'super_admin':
        flash("Access denied.", "danger")
        return redirect(url_for("login"))

    setting = GlobalSettings.query.first()
    if not setting:
        setting = GlobalSettings(allow_student_python_lab=True)
        db.session.add(setting)
        db.session.commit()

    # Toggle the flag
    setting.allow_student_python_lab = not setting.allow_student_python_lab
    db.session.commit()

    status = "enabled" if setting.allow_student_python_lab else "disabled"
    flash(f"Python Programming Lab access for students {status}.", "success")
    return redirect(url_for("admin_dashboard"))


@app.route('/mid_exam_cos', methods=['GET', 'POST'])
def mid_exam_cos():
    if 'loggedin' not in session or session.get('role') not in ['admin', 'super_admin']:
        flash('Access denied.', 'danger')
        return redirect(url_for('login'))
    if request.method == 'POST':
        co_number = int(request.form['co_number'])
        co_description = request.form['co_description']
        co = MidExamCO(co_number=co_number, co_description=co_description)
        db.session.add(co)
        db.session.commit()
        flash('CO added!', 'success')
    current_admin = User.query.filter_by(
        rollnumber=session['rollnumber']).first()
    cos = MidExamCO.query.all()
    return render_template('mid_exam_cos.html', cos=cos, current_admin=current_admin)
# --- CO Utility Functions ---


@app.route('/create_mid_exam_question_paper', methods=['GET', 'POST'])
def create_mid_exam_question_paper():
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Access denied. Only admins can create mid exam questions.", "error")
        return redirect(url_for("admin_dashboard"))
    # if 'loggedin' not in session or session.get('role') != 'super_admin':
     #   flash('Access denied.', 'danger')
      #  return redirect(url_for('login'))'''

    current_admin = User.query.filter_by(
        rollnumber=session['rollnumber']).first()

    if request.method == 'POST':
        mid_type = request.form['mid_type']
        cos_input = request.form['cos']
        cos_list = [c.strip() for c in cos_input.split(',') if c.strip()]

        co_details = []
        for co in cos_list:
            main_is_single = (request.form.get(
                f'main_subdivisions_{co}', '1') == '1')
            choice_is_single = (request.form.get(
                f'choice_subdivisions_{co}', '1') == '1')
            main_marks = int(request.form.get(f'main_marks_{co}', 0))
            choice_marks = int(request.form.get(f'choice_marks_{co}', 0))
            co_details.append({
                'co': co,
                'main_is_single': main_is_single,
                'main_marks': main_marks,
                'choice_is_single': choice_is_single,
                'choice_marks': choice_marks
            })

        # ✅ Effective total = sum of max(main, choice) per CO
        effective_total = sum(
            max(d['main_marks'], d['choice_marks']) for d in co_details)
        if effective_total > 40:
            flash(
                f'Effective total (sum of per-unit max) is {effective_total} (>40). Please reduce marks.', 'danger')
            return render_template('create_mid_exam_question_paper.html',
                                   current_admin=current_admin)

        from itertools import combinations
        import random

        def pick_exact_sum_questions(candidates, target, require_single, exclude_ids=None):
            """
            Pick a subset of questions whose co_marks == target.
            Restriction: max 4 questions (1.a to 1.d).
            """
            exclude_ids = set(exclude_ids or [])
            usable = [c for c in candidates if c["id"] not in exclude_ids]

            if target <= 0 or not usable:
                return None

            if require_single:
                matches = [c for c in usable if c["marks"] == target]
                return [random.choice(matches)] if matches else None

            valid_sets = []
            for r in range(2, min(len(usable), 4) + 1):  # ✅ allow max 4
                for combo in combinations(usable, r):
                    if sum(x["marks"] for x in combo) == target:
                        valid_sets.append(combo)
            if valid_sets:
                return list(random.choice(valid_sets))
            return None

        preview_questions = []
        for detail in co_details:
            co_num = int(detail['co'])
            # Get candidates for this CO
            mappings = (db.session.query(MidExamQuestion, MidExamQuestionCO)
                        .join(MidExamQuestionCO, MidExamQuestion.id == MidExamQuestionCO.question_id)
                        .filter(MidExamQuestionCO.co_number == co_num)
                        .all())

            candidates = [{
                "id": q.id,
                "question_text": q.question_text,
                "marks": maprow.co_marks,
                "bloom_level": q.bloom_level,
                "pi_level": q.pi_level
            } for q, maprow in mappings]

            main_sel = pick_exact_sum_questions(
                candidates, detail['main_marks'], require_single=detail['main_is_single']
            )
            exclude_for_choice = {c["id"] for c in (main_sel or [])}
            choice_sel = pick_exact_sum_questions(
                candidates, detail['choice_marks'],
                require_single=detail['choice_is_single'],
                exclude_ids=exclude_for_choice
            )

            preview_questions.append({
                'co': detail['co'],
                'main_marks': detail['main_marks'],
                'choice_marks': detail['choice_marks'],
                'main': main_sel if main_sel else [
                    {"question_text": "No valid combination found",
                        "marks": "", "bloom_level": "",  "pi_level": ""}
                ],
                'choice': choice_sel if choice_sel else [
                    {"question_text": "No valid combination found",
                        "marks": "", "bloom_level": "", "pi_level": ""}
                ]
            })

        # Save only IDs for download
        session['preview_data'] = {
            'mid_type': mid_type,
            'cos': cos_list,
            'questions': [{
                'co': q['co'],
                'main': [c["id"] for c in q['main'] if c["marks"] != ""],
                'choice': [c["id"] for c in q['choice'] if c["marks"] != ""]
            } for q in preview_questions]
        }

        return render_template(
            'create_mid_exam_question_paper.html',
            current_admin=current_admin,
            preview_data={'mid_type': mid_type,
                          'cos': cos_list,
                          'questions': preview_questions}
        )

    # GET
    return render_template('create_mid_exam_question_paper.html', current_admin=current_admin)


@app.route('/download_mid_exam_question_paper', methods=['POST'])
def download_mid_exam_question_paper():
    preview_data = session.get('preview_data')
    if not preview_data:
        flash("No preview data found. Please generate a preview first.", "danger")
        return redirect(url_for('create_mid_exam_question_paper'))

    mid_type = preview_data['mid_type']
    questions = []

    # Fetch questions by IDs from DB
    for q in preview_data['questions']:
        main_qs = (db.session.query(MidExamQuestion, MidExamQuestionCO)
                   .join(MidExamQuestionCO, MidExamQuestion.id == MidExamQuestionCO.question_id)
                   .filter(MidExamQuestion.id.in_(q['main']))
                   .all())
        choice_qs = (db.session.query(MidExamQuestion, MidExamQuestionCO)
                     .join(MidExamQuestionCO, MidExamQuestion.id == MidExamQuestionCO.question_id)
                     .filter(MidExamQuestion.id.in_(q['choice']))
                     .all())

        questions.append({
            'co': q['co'],
            'main': [{"id": m.id, "question_text": m.question_text,
                      "marks": maprow.co_marks, "bloom_level": m.bloom_level, "pi_level": m.pi_level}
                     for m, maprow in main_qs],
            'choice': [{"id": c.id, "question_text": c.question_text,
                        "marks": maprow.co_marks, "bloom_level": c.bloom_level, "pi_level": c.pi_level}
                       for c, maprow in choice_qs]
        })

    # ---- DOCX generation ----
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    # A4 margins
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # ✅ Calculate usable width globally (page width - margins)
    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_width_in = usable_width / 914400   # convert EMUs → inches

    # Logo
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    # Increase image size to fit page width (A4 width is about 8.27 inches, so use ~6.5 for margins)
    run.add_picture("/home/DOMASSIGNMENT/static/midlogo.png",
                    width=Inches(7.0))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Exam info
    # doc.add_paragraph("Reg.No.")
    '''p = doc.add_paragraph(
        #f"II/III/IV B.Tech, Mech (Section-A,B,C & D ) SEM- I/II, {mid_type}")
        f" Department of Mechanical Engineering
        f"III/IV B.Tech, Mech (Section-A,B,C & D ) SEM- I, {mid_type}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER'''
    p1 = doc.add_paragraph()
    run1 = p1.add_run("Department of Mechanical Engineering")
    run1.bold = True
    run1.font.size = Pt(14)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Second line - normal font
    p2 = doc.add_paragraph(
        f"III/IV B.Tech, Mech (Section-A,B,C & D) SEM- I, {mid_type}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # --- Subject and Max Marks ---
    sub_table = doc.add_table(rows=1, cols=2)
    sub_table.autofit = True
    sub_table.columns[0].width = Inches(3)
    sub_table.columns[1].width = Inches(3)

    sub_table.cell(0, 0).text = "Sub : Dynamics of Machinery"
    sub_table.cell(0, 1).text = "Max. Marks: 40"

    # Alignments
    for para in sub_table.cell(0, 0).paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for para in sub_table.cell(0, 1).paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- Date and Time ---
    dt_table = doc.add_table(rows=1, cols=2)
    dt_table.autofit = True
    dt_table.columns[0].width = Inches(3)
    dt_table.columns[1].width = Inches(3)

    dt_table.cell(0, 0).text = "Date:"
    dt_table.cell(0, 1).text = "Time:"

    # Alignments
    for para in dt_table.cell(0, 0).paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for para in dt_table.cell(0, 1).paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- Note ---
    note = doc.add_paragraph("Note:  Answer any 1 Question from Each UNIT.")
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT

    letters = ["a", "b", "c", "d"]  # max 4 sub-questions
    q_counter = 1
    co_alloc = {}

    for unit_idx, q in enumerate(questions, 1):
        # ✅ UNIT heading centered + bold + larger font
        # Roman numerals for UNIT heading
        roman_numerals = ["I", "II", "III", "IV",
                          "V", "VI", "VII", "VIII", "IX", "X"]
        unit_label = roman_numerals[unit_idx -
                                    1] if unit_idx <= len(roman_numerals) else str(unit_idx)
        p = doc.add_paragraph(f"UNIT {unit_label}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(12)

        # Table
        # ---- Table ----
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.autofit = False
        table.allow_autofit = False   # ✅ fully prevent resizing

        # ✅ Use already calculated usable_width_in
        col_widths_in = [
            usable_width_in * 0.07,  # Q.No
            usable_width_in * 0.60,  # Questions
            usable_width_in * 0.08,  # M
            usable_width_in * 0.08,  # CO
            usable_width_in * 0.08,  # BL
            usable_width_in * 0.09   # PI
        ]

        for i, width_in in enumerate(col_widths_in):
            for cell in table.columns[i].cells:
                cell.width = Inches(width_in)   # ✅ force width column-wise

        # Header
        hdr = table.rows[0].cells
        hdr[0].text = "Q.No"
        hdr[1].text = "Questions"
        hdr[2].text = "M"
        hdr[3].text = "CO"
        hdr[4].text = "BL"
        hdr[5].text = "PI"

        # Center header except Questions, enlarge header font
        for i in [0, 2, 3, 4, 5]:
            for para in hdr[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(14)
                    run.bold = True
        # Questions header font enlarged
        for run in hdr[1].paragraphs[0].runs:
            run.font.size = Pt(14)
            run.bold = True

        sum_main, sum_choice = 0, 0

        # --- Main rows ---
        if q['main']:
            for sub_idx, mq in enumerate(q['main'], start=1):
                row = table.add_row().cells
                row[0].text = f"{q_counter}.{letters[sub_idx-1]}"
                row[1].text = mq["question_text"]
                row[2].text = str(mq["marks"])
                row[3].text = f"CO-{q['co']}"
                row[4].text = str(mq["bloom_level"])
                row[5].text = str(mq["pi_level"])
                # row[6].text = ""
                sum_main += mq["marks"]

                # ✅ Center Q.No, M, CO, BL, PI
                for i in [0, 2, 3, 4, 5]:
                    for p in row[i].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- OR row ---
        if q['main'] and q['choice']:
            or_cells = table.add_row().cells
            merged = or_cells[0]
            for i in range(1, 6):
                merged = merged.merge(or_cells[i])
            para = merged.paragraphs[0]
            run = para.add_run("(OR)")
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Choice rows ---
        if q['choice']:
            for sub_idx, cq in enumerate(q['choice'], start=1):
                row = table.add_row().cells
                row[0].text = f"{q_counter+1}.{letters[sub_idx-1]}"
                row[1].text = cq["question_text"]
                row[2].text = str(cq["marks"])
                row[3].text = f"CO-{q['co']}"
                row[4].text = str(cq["bloom_level"])
                row[5].text = str(cq["pi_level"])
                # row[6].text = ""
                sum_choice += cq["marks"]

                # ✅ Center Q.No, M, CO, BL, PI
                for i in [0, 2, 3, 4, 5]:
                    for p in row[i].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # For charts
        try:
            co_key = int(q['co'])
        except Exception:
            co_key = q['co']
        co_alloc[co_key] = max(sum_main, sum_choice)

        q_counter += 2

        # Charts
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        # --- Bar chart: CO-wise allocation ---
        sorted_items = sorted(co_alloc.items(), key=lambda x: x[0])
        labels = [f"CO{k}" for k, _ in sorted_items]
        values = [v for _, v in sorted_items]

        fig1, ax1 = plt.subplots(figsize=(3.0, 3.0))
        ax1.bar(labels, values)
        ax1.set_title("CO-wise Marks Allocation")
        ax1.set_ylabel("Marks")
        ax1.set_xlabel("CO")
        for i, v in enumerate(values):
            ax1.text(i, v + 0.5, str(v), ha='center')
        buf_bar = io.BytesIO()
        fig1.tight_layout()
        fig1.savefig(buf_bar, format='png', dpi=200)
        plt.close(fig1)
        buf_bar.seek(0)

        # --- Pie chart: Bloom’s Level wise distribution ---
        bl_totals = {}
        bl_map = {
            "1": "Remember",
            "2": "Understand",
            "3": "Apply",
            "4": "Analyze",
            "5": "Evaluate",
            "6": "Create"
        }
        for q in questions:
            for mq in q['main']:
                if mq["bloom_level"]:
                    label = bl_map.get(
                        str(mq["bloom_level"]), str(mq["bloom_level"]))
                    bl_totals[label] = bl_totals.get(label, 0) + mq['marks']
            for cq in q['choice']:
                if cq["bloom_level"]:
                    label = bl_map.get(
                        str(cq["bloom_level"]), str(cq["bloom_level"]))
                    bl_totals[label] = bl_totals.get(label, 0) + cq['marks']

        blooms_order = ["Remember", "Understand",
                        "Apply", "Analyze", "Evaluate", "Create"]
        bl_labels = [b for b in blooms_order if b in bl_totals]
        bl_values = [bl_totals[b] for b in bl_labels]

        fig2, ax2 = plt.subplots(figsize=(3.0, 3.0))
        if sum(bl_values) > 0:
            ax2.pie(bl_values,
                    labels=bl_labels,     # ✅ only Bloom’s names
                    autopct='%1.0f%%',    # ✅ only percentages
                    startangle=90)
        ax2.set_title("Bloom’s Level Distribution")
        buf_pie = io.BytesIO()
        fig2.tight_layout()
        fig2.savefig(buf_pie, format='png', dpi=200)
        plt.close(fig2)
        buf_pie.seek(0)

        # --- Insert side by side ---
        doc.add_paragraph("\nCO & Bloom’s Level Analysis")
        chart_table = doc.add_table(rows=1, cols=2)
        chart_table.autofit = True

        cell1 = chart_table.rows[0].cells[0]
        run1 = cell1.paragraphs[0].add_run()
        run1.add_picture(buf_bar, width=Inches(3.0))

        cell2 = chart_table.rows[0].cells[1]
        run2 = cell2.paragraphs[0].add_run()
        run2.add_picture(buf_pie, width=Inches(3.0))

    except Exception:
        pass

    doc.add_paragraph("\n\n\nSignature of\nThe Course Coordinator")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name="MID_Exam_Question_Paper.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def initialize_co_mappings():
    """Initialize default CO mappings if they don't exist"""
    default_mappings = {
        1: 'CO1', 2: 'CO2', 3: 'CO1', 4: 'CO1', 5: 'CO2',
        6: 'CO2', 7: 'CO5', 8: 'CO5', 9: 'CO3', 10: 'CO4', 11: 'CO3'
    }

    for question_num, co in default_mappings.items():
        existing = COMapping.query.filter_by(
            question_number=question_num).first()
        if not existing:
            mapping = COMapping(
                question_number=question_num,
                co_number=co,
                max_marks=10.0  # Default 10 marks per question
            )
            db.session.add(mapping)

    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()


def update_co_marks(user_id):
    """Calculate and update CO marks for a user based on their problem scores"""
    try:
        user = User.query.get(user_id)
        if not user:
            return

        # Get all CO mappings
        co_mappings = COMapping.query.all()
        if not co_mappings:
            return

        co_data = {}

        # Initialize CO data
        for mapping in co_mappings:
            co = mapping.co_number
            if co not in co_data:
                co_data[co] = {'obtained': 0.0, 'max': 0.0}

        # Calculate marks for each CO
        problem_scores = [
            user.p1_score or 0, user.p2_score or 0, user.p3_score or 0, user.p4_score or 0, user.p5_score or 0,
            user.p6_score or 0, user.p7_score or 0, user.p8_score or 0, user.p9_score or 0, user.p10_score or 0, user.p11_score or 0
        ]

        for mapping in co_mappings:
            question_index = mapping.question_number - 1  # Convert to 0-based index
            if question_index < len(problem_scores):
                obtained_marks = float(problem_scores[question_index])
                co_data[mapping.co_number]['obtained'] += obtained_marks
                co_data[mapping.co_number]['max'] += mapping.max_marks

        # Update or create CO marks records
        for co_number, marks in co_data.items():
            co_mark = COMarks.query.filter_by(
                user_id=user_id, co_number=co_number).first()

            percentage = (marks['obtained'] / marks['max']
                          * 100) if marks['max'] > 0 else 0

            if co_mark:
                co_mark.marks_obtained = marks['obtained']
                co_mark.max_marks = marks['max']
                co_mark.percentage = percentage
            else:
                co_mark = COMarks(
                    user_id=user_id,
                    co_number=co_number,
                    marks_obtained=marks['obtained'],
                    max_marks=marks['max'],
                    percentage=percentage
                )
                db.session.add(co_mark)

        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()

    except Exception as e:
        pass

    def __repr__(self):
        if self.assignment_type == 'range':
            return f'<AdminRollAssignment {self.admin_id}: {self.roll_start}-{self.roll_end}>'
        else:
            return f'<AdminRollAssignment {self.admin_id}: {self.roll_numbers}>'


# --- Problem Visibility Model ---
class ProblemVisibility(db.Model):
    __tablename__ = 'problem_visibility'
    __table_args__ = {'extend_existing': True}

    id = db.Column(db.Integer, primary_key=True)
    problem_number = db.Column(db.Integer, nullable=False)
    is_released = db.Column(db.Boolean, default=False)
    release_date = db.Column(db.DateTime)
    # NULL means global release (for super admin)
    admin_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True)

    # Ensure unique combination of problem_number and admin_id
    __table_args__ = (db.UniqueConstraint('problem_number',
                      'admin_id', name='_problem_admin_uc'),)

    def __repr__(self):
        return f'<ProblemVisibility {self.problem_number}: {self.is_released} (Admin: {self.admin_id})>'


def get_problem_visibility(problem_number):
    """Get visibility status for a specific problem (legacy function - use get_problem_visibility_for_user instead)"""
    visibility = ProblemVisibility.query.filter_by(
        problem_number=problem_number, admin_id=None).first()
    if not visibility:
        # Create default global entry if doesn't exist
        visibility = ProblemVisibility(
            problem_number=problem_number, is_released=True, admin_id=None)
        db.session.add(visibility)
        db.session.commit()
    return visibility.is_released


def get_problem_visibility_for_user(problem_number, user_roll_number):
    """Get visibility status for a specific problem for a specific user"""
    user = User.query.filter_by(rollnumber=user_roll_number).first()
    if not user:
        return False

    # If user is an admin, they can see all problems
    if user.role == "admin":
        return True

    # For students, check which admin manages them and if that admin has released the problem
    # First check if any admin who manages this student has released the problem

    # Get all admins who have this student in their assigned roll numbers
    all_admins = User.query.filter_by(role="admin").all()

    for admin in all_admins:
        if admin.admin_level == 'super_admin':
            # Check if super admin has released it globally
            visibility = ProblemVisibility.query.filter_by(
                problem_number=problem_number, admin_id=None).first()
            if visibility and visibility.is_released:
                return True
        else:
            # Check if this sub admin manages the student and has released the problem
            assigned_rolls = get_assigned_roll_numbers(admin.id)
            if user_roll_number in assigned_rolls:
                visibility = ProblemVisibility.query.filter_by(
                    problem_number=problem_number, admin_id=admin.id).first()
                if visibility and visibility.is_released:
                    return True

    return False


def set_problem_visibility(problem_number, is_released):
    """Set visibility status for a specific problem (legacy function - creates global visibility)"""
    visibility = ProblemVisibility.query.filter_by(
        problem_number=problem_number, admin_id=None).first()
    if not visibility:
        visibility = ProblemVisibility(
            problem_number=problem_number, is_released=is_released, admin_id=None)
        db.session.add(visibility)
    else:
        visibility.is_released = is_released
    db.session.commit()
    return visibility


def check_problem_access(problem_number):
    """Check if student can access a specific problem"""
    if "rollnumber" not in session:
        return False, redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user:
        return False, redirect(url_for("login"))

    # Admin and super_admin users can access all problems regardless of release status
    if user.role in ["admin", "super_admin"]:
        return True, user

    # Check if problem is released for this specific user
    if not get_problem_visibility_for_user(problem_number, user.rollnumber):
        flash(
            f"Problem {problem_number} has not been released yet.", "warning")
        return False, redirect(url_for("student_dashboard"))

    # Removed sequential restriction - students can now access any released problem
    # No need to complete previous problems first

    return True, user


# --- Quiz Models ---

QUIZ_CO_OPTIONS = ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']
DEFAULT_QUIZ_NAME = 'default_quiz'
DOM_CONCEPTUAL_SYSTEM_INITIALIZED = False


class QuizBank(db.Model):
    __tablename__ = 'quiz_banks'
    __table_args__ = {'extend_existing': True}

    id = db.Column(db.Integer, primary_key=True)
    quiz_name = db.Column(db.String(100), unique=True, nullable=False)
    quiz_title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=True)
    covered_cos = db.Column(db.JSON, nullable=True)
    questions_per_attempt = db.Column(db.Integer, nullable=True, default=10)
    is_shared_with_all_admins = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=db.func.current_timestamp())
    created_by = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)

    creator = db.relationship('User', backref='created_quiz_banks')

    def __repr__(self):
        return f'<QuizBank {self.quiz_name}>'


class QuizQuestionAssignment(db.Model):
    __tablename__ = 'quiz_question_assignments'
    __table_args__ = (
        db.UniqueConstraint('quiz_id', 'question_id', name='uq_quiz_question_assignment'),
        {'extend_existing': True}
    )

    id = db.Column(db.Integer, primary_key=True)
    quiz_id = db.Column(db.Integer, db.ForeignKey('quiz_banks.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('quiz_questions.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=db.func.current_timestamp())

    quiz = db.relationship('QuizBank', backref='question_assignments')
    question = db.relationship('QuizQuestion', backref='quiz_assignments')


class QuizAdminAccess(db.Model):
    __tablename__ = 'quiz_admin_access'
    __table_args__ = (
        db.UniqueConstraint('quiz_id', 'admin_id', name='uq_quiz_admin_access'),
        {'extend_existing': True}
    )

    id = db.Column(db.Integer, primary_key=True)
    quiz_id = db.Column(db.Integer, db.ForeignKey('quiz_banks.id'), nullable=False)
    admin_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    granted_at = db.Column(db.DateTime, default=db.func.current_timestamp())
    granted_by = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True)

    quiz = db.relationship('QuizBank', backref='admin_access_entries')
    admin = db.relationship('User', foreign_keys=[admin_id], backref='accessible_quiz_entries')
    grantor = db.relationship('User', foreign_keys=[granted_by])

class QuizQuestion(db.Model):
    __tablename__ = 'quiz_questions'
    __table_args__ = {'extend_existing': True}

    id = db.Column(db.Integer, primary_key=True)
    question_text = db.Column(db.Text, nullable=False)
    co_number = db.Column(db.String(10), nullable=True)  # CO1, CO2, CO3, etc.
    points = db.Column(db.Integer, nullable=False, default=1)
    answer_type = db.Column(db.Enum(
        'single', 'multiple', name='answer_type_enum'), nullable=False, default='single')

    # Store choices as JSON array: ["Choice A", "Choice B", "Choice C", "Choice D"]
    choices = db.Column(db.JSON, nullable=False)

    # Store correct answer indices as JSON: [0] for single, [0, 2] for multiple
    correct_answers = db.Column(db.JSON, nullable=False)

    # Explanation for the correct answer
    explanation = db.Column(db.Text, nullable=True)

    created_at = db.Column(db.DateTime, default=db.func.current_timestamp())
    created_by = db.Column(
        db.Integer, db.ForeignKey('users.id'), nullable=False)

    # Relationship to creator (super admin)
    creator = db.relationship('User', backref='created_quiz_questions')

    def __repr__(self):
        return f'<QuizQuestion {self.id}: {self.co_number} - {self.points} points>'


class QuizAttempt(db.Model):
    __tablename__ = 'quiz_attempts'
    __table_args__ = {'extend_existing': True}

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    quiz_name = db.Column(db.String(100), default='default_quiz')
    started_at = db.Column(db.DateTime, default=db.func.current_timestamp())
    completed_at = db.Column(db.DateTime, nullable=True)
    auto_submitted = db.Column(db.Boolean, default=False)
    score = db.Column(db.Integer, default=0)
    total_points = db.Column(db.Integer, default=0)
    attempt_number = db.Column(db.Integer, nullable=False, default=1)

    # Relationships
    user = db.relationship('User', backref='quiz_attempts')

    def __repr__(self):
        return f'<QuizAttempt {self.id}: User {self.user_id} - Attempt {self.attempt_number}>'


class QuizResponse(db.Model):
    __tablename__ = 'quiz_responses'
    __table_args__ = {'extend_existing': True}

    id = db.Column(db.Integer, primary_key=True)
    attempt_id = db.Column(db.Integer, db.ForeignKey(
        'quiz_attempts.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey(
        'quiz_questions.id'), nullable=False)

    # Store user's selected answer as JSON (matching database column name)
    selected_answer = db.Column(db.JSON, nullable=True)
    response_data = db.Column(db.JSON, nullable=True)

    is_correct = db.Column(db.Boolean, nullable=False, default=False)
    points_earned = db.Column(db.Integer, nullable=False, default=0)

    answered_at = db.Column(db.DateTime, default=db.func.current_timestamp())

    # Relationships
    attempt = db.relationship('QuizAttempt', backref='responses')
    question = db.relationship('QuizQuestion', backref='responses')

    def __repr__(self):
        return f'<QuizResponse {self.id}: Q{self.question_id} - {self.points_earned} points>'


class QuizVisibility(db.Model):
    __tablename__ = 'quiz_visibility'
    __table_args__ = {'extend_existing': True}

    id = db.Column(db.Integer, primary_key=True)
    quiz_name = db.Column(db.String(100), nullable=False,
                          default='default_quiz')
    is_released = db.Column(db.Boolean, default=False)
    released_at = db.Column(db.DateTime, nullable=True)
    released_by = db.Column(
        db.Integer, db.ForeignKey('users.id'), nullable=True)
    # NULL means global release (for super admin), specific admin_id means per-admin release
    admin_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True)

    # Relationship to admin who released it
    admin = db.relationship(
        'User', backref='quiz_releases', foreign_keys=[released_by])
    # Relationship to admin who owns this visibility setting
    owner_admin = db.relationship('User', foreign_keys=[admin_id])

    def __repr__(self):
        admin_str = f"Admin {self.admin_id}" if self.admin_id else "Global"
        return f'<QuizVisibility {self.quiz_name}: {"Released" if self.is_released else "Locked"} ({admin_str})>'


class DOMSubjectStudentQuizRelease(db.Model):
    __tablename__ = 'dom_subject_student_quiz_releases'

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'))
    student_id = db.Column(db.Integer, db.ForeignKey(
        'users.id', ondelete='CASCADE'))
    is_released = db.Column(db.Boolean, default=False)
    created_at = db.Column(
        db.TIMESTAMP, server_default=db.func.current_timestamp())
    updated_at = db.Column(db.TIMESTAMP, server_default=db.func.current_timestamp(
    ), onupdate=db.func.current_timestamp())

    # Ensure one release record per admin-student pair
    __table_args__ = (db.UniqueConstraint('admin_id', 'student_id'),)

    # Relationships
    admin = db.relationship('User', foreign_keys=[
                            admin_id], backref='dom_quiz_releases_created')
    student = db.relationship('User', foreign_keys=[
                              student_id], backref='dom_quiz_releases_received')

    def __repr__(self):
        return f'<DOMSubjectStudentQuizRelease Admin:{self.admin_id} Student:{self.student_id} Released:{self.is_released}>'


class DOMConceptualQuizSession(db.Model):
    __tablename__ = 'dom_conceptual_quiz_sessions'
    __table_args__ = {'extend_existing': True}

    id = db.Column(db.Integer, primary_key=True)
    admin_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('quiz_questions.id'), nullable=False)
    is_active = db.Column(db.Boolean, default=True)
    posted_at = db.Column(db.DateTime, default=db.func.current_timestamp())
    closed_at = db.Column(db.DateTime, nullable=True)

    admin = db.relationship('User', foreign_keys=[admin_id], backref='dom_conceptual_sessions')
    question = db.relationship('QuizQuestion', backref='dom_conceptual_sessions')


class DOMConceptualQuizResponse(db.Model):
    __tablename__ = 'dom_conceptual_quiz_responses'
    __table_args__ = (
        db.UniqueConstraint('session_id', 'student_id', name='uq_dom_conceptual_session_student'),
        {'extend_existing': True}
    )

    id = db.Column(db.Integer, primary_key=True)
    session_id = db.Column(db.Integer, db.ForeignKey('dom_conceptual_quiz_sessions.id'), nullable=False)
    student_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    selected_answer = db.Column(db.JSON, nullable=True)
    is_correct = db.Column(db.Boolean, nullable=False, default=False)
    answered_at = db.Column(db.DateTime, default=db.func.current_timestamp())

    session = db.relationship('DOMConceptualQuizSession', backref='responses')
    student = db.relationship('User', foreign_keys=[student_id], backref='dom_conceptual_responses')


# --- Quiz Utility Functions ---

def get_quiz_visibility_for_user(user_roll_number):
    """Get quiz visibility status for a specific user"""
    try:
        user = User.query.filter_by(rollnumber=user_roll_number).first()
        if not user:
            print(
                f"⚠️ get_quiz_visibility_for_user: User {user_roll_number} not found")
            return False

        # If user is an admin or super_admin, they can always access quiz
        if user.role in ["admin", "super_admin"]:
            print(
                f"✅ get_quiz_visibility_for_user: Admin {user_roll_number} always has access")
            return True

        # For students, check if ANY admin who manages them has released the quiz
        return check_quiz_visibility_for_student(user)
    except Exception as e:
        # Quiz tables don't exist yet or other error
        print(f"⚠️ Error checking quiz visibility: {e}")
        import traceback
        traceback.print_exc()
        return False


def check_quiz_visibility_for_student(user):
    """Check if the quiz is visible for a specific student based on their assigned admins"""
    try:
        # Handle None user case
        if not user:
            print("⚠️ check_quiz_visibility_for_student: user is None")
            return False

        # First check if super admin has released it globally (admin_id=None)
        global_visibility = QuizVisibility.query.filter_by(
            quiz_name='default_quiz', admin_id=None).first()
        if global_visibility and global_visibility.is_released:
            print(
                f"✅ Quiz globally released by super admin for student {user.rollnumber}")
            return True

        # Check if any admin who manages this student (based on roll number assignments) has released the quiz
        # Get all admins and check if this student's roll number is in their assigned range
        try:
            all_admins = User.query.filter(User.admin_level == 'admin').all()
            for admin in all_admins:
                try:
                    assigned_roll_numbers = get_assigned_roll_numbers(admin.id)
                    if user.rollnumber in assigned_roll_numbers:
                        # First check for individual DOM subject quiz release
                        dom_individual_release = DOMSubjectStudentQuizRelease.query.filter_by(
                            admin_id=admin.id, student_id=user.id, is_released=True).first()
                        if dom_individual_release:
                            print(
                                f"✅ DOM Subject quiz individually released by admin {admin.id} for student {user.rollnumber}")
                            return True

                        # Then check bulk quiz release by this admin
                        admin_visibility = QuizVisibility.query.filter_by(
                            quiz_name='default_quiz', admin_id=admin.id).first()
                        if admin_visibility and admin_visibility.is_released:
                            print(
                                f"✅ Quiz bulk released by admin {admin.id} ({admin.rollnumber}) for student {user.rollnumber}")
                            return True
                except Exception as admin_error:
                    print(f"⚠️ Error checking admin {admin.id}: {admin_error}")
                    continue

        except Exception as admins_error:
            print(f"⚠️ Error getting admins list: {admins_error}")

        # Also check if any super admin has individually released the quiz for this specific student
        try:
            super_admins = User.query.filter(
                User.admin_level == 'super_admin').all()
            for super_admin in super_admins:
                dom_individual_release = DOMSubjectStudentQuizRelease.query.filter_by(
                    admin_id=super_admin.id, student_id=user.id, is_released=True).first()
                if dom_individual_release:
                    print(
                        f"✅ DOM Subject quiz individually released by super admin {super_admin.id} for student {user.rollnumber}")
                    return True
        except Exception as super_admin_error:
            print(
                f"⚠️ Error checking super admin individual releases: {super_admin_error}")

        print(f"❌ Quiz not released for student {user.rollnumber}")
        return False
    except Exception as e:
        print(f"⚠️ Error checking quiz visibility for student: {e}")
        import traceback
        traceback.print_exc()
        return False


def get_actual_quiz_release_status():
    """Get the actual quiz release status from database (for students) - DEPRECATED"""
    # This function is deprecated - use check_quiz_visibility_for_student instead
    try:
        visibility = QuizVisibility.query.filter_by(
            quiz_name='default_quiz').first()
        return visibility.is_released if visibility else False
    except Exception as e:
        print(f"⚠️ Error getting quiz release status: {e}")
        return False


def set_quiz_visibility(admin_id, is_released):
    """Set quiz visibility status for a specific admin (None for global)"""
    quiz_name = 'default_quiz'

    # Get current admin user to determine admin level
    admin_user = User.query.get(admin_id) if admin_id else None

    if admin_user and admin_user.admin_level == 'super_admin':
        # Super admin creates/updates global visibility (admin_id=None)
        visibility = QuizVisibility.query.filter_by(
            quiz_name=quiz_name, admin_id=None).first()
        if not visibility:
            visibility = QuizVisibility(
                quiz_name=quiz_name,
                is_released=is_released,
                released_at=datetime.now() if is_released else None,
                released_by=admin_id if is_released else None,
                admin_id=None  # Global release
            )
            db.session.add(visibility)
            print(
                f"🌍 Super admin creating global quiz visibility: {is_released}")
        else:
            visibility.is_released = is_released
            visibility.released_at = datetime.now() if is_released else None
            visibility.released_by = admin_id if is_released else None
            print(
                f"🌍 Super admin updating global quiz visibility: {is_released}")
    else:
        # Sub admin creates/updates their own visibility record
        visibility = QuizVisibility.query.filter_by(
            quiz_name=quiz_name, admin_id=admin_id).first()
        if not visibility:
            visibility = QuizVisibility(
                quiz_name=quiz_name,
                is_released=is_released,
                released_at=datetime.now() if is_released else None,
                released_by=admin_id if is_released else None,
                admin_id=admin_id  # Admin-specific release
            )
            db.session.add(visibility)
            print(
                f"👤 Sub admin {admin_id} creating quiz visibility: {is_released}")
        else:
            visibility.is_released = is_released
            visibility.released_at = datetime.now() if is_released else None
            visibility.released_by = admin_id if is_released else None
            print(
                f"👤 Sub admin {admin_id} updating quiz visibility: {is_released}")

    try:
        db.session.commit()
        print(f"💾 Quiz visibility saved successfully")
    except Exception as e:
        db.session.rollback()
        print(f"💥 Error saving quiz visibility: {e}")
        raise e

    return visibility


def get_user_quiz_attempts(user_id):
    """Get number of quiz attempts for a user"""
    return QuizAttempt.query.filter_by(user_id=user_id).count()


def can_user_take_quiz(user_roll_number):
    """Check if user can take quiz (visibility + attempt limits)"""
    user = User.query.filter_by(rollnumber=user_roll_number).first()
    if not user:
        print(f"❌ can_user_take_quiz: User {user_roll_number} not found")
        return False, "User not found"

    print(
        f"🔍 can_user_take_quiz: User {user_roll_number} has role '{user.role}'")

    # Admins and super_admins can always take the quiz for testing purposes
    if user.role in ["admin", "super_admin"]:
        print(
            f"✅ can_user_take_quiz: Admin {user_roll_number} can always take quiz")
        return True, "Admin can always take quiz"

    # For students, check visibility first
    quiz_visible = get_quiz_visibility_for_user(user_roll_number)
    print(
        f"🔍 can_user_take_quiz: Student {user_roll_number} quiz_visible = {quiz_visible}")

    if not quiz_visible:
        return False, "Quiz has not been released yet"

    # Check attempt limits for students only
    attempts = get_user_quiz_attempts(user.id)
    print(
        f"🔍 can_user_take_quiz: Student {user_roll_number} has {attempts} attempts")

    if attempts >= 2:
        return False, "Maximum 2 attempts exceeded"

    return True, "Can take quiz"


def normalize_quiz_name(value):
    value = (value or "").strip().lower()
    value = re.sub(r'[^a-z0-9]+', '_', value)
    return value.strip('_')[:100] or DEFAULT_QUIZ_NAME


def parse_quiz_cos(values):
    if isinstance(values, str):
        values = [item.strip() for item in values.split(',')]
    return [co for co in values if co in QUIZ_CO_OPTIONS]


def ensure_default_quiz_bank():
    quiz_bank = QuizBank.query.filter_by(quiz_name=DEFAULT_QUIZ_NAME).first()
    if not quiz_bank:
        creator = User.query.filter(User.admin_level == 'super_admin').first() or User.query.filter(
            User.role.in_(['admin', 'super_admin'])).first()
        creator_id = creator.id if creator else 1
        quiz_bank = QuizBank(
            quiz_name=DEFAULT_QUIZ_NAME,
            quiz_title="Default Quiz",
            description="Migrated default quiz bank",
            covered_cos=QUIZ_CO_OPTIONS,
            questions_per_attempt=10,
            is_shared_with_all_admins=True,
            created_by=creator_id
        )
        db.session.add(quiz_bank)
        db.session.flush()

    existing_ids = {
        row.question_id for row in QuizQuestionAssignment.query.filter_by(quiz_id=quiz_bank.id).all()
    }
    for question in QuizQuestion.query.all():
        if question.id not in existing_ids:
            db.session.add(QuizQuestionAssignment(
                quiz_id=quiz_bank.id, question_id=question.id))

    return quiz_bank


def ensure_quiz_bank_schema():
    inspector = inspect(db.engine)
    quiz_bank_columns = {column["name"] for column in inspector.get_columns("quiz_banks")}
    if "questions_per_attempt" not in quiz_bank_columns:
        db.session.execute(
            text("ALTER TABLE quiz_banks ADD COLUMN questions_per_attempt INTEGER NULL DEFAULT 10")
        )
        db.session.commit()


def ensure_quiz_attempt_schema():
    inspector = inspect(db.engine)
    quiz_attempt_columns = {column["name"] for column in inspector.get_columns("quiz_attempts")}
    if "auto_submitted" not in quiz_attempt_columns:
        db.session.execute(
            text("ALTER TABLE quiz_attempts ADD COLUMN auto_submitted BOOLEAN DEFAULT 0")
        )
        db.session.commit()


def initialize_dom_conceptual_testing_system():
    global DOM_CONCEPTUAL_SYSTEM_INITIALIZED
    if DOM_CONCEPTUAL_SYSTEM_INITIALIZED:
        return
    db.create_all()
    DOM_CONCEPTUAL_SYSTEM_INITIALIZED = True


def initialize_quiz_bank_system():
    db.create_all()
    ensure_quiz_bank_schema()
    ensure_quiz_attempt_schema()
    quiz_bank = ensure_default_quiz_bank()
    db.session.commit()
    return quiz_bank


def get_quiz_bank_by_name(quiz_name):
    initialize_quiz_bank_system()
    return QuizBank.query.filter_by(quiz_name=quiz_name or DEFAULT_QUIZ_NAME).first()


def get_choice_text(choice):
    if isinstance(choice, dict):
        return choice.get('text') or choice.get('value') or str(choice)
    return str(choice)


def get_conceptual_question_payload(question):
    if not question:
        return None
    return {
        "id": question.id,
        "co_number": question.co_number or "",
        "question_text": question.question_text,
        "answer_type": question.answer_type,
        "choices": [
            {"index": index, "text": get_choice_text(choice)}
            for index, choice in enumerate(question.choices or [])
        ],
    }


def get_conceptual_students_for_admin(admin_user):
    if admin_user.admin_level == 'super_admin':
        return User.query.filter_by(role='student').all()
    return get_students_for_admin(admin_user)


def get_dom_conceptual_active_session_for_admin(admin_user):
    initialize_dom_conceptual_testing_system()
    return (
        DOMConceptualQuizSession.query
        .filter_by(admin_id=admin_user.id, is_active=True)
        .order_by(DOMConceptualQuizSession.posted_at.desc(), DOMConceptualQuizSession.id.desc())
        .first()
    )


def get_dom_conceptual_active_session_for_student(student):
    initialize_dom_conceptual_testing_system()
    admin_ids = []
    responsible_admin = find_admin_for_student(student.rollnumber)
    if responsible_admin:
        admin_ids.append(responsible_admin.id)

    admin_ids.extend([
        row[0] for row in db.session.query(User.id)
        .filter(User.role.in_(['admin', 'super_admin']), User.admin_level == 'super_admin')
        .all()
    ])
    admin_ids = list(dict.fromkeys(admin_ids))
    if not admin_ids:
        return None

    return (
        DOMConceptualQuizSession.query
        .filter(
            DOMConceptualQuizSession.admin_id.in_(admin_ids),
            DOMConceptualQuizSession.is_active.is_(True)
        )
        .order_by(DOMConceptualQuizSession.posted_at.desc(), DOMConceptualQuizSession.id.desc())
        .first()
    )


def get_dom_conceptual_response(session_id, student_id):
    return DOMConceptualQuizResponse.query.filter_by(
        session_id=session_id,
        student_id=student_id
    ).first()


def get_dom_conceptual_stats(conceptual_session, admin_user):
    students = get_conceptual_students_for_admin(admin_user)
    student_ids = [student.id for student in students]
    responses = []
    if conceptual_session and student_ids:
        responses = (
            DOMConceptualQuizResponse.query
            .filter(
                DOMConceptualQuizResponse.session_id == conceptual_session.id,
                DOMConceptualQuizResponse.student_id.in_(student_ids)
            )
            .all()
        )

    correct_count = sum(1 for response in responses if response.is_correct)
    incorrect_responses = [response for response in responses if not response.is_correct]
    wrong_rollnumbers = sorted(
        response.student.rollnumber
        for response in incorrect_responses
        if response.student
    )

    return {
        "total_students": len(students),
        "answered_count": len(responses),
        "correct_count": correct_count,
        "incorrect_count": len(incorrect_responses),
        "pending_count": max(0, len(students) - len(responses)),
        "wrong_rollnumbers": wrong_rollnumbers,
    }


def get_admin_quiz_banks(user):
    initialize_quiz_bank_system()
    if not user or user.role not in ["admin", "super_admin"]:
        return []
    if user.admin_level == 'super_admin':
        return QuizBank.query.order_by(QuizBank.created_at.desc(), QuizBank.quiz_title.asc()).all()

    return QuizBank.query.outerjoin(
        QuizAdminAccess, QuizAdminAccess.quiz_id == QuizBank.id
    ).filter(
        or_(
            QuizBank.created_by == user.id,
            QuizBank.is_shared_with_all_admins.is_(True),
            QuizAdminAccess.admin_id == user.id
        )
    ).distinct().order_by(QuizBank.created_at.desc(), QuizBank.quiz_title.asc()).all()


def can_admin_access_quiz_bank(user, quiz_bank):
    if not user or not quiz_bank or user.role not in ["admin", "super_admin"]:
        return False
    if user.admin_level == 'super_admin':
        return True
    if quiz_bank.created_by == user.id or quiz_bank.is_shared_with_all_admins:
        return True
    return QuizAdminAccess.query.filter_by(quiz_id=quiz_bank.id, admin_id=user.id).first() is not None


def can_admin_edit_quiz_bank(user, quiz_bank):
    return bool(user and quiz_bank and (user.admin_level == 'super_admin' or quiz_bank.created_by == user.id))


def get_quiz_questions(quiz_bank):
    if not quiz_bank:
        return []
    return QuizQuestion.query.join(
        QuizQuestionAssignment, QuizQuestionAssignment.question_id == QuizQuestion.id
    ).filter(
        QuizQuestionAssignment.quiz_id == quiz_bank.id
    ).order_by(QuizQuestion.co_number, QuizQuestion.id).all()


def get_quiz_question_counts(quiz_bank):
    counts = {co: 0 for co in QUIZ_CO_OPTIONS}
    for question in get_quiz_questions(quiz_bank):
        if question.co_number in counts:
            counts[question.co_number] += 1
    return counts


def get_quiz_questions_per_attempt(quiz_bank, available_count=None):
    total_available = available_count if available_count is not None else len(get_quiz_questions(quiz_bank))
    configured = getattr(quiz_bank, "questions_per_attempt", None) if quiz_bank else None

    try:
        configured = int(configured)
    except (TypeError, ValueError):
        configured = 10

    configured = max(1, configured)
    if total_available <= 0:
        return configured
    return min(configured, total_available)


def select_random_questions_for_quiz(quiz_bank):
    quiz_questions = get_quiz_questions(quiz_bank)
    if not quiz_questions:
        return []

    target_count = get_quiz_questions_per_attempt(quiz_bank, available_count=len(quiz_questions))
    covered_cos = parse_quiz_cos(getattr(quiz_bank, "covered_cos", None) or [])
    if not covered_cos:
        covered_cos = sorted({
            question.co_number
            for question in quiz_questions
            if question.co_number in QUIZ_CO_OPTIONS
        })

    questions_by_co = {}
    for question in quiz_questions:
        if covered_cos and question.co_number not in covered_cos:
            continue
        questions_by_co.setdefault(question.co_number or "UNSPECIFIED", []).append(question)

    for pool in questions_by_co.values():
        random.shuffle(pool)

    co_keys = [co for co in covered_cos if questions_by_co.get(co)]
    if not co_keys:
        co_keys = [co for co, pool in questions_by_co.items() if pool]

    selected_questions = []
    base_quota = target_count // len(co_keys) if co_keys else 0
    remainder = target_count % len(co_keys) if co_keys else 0
    quotas = {co: base_quota for co in co_keys}

    remainder_cos = co_keys[:]
    random.shuffle(remainder_cos)
    for co in remainder_cos:
        if remainder <= 0:
            break
        quotas[co] += 1
        remainder -= 1

    for co in co_keys:
        pool = questions_by_co.get(co) or []
        take_count = min(quotas.get(co, 0), len(pool))
        for _ in range(take_count):
            selected_questions.append(pool.pop())

    if len(selected_questions) < target_count:
        remaining_questions = []
        for co in co_keys:
            remaining_questions.extend(questions_by_co.get(co) or [])
        random.shuffle(remaining_questions)
        needed = target_count - len(selected_questions)
        selected_questions.extend(remaining_questions[:needed])

    if len(selected_questions) < target_count:
        fallback_questions = [
            question for question in quiz_questions
            if question not in selected_questions
        ]
        random.shuffle(fallback_questions)
        selected_questions.extend(fallback_questions[:target_count - len(selected_questions)])

    random.shuffle(selected_questions)
    return selected_questions


def get_quiz_co_question_share(quiz_bank):
    covered_cos = parse_quiz_cos(getattr(quiz_bank, "covered_cos", None) or [])
    if not covered_cos:
        covered_cos = sorted({
            question.co_number
            for question in get_quiz_questions(quiz_bank)
            if question.co_number in QUIZ_CO_OPTIONS
        })

    if not covered_cos:
        return {}

    target_count = get_quiz_questions_per_attempt(quiz_bank)
    base_quota = target_count // len(covered_cos)
    remainder = target_count % len(covered_cos)
    shares = {co: float(base_quota) for co in covered_cos}
    for co in covered_cos:
        if remainder <= 0:
            break
        shares[co] += 1.0
        remainder -= 1
    return shares


def get_best_attempts_by_quiz(attempts):
    best_by_quiz = {}
    for attempt in attempts:
        quiz_name = attempt.quiz_name or DEFAULT_QUIZ_NAME
        current_best = best_by_quiz.get(quiz_name)
        attempt_score = attempt.score or 0
        attempt_total = attempt.total_points or 0
        attempt_percentage = (attempt_score / attempt_total * 100) if attempt_total > 0 else 0

        if current_best:
            current_score = current_best.score or 0
            current_total = current_best.total_points or 0
            current_percentage = (current_score / current_total * 100) if current_total > 0 else 0
        else:
            current_score = -1
            current_percentage = -1

        if (
            current_best is None
            or attempt_score > current_score
            or (attempt_score == current_score and attempt_percentage > current_percentage)
        ):
            best_by_quiz[quiz_name] = attempt
    return best_by_quiz


def assign_question_to_quiz(quiz_bank, question):
    existing = QuizQuestionAssignment.query.filter_by(
        quiz_id=quiz_bank.id, question_id=question.id).first()
    if not existing:
        db.session.add(QuizQuestionAssignment(
            quiz_id=quiz_bank.id, question_id=question.id))


def get_quiz_release_status_for_admin(user, quiz_name):
    if not user or user.role not in ["admin", "super_admin"]:
        return False
    if user.admin_level == 'super_admin':
        visibility = QuizVisibility.query.filter_by(
            quiz_name=quiz_name, admin_id=None).first()
        return visibility.is_released if visibility else False

    global_visibility = QuizVisibility.query.filter_by(
        quiz_name=quiz_name, admin_id=None).first()
    own_visibility = QuizVisibility.query.filter_by(
        quiz_name=quiz_name, admin_id=user.id).first()
    return bool((global_visibility and global_visibility.is_released) or (own_visibility and own_visibility.is_released))


def get_available_quiz_banks_for_user(user):
    initialize_quiz_bank_system()
    if not user:
        return []
    if user.role in ["admin", "super_admin"]:
        return get_admin_quiz_banks(user)

    quiz_names = {
        row.quiz_name for row in QuizVisibility.query.filter_by(admin_id=None, is_released=True).all()
    }

    assigned_admins = []
    for admin in User.query.filter(User.admin_level == 'admin').all():
        try:
            if user.rollnumber in get_assigned_roll_numbers(admin.id):
                assigned_admins.append(admin)
        except Exception:
            continue

    for admin in assigned_admins:
        if DOMSubjectStudentQuizRelease.query.filter_by(
            admin_id=admin.id, student_id=user.id, is_released=True
        ).first():
            for quiz_bank in get_admin_quiz_banks(admin):
                quiz_names.add(quiz_bank.quiz_name)
        for row in QuizVisibility.query.filter_by(admin_id=admin.id, is_released=True).all():
            quiz_names.add(row.quiz_name)

    for super_admin in User.query.filter(User.admin_level == 'super_admin').all():
        if DOMSubjectStudentQuizRelease.query.filter_by(
            admin_id=super_admin.id, student_id=user.id, is_released=True
        ).first():
            for quiz_bank in QuizBank.query.all():
                quiz_names.add(quiz_bank.quiz_name)

    if not quiz_names:
        return []

    return QuizBank.query.filter(QuizBank.quiz_name.in_(quiz_names)).order_by(QuizBank.quiz_title.asc()).all()


def get_student_quiz_summaries(user):
    if not user:
        return []

    available_quizzes = get_available_quiz_banks_for_user(user)
    quiz_bank_map = {quiz.quiz_name: quiz for quiz in available_quizzes}
    attempted_quiz_names = [
        row[0]
        for row in db.session.query(QuizAttempt.quiz_name)
        .filter_by(user_id=user.id)
        .distinct()
        .all()
        if row[0]
    ]

    quiz_names = sorted(set(quiz_bank_map.keys()) | set(attempted_quiz_names))
    summaries = []

    for quiz_name in quiz_names:
        quiz_bank = quiz_bank_map.get(quiz_name) or get_quiz_bank_by_name(quiz_name)
        if not quiz_bank:
            continue

        attempts = (
            QuizAttempt.query.filter_by(user_id=user.id, quiz_name=quiz_name)
            .order_by(QuizAttempt.attempt_number.asc(), QuizAttempt.id.asc())
            .all()
        )
        completed_attempts = [attempt for attempt in attempts if attempt.completed_at]
        best_attempt = max(completed_attempts, key=lambda attempt: attempt.score) if completed_attempts else None
        latest_attempt = completed_attempts[-1] if completed_attempts else None
        active_attempt = next((attempt for attempt in reversed(attempts) if not attempt.completed_at), None)
        released = get_quiz_visibility_for_user(user.rollnumber, quiz_name=quiz_name)

        summaries.append({
            "quiz_name": quiz_name,
            "quiz_title": quiz_bank.quiz_title or quiz_name,
            "released": released,
            "total_attempts": len(attempts),
            "completed_attempts": len(completed_attempts),
            "remaining_attempts": max(0, 2 - len(attempts)) if user.role == "student" else None,
            "best_attempt": best_attempt,
            "latest_attempt": latest_attempt,
            "active_attempt": active_attempt,
            "can_take": bool(active_attempt) or can_user_take_quiz(user.rollnumber, quiz_name=quiz_name)[0],
        })

    return summaries


def get_quiz_visibility_for_user(user_roll_number, quiz_name=None):
    try:
        user = User.query.filter_by(rollnumber=user_roll_number).first()
        if not user:
            return False
        available_quizzes = get_available_quiz_banks_for_user(user)
        if quiz_name:
            return any(quiz.quiz_name == quiz_name for quiz in available_quizzes)
        return bool(available_quizzes)
    except Exception as e:
        print(f"Error checking quiz visibility: {e}")
        return False


def get_actual_quiz_release_status():
    try:
        visibility = QuizVisibility.query.filter_by(quiz_name=DEFAULT_QUIZ_NAME).first()
        return visibility.is_released if visibility else False
    except Exception as e:
        print(f"Error getting quiz release status: {e}")
        return False


def set_quiz_visibility(admin_id, is_released, quiz_name=DEFAULT_QUIZ_NAME):
    admin_user = User.query.get(admin_id) if admin_id else None

    if admin_user and admin_user.admin_level == 'super_admin':
        visibility = QuizVisibility.query.filter_by(
            quiz_name=quiz_name, admin_id=None).first()
        if not visibility:
            visibility = QuizVisibility(quiz_name=quiz_name, admin_id=None)
            db.session.add(visibility)
    else:
        visibility = QuizVisibility.query.filter_by(
            quiz_name=quiz_name, admin_id=admin_id).first()
        if not visibility:
            visibility = QuizVisibility(quiz_name=quiz_name, admin_id=admin_id)
            db.session.add(visibility)

    visibility.is_released = is_released
    visibility.released_at = datetime.now() if is_released else None
    visibility.released_by = admin_id if is_released else None
    db.session.commit()
    return visibility


def get_user_quiz_attempts(user_id, quiz_name=None):
    query = QuizAttempt.query.filter_by(user_id=user_id)
    if quiz_name:
        query = query.filter_by(quiz_name=quiz_name)
    return query.count()


def can_user_take_quiz(user_roll_number, quiz_name=None):
    user = User.query.filter_by(rollnumber=user_roll_number).first()
    if not user:
        return False, "User not found"
    if user.role in ["admin", "super_admin"]:
        return True, "Admin can always take quiz"
    if not get_quiz_visibility_for_user(user_roll_number, quiz_name=quiz_name):
        return False, "Quiz has not been released yet"
    if quiz_name and get_user_quiz_attempts(user.id, quiz_name=quiz_name) >= 2:
        return False, "Maximum 2 attempts exceeded for this quiz"
    return True, "Can take quiz"


def update_score(user, problem_no, score):
    attempts_col = f"p{problem_no}_attempts"
    score_col = f"p{problem_no}_score"

    attempts = getattr(user, attempts_col, 0)
    best_score = getattr(user, score_col, 0)

    # Admin users have unlimited attempts
    if attempts >= 2 and user.role != "admin":
        # No more attempts allowed for non-admin users
        return False

    # Increment attempts
    setattr(user, attempts_col, attempts + 1)

    # Update best score
    setattr(user, score_col, max(best_score, score))

    # Mark completed if used 2 attempts OR got full marks
    if problem_no in [1, 2, 3]:
        full_marks = 5
    elif problem_no == 8:
        full_marks = 15
    else:
        full_marks = 10  # Problems 4, 5, 6, 7, 9

    # Update the completed count by counting all problems with attempts > 0
    # This counts problems that student has attempted at least once
    completed_count = 0
    for i in range(1, 12):
        if getattr(user, f"p{i}_attempts", 0) > 0:
            completed_count += 1
    user.completed = completed_count

    # Update total marks = sum of all pX_score
    user.marks = sum(getattr(user, f"p{i}_score", 0) for i in range(1, 12))

    db.session.commit()

    # Update CO marks after score update
    update_co_marks(user.id)

    return True


# ========================= KDM LAB INTERNAL QUIZ SYSTEM =========================

@app.route("/kdm_lab_upload_quiz_questions", methods=["GET", "POST"])
def kdm_lab_upload_quiz_questions():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    if request.method == "POST":
        if 'quiz_file' not in request.files:
            flash("No file selected!", "danger")
            return redirect(request.url)

        file = request.files['quiz_file']
        if file.filename == '':
            flash("No file selected!", "danger")
            return redirect(request.url)

        # Get upload mode (replace or add)
        upload_mode = request.form.get('upload_mode', 'replace')

        if file and file.filename.lower().endswith('.csv'):
            try:
                import pandas as pd
                import io

                # Read CSV file
                stream = io.StringIO(
                    file.stream.read().decode("UTF8"), newline=None)
                df = pd.read_csv(stream)

                # Normalize column names to handle different formats
                df.columns = df.columns.str.strip().str.lower()

                # Create mapping for different column name variations
                column_mapping = {}

                # Map SlNo variations
                for col in df.columns:
                    if col in ['slno', 'sl no', 'sl.no', 'serial no', 'sno', 's.no']:
                        column_mapping['slno'] = col

                # Map Question variations
                for col in df.columns:
                    if col in ['question', 'questions', 'question text']:
                        column_mapping['questions'] = col

                # Map Option variations
                for col in df.columns:
                    if col in ['option a', 'optiona', 'option_a', 'a']:
                        column_mapping['optiona'] = col
                    elif col in ['option b', 'optionb', 'option_b', 'b']:
                        column_mapping['optionb'] = col
                    elif col in ['option c', 'optionc', 'option_c', 'c']:
                        column_mapping['optionc'] = col
                    elif col in ['option d', 'optiond', 'option_d', 'd']:
                        column_mapping['optiond'] = col

                # Map Correct Answer variations
                for col in df.columns:
                    if col in ['correct answer', 'correct_answer', 'answer', 'correct', 'ans']:
                        column_mapping['correct_answer'] = col

                # Check if all required columns are found
                required_columns = ['slno', 'questions', 'optiona',
                                    'optionb', 'optionc', 'optiond', 'correct_answer']
                missing_columns = [
                    col for col in required_columns if col not in column_mapping]

                if missing_columns:
                    flash(
                        f"Could not find columns for: {', '.join(missing_columns)}. Available columns: {', '.join(df.columns)}", "danger")
                    return redirect(request.url)

                # Clear existing questions for this admin only if replace mode
                if upload_mode == 'replace':
                    KDMLabQuizQuestion.query.filter_by(
                        admin_id=user.id).delete()
                    action_text = "replaced with"
                else:
                    action_text = "added"

                # Insert new questions
                questions_added = 0
                for _, row in df.iterrows():
                    try:
                        # Use column mapping to get values
                        slno_val = row[column_mapping['slno']
                                       ] if column_mapping.get('slno') else None
                        question_val = row[column_mapping['questions']]
                        option_a_val = row[column_mapping['optiona']]
                        option_b_val = row[column_mapping['optionb']]
                        option_c_val = row[column_mapping['optionc']]
                        option_d_val = row[column_mapping['optiond']]
                        correct_answer_val = row[column_mapping['correct_answer']]

                        question = KDMLabQuizQuestion(
                            admin_id=user.id,
                            slno=int(slno_val) if pd.notna(slno_val) else None,
                            question=str(question_val).strip(),
                            option_a=str(option_a_val).strip(),
                            option_b=str(option_b_val).strip(),
                            option_c=str(option_c_val).strip(),
                            option_d=str(option_d_val).strip(),
                            correct_answer=str(
                                correct_answer_val).strip().upper()
                        )
                        db.session.add(question)
                        questions_added += 1
                    except Exception as e:
                        print(
                            f"Error adding question {row.get(column_mapping.get('slno', 'slno'), 'unknown')}: {e}")
                        continue

                db.session.commit()
                flash(
                    f"Successfully {action_text} {questions_added} quiz questions!", "success")

            except Exception as e:
                db.session.rollback()
                flash(f"Error uploading file: {str(e)}", "danger")
        else:
            flash("Please upload a CSV file only!", "danger")

    # Get existing questions count - Super admin sees all, regular admin sees only their own
    if user.admin_level == 'super_admin':
        questions_count = KDMLabQuizQuestion.query.count()
    else:
        questions_count = KDMLabQuizQuestion.query.filter_by(
            admin_id=user.id).count()

    return render_template("kdm_lab_upload_quiz_questions.html",
                           questions_count=questions_count,
                           current_admin=user)


from functools import lru_cache

@app.route("/kdm_lab_manage_internal_quiz")
def kdm_lab_manage_internal_quiz():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # --- Get or create quiz settings ---
    quiz_settings = KDMLabQuizSettings.query.filter_by(admin_id=user.id).first()
    if not quiz_settings:
        quiz_settings = KDMLabQuizSettings(admin_id=user.id)
        db.session.add(quiz_settings)
        db.session.commit()

    # --- Global question count ---
    questions_count = KDMLabQuizQuestion.query.count()

    # --- Optimized attempt count logic ---
    attempts_count = 0
    try:
        if user.role == "super_admin":
            # ✅ Super admin sees all attempts
            attempts_count = KDMLabQuizAttempt.query.count()
        else:
            # ✅ Cache admin lookups to avoid repeated DB hits
            @lru_cache(maxsize=None)
            def get_admin_for_roll(roll):
                try:
                    admin_user = find_admin_for_student(roll)
                    return admin_user.id if admin_user else None
                except Exception:
                    return None

            # Get all student IDs once
            all_students = User.query.with_entities(User.id, User.rollnumber).filter_by(role="student").all()

            # Filter students belonging to this admin using cached lookup
            student_ids = [
                s.id for s in all_students
                if get_admin_for_roll(s.rollnumber) == user.id
            ]

            if student_ids:
                attempts_count = (
                    KDMLabQuizAttempt.query
                    .filter(KDMLabQuizAttempt.user_id.in_(student_ids))
                    .count()
                )
            else:
                attempts_count = 0
    except Exception as e:
        print(f"⚠️ Error calculating attempt count: {e}")
        attempts_count = 0

    return render_template(
        "kdm_lab_manage_internal_quiz.html",
        quiz_settings=quiz_settings,
        questions_count=questions_count,
        attempts_count=attempts_count,
        current_admin=user
    )


@app.route("/kdm_lab_quiz_settings", methods=["GET", "POST"])
def kdm_lab_quiz_settings():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Get or create quiz settings
    quiz_settings = KDMLabQuizSettings.query.filter_by(
        admin_id=user.id).first()
    if not quiz_settings:
        quiz_settings = KDMLabQuizSettings(admin_id=user.id)
        db.session.add(quiz_settings)
        db.session.commit()

    if request.method == "POST":
        try:
            questions_per_quiz = int(
                request.form.get("questions_per_quiz", 10))
            duration_minutes = int(request.form.get("duration_minutes", 10))

            # Validate values
            if questions_per_quiz < 1 or questions_per_quiz > 100:
                flash("Number of questions must be between 1 and 100!", "danger")
                return redirect(url_for("kdm_lab_quiz_settings"))

            if duration_minutes < 1 or duration_minutes > 180:
                flash("Quiz duration must be between 1 and 180 minutes!", "danger")
                return redirect(url_for("kdm_lab_quiz_settings"))

            # Check if we have enough questions - Super admin sees all, regular admin sees only their own
            if user.admin_level == 'super_admin':
                questions_count = KDMLabQuizQuestion.query.count()
            else:
                questions_count = KDMLabQuizQuestion.query.filter_by(
                    admin_id=user.id).count()

            if questions_count < questions_per_quiz:
                flash(
                    f"You only have {questions_count} questions. Upload more questions before setting {questions_per_quiz} questions per quiz!", "warning")
                return redirect(url_for("kdm_lab_quiz_settings"))

            # Update settings
            quiz_settings.questions_per_quiz = questions_per_quiz
            quiz_settings.quiz_duration_minutes = duration_minutes
            db.session.commit()

            flash("Quiz settings updated successfully!", "success")
            return redirect(url_for("kdm_lab_manage_internal_quiz"))

        except ValueError:
            flash("Please enter valid numbers for questions and duration!", "danger")
        except Exception as e:
            db.session.rollback()
            flash("Error updating quiz settings. Please try again.", "danger")

    # Get current questions count for validation - Super admin sees all, regular admin sees only their own
    if user.admin_level == 'super_admin':
        questions_count = KDMLabQuizQuestion.query.count()
    else:
        questions_count = KDMLabQuizQuestion.query.filter_by(
            admin_id=user.id).count()

    return render_template("kdm_lab_quiz_settings.html",
                           quiz_settings=quiz_settings,
                           questions_count=questions_count,
                           current_admin=user)


@app.route("/kdm_lab_admin_preview_quiz")
def kdm_lab_admin_preview_quiz():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Get quiz settings
    quiz_settings = KDMLabQuizSettings.query.filter_by(
        admin_id=user.id).first()
    if not quiz_settings:
        quiz_settings = KDMLabQuizSettings(admin_id=user.id)
        db.session.add(quiz_settings)
        db.session.commit()

    questions_per_quiz = quiz_settings.questions_per_quiz
    duration_minutes = quiz_settings.quiz_duration_minutes

    # Get all questions - Super admin sees all, regular admin sees only their own
    if user.admin_level == 'super_admin':
        all_questions = KDMLabQuizQuestion.query.all()
    else:
        all_questions = KDMLabQuizQuestion.query.filter_by(
            admin_id=user.id).all()

    if len(all_questions) < questions_per_quiz:
        flash(
            f"You need at least {questions_per_quiz} questions to preview the quiz!", "warning")
        return redirect(url_for("kdm_lab_manage_internal_quiz"))

    # Get random questions for preview based on settings
    import random
    quiz_questions = random.sample(all_questions, questions_per_quiz)

    return render_template("kdm_lab_admin_preview_quiz.html",
                           quiz_questions=quiz_questions,
                           duration_minutes=duration_minutes,
                           current_admin=user)


@app.route("/kdm_lab_admin_submit_preview", methods=["POST"])
def kdm_lab_admin_submit_preview():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Get submitted answers
    question_ids = []
    admin_answers = {}

    for key in request.form.keys():
        if key.startswith('question_'):
            question_id = key.replace('question_', '')
            question_ids.append(question_id)
            admin_answers[question_id] = request.form[key]

    # Calculate score
    score = 0
    total_questions = len(question_ids)
    results = []

    for q_id in question_ids:
        question = KDMLabQuizQuestion.query.get(int(q_id))
        if question:
            admin_answer = admin_answers.get(q_id, 'Not Answered')
            is_correct = admin_answer == question.correct_answer
            if is_correct:
                score += 1

            results.append({
                'question': question,
                'admin_answer': admin_answer,
                'is_correct': is_correct
            })

    # Flash score message (don't save to database)
    percentage = (score / total_questions) * 100 if total_questions > 0 else 0
    flash(
        f"Preview Quiz Complete! Score: {score}/{total_questions} ({percentage:.1f}%)", "info")

    return render_template("kdm_lab_admin_preview_result.html",
                           score=score,
                           total_questions=total_questions,
                           percentage=percentage,
                           results=results,
                           current_admin=user)


@app.route("/kdm_lab_toggle_quiz_release", methods=["POST"])
def kdm_lab_toggle_quiz_release():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Super admin: toggle release for all admins
    if user.admin_level == "super_admin":
        all_quiz_settings = KDMLabQuizSettings.query.all()
        # If any are released, lock all; if all locked, release all
        any_released = any(qs.is_quiz_released for qs in all_quiz_settings)
        new_status = not any_released
        for qs in all_quiz_settings:
            qs.is_quiz_released = new_status
        db.session.commit()
        status = "released" if new_status else "locked"
        flash(f"Internal quiz has been {status} for ALL admins!", "success")
    else:
        # Regular admin: toggle only their own
        quiz_settings = KDMLabQuizSettings.query.filter_by(
            admin_id=user.id).first()
        if not quiz_settings:
            quiz_settings = KDMLabQuizSettings(admin_id=user.id)
            db.session.add(quiz_settings)
        quiz_settings.is_quiz_released = not quiz_settings.is_quiz_released
        db.session.commit()
        status = "released" if quiz_settings.is_quiz_released else "locked"
        flash(f"Internal quiz has been {status}!", "success")

    return redirect(url_for("kdm_lab_manage_internal_quiz"))


@app.route("/kdm_lab_individual_quiz_releases")
def kdm_lab_individual_quiz_releases():
    """Manage individual student quiz releases for KDM Lab Internal Quiz"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    # Ensure only admin or super_admin can access
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    # Get students based on admin level
    if current_admin.admin_level == "super_admin":
        students = User.query.filter_by(role="student").all()
    else:
        students = get_students_for_admin(current_admin)

    # Sort students by roll number
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # Get existing individual quiz releases
    individual_releases = {}
    releases = KDMLabStudentQuizRelease.query.filter_by(
        admin_id=current_admin.id
    ).all()

    for release in releases:
        individual_releases[release.student_id] = release.is_released

    # Check if quiz is bulk released
    bulk_released = False
    if current_admin.admin_level == "super_admin":
        # Check if any admin has released the quiz
        bulk_released = KDMLabQuizSettings.query.filter_by(
            is_quiz_released=True).first() is not None
    else:
        # Check if this admin has released the quiz in bulk
        quiz_settings = KDMLabQuizSettings.query.filter_by(
            admin_id=current_admin.id,
            is_quiz_released=True
        ).first()
        bulk_released = quiz_settings is not None

    return render_template(
        "kdm_lab_individual_quiz_releases.html",
        students=students,
        individual_releases=individual_releases,
        bulk_released=bulk_released,
        current_admin=current_admin
    )


@app.route("/kdm_lab_toggle_individual_quiz_release", methods=["POST"])
def kdm_lab_toggle_individual_quiz_release():
    """Toggle individual quiz release for a specific student"""
    if "loggedin" not in session:
        return jsonify({"error": "Not logged in"}), 401

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        return jsonify({"error": "Admin access required"}), 403

    try:
        data = request.get_json()
        student_id = data.get('student_id')
        is_released = data.get('is_released')

        # Find existing release record
        release_record = KDMLabStudentQuizRelease.query.filter_by(
            admin_id=current_admin.id,
            student_id=student_id
        ).first()

        if release_record:
            release_record.is_released = is_released
        else:
            # Create new release record
            release_record = KDMLabStudentQuizRelease(
                admin_id=current_admin.id,
                student_id=student_id,
                is_released=is_released
            )
            db.session.add(release_record)

        db.session.commit()
        return jsonify({"success": True, "is_released": is_released})

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@app.route("/kdm_lab_take_internal_quiz")
def kdm_lab_take_internal_quiz():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role != "student":
        flash("Student access required!", "danger")
        return redirect(url_for("login"))

    # Find which admin manages this student and check if quiz is released
    student_admin = find_admin_for_student(user.rollnumber)

    if not student_admin:
        flash("Could not find your assigned admin for the internal quiz!", "danger")
        return redirect(url_for("kdm_lab_student_dashboard"))

    # Check if quiz is released (bulk or individual)
    quiz_settings = KDMLabQuizSettings.query.filter_by(
        admin_id=student_admin.id).first()

    # Check bulk release
    bulk_released = quiz_settings and quiz_settings.is_quiz_released

    # Check individual release
    individual_release = KDMLabStudentQuizRelease.query.filter_by(
        admin_id=student_admin.id,
        student_id=user.id,
        is_released=True
    ).first()
    individual_released = individual_release is not None

    # Debug information
    print(
        f"🔍 DEBUG: Student {user.rollnumber} -> Admin {student_admin.rollnumber} ({student_admin.admin_level})")
    print(f"🔍 DEBUG: Quiz settings exists: {quiz_settings is not None}")
    if quiz_settings:
        print(f"🔍 DEBUG: Bulk released: {bulk_released}")
    print(f"🔍 DEBUG: Individual released: {individual_released}")

    # Student can access quiz if either bulk released OR individually released
    quiz_accessible = bulk_released or individual_released

    if not quiz_accessible:
        if not quiz_settings:
            flash(
                f"No quiz settings found for your admin ({student_admin.rollnumber})!", "warning")
        else:
            flash(
                f"Internal quiz is not available to you. Contact your admin ({student_admin.rollnumber}) for access!", "warning")
        return redirect(url_for("kdm_lab_student_dashboard"))

    # Check student's attempt history (maximum 2 attempts)
    completed_attempts = KDMLabQuizAttempt.query.filter_by(
        student_id=user.id,
        admin_id=student_admin.id,
        is_completed=True
    ).count()

    if completed_attempts >= 2:
        flash("You have already used your maximum 2 attempts for this quiz!", "warning")
        return redirect(url_for("kdm_lab_student_dashboard"))

    # Check if student has an ongoing attempt
    ongoing_attempt = KDMLabQuizAttempt.query.filter_by(
        student_id=user.id,
        admin_id=student_admin.id,
        is_completed=False
    ).first()

    # Get quiz settings for number of questions
    questions_per_quiz = quiz_settings.questions_per_quiz
    duration_minutes = quiz_settings.quiz_duration_minutes

    # Get random questions based on admin settings
    all_questions = KDMLabQuizQuestion.query.all()

    if len(all_questions) < questions_per_quiz:
        flash(
            f"Not enough questions available for the quiz! Need at least {questions_per_quiz} questions.", "danger")
        return redirect(url_for("kdm_lab_student_dashboard"))

    import random
    quiz_questions = random.sample(all_questions, questions_per_quiz)

    # Use ongoing attempt or create new one
    if ongoing_attempt:
        # Resume existing attempt
        attempt = ongoing_attempt
        question_ids = attempt.quiz_questions.split(',')
        quiz_questions = []
        for q_id in question_ids:
            question = KDMLabQuizQuestion.query.get(int(q_id))
            if question:
                quiz_questions.append(question)
    else:
        # Create new attempt record
        attempt_number = completed_attempts + 1
        attempt = KDMLabQuizAttempt(
            student_id=user.id,
            admin_id=student_admin.id,
            quiz_questions=','.join([str(q.id) for q in quiz_questions]),
            total_questions=questions_per_quiz,
            attempt_number=attempt_number
        )
        db.session.add(attempt)
        db.session.commit()

    return render_template("kdm_lab_take_internal_quiz.html",
                           quiz_questions=quiz_questions,
                           attempt=attempt,
                           duration_minutes=duration_minutes)


@app.route("/kdm_lab_submit_internal_quiz", methods=["POST"])
def kdm_lab_submit_internal_quiz():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role != "student":
        flash("Student access required!", "danger")
        return redirect(url_for("login"))

    attempt_id = request.form.get("attempt_id")
    attempt = KDMLabQuizAttempt.query.get_or_404(attempt_id)

    if attempt.student_id != user.id:
        flash("Invalid attempt!", "danger")
        return redirect(url_for("kdm_lab_student_dashboard"))

    if attempt.is_completed:
        flash("Quiz already completed!", "info")
        return redirect(url_for("kdm_lab_student_dashboard"))

    # Get student answers
    student_answers = {}
    question_ids = attempt.quiz_questions.split(',')

    for q_id in question_ids:
        answer = request.form.get(f"question_{q_id}")
        if answer:
            student_answers[q_id] = answer

    # Calculate score
    score = 0
    total_questions = len(question_ids)

    for q_id in question_ids:
        question = KDMLabQuizQuestion.query.get(int(q_id))
        if question and student_answers.get(q_id) == question.correct_answer:
            score += 1

    # Update attempt record
    attempt.student_answers = str(student_answers)
    attempt.score = score
    attempt.end_time = db.func.current_timestamp()
    attempt.is_completed = True
    db.session.commit()

    flash(f"Quiz completed! Your score: {score}/{total_questions}", "success")
    return redirect(url_for("kdm_lab_internal_quiz_result", attempt_id=attempt.id))


@app.route("/kdm_lab_internal_quiz_result/<int:attempt_id>")
def kdm_lab_internal_quiz_result(attempt_id):
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user:
        flash("User not found!", "danger")
        return redirect(url_for("login"))

    attempt = KDMLabQuizAttempt.query.get_or_404(attempt_id)

    # Check permissions
    if user.role == "student" and attempt.student_id != user.id:
        flash("Access denied!", "danger")
        return redirect(url_for("login"))
    elif user.role in ["admin", "super_admin"] and attempt.admin_id != user.id:
        flash("Access denied!", "danger")
        return redirect(url_for("login"))

    # Get questions and answers
    question_ids = attempt.quiz_questions.split(',')
    questions = []
    student_answers = eval(
        attempt.student_answers) if attempt.student_answers else {}

    for q_id in question_ids:
        question = KDMLabQuizQuestion.query.get(int(q_id))
        if question:
            questions.append({
                'question': question,
                'student_answer': student_answers.get(q_id, 'Not Answered'),
                'is_correct': student_answers.get(q_id) == question.correct_answer
            })

    return render_template("kdm_lab_internal_quiz_result.html",
                           attempt=attempt,
                           questions=questions,
                           user=user)


@app.route("/kdm_lab_quiz_question_management")
def kdm_lab_quiz_question_management():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Get all questions for this admin with pagination
    page = request.args.get('page', 1, type=int)
    per_page = 10

    # Get questions - Super admin sees all, regular admin sees only their own
    if user.admin_level == 'super_admin':
        questions = KDMLabQuizQuestion.query\
            .order_by(KDMLabQuizQuestion.slno.asc())\
            .paginate(page=page, per_page=per_page, error_out=False)
    else:
        questions = KDMLabQuizQuestion.query.filter_by(admin_id=user.id)\
            .order_by(KDMLabQuizQuestion.slno.asc())\
            .paginate(page=page, per_page=per_page, error_out=False)

    return render_template("kdm_lab_quiz_question_management.html",
                           questions=questions,
                           current_admin=user)


@app.route("/kdm_lab_add_quiz_question", methods=["GET", "POST"])
def kdm_lab_add_quiz_question():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    if request.method == "POST":
        try:
            question = KDMLabQuizQuestion(
                admin_id=user.id,
                slno=int(request.form.get('slno', 0)),
                question=request.form.get('question').strip(),
                option_a=request.form.get('option_a').strip(),
                option_b=request.form.get('option_b').strip(),
                option_c=request.form.get('option_c').strip(),
                option_d=request.form.get('option_d').strip(),
                correct_answer=request.form.get(
                    'correct_answer').strip().upper()
            )
            db.session.add(question)
            db.session.commit()
            flash("Question added successfully!", "success")
            return redirect(url_for("kdm_lab_quiz_question_management"))
        except Exception as e:
            db.session.rollback()
            flash(f"Error adding question: {str(e)}", "danger")

    return render_template("kdm_lab_add_quiz_question.html", current_admin=user)


@app.route("/kdm_lab_edit_quiz_question/<int:question_id>", methods=["GET", "POST"])
def kdm_lab_edit_quiz_question(question_id):
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    question = KDMLabQuizQuestion.query.get_or_404(question_id)

    # Check if this admin owns this question
    if question.admin_id != user.id:
        flash("You can only edit your own questions!", "danger")
        return redirect(url_for("kdm_lab_quiz_question_management"))

    if request.method == "POST":
        try:
            question.slno = int(request.form.get('slno', 0))
            question.question = request.form.get('question').strip()
            question.option_a = request.form.get('option_a').strip()
            question.option_b = request.form.get('option_b').strip()
            question.option_c = request.form.get('option_c').strip()
            question.option_d = request.form.get('option_d').strip()
            question.correct_answer = request.form.get(
                'correct_answer').strip().upper()

            db.session.commit()
            flash("Question updated successfully!", "success")
            return redirect(url_for("kdm_lab_quiz_question_management"))
        except Exception as e:
            db.session.rollback()
            flash(f"Error updating question: {str(e)}", "danger")

    return render_template("kdm_lab_edit_quiz_question.html",
                           question=question,
                           current_admin=user)


@app.route("/kdm_lab_delete_quiz_question/<int:question_id>", methods=["POST"])
def kdm_lab_delete_quiz_question(question_id):
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    question = KDMLabQuizQuestion.query.get_or_404(question_id)

    # Check if this admin owns this question
    if question.admin_id != user.id:
        flash("You can only delete your own questions!", "danger")
        return redirect(url_for("kdm_lab_quiz_question_management"))

    try:
        db.session.delete(question)
        db.session.commit()
        flash("Question deleted successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting question: {str(e)}", "danger")

    return redirect(url_for("kdm_lab_quiz_question_management"))


# ==========================================================
# ✅ PYTHON LAB ROUTES (Python Programming Lab Quiz System)
# ==========================================================

@app.route("/python_lab_admin_dashboard")
def python_lab_admin_dashboard():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    # Ensure only admin or super_admin can access
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    # Get quiz settings for this admin
    quiz_settings = PythonLabQuizSettings.query.filter_by(
        admin_id=current_admin.id).first()

    # Count questions per module (all questions from all admins)
    module_counts = {}
    for module_num in range(1, 6):  # Modules 1-5
        # All admins see all questions from all admins (shared question pool)
        count = PythonLabQuizQuestion.query.filter_by(
            module_number=module_num).count()
        module_counts[f"module_{module_num}"] = count

    total_questions = sum(module_counts.values())

    # Get students based on admin level
    if current_admin.admin_level == "super_admin":
        students = User.query.filter_by(role="student").all()
    else:
        students = get_students_for_admin(current_admin)

    # Sort students by roll number in ascending order
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # Add quiz attempt data and release status for each student
    for student in students:
        student_attempts = PythonLabQuizAttempt.query.filter_by(
            student_id=student.id, is_completed=True).all()
        student.quiz_attempts_count = len(student_attempts)
        if student_attempts:
            student.best_quiz_score = max(
                attempt.score for attempt in student_attempts)
            student.has_attempted = True
        else:
            student.best_quiz_score = 0
            student.has_attempted = False

        # Check individual release status for this student
        release_record = PythonLabStudentQuizRelease.query.filter_by(
            admin_id=current_admin.id, student_id=student.id).first()
        student.is_individually_released = release_record.is_released if release_record else False

    # Get quiz attempts for analysis
    if current_admin.admin_level == "super_admin":
        quiz_attempts = PythonLabQuizAttempt.query.filter_by(
            is_completed=True).all()
    else:
        quiz_attempts = PythonLabQuizAttempt.query.filter_by(
            admin_id=current_admin.id, is_completed=True).all()

    # Calculate statistics
    total_attempts = len(quiz_attempts)
    if total_attempts > 0:
        scores = [attempt.score for attempt in quiz_attempts]
        avg_score = sum(scores) / len(scores)
        max_score = max(scores)
        min_score = min(scores)
    else:
        avg_score = max_score = min_score = 0

    # Module definitions for display
    modules = {
        1: "Python Basics",
        2: "Operators and Control Flow",
        3: "Data Structures and Manipulation",
        4: "Functions and Modules",
        5: "Object-Oriented Programming"
    }

    # Get Python Lab settings
    lab_settings = PythonLabSettings.query.first()
    if not lab_settings:
        lab_settings = PythonLabSettings(num_experiments=12, num_criteria=10,
                                         student_visibility=False, admin_edit_locked=False)
        db.session.add(lab_settings)
        db.session.commit()

    # Get experiments
    experiments = PythonLabExperiment.query.order_by(
        PythonLabExperiment.experiment_number).all()

    # Get global release status (super admin releases)
    global_releases = {}
    for exp in experiments:
        global_releases[exp.id] = exp.is_released

    # Get admin-specific releases
    admin_releases = {}
    if current_admin.admin_level != 'super_admin':
        admin_release_records = PythonLabAdminRelease.query.filter_by(
            admin_id=current_admin.id).all()
        for rel in admin_release_records:
            admin_releases[rel.experiment_id] = rel.is_released

    return render_template("python_lab_admin_dashboard.html",
                           current_admin=current_admin,
                           quiz_settings=quiz_settings,
                           lab_settings=lab_settings,
                           experiments=experiments,
                           global_releases=global_releases,
                           admin_releases=admin_releases,
                           module_counts=module_counts,
                           total_questions=total_questions,
                           students=students,
                           total_attempts=total_attempts,
                           avg_score=avg_score,
                           max_score=max_score,
                           min_score=min_score,
                           modules=modules)


# ==========================================================
# ✅ PYTHON LAB EXPERIMENTS ROUTES (Similar to KDM Lab)
# ==========================================================

@app.route("/set_python_lab_settings", methods=["GET", "POST"])
def set_python_lab_settings():
    """Set Python Lab settings (Super Admin only)"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.admin_level != "super_admin":
        flash("Super Admin access required!", "danger")
        return redirect(url_for("python_lab_admin_dashboard"))

    lab_settings = PythonLabSettings.query.first()
    if not lab_settings:
        lab_settings = PythonLabSettings(num_experiments=12, num_criteria=10)
        db.session.add(lab_settings)
        db.session.commit()

    if request.method == "POST":
        num_experiments = int(request.form.get("num_experiments", 12))
        num_criteria = int(request.form.get("num_criteria", 10))

        lab_settings.num_experiments = num_experiments
        lab_settings.num_criteria = num_criteria
        db.session.commit()

        # Ensure experiments exist
        _ensure_python_experiments_exist(num_experiments, num_criteria)

        flash("Python Lab settings updated successfully!", "success")
        return redirect(url_for("python_lab_admin_dashboard"))

    # Get current experiments for display
    experiments = PythonLabExperiment.query.order_by(
        PythonLabExperiment.experiment_number).all()

    return render_template("set_python_lab_settings.html",
                           lab_settings=lab_settings,
                           experiments=experiments)


@app.route("/toggle_python_lab_visibility", methods=["POST"])
def toggle_python_lab_visibility():
    """Toggle Python Lab visibility for students"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.admin_level != "super_admin":
        flash("Super Admin access required!", "danger")
        return redirect(url_for("python_lab_admin_dashboard"))

    lab_settings = PythonLabSettings.query.first()
    if lab_settings:
        lab_settings.student_visibility = not lab_settings.student_visibility
        db.session.commit()
        status = "visible" if lab_settings.student_visibility else "hidden"
        flash(f"Python Lab is now {status} to students!", "success")

    return redirect(url_for("python_lab_admin_dashboard"))


@app.route("/toggle_python_lab_edit_lock", methods=["POST"])
def toggle_python_lab_edit_lock():
    """Toggle Python Lab admin edit lock"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.admin_level != "super_admin":
        flash("Super Admin access required!", "danger")
        return redirect(url_for("python_lab_admin_dashboard"))

    lab_settings = PythonLabSettings.query.first()
    if lab_settings:
        lab_settings.admin_edit_locked = not lab_settings.admin_edit_locked
        db.session.commit()
        status = "locked" if lab_settings.admin_edit_locked else "unlocked"
        flash(f"Python Lab admin editing is now {status}!", "success")

    return redirect(url_for("python_lab_admin_dashboard"))


@app.route("/toggle_python_experiment_release/<int:experiment_id>", methods=["POST"])
def toggle_python_experiment_release(experiment_id):
    """Toggle Python Lab experiment release status"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required!", "danger")
        return redirect(url_for("login"))

    experiment = PythonLabExperiment.query.get_or_404(experiment_id)

    if current_admin.admin_level == "super_admin":
        # Super admin toggles global release
        experiment.is_released = not experiment.is_released
        db.session.commit()
        status = "released" if experiment.is_released else "locked"
        flash(f"{experiment.title} globally {status}!", "success")
    else:
        # Sub admin toggles their specific release
        admin_release = PythonLabAdminRelease.query.filter_by(
            admin_id=current_admin.id, experiment_id=experiment_id
        ).first()

        if not admin_release:
            admin_release = PythonLabAdminRelease(
                admin_id=current_admin.id,
                experiment_id=experiment_id,
                is_released=True
            )
            db.session.add(admin_release)
        else:
            admin_release.is_released = not admin_release.is_released

        db.session.commit()
        status = "released" if admin_release.is_released else "locked"
        flash(f"{experiment.title} {status} for your students!", "success")

    return redirect(url_for("python_lab_admin_dashboard"))


@app.route("/python_lab_individual_releases/<int:experiment_id>")
def python_lab_individual_releases(experiment_id):
    """Manage individual student releases for Python Lab experiment"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required!", "danger")
        return redirect(url_for("login"))

    experiment = PythonLabExperiment.query.get_or_404(experiment_id)

    # Get students for this admin
    if current_admin.admin_level == "super_admin":
        students = User.query.filter_by(role="student").all()
    else:
        students = get_students_for_admin(current_admin)

    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # Get individual release records
    individual_releases = {}
    release_records = PythonLabStudentExperimentRelease.query.filter_by(
        admin_id=current_admin.id, experiment_id=experiment_id
    ).all()

    for record in release_records:
        individual_releases[record.student_id] = record.is_released

    return render_template("python_lab_individual_releases.html",
                           experiment=experiment,
                           students=students,
                           individual_releases=individual_releases,
                           current_admin=current_admin)


@app.route("/python_lab_toggle_individual_release/<int:experiment_id>/<int:student_id>", methods=["POST"])
def python_lab_toggle_individual_release(experiment_id, student_id):
    """Toggle individual student release for Python Lab experiment"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required!", "danger")
        return redirect(url_for("login"))

    # Check if release record exists
    release_record = PythonLabStudentExperimentRelease.query.filter_by(
        admin_id=current_admin.id,
        student_id=student_id,
        experiment_id=experiment_id
    ).first()

    if not release_record:
        release_record = PythonLabStudentExperimentRelease(
            admin_id=current_admin.id,
            student_id=student_id,
            experiment_id=experiment_id,
            is_released=True
        )
        db.session.add(release_record)
    else:
        release_record.is_released = not release_record.is_released

    db.session.commit()
    return redirect(url_for("python_lab_individual_releases", experiment_id=experiment_id))


@app.route("/python_lab_bulk_release_experiment/<int:experiment_id>/<action>", methods=["POST"])
def python_lab_bulk_release_experiment(experiment_id, action):
    """Bulk release/lock Python Lab experiment for all students"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required!", "danger")
        return redirect(url_for("login"))

    experiment = PythonLabExperiment.query.get_or_404(experiment_id)
    is_released = (action == "release")

    # Get students for this admin
    if current_admin.admin_level == "super_admin":
        students = User.query.filter_by(role="student").all()
    else:
        students = get_students_for_admin(current_admin)

    count = 0
    for student in students:
        release_record = PythonLabStudentExperimentRelease.query.filter_by(
            admin_id=current_admin.id,
            student_id=student.id,
            experiment_id=experiment_id
        ).first()

        if not release_record:
            release_record = PythonLabStudentExperimentRelease(
                admin_id=current_admin.id,
                student_id=student.id,
                experiment_id=experiment_id,
                is_released=is_released
            )
            db.session.add(release_record)
        else:
            release_record.is_released = is_released
        count += 1

    db.session.commit()
    action_text = "released" if is_released else "locked"
    flash(f"{experiment.title} {action_text} for {count} students!", "success")

    return redirect(url_for("python_lab_individual_releases", experiment_id=experiment_id))


# Helper function for Python Lab experiments
def _ensure_python_experiments_exist(num_experiments, num_criteria):
    """Ensure Python Lab experiments and criteria exist"""
    existing = {e.experiment_number: e for e in PythonLabExperiment.query.all()}

    # Create missing experiments
    for i in range(1, num_experiments + 1):
        if i not in existing:
            ex = PythonLabExperiment(
                experiment_number=i, title=f"EX-{i}", is_released=False)
            db.session.add(ex)
            db.session.commit()

            # Create criteria for this experiment
            for j in range(1, num_criteria + 1):
                criteria = PythonLabCriteria(
                    experiment_id=ex.id,
                    criteria_number=j,
                    question_text=f"Criteria {j} for {ex.title}",
                    option_a="Option A",
                    option_b="Option B",
                    option_c="Option C",
                    option_d="Option D",
                    marks_a=4.0, marks_b=3.0, marks_c=2.0, marks_d=0.0,
                    max_marks=4.0,
                    withheld=False
                )
                db.session.add(criteria)

    # Remove extra experiments
    for ex in PythonLabExperiment.query.all():
        if ex.experiment_number > num_experiments:
            db.session.delete(ex)

    db.session.commit()


@app.route("/python_lab_release_all_experiments", methods=["POST"])
def python_lab_release_all_experiments():
    """Release all Python Lab experiments for all students (bulk action)"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required!", "danger")
        return redirect(url_for("login"))

    # Release all experiments globally
    experiments = PythonLabExperiment.query.all()
    count = 0
    for experiment in experiments:
        if not experiment.is_released:
            experiment.is_released = True
            count += 1

    db.session.commit()
    flash(f"Released {count} experiments for all students!", "success")
    return redirect(url_for("python_lab_admin_dashboard"))


@app.route("/python_lab_lock_all_experiments", methods=["POST"])
def python_lab_lock_all_experiments():
    """Lock all Python Lab experiments for all students (bulk action)"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required!", "danger")
        return redirect(url_for("login"))

    # Lock all experiments globally
    experiments = PythonLabExperiment.query.all()
    count = 0
    for experiment in experiments:
        if experiment.is_released:
            experiment.is_released = False
            count += 1

    db.session.commit()
    flash(f"Locked {count} experiments for all students!", "success")
    return redirect(url_for("python_lab_admin_dashboard"))


@app.route("/create_python_lab_quiz", methods=["GET", "POST"])
def create_python_lab_quiz():
    """Create/Edit Python Lab experiment criteria (quiz questions) similar to KDM Lab"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    lab_settings = PythonLabSettings.query.first()
    if not lab_settings:
        flash("Python Lab settings not found. Please set them first.", "warning")
        return redirect(url_for("python_lab_admin_dashboard"))

    experiments = PythonLabExperiment.query.order_by(
        PythonLabExperiment.experiment_number).all()

    selected_experiment_number = None
    experiment = None
    existing_criteria = []

    # ------------------- POST -------------------
    if request.method == "POST":
        try:
            selected_experiment_number = int(
                request.form.get("experiment_number"))
            title = request.form.get("title").strip() if request.form.get(
                "title") else f"EX-{selected_experiment_number}"
            criteria_count = int(request.form.get(
                "criteria_count") or lab_settings.num_criteria)

            # Check if we have any criteria data to save
            has_criteria_data = any(request.form.get(
                f"question_{i}") for i in range(1, criteria_count + 1))

            if not has_criteria_data:
                flash(
                    "No criteria data found - please fill in at least one question before saving.", "warning")
                return redirect(url_for("create_python_lab_quiz", experiment=selected_experiment_number))

            experiment = PythonLabExperiment.query.filter_by(
                experiment_number=selected_experiment_number).first()
            if not experiment:
                experiment = PythonLabExperiment(
                    experiment_number=selected_experiment_number,
                    title=title,
                    is_released=False
                )
                db.session.add(experiment)
                db.session.flush()  # Get the experiment ID
            else:
                experiment.title = title

            db.session.commit()

            # --- Save / update criteria ---
            saved_count = 0
            for i in range(1, criteria_count + 1):
                q_text = request.form.get(f"question_{i}")
                if not q_text:
                    continue

                existing = PythonLabCriteria.query.filter_by(
                    experiment_id=experiment.id, criteria_number=i
                ).first()

                if not existing:
                    existing = PythonLabCriteria(
                        experiment_id=experiment.id, criteria_number=i)

                existing.question_text = q_text
                existing.option_a = request.form.get(f"optionA_{i}")
                existing.option_b = request.form.get(f"optionB_{i}")
                existing.option_c = request.form.get(f"optionC_{i}")
                existing.option_d = request.form.get(f"optionD_{i}")
                existing.marks_a = float(request.form.get(f"marksA_{i}") or 0)
                existing.marks_b = float(request.form.get(f"marksB_{i}") or 0)
                existing.marks_c = float(request.form.get(f"marksC_{i}") or 0)
                existing.marks_d = float(request.form.get(f"marksD_{i}") or 0)
                existing.max_marks = float(
                    request.form.get(f"max_marks_{i}") or 4)
                # --- Handle Withhold (Admin Only) checkbox ---
                checkbox_name = f"withhold_{i}"
                if checkbox_name in request.form:
                    existing.withheld = True
                else:
                    existing.withheld = False

                db.session.add(existing)
                saved_count += 1

            db.session.commit()
            flash(
                f"Python Lab Quiz saved successfully for EX-{selected_experiment_number}.", "success")

            # ✅ Redirect to same page with ?experiment=<saved_experiment_number>
            return redirect(url_for("create_python_lab_quiz", experiment=selected_experiment_number))

        except ValueError as e:
            flash(f"Invalid form data error: {str(e)}", "danger")
            return redirect(url_for("create_python_lab_quiz"))
        except Exception as e:
            db.session.rollback()
            flash(f"Error saving quiz: {str(e)}", "danger")
            return redirect(url_for("create_python_lab_quiz", experiment=selected_experiment_number))

    # ------------------- GET -------------------
    selected_experiment_number = request.args.get("experiment")
    if selected_experiment_number:
        experiment = PythonLabExperiment.query.filter_by(
            experiment_number=int(selected_experiment_number)
        ).first()
        if experiment:
            existing_criteria = PythonLabCriteria.query.filter_by(
                experiment_id=experiment.id
            ).order_by(PythonLabCriteria.criteria_number).all()

    return render_template(
        "create_python_lab_quiz.html",
        current_admin=current_admin,
        lab_settings=lab_settings,
        experiments=experiments,
        experiment=experiment,
        existing_criteria=existing_criteria
    )


@app.route("/python_lab_reset_all_users", methods=["POST"])
def python_lab_reset_all_users():
    """Reset all Python Lab experiment data for all users (Super Admin only)"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.admin_level != "super_admin":
        flash("Super admin access required!", "danger")
        return redirect(url_for("login"))

    try:
        # Delete all experiment responses and attempts
        PythonLabResponse.query.delete()
        PythonLabAttempt.query.delete()
        PythonLabManualMarks.query.delete()
        PythonLabStudentExperimentRelease.query.delete()

        db.session.commit()
        flash("All Python Lab experiment data has been reset!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error resetting data: {str(e)}", "danger")

    return redirect(url_for("python_lab_admin_dashboard"))


@app.route("/kdm_lab_reset_all_users", methods=["POST"])
def kdm_lab_reset_all_users():
    if not _require_login():
        return redirect(url_for("login"))
    current_admin = _current_admin_user()
    if not _is_super_admin(current_admin):
        flash("Only super admin can reset all lab users.", "danger")
        return redirect(url_for("kdm_lab_admin_dashboard"))

    try:
        KDMLabResponse.query.delete()
        KDMLabManualMarks.query.delete()
        # Keep student rows but zero out totals
        for s in KDMLabStudent.query.all():
            s.total_lab_score = 0.0
            s.completed_experiments = 0
        db.session.commit()
        flash("All KDM LAB student data reset.", "success")
    except Exception as e:
        db.session.rollback()
        print("❌ Error resetting lab data:", e)
        flash("Error resetting lab data.", "danger")

    return redirect(url_for("kdm_lab_admin_dashboard"))


@app.route("/reset_python_lab_experiment/<int:experiment_id>", methods=["POST"])
def reset_python_lab_experiment(experiment_id):
    """Reset specific Python Lab experiment data for all users"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required!", "danger")
        return redirect(url_for("login"))

    experiment = PythonLabExperiment.query.get_or_404(experiment_id)

    try:
        # Delete all data for this specific experiment
        PythonLabResponse.query.filter_by(experiment_id=experiment_id).delete()
        PythonLabAttempt.query.filter_by(experiment_id=experiment_id).delete()
        PythonLabManualMarks.query.filter_by(
            experiment_id=experiment_id).delete()
        PythonLabStudentExperimentRelease.query.filter_by(
            experiment_id=experiment_id).delete()

        db.session.commit()
        flash(f"All data for {experiment.title} has been reset!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error resetting {experiment.title}: {str(e)}", "danger")

    return redirect(url_for("python_lab_admin_dashboard"))


@app.route("/delete_python_lab_experiment", methods=["POST"])
def delete_python_lab_experiment():
    """Delete a specific Python Lab experiment permanently (Super Admin only)"""
    if "loggedin" not in session:
        return {"success": False, "message": "Not logged in"}

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.admin_level != "super_admin":
        return {"success": False, "message": "Super Admin access required"}

    try:
        data = request.get_json()
        experiment_id = data.get("experiment_id")

        if not experiment_id:
            return {"success": False, "message": "Experiment ID required"}

        experiment = PythonLabExperiment.query.get(experiment_id)
        if not experiment:
            return {"success": False, "message": "Experiment not found"}

        experiment_title = experiment.title

        # Delete all related data first (to avoid foreign key constraints)
        PythonLabCriteria.query.filter_by(experiment_id=experiment_id).delete()
        PythonLabResponse.query.filter_by(experiment_id=experiment_id).delete()
        PythonLabAttempt.query.filter_by(experiment_id=experiment_id).delete()
        PythonLabManualMarks.query.filter_by(
            experiment_id=experiment_id).delete()
        PythonLabStudentExperimentRelease.query.filter_by(
            experiment_id=experiment_id).delete()
        PythonLabAdminRelease.query.filter_by(
            experiment_id=experiment_id).delete()

        # Delete the experiment itself
        db.session.delete(experiment)
        db.session.commit()

        return {
            "success": True,
            "message": f"Successfully deleted {experiment_title} and all related data!"
        }

    except Exception as e:
        db.session.rollback()
        return {
            "success": False,
            "message": f"Error deleting experiment: {str(e)}"
        }


@app.route("/delete_all_python_lab_experiments", methods=["POST"])
def delete_all_python_lab_experiments():
    """Delete ALL Python Lab experiments permanently (Super Admin only)"""
    if "loggedin" not in session:
        return {"success": False, "message": "Not logged in"}

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.admin_level != "super_admin":
        return {"success": False, "message": "Super Admin access required"}

    try:
        # Delete all related data first (to avoid foreign key constraints)
        PythonLabCriteria.query.delete()
        PythonLabResponse.query.delete()
        PythonLabAttempt.query.delete()
        PythonLabManualMarks.query.delete()
        PythonLabStudentExperimentRelease.query.delete()
        PythonLabAdminRelease.query.delete()

        # Get count before deletion
        experiment_count = PythonLabExperiment.query.count()

        # Delete all experiments
        PythonLabExperiment.query.delete()
        db.session.commit()

        return {
            "success": True,
            "message": f"Successfully deleted all {experiment_count} experiments and related data!"
        }

    except Exception as e:
        db.session.rollback()
        return {
            "success": False,
            "message": f"Error deleting all experiments: {str(e)}"
        }


@app.route("/delete_kdm_lab_experiment/<int:experiment_id>", methods=["DELETE"])
def delete_kdm_lab_experiment(experiment_id):
    """Delete a specific KDM Lab experiment permanently (Super Admin only)"""
    if "loggedin" not in session:
        return {"success": False, "message": "Not logged in"}

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.admin_level != "super_admin":
        return {"success": False, "message": "Super Admin access required"}

    try:
        experiment = KDMLabExperiment.query.get(experiment_id)
        if not experiment:
            return {"success": False, "message": "Experiment not found"}

        experiment_title = experiment.title

        # Delete all related data first (to avoid foreign key constraints)
        KDMLabCriteria.query.filter_by(experiment_id=experiment_id).delete()
        KDMLabResponse.query.filter_by(experiment_id=experiment_id).delete()
        KDMLabAttempts.query.filter_by(experiment_id=experiment_id).delete()
        KDMLabManualMarks.query.filter_by(experiment_id=experiment_id).delete()
        KDMLabStudentExperimentRelease.query.filter_by(
            experiment_id=experiment_id).delete()
        KDMLabAdminRelease.query.filter_by(
            experiment_id=experiment_id).delete()

        # Delete the experiment itself
        db.session.delete(experiment)
        db.session.commit()

        return {
            "success": True,
            "message": f"Successfully deleted {experiment_title} and all related data!"
        }

    except Exception as e:
        db.session.rollback()
        return {
            "success": False,
            "message": f"Error deleting experiment: {str(e)}"
        }


@app.route("/delete_all_kdm_lab_experiments", methods=["DELETE"])
def delete_all_kdm_lab_experiments():
    """Delete ALL KDM Lab experiments permanently (Super Admin only)"""
    if "loggedin" not in session:
        return {"success": False, "message": "Not logged in"}

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.admin_level != "super_admin":
        return {"success": False, "message": "Super Admin access required"}

    try:
        # Delete all related data first (to avoid foreign key constraints)
        KDMLabCriteria.query.delete()
        KDMLabResponse.query.delete()
        KDMLabAttempts.query.delete()
        KDMLabManualMarks.query.delete()
        KDMLabStudentExperimentRelease.query.delete()
        KDMLabAdminRelease.query.delete()

        # Get count before deletion
        experiment_count = KDMLabExperiment.query.count()

        # Delete all experiments
        KDMLabExperiment.query.delete()
        db.session.commit()

        return {
            "success": True,
            "message": f"Successfully deleted all {experiment_count} KDM Lab experiments and related data!"
        }

    except Exception as e:
        db.session.rollback()
        return {
            "success": False,
            "message": f"Error deleting all experiments: {str(e)}"
        }


@app.route("/view_python_lab_submissions/<int:experiment_id>")
def view_python_lab_submissions(experiment_id):
    """View all student submissions for a specific Python Lab experiment"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required!", "danger")
        return redirect(url_for("login"))

    experiment = PythonLabExperiment.query.get_or_404(experiment_id)

    # Get all student responses for this experiment
    submissions = db.session.query(
        User.rollnumber, User.email,
        PythonLabResponse.user_id,
        db.func.count(PythonLabResponse.id).label('total_responses'),
        db.func.sum(PythonLabResponse.marks_earned).label('total_marks')
    ).join(
        PythonLabResponse, User.id == PythonLabResponse.user_id
    ).filter(
        PythonLabResponse.experiment_id == experiment_id,
        User.role == 'student'
    ).group_by(
        User.id, User.rollnumber, User.email, PythonLabResponse.user_id
    ).all()

    return render_template("view_python_lab_submissions.html",
                           current_admin=current_admin,
                           experiment=experiment,
                           submissions=submissions)


@app.route("/fill_python_lab_manual_marks/<int:experiment_id>", methods=["GET", "POST"])
def fill_python_lab_manual_marks(experiment_id):
    """Fill manual marks for a specific Python Lab experiment"""
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin:
        flash("Admin not found.", "danger")
        return redirect(url_for("python_lab_admin_dashboard"))

    experiment = PythonLabExperiment.query.get_or_404(experiment_id)

    # ✅ Get withheld criteria only
    withheld_criteria = PythonLabCriteria.query.filter_by(
        experiment_id=experiment.id, withheld=True
    ).order_by(PythonLabCriteria.criteria_number).all()

    # ✅ Get students in admin's scope
    if current_admin.admin_level == "super_admin":
        students = User.query.filter_by(role="student").all()
    else:
        student_rolls = [
            s.rollnumber for s in get_students_for_admin(current_admin)]
        students = User.query.filter(User.rollnumber.in_(student_rolls)).all()

    # Sort students by roll number in ascending order
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # ✅ Load existing manual marks (by student_id)
    existing_marks = {}
    manual_marks = PythonLabManualMarks.query.filter_by(
        experiment_id=experiment.id).all()
    for mark in manual_marks:
        existing_marks[(mark.student_id, mark.criteria_id)] = mark.marks_given

    # ✅ Handle form submission
    if request.method == "POST":
        bulk_criteria_id = request.form.get("bulk_criteria_id")
        bulk_mark = request.form.get("bulk_mark")

        # --- Bulk marks applied to all students ---
        if bulk_criteria_id and bulk_mark:
            for student in students:
                existing = PythonLabManualMarks.query.filter_by(
                    student_id=student.id,
                    experiment_id=experiment.id,
                    criteria_id=bulk_criteria_id
                ).first()
                if existing:
                    existing.marks_given = float(bulk_mark)
                else:
                    db.session.add(PythonLabManualMarks(
                        student_id=student.id,
                        experiment_id=experiment.id,
                        criteria_id=bulk_criteria_id,
                        marks_given=float(bulk_mark)
                    ))
            db.session.commit()
            flash("✅ Bulk marks applied successfully!", "success")
            return redirect(url_for("fill_python_lab_manual_marks", experiment_id=experiment.id))

        # --- Individual marks per student ---
        for student in students:
            for c in withheld_criteria:
                field_name = f"mark_{student.id}_{c.id}"
                mark_val = request.form.get(field_name)
                if mark_val is not None and mark_val.strip() != "":
                    existing = PythonLabManualMarks.query.filter_by(
                        student_id=student.id,
                        experiment_id=experiment.id,
                        criteria_id=c.id
                    ).first()
                    if existing:
                        existing.marks_given = float(mark_val)
                    else:
                        db.session.add(PythonLabManualMarks(
                            student_id=student.id,
                            experiment_id=experiment.id,
                            criteria_id=c.id,
                            marks_given=float(mark_val)
                        ))

        db.session.commit()
        flash("✅ Manual marks saved successfully!", "success")
        return redirect(url_for("fill_python_lab_manual_marks", experiment_id=experiment.id))

    # ✅ Render page
    return render_template(
        "fill_python_lab_manual_marks.html",
        experiment=experiment,
        withheld_criteria=withheld_criteria,
        students=students,
        manual_marks=existing_marks,
    )


@app.route("/fill_python_lab_manual_marks_all_experiments", methods=["GET", "POST"])
def fill_python_lab_manual_marks_all_experiments():
    """Fill manual marks for all Python Lab experiments"""
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin:
        flash("Admin not found.", "danger")
        return redirect(url_for("python_lab_admin_dashboard"))

    # Get all experiments
    experiments = PythonLabExperiment.query.order_by(
        PythonLabExperiment.experiment_number).all()

    # Get all withheld criteria across all experiments
    all_withheld_criteria = []
    for exp in experiments:
        criteria = PythonLabCriteria.query.filter_by(
            experiment_id=exp.id, withheld=True
        ).order_by(PythonLabCriteria.criteria_number).all()
        for c in criteria:
            all_withheld_criteria.append({
                'id': c.id,
                'criteria': c,
                'experiment': exp
            })

    # Get students in admin's scope
    if current_admin.admin_level == "super_admin":
        students = User.query.filter_by(role="student").all()
    else:
        student_rolls = [
            s.rollnumber for s in get_students_for_admin(current_admin)]
        students = User.query.filter(User.rollnumber.in_(student_rolls)).all()

    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # Load existing manual marks
    manual_marks = {}
    marks = PythonLabManualMarks.query.all()
    for mark in marks:
        manual_marks[(mark.student_id, mark.criteria_id)] = mark.marks_given

    # Handle POST request
    if request.method == "POST":
        try:
            # Process each student's marks for each criteria
            for student in students:
                for item in all_withheld_criteria:
                    field_name = f"mark_{student.id}_{item['id']}"
                    mark_val = request.form.get(field_name)

                    if mark_val is not None and mark_val.strip() != "":
                        existing = PythonLabManualMarks.query.filter_by(
                            student_id=student.id,
                            experiment_id=item['experiment'].id,
                            criteria_id=item['id']
                        ).first()

                        if existing:
                            existing.marks_given = float(mark_val)
                        else:
                            db.session.add(PythonLabManualMarks(
                                student_id=student.id,
                                experiment_id=item['experiment'].id,
                                criteria_id=item['id'],
                                marks_given=float(mark_val)
                            ))

            db.session.commit()
            flash("✅ All manual marks saved successfully!", "success")
            return redirect(url_for("fill_python_lab_manual_marks_all_experiments"))

        except Exception as e:
            db.session.rollback()
            flash(f"❌ Error saving marks: {str(e)}", "danger")

    return render_template("fill_python_lab_manual_marks_all_experiments.html",
                           current_admin=current_admin,
                           experiments=experiments,
                           all_withheld_criteria=all_withheld_criteria,
                           students=students,
                           manual_marks=manual_marks)


@app.route("/attempt_python_lab_experiment/<int:experiment_id>")
def attempt_python_lab_experiment(experiment_id):
    """Student and Admin interface to attempt Python Lab experiment"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["student", "admin", "super_admin"]:
        flash("Access denied!", "danger")
        return redirect(url_for("login"))

    experiment = PythonLabExperiment.query.get_or_404(experiment_id)

    # For students: Check if experiment is available
    if user.role == "student":
        if not experiment.is_released:
            # Check individual releases
            individual_release = PythonLabStudentExperimentRelease.query.filter_by(
                student_id=user.id, experiment_id=experiment_id, is_released=True
            ).first()
            if not individual_release:
                flash("This experiment is not available yet.", "warning")
                return redirect(url_for("python_lab_student_dashboard"))

    # For admins: No restrictions, they can attempt any experiment for testing/preview

    # Get experiment criteria (exclude withheld criteria - they are marked separately by instructor)
    criteria = PythonLabCriteria.query.filter_by(
        experiment_id=experiment_id, withheld=False
    ).order_by(PythonLabCriteria.criteria_number).all()

    # Shuffle options for each criteria (like KDM Lab)
    import random
    randomized_criteria = []
    for c in criteria:
        options = [
            {"key": "A", "text": c.option_a or "", "marks": c.marks_a or 0},
            {"key": "B", "text": c.option_b or "", "marks": c.marks_b or 0},
            {"key": "C", "text": c.option_c or "", "marks": c.marks_c or 0},
            {"key": "D", "text": c.option_d or "", "marks": c.marks_d or 0},
        ]
        options = [opt for opt in options if opt["text"].strip()]
        random.shuffle(options)
        randomized_criteria.append({"criteria": c, "options": options})

    # Get existing responses for this user and experiment (if any)
    existing_responses = {}
    responses = PythonLabResponse.query.filter_by(
        user_id=user.id, experiment_id=experiment_id
    ).all()
    for response in responses:
        existing_responses[response.criteria_id] = response.selected_option

    # Get attempt record if it exists
    attempt_record = PythonLabAttempt.query.filter_by(
        student_id=user.id, experiment_id=experiment_id
    ).first()

    # Enforce single attempt for students
    if user.role == "student" and attempt_record and attempt_record.attempt_count >= 1:
        flash("You have already attempted this experiment. Only one attempt is allowed.", "warning")
        return redirect(url_for("view_my_python_lab_submission", experiment_id=experiment_id))

    return render_template("attempt_python_lab_experiment.html",
                           user=user,
                           experiment=experiment,
                           criteria=randomized_criteria,
                           existing_responses=existing_responses,
                           attempt_record=attempt_record)


@app.route("/submit_python_lab_experiment/<int:experiment_id>", methods=["POST"])
def submit_python_lab_experiment(experiment_id):
    """Submit Python Lab experiment responses"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["student", "admin", "super_admin"]:
        flash("Access denied!", "danger")
        return redirect(url_for("login"))

    experiment = PythonLabExperiment.query.get_or_404(experiment_id)

    try:
        # Delete existing responses for this experiment
        PythonLabResponse.query.filter_by(
            user_id=user.id, experiment_id=experiment_id
        ).delete()

        # Process new responses
        total_marks = 0
        criteria_responses = 0

        for key, value in request.form.items():
            if key.startswith("criteria_"):
                criteria_id = int(key.split("_")[1])
                selected_option = value

                criteria = PythonLabCriteria.query.get(criteria_id)
                if criteria:
                    # Calculate marks based on selected option (weighted calculation like KDM Lab)
                    marks_map = {
                        'A': criteria.marks_a or 0,
                        'B': criteria.marks_b or 0,
                        'C': criteria.marks_c or 0,
                        'D': criteria.marks_d or 0
                    }
                    obtained_points = marks_map.get(selected_option, 0)
                    marks_earned = (criteria.max_marks or 4) * \
                        (obtained_points / 4)

                    # Create response record
                    response = PythonLabResponse(
                        user_id=user.id,
                        experiment_id=experiment_id,
                        criteria_id=criteria_id,
                        selected_option=selected_option,
                        obtained_points=obtained_points,
                        marks_earned=marks_earned,
                        attempt_number=1
                    )
                    db.session.add(response)

                    total_marks += marks_earned
                    criteria_responses += 1

        # Update or create attempt record
        attempt = PythonLabAttempt.query.filter_by(
            student_id=user.id, experiment_id=experiment_id
        ).first()

        if attempt:
            attempt.attempt_count += 1
        else:
            attempt = PythonLabAttempt(
                student_id=user.id,
                experiment_id=experiment_id,
                attempt_count=1
            )
            db.session.add(attempt)

        db.session.commit()
        flash(
            f"Successfully submitted {experiment.title}! Total marks: {total_marks}", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error submitting experiment: {str(e)}", "danger")

    # Redirect based on user role
    if user.role in ["admin", "super_admin"]:
        return redirect(url_for("python_lab_admin_dashboard"))
    else:
        return redirect(url_for("python_lab_student_dashboard"))


@app.route("/view_my_python_lab_submission/<int:experiment_id>")
def view_my_python_lab_submission(experiment_id):
    """View student's own submission for Python Lab experiment"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role != "student":
        flash("Student access required!", "danger")
        return redirect(url_for("login"))

    experiment = PythonLabExperiment.query.get_or_404(experiment_id)

    # Get student's responses for this experiment
    responses = PythonLabResponse.query.filter_by(
        user_id=user.id, experiment_id=experiment_id
    ).join(PythonLabCriteria).order_by(PythonLabCriteria.criteria_number).all()

    if not responses:
        flash("No submission found for this experiment.", "warning")
        return redirect(url_for("python_lab_student_dashboard"))

    # Calculate quiz (auto) and manual (withheld) marks
    # Quiz marks: non-withheld criteria
    quiz_criteria = PythonLabCriteria.query.filter_by(
        experiment_id=experiment_id, withheld=False).all()
    quiz_criteria_ids = [c.id for c in quiz_criteria]
    quiz_responses = [
        r for r in responses if r.criteria_id in quiz_criteria_ids]
    total_quiz_marks = sum(r.marks_earned for r in quiz_responses)
    total_max_quiz_marks = sum(c.max_marks or 4 for c in quiz_criteria)

    # Manual marks: withheld criteria
    manual_criteria = PythonLabCriteria.query.filter_by(
        experiment_id=experiment_id, withheld=True).all()
    manual_criteria_ids = [c.id for c in manual_criteria]
    manual_marks_records = PythonLabManualMarks.query.filter_by(
        student_id=user.id, experiment_id=experiment_id).all()
    manual_marks_dict = {
        m.criteria_id: m.marks_given for m in manual_marks_records}
    # Manual marks obtained should be weighted by the criteria's max_marks
    total_manual_marks = 0
    total_max_manual_marks = 0
    for c in manual_criteria:
        given = manual_marks_dict.get(c.id)
        # max for this criteria
        crit_max = c.max_marks or 4
        total_max_manual_marks += crit_max
        if given is not None:
            # instructor gives a value out of 4; convert to weighted marks
            try:
                given_val = float(given)
            except Exception:
                given_val = 0
            weighted = crit_max * (given_val / 4.0)
            total_manual_marks += weighted

    # For backward compatibility
    total_marks = total_quiz_marks + total_manual_marks
    max_marks = total_max_quiz_marks + total_max_manual_marks
    percentage = (total_marks / max_marks * 100) if max_marks > 0 else 0

    return render_template("view_my_python_lab_submission.html",
                           user=user,
                           student=user,
                           experiment=experiment,
                           responses=responses,
                           total_marks=total_marks,
                           max_marks=max_marks,
                           percentage=percentage,
                           total_quiz_marks=total_quiz_marks,
                           total_manual_marks=total_manual_marks,
                           total_max_quiz_marks=total_max_quiz_marks,
                           total_max_manual_marks=total_max_manual_marks)


@app.route("/python_lab_toggle_student_visibility", methods=["POST"])
def python_lab_toggle_student_visibility():
    """Toggle visibility of Python Lab questions for students (Super Admin only)"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level != "super_admin":
        flash("Super Admin access required!", "danger")
        return redirect(url_for("python_lab_admin_dashboard"))

    # Get current visibility status from any settings record (they should all be the same)
    settings = PythonLabQuizSettings.query.first()
    current_status = settings.questions_visible_to_students if settings else False

    # Toggle the visibility for all settings records
    new_status = not current_status
    PythonLabQuizSettings.query.update(
        {'questions_visible_to_students': new_status})
    db.session.commit()

    status_text = "visible" if new_status else "hidden"
    flash(
        f"Python Lab questions are now {status_text} to students!", "success")
    return redirect(url_for("python_lab_admin_dashboard"))


@app.route("/python_lab_upload_quiz_questions", methods=["GET", "POST"])
def python_lab_upload_quiz_questions():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    modules = {
        1: "Python Basics",
        2: "Operators and Control Flow",
        3: "Data Structures and Manipulation",
        4: "Functions and Modules",
        5: "Object-Oriented Programming"
    }

    if request.method == "POST":
        try:
            module_number = int(request.form.get("module_number"))
            upload_mode = request.form.get(
                "upload_mode", "add")  # "add" or "replace"

            if module_number not in range(1, 6):
                flash("Invalid module number!", "danger")
                return redirect(url_for("python_lab_upload_quiz_questions"))

            if "csv_file" not in request.files:
                flash("No file uploaded!", "danger")
                return redirect(url_for("python_lab_upload_quiz_questions"))

            file = request.files["csv_file"]
            if file.filename == "":
                flash("No file selected!", "danger")
                return redirect(url_for("python_lab_upload_quiz_questions"))

            if not file.filename.lower().endswith('.csv'):
                flash("Please upload a CSV file!", "danger")
                return redirect(url_for("python_lab_upload_quiz_questions"))

            # Read and parse CSV file
            csv_content = file.read().decode('utf-8')
            csv_reader = csv.DictReader(StringIO(csv_content))

            # Validate required columns
            required_columns = ['slno', 'question', 'option_a',
                                'option_b', 'option_c', 'option_d', 'correct_answer']
            if not all(col in csv_reader.fieldnames for col in required_columns):
                flash(
                    f"CSV must contain columns: {', '.join(required_columns)}", "danger")
                return redirect(url_for("python_lab_upload_quiz_questions"))

            # If replace mode, delete existing questions for this module and admin
            if upload_mode == "replace":
                if user.admin_level == "super_admin":
                    # Super admin replaces all questions for this module
                    PythonLabQuizQuestion.query.filter_by(
                        module_number=module_number).delete()
                else:
                    # Regular admin only replaces their own questions for this module
                    PythonLabQuizQuestion.query.filter_by(
                        admin_id=user.id, module_number=module_number).delete()

            # Process CSV rows
            questions_added = 0
            for row in csv_reader:
                try:
                    # Validate correct answer
                    correct_answer = row['correct_answer'].upper().strip()
                    if correct_answer not in ['A', 'B', 'C', 'D']:
                        flash(
                            f"Invalid correct answer '{correct_answer}' in row {row.get('slno', '?')}. Must be A, B, C, or D.", "danger")
                        continue

                    question = PythonLabQuizQuestion(
                        admin_id=user.id,
                        module_number=module_number,
                        module_name=modules[module_number],
                        slno=int(row['slno']) if row['slno'].strip() else None,
                        question=row['question'].strip(),
                        option_a=row['option_a'].strip(),
                        option_b=row['option_b'].strip(),
                        option_c=row['option_c'].strip(),
                        option_d=row['option_d'].strip(),
                        correct_answer=correct_answer
                    )
                    db.session.add(question)
                    questions_added += 1

                except ValueError as e:
                    flash(
                        f"Invalid data in row {row.get('slno', '?')}: {str(e)}", "warning")
                    continue

            db.session.commit()

            action = "replaced" if upload_mode == "replace" else "added"
            flash(
                f"Successfully {action} {questions_added} questions for Module {module_number}: {modules[module_number]}!", "success")

        except Exception as e:
            db.session.rollback()
            flash(f"Error uploading questions: {str(e)}", "danger")

        return redirect(url_for("python_lab_upload_quiz_questions"))

    # GET request - show upload form
    # Count existing questions per module (from shared question pool)
    module_counts = {}
    for module_num in range(1, 6):
        # All admins see all questions from all admins (shared question pool)
        count = PythonLabQuizQuestion.query.filter_by(
            module_number=module_num).count()
        module_counts[module_num] = count

    return render_template("python_lab_upload_quiz_questions.html",
                           current_admin=user,
                           modules=modules,
                           module_counts=module_counts)


@app.route("/download_python_lab_template")
def download_python_lab_template():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Create sample CSV data
    sample_data = [
        {
            'slno': 1,
            'question': 'Which of the following is the correct way to create a variable in Python?',
            'option_a': 'var x = 10',
            'option_b': 'x = 10',
            'option_c': 'int x = 10',
            'option_d': 'variable x = 10',
            'correct_answer': 'B'
        },
        {
            'slno': 2,
            'question': 'What is the output of print(type(5.0))?',
            'option_a': '<class "int">',
            'option_b': '<class "float">',
            'option_c': '<class "number">',
            'option_d': '<class "decimal">',
            'correct_answer': 'B'
        },
        {
            'slno': 3,
            'question': 'Which method is used to add an element at the end of a list?',
            'option_a': 'add()',
            'option_b': 'insert()',
            'option_c': 'append()',
            'option_d': 'extend()',
            'correct_answer': 'C'
        }
    ]

    # Create CSV content
    output = StringIO()
    fieldnames = ['slno', 'question', 'option_a',
                  'option_b', 'option_c', 'option_d', 'correct_answer']
    writer = csv.DictWriter(output, fieldnames=fieldnames)

    writer.writeheader()
    for row in sample_data:
        writer.writerow(row)

    # Create response
    csv_content = output.getvalue()
    output.close()

    response = make_response(csv_content)
    response.headers['Content-Type'] = 'text/csv'
    response.headers['Content-Disposition'] = 'attachment; filename=python_lab_questions_template.csv'

    return response


@app.route("/python_lab_quiz_settings", methods=["GET", "POST"])
def python_lab_quiz_settings():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Get quiz settings for all modules for this admin
    all_settings = PythonLabQuizSettings.query.filter_by(
        admin_id=user.id).all()

    # Create a settings dict by module number for easy access
    settings_by_module = {
        setting.module_number: setting for setting in all_settings}

    if request.method == "POST":
        try:
            quiz_duration = int(request.form.get("quiz_duration", 60))
            max_attempts = int(request.form.get("max_attempts", 3))
            questions_per_quiz = int(
                request.form.get("questions_per_quiz", 20))

            # Calculate questions per module from total questions
            if questions_per_quiz % 5 != 0:
                flash(
                    "Total questions must be divisible by 5 to distribute equally across all modules!", "danger")
                return redirect(url_for("python_lab_quiz_settings"))

            questions_per_module = questions_per_quiz // 5

            # Validate inputs
            if quiz_duration < 5 or quiz_duration > 120:
                flash("Quiz duration must be between 5 and 120 minutes!", "danger")
                return redirect(url_for("python_lab_quiz_settings"))

            if max_attempts < 1 or max_attempts > 10:
                flash("Max attempts must be between 1 and 10!", "danger")
                return redirect(url_for("python_lab_quiz_settings"))

            if questions_per_quiz < 5 or questions_per_quiz > 100:
                flash("Total questions must be between 5 and 100!", "danger")
                return redirect(url_for("python_lab_quiz_settings"))

            if questions_per_module < 1 or questions_per_module > 20:
                flash("Questions per module must be between 1 and 20!", "danger")
                return redirect(url_for("python_lab_quiz_settings"))

            # Update or create settings for all 5 modules
            for module_num in range(1, 6):
                existing_setting = settings_by_module.get(module_num)

                if existing_setting:
                    # Update existing settings
                    existing_setting.quiz_duration = quiz_duration
                    existing_setting.max_attempts = max_attempts
                    existing_setting.questions_per_module = questions_per_module
                else:
                    # Create new settings for this module
                    new_setting = PythonLabQuizSettings(
                        admin_id=user.id,
                        module_number=module_num,
                        quiz_duration=quiz_duration,
                        max_attempts=max_attempts,
                        questions_per_module=questions_per_module,
                        is_active=False,
                        is_quiz_released=False
                    )
                    db.session.add(new_setting)

            db.session.commit()
            flash("Quiz settings updated successfully for all modules!", "success")

        except ValueError:
            flash("Invalid input values! Please enter valid numbers.", "danger")
        except Exception as e:
            db.session.rollback()
            flash(f"Error updating settings: {str(e)}", "danger")

        return redirect(url_for("python_lab_quiz_settings"))

    # If no settings exist, create default ones for all modules
    if not all_settings:
        for module_num in range(1, 6):
            default_setting = PythonLabQuizSettings(
                admin_id=user.id,
                module_number=module_num,
                quiz_duration=60,
                max_attempts=3,
                questions_per_module=4,
                is_active=False,
                is_quiz_released=False
            )
            db.session.add(default_setting)
        db.session.commit()

        # Refresh settings after creating defaults
        all_settings = PythonLabQuizSettings.query.filter_by(
            admin_id=user.id).all()
        settings_by_module = {
            setting.module_number: setting for setting in all_settings}

    # Count questions per module for validation (all questions from all admins)
    module_counts = {}
    for module_num in range(1, 6):
        # All admins see all questions from all admins (shared question pool)
        count = PythonLabQuizQuestion.query.filter_by(
            module_number=module_num).count()
        module_counts[module_num] = count

    # Get the first setting for general settings (they should all have same duration/attempts)
    general_settings = all_settings[0] if all_settings else None

    return render_template("python_lab_quiz_settings.html",
                           current_admin=user,
                           quiz_settings=general_settings,
                           all_settings=all_settings,
                           settings_by_module=settings_by_module,
                           module_counts=module_counts)


@app.route("/python_lab_student_dashboard")
def python_lab_student_dashboard():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role != "student":
        flash("Student access required.", "danger")
        return redirect(url_for("login"))

    # Get quiz settings from any admin (for students to see available quizzes)
    active_settings = PythonLabQuizSettings.query.filter_by(
        is_quiz_released=True).all()

    # Fix any legacy data where is_released=True but is_active=False
    for settings in active_settings:
        if not settings.is_active:
            settings.is_active = True
    if active_settings:  # Only commit if there are settings to update
        db.session.commit()

    # Also check for individual student releases
    individual_releases = PythonLabStudentQuizRelease.query.filter_by(
        student_id=user.id, is_released=True).all()

    # Get the first active setting for display purposes, preferring bulk over individual
    quiz_settings = None
    if active_settings:
        quiz_settings = active_settings[0]
    elif individual_releases:
        # If individually released but not bulk released, get settings from releasing admin
        release_admin = User.query.get(individual_releases[0].admin_id)
        if release_admin:
            quiz_settings = PythonLabQuizSettings.query.filter_by(
                admin_id=release_admin.id).first()

            # If admin doesn't have settings yet, create default ones
            if not quiz_settings:
                quiz_settings = PythonLabQuizSettings(
                    admin_id=release_admin.id,
                    module_number=1,
                    is_quiz_released=False,  # Individual release, not bulk
                    max_attempts=3,
                    quiz_duration=60,
                    questions_per_module=2  # Changed from 4 to 2 for 10 total questions
                )
                db.session.add(quiz_settings)
                db.session.commit()

    # Get any settings record to check question visibility (simple toggle by super admin)
    any_settings = PythonLabQuizSettings.query.first()
    questions_visible = any_settings.questions_visible_to_students if any_settings else False

    # Get student's quiz attempts
    student_attempts = PythonLabQuizAttempt.query.filter_by(
        student_id=user.id, is_completed=True).order_by(
        PythonLabQuizAttempt.attempt_number.desc()).all()

    # Calculate statistics
    total_attempts = len(student_attempts)
    if total_attempts > 0:
        scores = [attempt.score for attempt in student_attempts]
        best_score = max(scores)
        avg_score = sum(scores) / len(scores)
        latest_attempt = student_attempts[0]
    else:
        best_score = avg_score = 0
        latest_attempt = None

    # Check if student can take quiz (based on max attempts)
    can_take_quiz = False
    remaining_attempts = 0

    if quiz_settings:
        max_attempts = quiz_settings.max_attempts
        attempts_used = PythonLabQuizAttempt.query.filter_by(
            student_id=user.id,
            admin_id=quiz_settings.admin_id,
            is_completed=True
        ).count()
        remaining_attempts = max(0, max_attempts - attempts_used)
        # Student can take quiz if either bulk released OR individually released
        quiz_available = len(active_settings) > 0 or len(
            individual_releases) > 0
        can_take_quiz = remaining_attempts > 0 and quiz_available

    else:
        quiz_available = False
        can_take_quiz = False
        remaining_attempts = 0

    # Module definitions
    modules = {
        1: "Python Basics",
        2: "Operators and Control Flow",
        3: "Data Structures and Manipulation",
        4: "Functions and Modules",
        5: "Object-Oriented Programming"
    }

    # Check if quiz is available (either bulk or individual release)
    quiz_available = (len(active_settings) > 0 or len(
        individual_releases) > 0) and quiz_settings is not None

    # Get Python Lab Experiments data
    # Get available experiments for this student (either globally released or individually released)
    available_experiments = []

    # Get globally released experiments
    global_experiments = PythonLabExperiment.query.filter_by(
        is_released=True).all()
    available_experiments.extend(global_experiments)

    # Get individually released experiments
    individual_experiment_releases = PythonLabStudentExperimentRelease.query.filter_by(
        student_id=user.id, is_released=True
    ).all()

    for release in individual_experiment_releases:
        experiment = PythonLabExperiment.query.get(release.experiment_id)
        if experiment and experiment not in available_experiments:
            available_experiments.append(experiment)

    # Sort experiments by experiment number
    available_experiments.sort(key=lambda x: x.experiment_number)

    # Get experiment criteria counts
    experiment_criteria_count = {}
    for experiment in available_experiments:
        count = PythonLabCriteria.query.filter_by(
            experiment_id=experiment.id).count()
        experiment_criteria_count[experiment.id] = count

    # Get student's experiment attempts
    student_experiment_attempts = {}
    for experiment in available_experiments:
        attempt = PythonLabAttempt.query.filter_by(
            student_id=user.id, experiment_id=experiment.id
        ).first()
        if attempt:
            student_experiment_attempts[experiment.id] = attempt

    # Get student's experiment marks (earned and max)
    student_experiment_marks = {}
    for experiment in available_experiments:
        responses = PythonLabResponse.query.filter_by(
            user_id=user.id, experiment_id=experiment.id
        ).all()

        if responses:
            marks_earned = sum(response.marks_earned for response in responses)
            # Calculate max marks from criteria for this experiment
            criteria = PythonLabCriteria.query.filter_by(
                experiment_id=experiment.id).all()
            max_marks = sum(c.max_marks for c in criteria)
            student_experiment_marks[experiment.id] = {
                'marks_earned': marks_earned,
                'max_marks': max_marks
            }

    return render_template("python_lab_student_dashboard.html",
                           user=user,
                           quiz_settings=quiz_settings,
                           quiz_available=quiz_available,
                           student_attempts=student_attempts,
                           total_attempts=total_attempts,
                           best_score=best_score,
                           avg_score=avg_score,
                           latest_attempt=latest_attempt,
                           can_take_quiz=can_take_quiz,
                           remaining_attempts=remaining_attempts,
                           questions_visible_to_students=questions_visible,
                           modules=modules,
                           available_experiments=available_experiments,
                           experiment_criteria_count=experiment_criteria_count,
                           student_experiment_attempts=student_experiment_attempts,
                           student_experiment_marks=student_experiment_marks)


@app.route("/python_lab_student_consolidated_marks")
def python_lab_student_consolidated_marks():
    """Show student their individual Python Lab consolidated marks"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role != 'student':
        flash("Student access required.", "danger")
        return redirect(url_for("login"))

    roll = user.rollnumber

    # Try to get student name from KDM Lab records (if available)
    kdm_student = KDMLabStudent.query.filter_by(rollnumber=roll).first()
    student_name = kdm_student.name if kdm_student and kdm_student.name else "N/A"

    # Find the admin responsible for this student
    responsible_admin = find_admin_for_student(roll)
    if not responsible_admin:
        flash("No admin assigned for your Python Lab section.", "warning")
        return redirect(url_for("python_lab_student_dashboard"))

    # Load uploaded marks data for the responsible admin
    import os
    import pandas as pd
    uploaded_data = []

    try:
        admin_id = responsible_admin.id
        upload_path = os.path.join(
            os.getcwd(), "tmp_uploads", f"python_lab_marks_{admin_id}.csv")
        if os.path.exists(upload_path):
            df = pd.read_csv(upload_path)
            uploaded_data = df.to_dict(orient="records")
    except Exception as e:
        print(f"⚠️ Unable to read Python Lab marks file: {e}")

    # Calculate student's marks
    # Default marks
    rubrics_marks = "N/A"  # 20 marks - from instructor
    script_marks = "N/A"   # 15 marks - from instructor
    attendance_marks = "N/A"  # 5 marks - from instructor

    # Load from CSV if available
    for row in uploaded_data:
        if str(row.get("Roll Number")).strip() == str(roll):
            rubrics_marks = row.get("Rubrics Marks (20)", "N/A")
            script_marks = row.get("Script Marks (15)", "N/A")
            attendance_marks = row.get("Attendance Marks (5)", "N/A")

            # Convert numeric 0 values to actual 0 instead of N/A
            if rubrics_marks == 0 or rubrics_marks == "0" or rubrics_marks == 0.0:
                rubrics_marks = 0
            if script_marks == 0 or script_marks == "0" or script_marks == 0.0:
                script_marks = 0
            if attendance_marks == 0 or attendance_marks == "0" or attendance_marks == 0.0:
                attendance_marks = 0
            break

    # Calculate Viva Marks (10) from best Python Lab quiz attempt
    viva_marks = 0.0
    best_attempt = PythonLabQuizAttempt.query.filter_by(
        student_id=user.id,
        admin_id=responsible_admin.id,
        is_completed=True
    ).order_by(PythonLabQuizAttempt.score.desc()).first()

    if best_attempt:
        # Convert percentage to marks out of 10
        viva_marks = round((best_attempt.score * 10) / 100, 1)

    # Calculate Total Internal Marks (25) = Script (15) + Viva (10)
    total_internal_marks = "N/A"
    try:
        script_numeric = None
        if script_marks != "N/A":
            if isinstance(script_marks, (int, float)) or str(script_marks).replace('.', '', 1).replace('-', '', 1).isdigit():
                script_numeric = float(script_marks)

        if script_numeric is not None and isinstance(viva_marks, (int, float)):
            total_internal_marks = round(script_numeric + viva_marks, 1)
    except Exception:
        total_internal_marks = "N/A"

    # Calculate Final Internal Marks (50) = Rubrics (20) + Total Internal (25) + Attendance (5)
    final_internal = "N/A"
    try:
        rubrics_numeric = None
        attendance_numeric = None

        if rubrics_marks != "N/A":
            if isinstance(rubrics_marks, (int, float)) or str(rubrics_marks).replace('.', '', 1).replace('-', '', 1).isdigit():
                rubrics_numeric = float(rubrics_marks)

        if attendance_marks != "N/A":
            if isinstance(attendance_marks, (int, float)) or str(attendance_marks).replace('.', '', 1).replace('-', '', 1).isdigit():
                attendance_numeric = float(attendance_marks)

        if (rubrics_numeric is not None and
            isinstance(total_internal_marks, (int, float)) and
                attendance_numeric is not None):
            total_marks = rubrics_numeric + total_internal_marks + attendance_numeric
            final_internal = round(total_marks, 1)
    except Exception:
        final_internal = "N/A"

    # Prepare student data
    student_marks = {
        "Roll Number": roll,
        "Rubrics Marks (20)": rubrics_marks,
        "Script Marks (15)": script_marks,
        "Viva Marks (10)": viva_marks,
        "Total Internal (25)": total_internal_marks,
        "Attendance Marks (5)": attendance_marks,
        "Final Internal Marks (50)": final_internal
    }

    # Get quiz attempt details for display
    quiz_attempts = PythonLabQuizAttempt.query.filter_by(
        student_id=user.id,
        admin_id=responsible_admin.id,
        is_completed=True
    ).order_by(PythonLabQuizAttempt.completed_at.desc()).all()

    # Mark type descriptions for display
    mark_types = {
        "Rubrics Marks (20)": "Lab rubrics evaluation by instructor",
        "Script Marks (15)": "Python programming script evaluation",
        "Viva Marks (10)": "Automatically calculated from quiz performance",
        "Total Internal (25)": "Combined script and viva marks",
        "Attendance Marks (5)": "Lab attendance evaluation",
        "Final Internal Marks (50)": "Total internal assessment marks"
    }

    return render_template("python_lab_student_consolidated_marks.html",
                           user=user,
                           student_marks=student_marks,
                           responsible_admin=responsible_admin,
                           quiz_attempts=quiz_attempts,
                           best_attempt=best_attempt,
                           mark_types=mark_types)


@app.route("/python_lab_view_all_questions")
def python_lab_view_all_questions():
    """Students can view all Python Lab questions with answers (if enabled by super admin)"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role != "student":
        flash("Student access required.", "danger")
        return redirect(url_for("login"))

    # Check if questions are visible to students
    settings = PythonLabQuizSettings.query.first()
    if not settings or not settings.questions_visible_to_students:
        flash("Question bank is not currently available to students.", "warning")
        return redirect(url_for("python_lab_student_dashboard"))

    # Get all questions from all admins, grouped by module
    all_questions = PythonLabQuizQuestion.query.order_by(
        PythonLabQuizQuestion.module_number, PythonLabQuizQuestion.slno).all()

    # Group questions by module
    questions_by_module = {}
    for question in all_questions:
        if question.module_number not in questions_by_module:
            questions_by_module[question.module_number] = []
        questions_by_module[question.module_number].append(question)

    modules = {
        1: "Python Basics",
        2: "Operators and Control Flow",
        3: "Data Structures and Manipulation",
        4: "Functions and Modules",
        5: "Object-Oriented Programming"
    }

    return render_template("python_lab_student_questions.html",
                           user=user,
                           questions_by_module=questions_by_module,
                           modules=modules,
                           total_questions=len(all_questions))


@app.route("/python_lab_take_quiz", methods=["GET", "POST"])
def python_lab_take_quiz():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role != "student":
        flash("Student access required.", "danger")
        return redirect(url_for("login"))

    # Get quiz settings - check if any module is released (bulk release)
    active_settings = PythonLabQuizSettings.query.filter_by(
        is_quiz_released=True).all()

    # Fix any legacy data where is_released=True but is_active=False
    for settings in active_settings:
        if not settings.is_active:
            settings.is_active = True
    if active_settings:  # Only commit if there are settings to update
        db.session.commit()

    # Also check for individual student release
    individual_releases = PythonLabStudentQuizRelease.query.filter_by(
        student_id=user.id, is_released=True).all()

    # Student can take quiz if either bulk released OR individually released
    if not active_settings and not individual_releases:
        flash("No quiz is currently available! Please contact your instructor.", "warning")
        return redirect(url_for("python_lab_student_dashboard"))

    # Determine quiz settings and admin
    if active_settings:
        # Use bulk release settings
        quiz_settings = active_settings[0]
        admin_user = User.query.get(quiz_settings.admin_id)
    else:
        # Use individual release settings - get settings from the releasing admin
        individual_release = individual_releases[0]
        admin_user = User.query.get(individual_release.admin_id)

        # Get quiz settings for this admin
        quiz_settings = PythonLabQuizSettings.query.filter_by(
            admin_id=admin_user.id).first()
        if not quiz_settings:
            # Create default settings if none exist
            quiz_settings = PythonLabQuizSettings(
                admin_id=admin_user.id,
                module_number=1,
                is_quiz_released=False,  # Keep as False since this is individual release
                max_attempts=3,
                quiz_duration=60,
                questions_per_module=4
            )
            db.session.add(quiz_settings)
            db.session.commit()

    if not admin_user:
        flash("Quiz administrator not found! Please contact support.", "danger")
        return redirect(url_for("python_lab_student_dashboard"))

    # Check if student has attempts left (only count completed attempts)
    attempts_used = PythonLabQuizAttempt.query.filter_by(
        student_id=user.id,
        admin_id=quiz_settings.admin_id,
        is_completed=True
    ).count()

    if attempts_used >= quiz_settings.max_attempts:
        flash("You have used all your quiz attempts!", "danger")
        return redirect(url_for("python_lab_student_dashboard"))

    if request.method == "POST":
        # Process quiz submission
        try:
            # Get the current attempt
            attempt_id = request.form.get("attempt_id")
            attempt = PythonLabQuizAttempt.query.get(attempt_id)

            if not attempt or attempt.student_id != user.id:
                flash("Invalid quiz attempt!", "danger")
                return redirect(url_for("python_lab_student_dashboard"))

            # Get choice mappings from session
            choice_mappings = session.get(
                f"python_quiz_mappings_{attempt.id}", {})

            # Get the actual shuffled questions that were shown to the student
            try:
                # Try to parse as JSON (new format with full question data)
                quiz_questions_data = json.loads(attempt.quiz_questions)
                quiz_questions = []

                # Convert stored data back to question-like objects for compatibility
                for q_data in quiz_questions_data:
                    # Create a simple object to hold question data
                    class QuestionObj:
                        def __init__(self, data):
                            self.id = data["id"]
                            self.module_number = data["module"]
                            self.question = data["question"]
                            self.option_a = data["option_a"]
                            self.option_b = data["option_b"]
                            self.option_c = data["option_c"]
                            self.option_d = data["option_d"]
                            self.correct_answer = data["correct_answer"]

                    quiz_questions.append(QuestionObj(q_data))

            except (json.JSONDecodeError, KeyError):
                # Fallback to old format (comma-separated IDs) for backward compatibility
                question_ids = attempt.quiz_questions.split(',')
                quiz_questions = PythonLabQuizQuestion.query.filter(
                    PythonLabQuizQuestion.id.in_(question_ids)).all()

                # Sort questions to maintain the same order as shown to student
                question_order = {int(qid): idx for idx,
                                  qid in enumerate(question_ids)}
                quiz_questions.sort(
                    key=lambda q: question_order.get(q.id, 999))

            student_answers = {}
            module_scores = {"module_1": 0, "module_2": 0,
                             "module_3": 0, "module_4": 0, "module_5": 0}
            total_correct = 0

            # Process each answer
            for question in quiz_questions:
                answer_key = f"question_{question.id}"
                student_answer = request.form.get(
                    answer_key, "").upper().strip()
                student_answers[str(question.id)] = student_answer

                # Check if correct - account for shuffled choices
                correct_answer = question.correct_answer
                question_mapping = choice_mappings.get(str(question.id), {})

                # If we have a mapping, convert student answer back to original position
                if question_mapping and student_answer in question_mapping:
                    original_answer = question_mapping[student_answer]
                    is_answer_correct = original_answer == correct_answer
                else:
                    # Fallback to direct comparison if no mapping found
                    is_answer_correct = student_answer == correct_answer
                    # Flash a warning if mapping is missing (shouldn't happen normally)
                    if student_answer and not question_mapping:
                        flash(
                            "Technical issue detected during quiz evaluation. Results may be affected. Please inform your instructor.", "warning")

                if is_answer_correct:
                    total_correct += 1
                    module_scores[f"module_{question.module_number}"] += 1

                # Save individual response
                response = PythonLabQuizResponse(
                    attempt_id=attempt.id,
                    question_id=question.id,
                    student_answer=student_answer if student_answer else None,
                    correct_answer=question.correct_answer,
                    is_correct=is_answer_correct,
                    module_number=question.module_number
                )
                db.session.add(response)

            # Calculate final score based on actual number of questions
            total_questions = len(quiz_questions)
            final_score = int((total_correct / total_questions)
                              * 100) if total_questions > 0 else 0

            # Update attempt
            attempt.score = final_score
            attempt.total_questions = total_questions
            attempt.completed_at = datetime.now()
            attempt.is_completed = True

            db.session.commit()

            # Clean up session data for this attempt
            session_key = f"python_quiz_mappings_{attempt.id}"
            if session_key in session:
                del session[session_key]

            flash(
                f"Quiz submitted successfully! Your score: {final_score}%", "success")
            return redirect(url_for("python_lab_view_attempt_result", attempt_id=attempt.id))

        except Exception as e:
            db.session.rollback()
            flash(f"Error submitting quiz: {str(e)}", "danger")
            return redirect(url_for("python_lab_student_dashboard"))

    # GET request - start new quiz attempt
    try:
        # Generate quiz questions (4 from each module = 20 total)
        quiz_questions = []

        questions_per_module = quiz_settings.questions_per_module if hasattr(
            quiz_settings, 'questions_per_module') else 4

        for module_num in range(1, 6):
            # Get questions for this module from shared question pool (all admins)
            module_questions = PythonLabQuizQuestion.query.filter_by(
                module_number=module_num).all()

            if len(module_questions) < questions_per_module:
                flash(
                    f"Insufficient questions in Module {module_num}. Contact your instructor.", "danger")
                return redirect(url_for("python_lab_student_dashboard"))

            # Randomly select questions from this module
            selected = random.sample(module_questions, questions_per_module)
            for q in selected:
                # Create list of choices for shuffling
                choices = [
                    {"label": "A", "text": q.option_a, "original": "A"},
                    {"label": "B", "text": q.option_b, "original": "B"},
                    {"label": "C", "text": q.option_c, "original": "C"},
                    {"label": "D", "text": q.option_d, "original": "D"}
                ]

                # Shuffle the choices
                random.shuffle(choices)

                # Create mapping of new positions to original labels
                choice_mapping = {}
                for i, choice in enumerate(choices):
                    new_label = chr(ord('A') + i)
                    choice_mapping[new_label] = choice["original"]

                quiz_questions.append({
                    "id": q.id,
                    "module": q.module_number,
                    "question": q.question,
                    "option_a": choices[0]["text"],
                    "option_b": choices[1]["text"],
                    "option_c": choices[2]["text"],
                    "option_d": choices[3]["text"],
                    "correct_answer": q.correct_answer,
                    "choice_mapping": choice_mapping  # Store the shuffle mapping
                })

        # Shuffle the final question order
        random.shuffle(quiz_questions)

        # Store choice mappings in session for answer validation
        choice_mappings = {}
        for q in quiz_questions:
            choice_mappings[str(q["id"])] = q["choice_mapping"]

        # Check for existing incomplete attempts and clean them up
        existing_incomplete = PythonLabQuizAttempt.query.filter_by(
            student_id=user.id,
            admin_id=quiz_settings.admin_id,
            is_completed=False
        ).all()

        if existing_incomplete:
            # Delete incomplete attempts and their responses
            for incomplete in existing_incomplete:
                # Delete associated responses
                PythonLabQuizResponse.query.filter_by(
                    attempt_id=incomplete.id).delete()
                # Delete the attempt
                db.session.delete(incomplete)
            db.session.commit()

        # Double-check for any remaining attempt with the same attempt number (safety check)
        potential_duplicate = PythonLabQuizAttempt.query.filter_by(
            student_id=user.id,
            admin_id=quiz_settings.admin_id,
            attempt_number=attempts_used + 1
        ).first()

        if potential_duplicate:
            # Delete associated responses
            PythonLabQuizResponse.query.filter_by(
                attempt_id=potential_duplicate.id).delete()
            # Delete the attempt
            db.session.delete(potential_duplicate)
            db.session.commit()

        # Create new attempt
        attempt_number = attempts_used + 1
        total_questions = len(quiz_questions)

        # Store complete shuffled question data as JSON for consistency
        shuffled_questions_data = []
        for q in quiz_questions:
            shuffled_questions_data.append({
                "id": q["id"],
                "module": q["module"],
                "question": q["question"],
                "option_a": q["option_a"],  # These are already shuffled
                "option_b": q["option_b"],
                "option_c": q["option_c"],
                "option_d": q["option_d"],
                # Original correct answer
                "correct_answer": q["correct_answer"],
                # Mapping for evaluation
                "choice_mapping": q["choice_mapping"]
            })

        try:
            attempt = PythonLabQuizAttempt(
                student_id=user.id,
                admin_id=quiz_settings.admin_id,
                attempt_number=attempt_number,
                total_questions=total_questions,
                # Store full question data
                quiz_questions=json.dumps(shuffled_questions_data)
            )
            db.session.add(attempt)
            db.session.commit()

        except Exception as create_error:
            db.session.rollback()

            # If still getting duplicate key error, find and use the next available attempt number
            if "Duplicate entry" in str(create_error) and "unique_student_attempt" in str(create_error):
                # Find the highest attempt number for this student-admin combination
                max_attempt = db.session.query(db.func.max(PythonLabQuizAttempt.attempt_number)).filter_by(
                    student_id=user.id,
                    admin_id=quiz_settings.admin_id
                ).scalar()

                next_attempt_number = (max_attempt or 0) + 1

                attempt = PythonLabQuizAttempt(
                    student_id=user.id,
                    admin_id=quiz_settings.admin_id,
                    attempt_number=next_attempt_number,
                    total_questions=total_questions,
                    quiz_questions=json.dumps(shuffled_questions_data)
                )
                db.session.add(attempt)
                db.session.commit()
                flash("Quiz attempt created successfully!", "success")
            else:
                # Re-raise if it's not a duplicate key error
                raise create_error

        # Store the choice mappings in session for this attempt
        session[f"python_quiz_mappings_{attempt.id}"] = choice_mappings

        # Module definitions for display
        modules = {
            1: "Python Basics",
            2: "Operators and Control Flow",
            3: "Data Structures and Manipulation",
            4: "Functions and Modules",
            5: "Object-Oriented Programming"
        }

        return render_template("python_lab_take_quiz.html",
                               user=user,
                               quiz_settings=quiz_settings,
                               quiz_questions=quiz_questions,
                               attempt=attempt,
                               attempt_number=attempt_number,
                               modules=modules)

    except Exception as e:
        db.session.rollback()
        flash(f"Error starting quiz: {str(e)}", "danger")
        return redirect(url_for("python_lab_student_dashboard"))


@app.route("/python_lab_view_attempt_result/<int:attempt_id>")
def python_lab_view_attempt_result(attempt_id):
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user:
        flash("User not found.", "danger")
        return redirect(url_for("login"))

    attempt = PythonLabQuizAttempt.query.get_or_404(attempt_id)

    # Allow access if user is the student who took the quiz OR an admin/super_admin
    if user.role == "student":
        if attempt.student_id != user.id:
            flash("Unauthorized access!", "danger")
            return redirect(url_for("python_lab_student_dashboard"))
    elif user.role not in ["admin", "super_admin"]:
        flash("Access denied. Student or admin access required.", "danger")
        return redirect(url_for("login"))
    else:
        # Admin/super_admin access - check if they have permission for this attempt
        if user.role == "admin" and attempt.admin_id != user.id:
            flash("You can only view attempts from your own quizzes.", "danger")
            return redirect(url_for("python_lab_admin_dashboard"))
        # super_admin can view all attempts

    if not attempt.is_completed:
        flash("This quiz attempt is not completed yet!", "warning")
        if user.role == "student":
            return redirect(url_for("python_lab_student_dashboard"))
        else:
            return redirect(url_for("python_lab_quiz_results"))

    # Get all responses for this attempt
    responses = PythonLabQuizResponse.query.filter_by(
        attempt_id=attempt_id).all()

    # Organize responses by question
    question_responses = {}
    for response in responses:
        question_responses[response.question_id] = response

    # Build questions dictionary for template (needed for detailed analysis)
    questions = {}
    quiz_questions_data = []

    try:
        # Try to parse the stored shuffled question data (new format)
        stored_questions = json.loads(attempt.quiz_questions)
        quiz_questions_data = stored_questions
        question_ids = [q["id"] for q in quiz_questions_data]
        questions = {q.id: q for q in PythonLabQuizQuestion.query.filter(
            PythonLabQuizQuestion.id.in_(question_ids)).all()}
    except (json.JSONDecodeError, KeyError):
        # Fallback for old format - get original questions from database
        question_ids = [response.question_id for response in responses]
        questions = {q.id: q for q in PythonLabQuizQuestion.query.filter(
            PythonLabQuizQuestion.id.in_(question_ids)).all()}

        # Create quiz_questions_data from original database questions
        for response in responses:
            question = questions.get(response.question_id)
            if question:
                quiz_questions_data.append({
                    "id": question.id,
                    "module": question.module_number,
                    "question": question.question,
                    "option_a": question.option_a,
                    "option_b": question.option_b,
                    "option_c": question.option_c,
                    "option_d": question.option_d,
                    "correct_answer": question.correct_answer
                })

    # Calculate module-wise performance from responses
    module_performance = {}

    for module_num in range(1, 6):
        module_questions = [
            q for q in quiz_questions_data if q["module"] == module_num]

        # Count correct answers for this module
        correct_answers = 0
        for response in responses:
            if response.question_id in [q["id"] for q in module_questions] and response.is_correct:
                correct_answers += 1

        total_questions = len(module_questions)

        module_performance[module_num] = {
            "correct": correct_answers,
            "total": total_questions,
            "percentage": (correct_answers / total_questions * 100) if total_questions > 0 else 0
        }

    # Module definitions
    modules = {
        1: "Python Basics",
        2: "Operators and Control Flow",
        3: "Data Structures and Manipulation",
        4: "Functions and Modules",
        5: "Object-Oriented Programming"
    }

    # Get student information for display (needed when admin views student's attempt)
    student = User.query.get(attempt.student_id)

    return render_template("python_lab_view_attempt_result.html",
                           user=user,
                           student=student,
                           attempt=attempt,
                           question_responses=question_responses,
                           quiz_questions_data=quiz_questions_data,
                           module_performance=module_performance,
                           modules=modules,
                           questions=questions)


@app.route("/python_lab_quiz_question_management")
def python_lab_quiz_question_management():
    """Manage Python Lab quiz questions"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    # Get all questions from all admins (shared question pool)
    questions = PythonLabQuizQuestion.query.order_by(
        PythonLabQuizQuestion.module_number, PythonLabQuizQuestion.slno).all()

    # Group questions by module
    questions_by_module = {}
    for question in questions:
        if question.module_number not in questions_by_module:
            questions_by_module[question.module_number] = []
        questions_by_module[question.module_number].append(question)

    modules = {
        1: "Python Basics",
        2: "Operators and Control Flow",
        3: "Data Structures and Manipulation",
        4: "Functions and Modules",
        5: "Object-Oriented Programming"
    }

    return render_template("python_lab_quiz_question_management.html",
                           current_admin=current_admin,
                           questions_by_module=questions_by_module,
                           modules=modules,
                           total_questions=len(questions))


@app.route("/python_lab_admin_preview_quiz")
def python_lab_admin_preview_quiz():
    """Preview Python Lab quiz"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    # Get quiz settings to determine how many questions per module
    quiz_settings = PythonLabQuizSettings.query.filter_by(
        admin_id=current_admin.id).first()

    questions_per_module = quiz_settings.questions_per_module if quiz_settings else 4

    # Generate sample quiz questions based on settings (from shared question pool)
    quiz_questions = []
    for module_num in range(1, 6):
        # All admins can preview questions from all admins (shared question pool)
        module_questions = PythonLabQuizQuestion.query.filter_by(
            module_number=module_num).limit(questions_per_module).all()
        quiz_questions.extend(module_questions)

    modules = {
        1: "Python Basics",
        2: "Operators and Control Flow",
        3: "Data Structures and Manipulation",
        4: "Functions and Modules",
        5: "Object-Oriented Programming"
    }

    return render_template("python_lab_admin_preview_quiz.html",
                           current_admin=current_admin,
                           quiz_questions=quiz_questions,
                           modules=modules,
                           quiz_settings=quiz_settings,
                           questions_per_module=questions_per_module,
                           total_questions=len(quiz_questions))


@app.route("/python_lab_quiz_results")
def python_lab_quiz_results():
    """View Python Lab quiz results"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    # Get completed attempts sorted by roll number
    if current_admin.admin_level == "super_admin":
        attempts = PythonLabQuizAttempt.query.filter_by(is_completed=True).join(
            User, PythonLabQuizAttempt.student_id == User.id).order_by(
            User.rollnumber.asc()).all()
    else:
        attempts = PythonLabQuizAttempt.query.filter_by(
            admin_id=current_admin.id, is_completed=True).join(
            User, PythonLabQuizAttempt.student_id == User.id).order_by(
            User.rollnumber.asc()).all()

    # Calculate marks for each attempt and include user data
    attempts_with_marks = []
    for attempt in attempts:
        student = User.query.get(attempt.student_id)
        marks_scored = (attempt.score * attempt.total_questions) // 100
        attempt_data = {
            'attempt': attempt,
            'student': student,
            'marks_scored': marks_scored,
            'total_marks': attempt.total_questions
        }
        attempts_with_marks.append(attempt_data)

    return render_template("python_lab_quiz_results.html",
                           current_admin=current_admin,
                           attempts=attempts,
                           attempts_with_marks=attempts_with_marks,
                           User=User)


@app.route("/python_lab_toggle_quiz_release", methods=["POST"])
def python_lab_toggle_quiz_release():
    """Toggle Python Lab quiz release status"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    # Get or create quiz settings
    quiz_settings = PythonLabQuizSettings.query.filter_by(
        admin_id=current_admin.id).first()
    if not quiz_settings:
        quiz_settings = PythonLabQuizSettings(
            admin_id=current_admin.id,
            module_number=1,  # Default module
            is_quiz_released=False
        )
        db.session.add(quiz_settings)

    # Toggle release status
    quiz_settings.is_quiz_released = not quiz_settings.is_quiz_released
    # Also set is_active when releasing the quiz
    quiz_settings.is_active = quiz_settings.is_quiz_released
    db.session.commit()

    status = "released" if quiz_settings.is_quiz_released else "unreleased"
    flash(f"Python Lab quiz {status} successfully!", "success")
    return redirect(url_for("python_lab_admin_dashboard"))


@app.route("/python_lab_cleanup_attempts", methods=["POST"])
def python_lab_cleanup_attempts():
    """Clean up incomplete quiz attempts - Admin only"""
    if "loggedin" not in session:
        return jsonify({"success": False, "message": "Login required."})

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        return jsonify({"success": False, "message": "Admin access required."})

    try:
        # Find incomplete attempts
        incomplete_attempts = PythonLabQuizAttempt.query.filter_by(
            is_completed=False).all()

        cleaned_count = 0
        for attempt in incomplete_attempts:
            # Delete associated responses
            PythonLabQuizResponse.query.filter_by(
                attempt_id=attempt.id).delete()
            # Delete the attempt
            db.session.delete(attempt)
            cleaned_count += 1

        db.session.commit()

        return jsonify({
            "success": True,
            "message": f"Cleaned up {cleaned_count} incomplete attempts."
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": f"Error: {str(e)}"})


@app.route("/python_lab_toggle_student_release", methods=["POST"])
def python_lab_toggle_student_release():
    """Toggle Python Lab quiz release for individual student"""
    if "loggedin" not in session:
        return jsonify({"success": False, "message": "Login required."})

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        return jsonify({"success": False, "message": "Admin access required."})

    try:
        data = request.get_json()
        student_id = data.get("student_id")

        if not student_id:
            return jsonify({"success": False, "message": "Student ID is required."})

        student = User.query.get(student_id)
        if not student or student.role != "student":
            return jsonify({"success": False, "message": "Student not found."})

        # Get or create release record
        release_record = PythonLabStudentQuizRelease.query.filter_by(
            admin_id=current_admin.id, student_id=student.id).first()

        if not release_record:
            release_record = PythonLabStudentQuizRelease(
                admin_id=current_admin.id,
                student_id=student.id,
                is_released=False
            )
            db.session.add(release_record)

        # Toggle release status
        release_record.is_released = not release_record.is_released
        db.session.commit()

        status = "released" if release_record.is_released else "locked"
        return jsonify({
            "success": True,
            "message": f"Quiz {status} for {student.rollnumber}",
            "is_released": release_record.is_released
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": f"Error: {str(e)}"})


@app.route("/python_lab_consolidated_marks")
def python_lab_consolidated_marks():
    """Python Lab Consolidated Internal Marks Page"""
    import os
    import pandas as pd

    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    current_admin = user

    # --- Fetch students based on admin level ---
    if current_admin.admin_level == "super_admin":
        students = User.query.filter(User.role == "student").all()
    else:
        students = get_students_for_admin(current_admin)
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # --- Load uploaded CSV (if exists) for Rubrics & Attendance marks ---
    uploaded_data = []
    try:
        if current_admin.admin_level == "super_admin":
            import glob
            upload_dir = os.path.join(os.getcwd(), "tmp_uploads")
            csv_files = glob.glob(os.path.join(upload_dir, "python_lab_marks_*.csv"))
            for csv_file in csv_files:
                try:
                    df = pd.read_csv(csv_file)
                    uploaded_data.extend(df.to_dict(orient="records"))
                except Exception as e:
                    print(f"⚠️ Error reading {csv_file}: {e}")
        else:
            upload_path = os.path.join(
                os.getcwd(), "tmp_uploads", f"python_lab_marks_{current_admin.id}.csv"
            )
            if os.path.exists(upload_path):
                df = pd.read_csv(upload_path)
                uploaded_data = df.to_dict(orient="records")
    except Exception as e:
        print(f"⚠️ Unable to read uploaded Python Lab marks file: {e}")

    # --- Compute all marks ---
    data = []
    for s in students:
        roll = s.rollnumber

        # Default Rubrics & Attendance marks
        rubrics_marks = "N/A"
        attendance_marks = "N/A"

        # If uploaded, read from CSV
        for row in uploaded_data:
            if str(row.get("Roll Number")).strip() == str(roll):
                rubrics_marks = row.get("Rubrics Marks (20)", "N/A")
                attendance_marks = row.get("Attendance Marks (5)", "N/A")
                if rubrics_marks in [0, "0", 0.0]:
                    rubrics_marks = 0
                if attendance_marks in [0, "0", 0.0]:
                    attendance_marks = 0

        # Script Marks (15)
        script_marks = "N/A"
        for row in uploaded_data:
            if str(row.get("Roll Number")).strip() == str(roll):
                script_marks = row.get("Script Marks (15)", "N/A")
                if script_marks in [0, "0", 0.0]:
                    script_marks = 0

        # Viva Marks (10) — Hybrid Logic (subject-specific + safe fallback)
        student_user = User.query.filter_by(rollnumber=roll).first()
        viva_marks = 0.0
        if student_user:
            # 1️⃣ Try to find attempt under current admin (subject-specific)
            best_attempt = PythonLabQuizAttempt.query.filter_by(
                student_id=student_user.id,
                admin_id=current_admin.id,
                is_completed=True,
            ).order_by(PythonLabQuizAttempt.score.desc()).first()

            # 2️⃣ Fallback: if no attempt found under this admin,
            # check if there’s only one attempt total (safe to use)
            if not best_attempt:
                all_attempts = PythonLabQuizAttempt.query.filter_by(
                    student_id=student_user.id, is_completed=True
                ).all()
                if len(all_attempts) == 1:
                    best_attempt = all_attempts[0]

            if best_attempt:
                viva_marks = round((best_attempt.score * 10) / 100, 1)

        # Total Internal (25) = Script + Viva
        total_internal_marks = "N/A"
        try:
            script_numeric = None
            if script_marks != "N/A":
                if isinstance(script_marks, (int, float)) or str(script_marks).replace(".", "", 1).isdigit():
                    script_numeric = float(script_marks)
            if script_numeric is not None and isinstance(viva_marks, (int, float)):
                total_internal_marks = round(script_numeric + viva_marks, 1)
        except Exception:
            total_internal_marks = "N/A"

        # Final Internal (50) = Rubrics + Total Internal + Attendance
        final_internal = "N/A"
        try:
            rubrics_numeric = None
            attendance_numeric = None
            if rubrics_marks != "N/A":
                if isinstance(rubrics_marks, (int, float)) or str(rubrics_marks).replace(".", "", 1).isdigit():
                    rubrics_numeric = float(rubrics_marks)
            if attendance_marks != "N/A":
                if isinstance(attendance_marks, (int, float)) or str(attendance_marks).replace(".", "", 1).isdigit():
                    attendance_numeric = float(attendance_marks)
            if (
                rubrics_numeric is not None
                and isinstance(total_internal_marks, (int, float))
                and attendance_numeric is not None
            ):
                total_marks = rubrics_numeric + total_internal_marks + attendance_numeric
                final_internal = round(total_marks, 1)
        except Exception:
            final_internal = "N/A"

        data.append(
            {
                "Roll Number": roll,
                "Rubrics Marks (20)": rubrics_marks,
                "Script Marks (15)": script_marks,
                "Viva Marks (10)": viva_marks,
                "Total Internal (25)": total_internal_marks,
                "Attendance Marks (5)": attendance_marks,
                "Final Internal Marks (50)": final_internal,
            }
        )

    mark_types = {
        "Rubrics Marks (20)": "Lab rubrics evaluation by instructor",
        "Script Marks (15)": "Python programming script evaluation",
        "Viva Marks (10)": "Automatically calculated from quiz performance",
        "Total Internal (25)": "Combined script and viva marks",
        "Attendance Marks (5)": "Lab attendance evaluation",
        "Final Internal Marks (50)": "Total internal assessment marks",
    }

    return render_template(
        "python_lab_consolidated_marks.html",
        data=data,
        current_admin=current_admin,
        mark_types=mark_types,
    )


@app.route("/download_python_lab_marks_template")
def download_python_lab_marks_template():
    """Download CSV template for Python Lab marks upload"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    current_admin = user

    # Get students for this admin
    if current_admin.admin_level == 'super_admin':
        students = User.query.filter(User.role == 'student').all()
    else:
        students = get_students_for_admin(current_admin)
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # Create CSV template
    import io
    import csv
    from flask import make_response

    output = io.StringIO()
    writer = csv.writer(output)

    # Write headers
    writer.writerow([
        'Roll Number',
        'Rubrics Marks (20)',
        'Script Marks (15)',
        'Attendance Marks (5)'
    ])

    # Write student data with empty marks
    for student in students:
        writer.writerow([
            student.rollnumber,
            '',  # Rubrics Marks (20) - to be filled
            '',  # Script Marks (15) - to be filled
            ''   # Attendance Marks (5) - to be filled
        ])

    output.seek(0)

    response = make_response(output.getvalue())
    response.headers['Content-Type'] = 'text/csv'
    response.headers[
        'Content-Disposition'] = f'attachment; filename=python_lab_marks_template_{current_admin.id}.csv'

    return response


@app.route("/upload_python_lab_marks", methods=["POST"])
def upload_python_lab_marks():
    """Upload Python Lab marks from CSV file"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    current_admin = user

    if 'file' not in request.files:
        flash("No file selected!", "danger")
        return redirect(url_for("python_lab_consolidated_marks"))

    file = request.files['file']
    if file.filename == '':
        flash("No file selected!", "danger")
        return redirect(url_for("python_lab_consolidated_marks"))

    try:
        import pandas as pd
        import os

        # Save uploaded file
        upload_dir = os.path.join(os.getcwd(), "tmp_uploads")
        if not os.path.exists(upload_dir):
            os.makedirs(upload_dir)

        admin_id = current_admin.id
        upload_path = os.path.join(
            upload_dir, f"python_lab_marks_{admin_id}.csv")

        # Read and validate file
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file)
        elif file.filename.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file)
        else:
            flash("Please upload a CSV or Excel file!", "danger")
            return redirect(url_for("python_lab_consolidated_marks"))

        # Validate required columns
        required_columns = [
            'Roll Number', 'Rubrics Marks (20)', 'Script Marks (15)', 'Attendance Marks (5)']
        missing_columns = [
            col for col in required_columns if col not in df.columns]

        if missing_columns:
            flash(
                f"Missing required columns: {', '.join(missing_columns)}", "danger")
            return redirect(url_for("python_lab_consolidated_marks"))

        # Save as CSV
        df.to_csv(upload_path, index=False)

        flash(
            f"Python Lab marks uploaded successfully! {len(df)} records processed.", "success")

    except Exception as e:
        flash(f"Error processing file: {str(e)}", "danger")

    return redirect(url_for("python_lab_consolidated_marks"))


@app.route("/reset_academic_year_data", methods=["POST"])
def reset_academic_year_data():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.admin_level != "super_admin":
        flash("Only the main admin can use this reset.", "danger")
        return redirect(url_for("admin_dashboard"))

    if session.get("impersonating"):
        flash("Stop impersonation before running the yearly reset.", "warning")
        return redirect(url_for("admin_dashboard"))

    try:
        students = User.query.filter(User.role == 'student').all()
        student_ids = [student.id for student in students]

        if student_ids:
            dom_attempt_ids = [
                row[0] for row in db.session.query(QuizAttempt.id)
                .filter(QuizAttempt.user_id.in_(student_ids)).all()
            ]
            if dom_attempt_ids:
                QuizResponse.query.filter(
                    QuizResponse.attempt_id.in_(dom_attempt_ids)
                ).delete(synchronize_session=False)

            python_attempt_ids = [
                row[0] for row in db.session.query(PythonLabQuizAttempt.id)
                .filter(PythonLabQuizAttempt.student_id.in_(student_ids)).all()
            ]
            if python_attempt_ids:
                PythonLabQuizResponse.query.filter(
                    PythonLabQuizResponse.attempt_id.in_(python_attempt_ids)
                ).delete(synchronize_session=False)

            COMarks.query.filter(COMarks.user_id.in_(student_ids)).delete(
                synchronize_session=False)
            DOMSubjectStudentQuizRelease.query.filter(
                DOMSubjectStudentQuizRelease.student_id.in_(student_ids)
            ).delete(synchronize_session=False)
            KDMLabResponse.query.filter(KDMLabResponse.user_id.in_(student_ids)).delete(
                synchronize_session=False)
            KDMLabManualMarks.query.filter(
                KDMLabManualMarks.student_id.in_(student_ids)
            ).delete(synchronize_session=False)
            KDMLabAttempt.query.filter(KDMLabAttempt.student_id.in_(student_ids)).delete(
                synchronize_session=False)
            KDMLabStudentExperimentRelease.query.filter(
                KDMLabStudentExperimentRelease.student_id.in_(student_ids)
            ).delete(synchronize_session=False)
            KDMLabQuizAttempt.query.filter(
                KDMLabQuizAttempt.student_id.in_(student_ids)
            ).delete(synchronize_session=False)
            KDMLabStudentQuizRelease.query.filter(
                KDMLabStudentQuizRelease.student_id.in_(student_ids)
            ).delete(synchronize_session=False)
            PythonLabResponse.query.filter(
                PythonLabResponse.user_id.in_(student_ids)
            ).delete(synchronize_session=False)
            PythonLabManualMarks.query.filter(
                PythonLabManualMarks.student_id.in_(student_ids)
            ).delete(synchronize_session=False)
            PythonLabAttempt.query.filter(
                PythonLabAttempt.student_id.in_(student_ids)
            ).delete(synchronize_session=False)
            PythonLabStudentExperimentRelease.query.filter(
                PythonLabStudentExperimentRelease.student_id.in_(student_ids)
            ).delete(synchronize_session=False)
            PythonLabStudentQuizRelease.query.filter(
                PythonLabStudentQuizRelease.student_id.in_(student_ids)
            ).delete(synchronize_session=False)
            QuizAttempt.query.filter(QuizAttempt.user_id.in_(student_ids)).delete(
                synchronize_session=False)
            PythonLabQuizAttempt.query.filter(
                PythonLabQuizAttempt.student_id.in_(student_ids)
            ).delete(synchronize_session=False)
            User.query.filter(User.id.in_(student_ids)).delete(
                synchronize_session=False)

        KDMLabStudent.query.delete(synchronize_session=False)
        PythonLabStudent.query.delete(synchronize_session=False)
        AdminRollAssignment.query.delete(synchronize_session=False)

        removed_files = []
        for filename in os.listdir(UPLOAD_FOLDER):
            if not filename.endswith(".csv"):
                continue
            if filename.startswith(("internal_marks_", "kdm_internal_marks_", "python_lab_marks_")):
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                if os.path.isfile(filepath):
                    os.remove(filepath)
                    removed_files.append(filename)

        db.session.commit()
        flash(
            f"Academic year data cleared. Removed {len(student_ids)} students and {len(removed_files)} uploaded marks files. Admins and question banks were kept.",
            "success"
        )
    except Exception as e:
        db.session.rollback()
        flash(f"Year reset failed: {str(e)}", "danger")

    return redirect(url_for("admin_dashboard"))


@app.route("/admin_dashboard")
def admin_dashboard():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    # Ensure only admin or super_admin can access
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Sorting setup
    sort_by = request.args.get('sort_by', 'rollnumber')
    sort_order = request.args.get('sort_order', 'asc')

    # Current admin
    current_admin = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    def get_completed_problems_count(user):
        count = 0
        for i in range(1, 12):
            if getattr(user, f"p{i}_attempts", 0) > 0:
                count += 1
        return count

    # Handle data based on admin type
    if current_admin.admin_level == 'super_admin':
        base_query = User.query.filter(User.role == 'student')
        users = base_query.all()
        for u in users:
            u.completed_problems_count = get_completed_problems_count(u)
        if sort_by == 'rollnumber':
            users.sort(key=lambda x: x.rollnumber, reverse=(sort_order == 'desc'))
        elif sort_by == 'marks':
            users.sort(key=lambda x: x.marks or 0, reverse=(sort_order == 'desc'))
        elif sort_by == 'completed':
            users.sort(key=lambda x: x.completed_problems_count, reverse=(sort_order == 'desc'))
        else:
            users.sort(key=lambda x: x.id, reverse=(sort_order == 'desc'))

        all_admins = User.query.filter(User.admin_level.in_(['admin', 'super_admin'])).all()

    else:
        students = get_students_for_admin(current_admin)
        for u in students:
            u.completed_problems_count = get_completed_problems_count(u)
        if sort_by == 'rollnumber':
            students.sort(key=lambda x: x.rollnumber, reverse=(sort_order == 'desc'))
        elif sort_by == 'marks':
            students.sort(key=lambda x: x.marks or 0, reverse=(sort_order == 'desc'))
        elif sort_by == 'completed':
            students.sort(key=lambda x: x.completed_problems_count, reverse=(sort_order == 'desc'))
        else:
            students.sort(key=lambda x: x.id, reverse=(sort_order == 'desc'))
        users = students
        all_admins = []

    # Problem visibility setup
    problem_visibility = {}
    for i in range(1, 12):
        if current_admin.admin_level == 'super_admin':
            visibility = ProblemVisibility.query.filter_by(problem_number=i, admin_id=None).first()
        else:
            global_visibility = ProblemVisibility.query.filter_by(problem_number=i, admin_id=None).first()
            own_visibility = ProblemVisibility.query.filter_by(problem_number=i, admin_id=current_admin.id).first()
            is_globally_released = global_visibility.is_released if global_visibility else False
            is_own_released = own_visibility.is_released if own_visibility else False

            class VisibilityStatus:
                def __init__(self, is_released):
                    self.is_released = is_released

            visibility = VisibilityStatus(is_globally_released or is_own_released)

        problem_visibility[i] = visibility.is_released if visibility else False

    # Quiz release check
    quiz_released = False
    try:
        if current_admin.admin_level == 'super_admin':
            quiz_visibility = QuizVisibility.query.filter_by(quiz_name='default_quiz', admin_id=None).first()
            quiz_released = quiz_visibility.is_released if quiz_visibility else False
        else:
            global_quiz_visibility = QuizVisibility.query.filter_by(quiz_name='default_quiz', admin_id=None).first()
            own_quiz_visibility = QuizVisibility.query.filter_by(quiz_name='default_quiz', admin_id=current_admin.id).first()
            is_globally_released = global_quiz_visibility.is_released if global_quiz_visibility else False
            is_own_released = own_quiz_visibility.is_released if own_quiz_visibility else False
            quiz_released = is_globally_released or is_own_released
    except Exception as e:
        quiz_released = False
        print(f"⚠️ Error getting quiz visibility: {e}")

    # Quiz info per student
    try:
        for user in users:
            quiz_attempts = QuizAttempt.query.filter_by(user_id=user.id, quiz_name=DEFAULT_QUIZ_NAME).all()
            if quiz_attempts:
                user.quiz_best_score = max(a.score for a in quiz_attempts)
                user.quiz_total_points = quiz_attempts[0].total_points if quiz_attempts else 0
                user.quiz_attempts_count = len(quiz_attempts)
                latest_attempt = max(quiz_attempts, key=lambda x: x.started_at)
                user.quiz_latest_score = latest_attempt.score
                user.quiz_completed_at = latest_attempt.completed_at
            else:
                user.quiz_best_score = user.quiz_total_points = 0
                user.quiz_attempts_count = user.quiz_latest_score = 0
                user.quiz_completed_at = None
    except Exception as e:
        print(f"⚠️ Error getting quiz scores: {e}")
        for user in users:
            user.quiz_best_score = user.quiz_total_points = 0
            user.quiz_attempts_count = user.quiz_latest_score = 0
            user.quiz_completed_at = None

    setting = GlobalSettings.query.first()
    allow_student_mid_exam_questions = setting.allow_student_mid_exam_questions if setting else False
    allow_student_python_lab = setting.allow_student_python_lab if setting else True
    allow_admin_add_question = get_config('allow_admin_add_question', '0') == '1'

    # ✅ NEW BLOCK: Dashboard visibility controls (for super admin only)
    if current_admin.admin_level == 'super_admin':
        from sqlalchemy import text
        visibility_records = db.session.query(AdminDashboardVisibility).all()
        visibility_map = {}
        admin_student_counts = {}
        for v in visibility_records:
            if v.admin_id not in visibility_map:
                visibility_map[v.admin_id] = {}
            visibility_map[v.admin_id][v.dashboard_type] = v.is_visible

        for admin in all_admins:
            admin.visibility_settings = visibility_map.get(admin.id, {
                'dom': True,
                'kdm': True,
                'python': True
            })
            if admin.admin_level == 'admin':
                assignments = AdminRollAssignment.query.filter_by(admin_id=admin.id).all()
                admin_student_counts[admin.id] = sum(
                    count_rolls_in_assignment(assignment) for assignment in assignments
                )
            else:
                admin_student_counts[admin.id] = len(users)
    else:
        admin_student_counts = {}

    # --- Apply dashboard visibility settings for sub-admins ---
    if current_admin.admin_level != 'super_admin':
        visibilities = AdminDashboardVisibility.query.filter_by(admin_id=current_admin.id).all()
        visibility_map = {v.dashboard_type: v.is_visible for v in visibilities}

        # Use safe lookups, prefer new names
        show_dom_dashboard = visibility_map.get("dom")
        show_kdm_dashboard = visibility_map.get("kdm")
        show_python_dashboard = visibility_map.get("python")

        # Default to True only if not present at all
        if show_dom_dashboard is None:
            show_dom_dashboard = visibility_map.get("domsubject", True)
        if show_kdm_dashboard is None:
            show_kdm_dashboard = visibility_map.get("kdmlab", True)
        if show_python_dashboard is None:
            show_python_dashboard = visibility_map.get("pythonlab", True)
    else:
        show_dom_dashboard = show_kdm_dashboard = show_python_dashboard = True

    return render_template(
        "admin_dashboard.html",
        users=users,
        problem_visibility=problem_visibility,
        quiz_released=quiz_released,
        current_sort=sort_by,
        current_order=sort_order,
        current_admin=current_admin,
        all_admins=all_admins,
        admin_student_counts=admin_student_counts,
        count_students_for_admin=count_students_for_admin,
        is_super_admin=(current_admin.admin_level == 'super_admin'),
        allow_admin_add_question=allow_admin_add_question,
        allow_student_mid_exam_questions=allow_student_mid_exam_questions,
        allow_student_python_lab=allow_student_python_lab,
        show_dom_dashboard=show_dom_dashboard,
        show_kdm_dashboard=show_kdm_dashboard,
        show_python_dashboard=show_python_dashboard

    )
@app.route("/update_admin_dashboard_visibility", methods=["POST"])
def update_admin_dashboard_visibility():
    try:
        data = request.get_json()
        admin_id = int(data.get("admin_id"))
        dashboard_type = str(data.get("dashboard_type")).strip()
        is_visible = bool(data.get("is_visible"))

        if not admin_id or not dashboard_type:
            return jsonify({"status": "error", "message": "Invalid data"}), 400

        # 🔍 Ensure only one record per (admin_id, dashboard_type)
        record = AdminDashboardVisibility.query.filter(
            AdminDashboardVisibility.admin_id == admin_id,
            AdminDashboardVisibility.dashboard_type == dashboard_type
        ).first()

        if record:
            record.is_visible = is_visible
        else:
            record = AdminDashboardVisibility(
                admin_id=admin_id,
                dashboard_type=dashboard_type,
                is_visible=is_visible
            )
            db.session.add(record)

        db.session.commit()
        return jsonify({"status": "success", "message": "Visibility updated"})

    except Exception as e:
        print(f"⚠️ Error updating visibility: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route("/dom_subject_admin_dashboard")
def dom_subject_admin_dashboard():
    """DOM Subject dedicated admin dashboard"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    # Ensure only admin or super_admin can access
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    current_admin = user

    # Get sort parameters
    sort_by = request.args.get('sort_by', 'rollnumber')
    sort_order = request.args.get('sort_order', 'asc')  # Default ascending

    # Get students data for the table
    def get_completed_problems_count(user):
        count = 0
        for i in range(1, 12):
            attempts = getattr(user, f"p{i}_attempts", 0)
            if attempts > 0:
                count += 1
        return count

    # Filter users based on admin level
    if current_admin.admin_level == 'super_admin':
        base_query = User.query.filter(User.role == 'student')
        users = base_query.all()
        for u in users:
            u.completed_problems_count = get_completed_problems_count(u)
        if sort_by == 'rollnumber':
            users.sort(key=lambda x: x.rollnumber,
                       reverse=(sort_order == 'desc'))
        elif sort_by == 'marks':
            users.sort(key=lambda x: x.marks or 0,
                       reverse=(sort_order == 'desc'))
        elif sort_by == 'completed':
            users.sort(key=lambda x: x.completed_problems_count,
                       reverse=(sort_order == 'desc'))
        else:
            users.sort(key=lambda x: x.id, reverse=(sort_order == 'desc'))
    else:
        students = get_students_for_admin(current_admin)
        for u in students:
            u.completed_problems_count = get_completed_problems_count(u)
        if sort_by == 'rollnumber':
            students.sort(key=lambda x: x.rollnumber,
                          reverse=(sort_order == 'desc'))
        elif sort_by == 'marks':
            students.sort(key=lambda x: x.marks or 0,
                          reverse=(sort_order == 'desc'))
        elif sort_by == 'completed':
            students.sort(key=lambda x: x.completed_problems_count,
                          reverse=(sort_order == 'desc'))
        else:
            students.sort(key=lambda x: x.id, reverse=(sort_order == 'desc'))
        users = students

    # Add quiz score information for each user
    try:
        for user in users:
            # Get quiz attempts for this user
            quiz_attempts = QuizAttempt.query.filter_by(user_id=user.id, quiz_name=DEFAULT_QUIZ_NAME).all()

            if quiz_attempts:
                # Get best score and total attempts
                user.quiz_best_score = max(
                    attempt.score for attempt in quiz_attempts)
                user.quiz_total_points = quiz_attempts[0].total_points if quiz_attempts else 0
                user.quiz_attempts_count = len(quiz_attempts)

                # Get latest attempt details
                latest_attempt = max(quiz_attempts, key=lambda x: x.started_at)
                user.quiz_latest_score = latest_attempt.score
                user.quiz_completed_at = latest_attempt.completed_at
            else:
                user.quiz_best_score = 0
                user.quiz_total_points = 0
                user.quiz_attempts_count = 0
                user.quiz_latest_score = 0
                user.quiz_completed_at = None
    except Exception as e:
        print(f"⚠️ Error getting quiz scores: {e}")
        # Set default values for all users
        for user in users:
            user.quiz_best_score = 0
            user.quiz_total_points = 0
            user.quiz_attempts_count = 0
            user.quiz_latest_score = 0
            user.quiz_completed_at = None

    # Get statistics for DOM Subject dashboard
    total_students = len(users)
    total_quiz_questions = 0
    total_exam_questions = 0

    try:
        # Count quiz questions
        total_quiz_questions = QuizQuestion.query.count()

        # Count exam questions
        total_exam_questions = MidExamQuestion.query.count()

    except Exception as e:
        print(f"⚠️ Error getting DOM subject statistics: {e}")
        # Set defaults if tables don't exist
        pass    # Check admin permissions
    allow_admin_add_question = get_config(
        'allow_admin_add_question', '0') == '1'

    # Fetch problem visibility settings for current admin
    problem_visibility = {}
    for i in range(1, 12):  # Problems 1-11
        if current_admin.admin_level == 'super_admin':
            # Super admin sees global visibility status
            visibility = ProblemVisibility.query.filter_by(
                problem_number=i, admin_id=None).first()
        else:
            # Sub admin sees BOTH their own AND global visibility status
            # First check for global release (admin_id=None)
            global_visibility = ProblemVisibility.query.filter_by(
                problem_number=i, admin_id=None).first()
            # Then check for their own release
            own_visibility = ProblemVisibility.query.filter_by(
                problem_number=i, admin_id=current_admin.id).first()

            # Problem is released if EITHER global OR own release exists and is released
            is_globally_released = global_visibility.is_released if global_visibility else False
            is_own_released = own_visibility.is_released if own_visibility else False

            # Use a dummy visibility object to maintain compatibility
            class VisibilityStatus:
                def __init__(self, is_released):
                    self.is_released = is_released

            visibility = VisibilityStatus(
                is_globally_released or is_own_released)

        problem_visibility[i] = visibility.is_released if visibility else False

    # Get DOM quiz release status
    quiz_released = False
    try:
        if current_admin.admin_level == 'super_admin':
            # Super admin sees global quiz visibility status
            quiz_visibility = QuizVisibility.query.filter_by(
                quiz_name='default_quiz', admin_id=None).first()
            quiz_released = quiz_visibility.is_released if quiz_visibility else False
        else:
            # Sub admin sees BOTH their own AND global quiz visibility status
            # First check for global release (admin_id=None)
            global_quiz_visibility = QuizVisibility.query.filter_by(
                quiz_name='default_quiz', admin_id=None).first()
            # Then check for their own release
            own_quiz_visibility = QuizVisibility.query.filter_by(
                quiz_name='default_quiz', admin_id=current_admin.id).first()

            # Quiz is released if EITHER global OR own release exists and is released
            is_globally_released = global_quiz_visibility.is_released if global_quiz_visibility else False
            is_own_released = own_quiz_visibility.is_released if own_quiz_visibility else False

            quiz_released = is_globally_released or is_own_released
    except Exception as e:
        # Quiz tables don't exist yet - set default value
        quiz_released = False
        print(f"⚠️ Error getting DOM quiz visibility: {e}")

    initialize_dom_conceptual_testing_system()
    conceptual_questions_by_co = {}
    for co in QUIZ_CO_OPTIONS:
        conceptual_questions_by_co[co] = [
            get_conceptual_question_payload(question)
            for question in QuizQuestion.query.filter_by(co_number=co).order_by(QuizQuestion.id.asc()).all()
        ]
    active_conceptual_session = get_dom_conceptual_active_session_for_admin(current_admin)
    conceptual_stats = get_dom_conceptual_stats(active_conceptual_session, current_admin)

    return render_template("dom_subject_admin_dashboard.html",
                           current_admin=current_admin,
                           allow_admin_add_question=allow_admin_add_question,
                           quiz_released=quiz_released,
                           quiz_co_options=QUIZ_CO_OPTIONS,
                           conceptual_questions_by_co=conceptual_questions_by_co,
                           active_conceptual_session=active_conceptual_session,
                           conceptual_stats=conceptual_stats,
                           users=users,
                           current_sort=sort_by,
                           current_order=sort_order,
                           problem_visibility=problem_visibility,
                           total_students=total_students,
                           total_quiz_questions=total_quiz_questions,
                           total_exam_questions=total_exam_questions)


@app.route("/dom_conceptual_post", methods=["POST"])
def dom_conceptual_post():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    admin_user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not admin_user or admin_user.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    initialize_dom_conceptual_testing_system()
    co_number = request.form.get("co_number")
    question_id = request.form.get("question_id", type=int)
    question = QuizQuestion.query.get(question_id) if question_id else None

    if co_number not in QUIZ_CO_OPTIONS or not question or question.co_number != co_number:
        flash("Please select one valid question under the selected CO.", "danger")
        return redirect(url_for("dom_subject_admin_dashboard"))

    DOMConceptualQuizSession.query.filter_by(
        admin_id=admin_user.id,
        is_active=True
    ).update({"is_active": False, "closed_at": datetime.now()})

    conceptual_session = DOMConceptualQuizSession(
        admin_id=admin_user.id,
        question_id=question.id,
        is_active=True
    )
    db.session.add(conceptual_session)
    db.session.commit()

    flash("Conceptual testing question posted to students.", "success")
    return redirect(url_for("dom_subject_admin_dashboard"))


@app.route("/dom_conceptual_close", methods=["POST"])
def dom_conceptual_close():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    admin_user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not admin_user or admin_user.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    active_session = get_dom_conceptual_active_session_for_admin(admin_user)
    if active_session:
        active_session.is_active = False
        active_session.closed_at = datetime.now()
        db.session.commit()
        flash("Conceptual testing question closed.", "success")
    else:
        flash("No active conceptual testing question to close.", "warning")

    return redirect(url_for("dom_subject_admin_dashboard"))


@app.route("/dom_conceptual_stats")
def dom_conceptual_stats():
    if "loggedin" not in session:
        return jsonify({"error": "Not logged in"}), 401

    admin_user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not admin_user or admin_user.role not in ["admin", "super_admin"]:
        return jsonify({"error": "Admin access required"}), 403

    active_session = get_dom_conceptual_active_session_for_admin(admin_user)
    stats = get_dom_conceptual_stats(active_session, admin_user)
    return jsonify({
        "active": bool(active_session),
        "session_id": active_session.id if active_session else None,
        **stats
    })


@app.route("/dom_conceptual_answer", methods=["POST"])
def dom_conceptual_answer():
    if "loggedin" not in session or "rollnumber" not in session:
        return redirect(url_for("login"))

    student = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not student or student.role != "student":
        flash("Only students can answer conceptual testing questions.", "danger")
        return redirect(url_for("login"))

    active_session = get_dom_conceptual_active_session_for_student(student)
    session_id = request.form.get("session_id", type=int)
    if not active_session or active_session.id != session_id:
        flash("This conceptual testing question is no longer active.", "warning")
        return redirect(url_for("dom_subject_student_dashboard"))

    if get_dom_conceptual_response(active_session.id, student.id):
        flash("You already answered this conceptual testing question.", "info")
        return redirect(url_for("dom_subject_student_dashboard"))

    question = active_session.question
    selected_indices = []
    for value in request.form.getlist("selected_answer"):
        try:
            selected_indices.append(int(value))
        except (TypeError, ValueError):
            continue

    if not selected_indices:
        flash("Please select an answer before submitting.", "warning")
        return redirect(url_for("dom_subject_student_dashboard"))

    choice_ids = ['A', 'B', 'C', 'D']
    selected_letter_ids = [
        choice_ids[index]
        for index in selected_indices
        if 0 <= index < len(choice_ids)
    ]
    question_score = calculate_question_score(question, selected_letter_ids)
    is_correct = question_score == question.points

    response = DOMConceptualQuizResponse(
        session_id=active_session.id,
        student_id=student.id,
        selected_answer=selected_indices,
        is_correct=is_correct
    )
    db.session.add(response)
    db.session.commit()

    flash("Your conceptual testing answer was submitted.", "success")
    return redirect(url_for("dom_subject_student_dashboard"))


@app.route("/dom_subject_individual_quiz_releases")
def dom_subject_individual_quiz_releases():
    """Manage individual student quiz releases for DOM Subject Quiz"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    # Ensure only admin or super_admin can access
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    # Get students based on admin level
    if current_admin.admin_level == "super_admin":
        students = User.query.filter_by(role="student").all()
    else:
        students = get_students_for_admin(current_admin)

    # Sort students by roll number
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # Get existing individual quiz releases
    individual_releases = {}
    releases = DOMSubjectStudentQuizRelease.query.filter_by(
        admin_id=current_admin.id
    ).all()

    for release in releases:
        individual_releases[release.student_id] = release.is_released

    # Check if quiz is bulk released
    bulk_released = False
    try:
        if current_admin.admin_level == "super_admin":
            # Check global quiz visibility
            quiz_visibility = QuizVisibility.query.filter_by(
                quiz_name='default_quiz', admin_id=None).first()
            bulk_released = quiz_visibility.is_released if quiz_visibility else False
        else:
            # Check if this admin has released the quiz in bulk
            quiz_visibility = QuizVisibility.query.filter_by(
                quiz_name='default_quiz', admin_id=current_admin.id).first()
            bulk_released = quiz_visibility.is_released if quiz_visibility else False
    except Exception as e:
        print(f"⚠️ Error checking bulk quiz release status: {e}")
        bulk_released = False

    return render_template(
        "dom_subject_individual_quiz_releases.html",
        students=students,
        individual_releases=individual_releases,
        bulk_released=bulk_released,
        current_admin=current_admin
    )


@app.route("/dom_subject_toggle_individual_quiz_release", methods=["POST"])
def dom_subject_toggle_individual_quiz_release():
    """Toggle individual quiz release for a specific student"""
    if "loggedin" not in session:
        return jsonify({"error": "Not logged in"}), 401

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        return jsonify({"error": "Admin access required"}), 403

    try:
        data = request.get_json()
        student_id = data.get('student_id')
        is_released = data.get('is_released')

        # Find existing release record
        release_record = DOMSubjectStudentQuizRelease.query.filter_by(
            admin_id=current_admin.id,
            student_id=student_id
        ).first()

        if release_record:
            release_record.is_released = is_released
        else:
            # Create new release record
            release_record = DOMSubjectStudentQuizRelease(
                admin_id=current_admin.id,
                student_id=student_id,
                is_released=is_released
            )
            db.session.add(release_record)

        db.session.commit()
        return jsonify({"success": True, "is_released": is_released})

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


def roll_sort_key(roll):
    if not roll:
        return (1, [])
    parts = re.split(r'(\d+)', roll)
    return (0, [int(p) if p.isdigit() else p.lower() for p in parts])


@app.route("/download_internal_marks_template")
def download_internal_marks_template():
    import pandas as pd
    from io import BytesIO
    from flask import send_file

    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Get students list
    if current_admin.admin_level == 'super_admin':
        students = (User.query.filter(User.role == 'student')
                    .order_by(User.rollnumber.asc()).all())
    else:
        students = get_students_for_admin(current_admin)
        students.sort(key=lambda s: s.rollnumber or "")

    # Prepare blank template
    rows = []
    for s in students:
        rows.append({
            "Roll Number": s.rollnumber,
            "MID-I Marks (20)": "",
            "MID-II Marks (20)": "",
            "Attendance Marks (5)": ""
        })

    import pandas as pd
    df = pd.DataFrame(rows)
    csv_data = df.to_csv(index=False)
    output = BytesIO(csv_data.encode("utf-8"))
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name="internal_marks_template.csv",
        mimetype="text/csv"
    )


@app.route("/kdm_lab_reset_individual_experiment", methods=["POST"])
def kdm_lab_reset_individual_experiment():
    """Reset all experiment attempts and marks for a specific student and experiment (admin-specific)"""
    if "loggedin" not in session:
        return jsonify({"error": "Not logged in"}), 401

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        return jsonify({"error": "Admin access required"}), 403

    try:
        data = request.get_json()
        student_id = data.get('student_id')
        experiment_id = data.get('experiment_id')
        if not student_id or not experiment_id:
            return jsonify({"error": "Missing student_id or experiment_id"}), 400

        # Delete all experiment attempts for this student and experiment for this admin
        attempts = KDMLabAttempt.query.filter_by(
            student_id=student_id, experiment_id=experiment_id).all()
        for attempt in attempts:
            db.session.delete(attempt)

        # Optionally, reset individual release (set is_released to False)
        release = KDMLabStudentExperimentRelease.query.filter_by(
            admin_id=current_admin.id, student_id=student_id, experiment_id=experiment_id).first()
        if release:
            release.is_released = False

        db.session.commit()
        return jsonify({"success": True})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@app.route("/upload_internal_marks", methods=["POST"])
def upload_internal_marks():
    import pandas as pd
    import os

    if "loggedin" not in session:
        return redirect(url_for("login"))

    file = request.files.get("file")
    if not file or file.filename == "":
        flash("Please select a CSV file to upload.", "warning")
        return redirect(url_for("admin_consolidated_marks"))

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin:
        flash("Unauthorized access.", "danger")
        return redirect(url_for("login"))

    # Save uploaded CSV file
    admin_id = current_admin.id
    filepath = os.path.join(UPLOAD_FOLDER, f"internal_marks_{admin_id}.csv")
    file.save(filepath)

    flash("Internal marks uploaded successfully.", "success")
    return redirect(url_for("admin_consolidated_marks"))


DOM_CONSOLIDATION_TOTAL = 15


def get_dom_consolidation_settings_key(admin_id):
    return f"dom_consolidation_settings_{admin_id}"


def get_default_dom_consolidation_settings(admin_user=None):
    quiz_count_default = 0
    if admin_user:
        try:
            quiz_count_default = len(get_admin_quiz_banks(admin_user))
        except Exception:
            quiz_count_default = 0

    return {
        "quiz_best_count": quiz_count_default,
        "quiz_selection_mode": "overall",
        "quiz_co_best_counts": {co: 0 for co in QUIZ_CO_OPTIONS},
        "quiz_target_marks": 5.0,
        "assignment_best_count": 11,
        "assignment_target_marks": 10.0,
    }


def get_dom_consolidation_settings(admin_user):
    settings = get_default_dom_consolidation_settings(admin_user)
    raw_value = get_config(get_dom_consolidation_settings_key(admin_user.id))
    if not raw_value:
        return settings

    try:
        stored = json.loads(raw_value)
        settings.update({
            "quiz_best_count": int(stored.get("quiz_best_count", settings["quiz_best_count"])),
            "quiz_selection_mode": stored.get("quiz_selection_mode", settings["quiz_selection_mode"]),
            "quiz_co_best_counts": {
                co: int(stored.get("quiz_co_best_counts", {}).get(co, 0))
                for co in QUIZ_CO_OPTIONS
            },
            "quiz_target_marks": float(stored.get("quiz_target_marks", settings["quiz_target_marks"])),
            "assignment_best_count": int(stored.get("assignment_best_count", settings["assignment_best_count"])),
            "assignment_target_marks": float(stored.get("assignment_target_marks", settings["assignment_target_marks"])),
        })
    except Exception:
        return settings

    if settings["quiz_selection_mode"] not in ["overall", "co_minimum"]:
        settings["quiz_selection_mode"] = "overall"

    return settings


def save_dom_consolidation_settings(admin_user, settings):
    set_config(
        get_dom_consolidation_settings_key(admin_user.id),
        json.dumps(settings, separators=(",", ":"))
    )


def get_dom_admin_quiz_names(admin_user):
    try:
        quiz_names = {
            quiz.quiz_name for quiz in get_admin_quiz_banks(admin_user)
            if getattr(quiz, "quiz_name", None)
        }
        if DEFAULT_QUIZ_NAME not in quiz_names:
            quiz_names.add(DEFAULT_QUIZ_NAME)
        return sorted(quiz_names)
    except Exception as e:
        print(f"⚠️ Error getting admin quiz names for consolidation: {e}")
        return [DEFAULT_QUIZ_NAME]


def get_dom_admin_quiz_bank_map(admin_user):
    try:
        return {
            quiz.quiz_name: quiz
            for quiz in get_admin_quiz_banks(admin_user)
            if getattr(quiz, "quiz_name", None)
        }
    except Exception as e:
        print(f"Error getting admin quiz banks for consolidation: {e}")
        return {}


def get_dom_quiz_cos(quiz_name, quiz_bank_map=None):
    quiz_bank_map = quiz_bank_map or {}
    quiz_bank = quiz_bank_map.get(quiz_name) or get_quiz_bank_by_name(quiz_name)
    cos = []
    if quiz_bank and quiz_bank.covered_cos:
        cos = parse_quiz_cos(quiz_bank.covered_cos)

    if not cos and quiz_bank:
        cos = sorted({
            assignment.question.co_number
            for assignment in getattr(quiz_bank, "question_assignments", [])
            if assignment.question and assignment.question.co_number in QUIZ_CO_OPTIONS
        })

    return cos


def load_uploaded_internal_marks(admin_id):
    uploaded_data = []
    uploaded_lookup = {}
    try:
        upload_path = os.path.join(
            os.getcwd(), "tmp_uploads", f"internal_marks_{admin_id}.csv")
        if os.path.exists(upload_path):
            df = pd.read_csv(upload_path)
            uploaded_data = df.to_dict(orient="records")
            for row in uploaded_data:
                roll = str(row.get("Roll Number", "")).strip()
                if roll:
                    uploaded_lookup[roll] = row
    except Exception as e:
        print(f"⚠️ Unable to read uploaded internal marks file for admin {admin_id}: {e}")

    return uploaded_data, uploaded_lookup


def parse_numeric_mark(value, default="N/A"):
    if value is None:
        return default
    text = str(value).strip()
    if not text or text.lower() == "nan":
        return default
    return value


def calculate_dom_assignment_component(student, best_count, target_marks):
    components = []
    for problem_no in range(1, 12):
        score = float(getattr(student, f"p{problem_no}_score", 0) or 0)
        max_marks = float(get_max_marks_for_problem(problem_no) or 0)
        percentage = (score / max_marks) if max_marks > 0 else 0
        components.append({
            "problem_no": problem_no,
            "score": score,
            "max_marks": max_marks,
            "percentage": percentage
        })

    selected = sorted(
        components,
        key=lambda item: (item["percentage"], item["score"], item["max_marks"], -item["problem_no"]),
        reverse=True
    )[:max(0, best_count)]

    total_scored = sum(item["score"] for item in selected)
    total_possible = sum(item["max_marks"] for item in selected)
    converted = round((total_scored / total_possible) * target_marks, 2) if total_possible > 0 else 0.0

    return {
        "selected_count": len(selected),
        "total_scored": total_scored,
        "total_possible": total_possible,
        "converted_marks": converted,
    }


def select_dom_quiz_components(components, settings):
    best_count = max(0, int(settings.get("quiz_best_count", 0) or 0))
    ranked = sorted(
        components,
        key=lambda item: (item["percentage"], item["score"], item["total_points"], item["quiz_name"]),
        reverse=True
    )

    if settings.get("quiz_selection_mode") != "co_minimum":
        return ranked[:best_count]

    selected = []
    selected_quizzes = set()
    co_counts = settings.get("quiz_co_best_counts", {}) or {}

    for co in QUIZ_CO_OPTIONS:
        co_quota = max(0, int(co_counts.get(co, 0) or 0))
        if co_quota <= 0:
            continue

        co_candidates = [
            item for item in ranked
            if item["quiz_name"] not in selected_quizzes and co in item.get("cos", [])
        ]
        for item in co_candidates[:co_quota]:
            selected.append(item)
            selected_quizzes.add(item["quiz_name"])

    remaining_slots = max(0, best_count - len(selected))
    if remaining_slots:
        for item in ranked:
            if item["quiz_name"] in selected_quizzes:
                continue
            selected.append(item)
            selected_quizzes.add(item["quiz_name"])
            remaining_slots -= 1
            if remaining_slots == 0:
                break

    return selected[:best_count]


def calculate_dom_quiz_component(student, quiz_names, target_marks, settings=None, quiz_bank_map=None):
    settings = settings or {"quiz_best_count": 0, "quiz_selection_mode": "overall", "quiz_co_best_counts": {}}
    quiz_bank_map = quiz_bank_map or {}
    components = []
    for quiz_name in quiz_names:
        attempts = (
            QuizAttempt.query.filter_by(user_id=student.id, quiz_name=quiz_name)
            .filter(QuizAttempt.completed_at.isnot(None))
            .all()
        )
        if not attempts:
            continue

        best_attempt = max(
            attempts,
            key=lambda attempt: ((attempt.score or 0), (attempt.total_points or 0), attempt.completed_at or attempt.started_at or datetime.min)
        )
        total_points = float(best_attempt.total_points or 0)
        score = float(best_attempt.score or 0)
        percentage = (score / total_points) if total_points > 0 else 0
        components.append({
            "quiz_name": quiz_name,
            "cos": get_dom_quiz_cos(quiz_name, quiz_bank_map),
            "score": score,
            "total_points": total_points,
            "percentage": percentage
        })

    selected = select_dom_quiz_components(components, settings)

    total_scored = sum(item["score"] for item in selected)
    total_possible = sum(item["total_points"] for item in selected)
    converted = round((total_scored / total_possible) * target_marks, 2) if total_possible > 0 else 0.0

    return {
        "selected_count": len(selected),
        "available_count": len(components),
        "total_scored": total_scored,
        "total_possible": total_possible,
        "converted_marks": converted,
        "selection_mode": settings.get("quiz_selection_mode", "overall"),
    }


def calculate_dom_mid_average(mid1, mid2):
    try:
        if str(mid1).replace('.', '', 1).isdigit() and str(mid2).replace('.', '', 1).isdigit():
            m1, m2 = float(mid1), float(mid2)
            best, second = sorted([m1, m2], reverse=True)
            return round((2 / 3) * best + (1 / 3) * second, 2)
    except Exception:
        pass
    return "N/A"


def build_dom_consolidated_row(student, current_admin, uploaded_lookup=None, settings=None, quiz_names=None, quiz_bank_map=None):
    uploaded_lookup = uploaded_lookup or {}
    settings = settings or get_dom_consolidation_settings(current_admin)
    quiz_names = quiz_names or get_dom_admin_quiz_names(current_admin)
    quiz_bank_map = quiz_bank_map or get_dom_admin_quiz_bank_map(current_admin)

    row = uploaded_lookup.get(str(student.rollnumber).strip(), {})
    mid1 = parse_numeric_mark(row.get("MID-I Marks (20)", "N/A"))
    mid2 = parse_numeric_mark(row.get("MID-II Marks (20)", "N/A"))
    attendance = parse_numeric_mark(row.get("Attendance Marks (5)", "N/A"))

    assignment_component = calculate_dom_assignment_component(
        student,
        best_count=settings["assignment_best_count"],
        target_marks=settings["assignment_target_marks"]
    )
    quiz_component = calculate_dom_quiz_component(
        student,
        quiz_names=quiz_names,
        target_marks=settings["quiz_target_marks"],
        settings=settings,
        quiz_bank_map=quiz_bank_map
    )

    avg_mid = calculate_dom_mid_average(mid1, mid2)
    attendance_value = 0.0
    if str(attendance).replace('.', '', 1).isdigit():
        attendance_value = float(attendance)

    final_internal = "N/A"
    if isinstance(avg_mid, (int, float)):
        final_internal = math.ceil(
            float(avg_mid)
            + assignment_component["converted_marks"]
            + quiz_component["converted_marks"]
            + attendance_value
        )

    return {
        "Roll Number": student.rollnumber,
        "MID-I Marks (20)": mid1,
        "MID-II Marks (20)": mid2,
        "Average of MIDs": avg_mid,
        "Assignment Marks (10)": assignment_component["converted_marks"],
        "Quiz Marks (5)": quiz_component["converted_marks"],
        "Attendance Marks (5)": attendance,
        "Final Internal Marks (40)": final_internal,
        "assignment_details": assignment_component,
        "quiz_details": quiz_component,
    }


@app.route("/admin_consolidated_marks", methods=["GET", "POST"])
def admin_consolidated_marks():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    current_admin = user

    if request.method == "POST":
        defaults = get_default_dom_consolidation_settings(current_admin)

        def clamp_int(name, fallback, minimum=0, maximum=100):
            raw_value = request.form.get(name, fallback)
            try:
                value = int(raw_value)
            except (TypeError, ValueError):
                value = fallback
            return max(minimum, min(maximum, value))

        def clamp_float(name, fallback, minimum=0.0, maximum=DOM_CONSOLIDATION_TOTAL):
            raw_value = request.form.get(name, fallback)
            try:
                value = float(raw_value)
            except (TypeError, ValueError):
                value = fallback
            return max(minimum, min(maximum, value))

        quiz_target_marks = clamp_float(
            "quiz_target_marks", defaults["quiz_target_marks"])
        assignment_target_marks = clamp_float(
            "assignment_target_marks", defaults["assignment_target_marks"])

        if round(quiz_target_marks + assignment_target_marks, 2) != DOM_CONSOLIDATION_TOTAL:
            assignment_target_marks = round(
                max(0.0, DOM_CONSOLIDATION_TOTAL - quiz_target_marks), 2)
            flash(
                f"Assignment conversion marks were adjusted to {assignment_target_marks} so quiz + assignment stays {DOM_CONSOLIDATION_TOTAL}.",
                "info"
            )

        quiz_selection_mode = request.form.get("quiz_selection_mode", "overall")
        if quiz_selection_mode not in ["overall", "co_minimum"]:
            quiz_selection_mode = "overall"

        quiz_best_count = clamp_int("quiz_best_count", defaults["quiz_best_count"])
        quiz_co_best_counts = {co: 0 for co in QUIZ_CO_OPTIONS}
        if quiz_selection_mode == "co_minimum":
            quiz_co_best_counts = {
                co: clamp_int(f"quiz_co_best_counts_{co}", 0)
                for co in QUIZ_CO_OPTIONS
            }
            co_minimum_total = sum(quiz_co_best_counts.values())
            if co_minimum_total > quiz_best_count:
                flash(
                    f"CO-wise quiz counts total {co_minimum_total}, but overall best quizzes is {quiz_best_count}. Increase overall count or reduce CO counts.",
                    "danger"
                )
                return redirect(url_for("admin_consolidated_marks"))

        settings = {
            "quiz_best_count": quiz_best_count,
            "quiz_selection_mode": quiz_selection_mode,
            "quiz_co_best_counts": quiz_co_best_counts,
            "quiz_target_marks": quiz_target_marks,
            "assignment_best_count": clamp_int("assignment_best_count", defaults["assignment_best_count"], maximum=11),
            "assignment_target_marks": assignment_target_marks,
        }
        save_dom_consolidation_settings(current_admin, settings)
        flash("Consolidation settings updated successfully.", "success")
        return redirect(url_for("admin_consolidated_marks"))

    if current_admin.admin_level == 'super_admin':
        students = User.query.filter(User.role == 'student').all()
    else:
        students = get_students_for_admin(current_admin)
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    settings = get_dom_consolidation_settings(current_admin)
    quiz_names = get_dom_admin_quiz_names(current_admin)
    quiz_bank_map = get_dom_admin_quiz_bank_map(current_admin)
    quiz_cos_summary = {
        quiz_name: get_dom_quiz_cos(quiz_name, quiz_bank_map)
        for quiz_name in quiz_names
    }
    _, uploaded_lookup = load_uploaded_internal_marks(current_admin.id)

    data = []
    for s in students:
        data.append(build_dom_consolidated_row(
            s,
            current_admin=current_admin,
            uploaded_lookup=uploaded_lookup,
            settings=settings,
            quiz_names=quiz_names,
            quiz_bank_map=quiz_bank_map
        ))

    return render_template(
        "admin_consolidated_marks.html",
        data=data,
        current_admin=current_admin,
        settings=settings,
        quiz_names=quiz_names,
        quiz_cos_summary=quiz_cos_summary,
        quiz_co_options=QUIZ_CO_OPTIONS,
        dom_consolidation_total=DOM_CONSOLIDATION_TOTAL
    )

@app.route("/download_consolidated_marks")
def download_consolidated_marks():
    from io import BytesIO

    if "loggedin" not in session:
        return redirect(url_for("login"))

    current_admin = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    if current_admin.admin_level == 'super_admin':
        students = (
            User.query.filter(User.role == 'student')
            .order_by(User.rollnumber.asc())
            .all()
        )
    else:
        students = get_students_for_admin(current_admin)
        students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    settings = get_dom_consolidation_settings(current_admin)
    quiz_names = get_dom_admin_quiz_names(current_admin)
    quiz_bank_map = get_dom_admin_quiz_bank_map(current_admin)
    _, uploaded_lookup = load_uploaded_internal_marks(current_admin.id)

    rows = []
    export_columns = [
        "Roll Number",
        "MID-I Marks (20)",
        "MID-II Marks (20)",
        "Average of MIDs",
        "Assignment Marks (10)",
        "Quiz Marks (5)",
        "Attendance Marks (5)",
        "Final Internal Marks (40)",
    ]
    for s in students:
        row = build_dom_consolidated_row(
            s,
            current_admin=current_admin,
            uploaded_lookup=uploaded_lookup,
            settings=settings,
            quiz_names=quiz_names,
            quiz_bank_map=quiz_bank_map
        )
        rows.append({column: row[column] for column in export_columns})

    df = pd.DataFrame(rows)
    csv_data = df.to_csv(index=False)
    output = BytesIO(csv_data.encode("utf-8"))
    output.seek(0)

    print(f"Final Consolidated CSV ready: {len(rows)} records")

    return send_file(
        output,
        as_attachment=True,
        download_name="final_internal_marks.csv",
        mimetype="text/csv"
    )

@app.route("/export_users_csv")
def export_users_csv():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    # Ensure only admin or super_admin can access
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    # Get current admin user
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    # Filter users based on admin level
    if current_admin.admin_level == 'super_admin':
        # Super admin sees all users
        users = User.query.all()
    else:
        # Sub admin sees only assigned students
        users = get_students_for_admin(current_admin)

    # Create CSV content
    output = []
    # CSV Header
    output.append("ID,Roll Number,Password,Email,Phone,Role,Completed,Total Score,P1 Score,P1 Attempts,P2 Score,P2 Attempts,P3 Score,P3 Attempts,P4 Score,P4 Attempts,P5 Score,P5 Attempts,P6 Score,P6 Attempts,P7 Score,P7 Attempts,P8 Score,P8 Attempts,P9 Score,P9 Attempts,P10 Score,P10 Attempts,P11 Score,P11 Attempts")

    # Add user data
    for u in users:
        row = f"{u.id},{u.rollnumber},{u.password},{u.email or ''},{u.phonenumber or ''},{u.role},{u.completed or 0},{u.marks or 0},{u.p1_score or 0},{u.p1_attempts or 0},{u.p2_score or 0},{u.p2_attempts or 0},{u.p3_score or 0},{u.p3_attempts or 0},{u.p4_score or 0},{u.p4_attempts or 0},{u.p5_score or 0},{u.p5_attempts or 0},{u.p6_score or 0},{u.p6_attempts or 0},{u.p7_score or 0},{u.p7_attempts or 0},{u.p8_score or 0},{u.p8_attempts or 0},{u.p9_score or 0},{u.p9_attempts or 0},{u.p10_score or 0},{u.p10_attempts or 0},{u.p11_score or 0},{u.p11_attempts or 0}"
        output.append(row)

    csv_content = "\n".join(output)
    response = make_response(csv_content)
    response.headers["Content-Disposition"] = "attachment; filename=users_export.csv"
    response.headers["Content-Type"] = "text/csv"
    return response


@app.route("/edit_user/<int:user_id>", methods=["GET", "POST"])
def edit_user(user_id):
    if "loggedin" not in session or session.get("role") != "admin":
        flash("Unauthorized access", "danger")
        return redirect(url_for("login"))

    # Get current admin user
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    user = User.query.get(user_id)
    if not user:
        flash("User not found", "danger")
        return redirect(url_for("admin_dashboard"))

    # Check if admin has permission to edit this user
    if current_admin.admin_level == 'admin':
        # Sub admin - check if user's roll number is in assigned range
        assigned_roll_numbers = get_assigned_roll_numbers(current_admin.id)
        if user.rollnumber not in assigned_roll_numbers:
            flash("You don't have permission to edit this user", "danger")
            return redirect(url_for("admin_dashboard"))

    if request.method == "POST":
        # Update basic fields from form
        user.rollnumber = request.form.get("rollnumber")
        user.email = request.form.get("email")
        user.phonenumber = request.form.get("phonenumber")
        user.marks = int(request.form.get("marks", 0))
        user.completed = int(request.form.get("completed", 0))

        # Role assignment: sub-admins can only assign student role
        if current_admin.admin_level == 'super_admin':
            # Super admin can set any role
            user.role = request.form.get("role")
        else:
            # Sub admin can only assign student role
            user.role = "student"

        # Update problem scores and attempts
        user.p1_score = int(request.form.get("p1_score", 0))
        user.p1_attempts = int(request.form.get("p1_attempts", 0))
        user.p2_score = int(request.form.get("p2_score", 0))
        user.p2_attempts = int(request.form.get("p2_attempts", 0))
        user.p3_score = int(request.form.get("p3_score", 0))
        user.p3_attempts = int(request.form.get("p3_attempts", 0))
        user.p4_score = int(request.form.get("p4_score", 0))
        user.p4_attempts = int(request.form.get("p4_attempts", 0))
        user.p5_score = int(request.form.get("p5_score", 0))
        user.p5_attempts = int(request.form.get("p5_attempts", 0))
        user.p6_score = int(request.form.get("p6_score", 0))
        user.p6_attempts = int(request.form.get("p6_attempts", 0))
        user.p7_score = int(request.form.get("p7_score", 0))
        user.p7_attempts = int(request.form.get("p7_attempts", 0))
        user.p8_score = int(request.form.get("p8_score", 0))
        user.p8_attempts = int(request.form.get("p8_attempts", 0))
        user.p9_score = int(request.form.get("p9_score", 0))
        user.p9_attempts = int(request.form.get("p9_attempts", 0))
        user.p10_score = int(request.form.get("p10_score", 0))
        user.p10_attempts = int(request.form.get("p10_attempts", 0))
        user.p11_score = int(request.form.get("p11_score", 0))
        user.p11_attempts = int(request.form.get("p11_attempts", 0))

        # Recalculate total marks based on individual scores
        user.marks = (user.p1_score + user.p2_score + user.p3_score +
                      user.p4_score + user.p5_score + user.p6_score + user.p7_score + user.p8_score + user.p9_score + user.p10_score + user.p11_score)

        db.session.commit()
        flash("✅ User updated successfully.", "success")
        return redirect(url_for("admin_dashboard"))

    return render_template("edit_user.html", user=user, current_admin=current_admin)


@app.route("/delete_user/<int:user_id>", methods=["POST"])
def delete_user(user_id):
    if "loggedin" not in session or session.get("role") != "admin":
        flash("Unauthorized access", "danger")
        return redirect(url_for("login"))

    # Get current admin user
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    user = User.query.get_or_404(user_id)

    # Check if admin has permission to delete this user
    if current_admin.admin_level == 'admin':
        # Sub admin - check if user's roll number is in assigned range
        assigned_roll_numbers = get_assigned_roll_numbers(current_admin.id)
        if user.rollnumber not in assigned_roll_numbers:
            flash("You don't have permission to delete this user", "danger")
            return redirect(url_for("admin_dashboard"))

    db.session.delete(user)
    db.session.commit()
    flash("User deleted successfully!", "success")
    return redirect(url_for("admin_dashboard"))


@app.route("/reset_user/<int:user_id>", methods=["POST"])
def reset_user(user_id):
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Unauthorized access", "danger")
        return redirect(url_for("login"))

    # Get current admin user
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    user = User.query.get_or_404(user_id)

    # Check if admin has permission to reset this user
    if current_admin.admin_level == 'admin':
        # Sub admin - check if user's roll number is in assigned range
        assigned_roll_numbers = get_assigned_roll_numbers(current_admin.id)
        if user.rollnumber not in assigned_roll_numbers:
            flash("You don't have permission to reset this user", "danger")
            return redirect(url_for("admin_dashboard"))

    # Reset all progress to initial values
    user.marks = 0
    user.completed = 0
    user.p1_attempts = 0
    user.p1_score = 0
    user.p2_attempts = 0
    user.p2_score = 0
    user.p3_attempts = 0
    user.p3_score = 0
    user.p4_attempts = 0
    user.p4_score = 0
    user.p5_attempts = 0
    user.p5_score = 0
    user.p6_attempts = 0
    user.p6_score = 0
    user.p7_attempts = 0
    user.p7_score = 0
    user.p8_attempts = 0
    user.p8_score = 0
    user.p9_attempts = 0
    user.p9_score = 0
    user.p10_attempts = 0
    user.p10_score = 0
    user.p11_attempts = 0
    user.p11_score = 0

    db.session.commit()
    flash(f"✅ User {user.rollnumber} has been reset successfully!", "success")
    return redirect(url_for("admin_dashboard"))


@app.route("/reset_user_quiz/<int:user_id>", methods=["POST"])
def reset_user_quiz(user_id):
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Unauthorized access", "danger")
        return redirect(url_for("login"))

    # Get current admin user
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    user = User.query.get_or_404(user_id)

    # Check if admin has permission to reset this user's quiz
    if current_admin.admin_level == 'admin':
        # Sub admin - check if user's roll number is in assigned range
        assigned_roll_numbers = get_assigned_roll_numbers(current_admin.id)
        if user.rollnumber not in assigned_roll_numbers:
            flash("You don't have permission to reset this user's quiz", "danger")
            return redirect(url_for("admin_dashboard"))

    try:
        # Delete all quiz attempts for this user
        quiz_attempts = QuizAttempt.query.filter_by(user_id=user_id, quiz_name=DEFAULT_QUIZ_NAME).all()
        for attempt in quiz_attempts:
            # Delete related quiz responses first
            QuizResponse.query.filter_by(attempt_id=attempt.id).delete()
            # Delete the attempt
            db.session.delete(attempt)

        db.session.commit()
        flash(
            f"✅ Quiz attempts for user {user.rollnumber} in {DEFAULT_QUIZ_NAME} have been reset successfully!", "success")

    except Exception as e:
        db.session.rollback()
        flash(
            f"❌ Error resetting quiz for user {user.rollnumber}: {str(e)}", "error")

    return redirect(url_for("dom_subject_admin_dashboard"))


@app.route("/reset_my_admin_progress", methods=["POST"])
def reset_my_admin_progress():
    """Reset current admin's progress and redirect back to assignment view"""
    if "loggedin" not in session or session.get("role") not in ["admin", "super_admin"]:
        flash("Unauthorized access", "danger")
        return redirect(url_for("login"))

    # Get current admin user
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    if not current_admin:
        flash("Admin user not found", "danger")
        return redirect(url_for("login"))

    # Reset all progress to initial values
    current_admin.marks = 0
    current_admin.completed = 0
    current_admin.p1_attempts = 0
    current_admin.p1_score = 0
    current_admin.p2_attempts = 0
    current_admin.p2_score = 0
    current_admin.p3_attempts = 0
    current_admin.p3_score = 0
    current_admin.p4_attempts = 0
    current_admin.p4_score = 0
    current_admin.p5_attempts = 0
    current_admin.p5_score = 0
    current_admin.p6_attempts = 0
    current_admin.p6_score = 0
    current_admin.p7_attempts = 0
    current_admin.p7_score = 0
    current_admin.p8_attempts = 0
    current_admin.p8_score = 0
    current_admin.p9_attempts = 0
    current_admin.p9_score = 0
    current_admin.p10_attempts = 0
    current_admin.p10_score = 0
    current_admin.p11_attempts = 0
    current_admin.p11_score = 0

    db.session.commit()
    flash(f"✅ Your progress has been reset successfully!", "success")
    return redirect(url_for("admin_assignment_view"))


@app.route("/reset_all_users", methods=["POST"])
def reset_all_users():
    if "loggedin" not in session or session.get("role") != "admin":
        flash("Unauthorized access", "danger")
        return redirect(url_for("login"))

    # Get current admin user
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    # Filter users based on admin level
    if current_admin.admin_level == 'super_admin':
        # Super admin can reset ALL users (students + admins, but not other super admins)
        students = User.query.filter_by(role="student").all()
        # Only regular admins, not super admins
        admins = User.query.filter_by(admin_level="admin").all()
        users = students + admins
    else:
        # Sub admin can reset assigned students + themselves
        students = get_students_for_admin(current_admin)
        users = students + [current_admin]  # Include the admin themselves

    reset_count = 0

    for user in users:
        user.marks = 0
        user.completed = 0
        user.p1_attempts = 0
        user.p1_score = 0
        user.p2_attempts = 0
        user.p2_score = 0
        user.p3_attempts = 0
        user.p3_score = 0
        user.p4_attempts = 0
        user.p4_score = 0
        user.p5_attempts = 0
        user.p5_score = 0
        user.p6_attempts = 0
        user.p6_score = 0
        user.p7_attempts = 0
        user.p7_score = 0
        user.p8_attempts = 0
        user.p8_score = 0
        user.p9_attempts = 0
        user.p9_score = 0
        user.p10_attempts = 0
        user.p10_score = 0
        user.p11_attempts = 0
        user.p11_score = 0

        # Delete all quiz attempts and responses for this user
        quiz_attempts = QuizAttempt.query.filter_by(user_id=user.id, quiz_name=DEFAULT_QUIZ_NAME).all()
        for attempt in quiz_attempts:
            QuizResponse.query.filter_by(attempt_id=attempt.id).delete()
            db.session.delete(attempt)
        reset_count += 1

    db.session.commit()

    # Create appropriate success message
    if current_admin.admin_level == 'super_admin':
        flash(
            f"✅ Reset {reset_count} accounts successfully (students + admins)!", "success")
    else:
        flash(
            f"✅ Reset {reset_count} accounts successfully (your assigned students + your account)!", "success")

    return redirect(url_for("admin_dashboard"))


@app.route("/reset_all_admins", methods=["POST"])
def reset_all_admins():
    """Reset only admin accounts (not students) - Super admin only"""
    if "loggedin" not in session or session.get("role") != "admin":
        flash("Unauthorized access", "danger")
        return redirect(url_for("login"))

    # Get current admin user
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    # Only super admins can reset other admins
    if current_admin.admin_level != 'super_admin':
        flash("Only super admins can reset admin accounts", "danger")
        return redirect(url_for("admin_dashboard"))

    # Get only admin accounts (not super admins, not students)
    admin_users = User.query.filter_by(admin_level="admin").all()

    reset_count = 0

    for user in admin_users:
        user.marks = 0
        user.completed = 0
        user.p1_attempts = 0
        user.p1_score = 0
        user.p2_attempts = 0
        user.p2_score = 0
        user.p3_attempts = 0
        user.p3_score = 0
        user.p4_attempts = 0
        user.p4_score = 0
        user.p5_attempts = 0
        user.p5_score = 0
        user.p6_attempts = 0
        user.p6_score = 0
        user.p7_attempts = 0
        user.p7_score = 0
        user.p8_attempts = 0
        user.p8_score = 0
        user.p9_attempts = 0
        user.p9_score = 0
        user.p10_attempts = 0
        user.p10_score = 0
        user.p11_attempts = 0
        user.p11_score = 0
        reset_count += 1

    db.session.commit()

    flash(f"✅ Reset {reset_count} admin accounts successfully!", "success")

    return redirect(url_for("admin_dashboard"))


@app.route("/toggle_problem_release/<int:problem_number>", methods=["POST"])
def toggle_problem_release(problem_number):
    if "loggedin" not in session or session.get("role") != "admin":
        flash("Unauthorized access", "danger")
        return redirect(url_for("login"))

    if problem_number < 1 or problem_number > 11:  # Updated to include problems 1-11
        flash("Invalid problem number", "danger")
        return redirect(url_for("admin_dashboard"))

    # Get current admin
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    # Determine admin_id for visibility record
    if current_admin.admin_level == 'super_admin':
        # Super admin can choose to release globally (admin_id = NULL) or for specific admin
        # For now, we'll make it global. Later we can add UI to choose specific admin
        admin_id_for_visibility = None
    else:
        # Sub admin releases only for their assigned students
        admin_id_for_visibility = current_admin.id

    visibility = ProblemVisibility.query.filter_by(
        problem_number=problem_number, admin_id=admin_id_for_visibility).first()

    if not visibility:
        visibility = ProblemVisibility(
            problem_number=problem_number,
            is_released=True,
            admin_id=admin_id_for_visibility)
        db.session.add(visibility)
    else:
        visibility.is_released = not visibility.is_released

    db.session.commit()

    status = "released" if visibility.is_released else "locked"
    scope = "globally" if admin_id_for_visibility is None else f"for your assigned students"
    flash(f"✅ Problem {problem_number} has been {status} {scope}!", "success")
    return redirect(url_for("admin_dashboard"))


@app.route("/release_all_problems_for_all_users", methods=["POST"])
def release_all_problems_for_all_users():
    """Release all problems for all users (or assigned students for sub-admins)"""
    if "loggedin" not in session or session.get("role") != "admin":
        flash("Unauthorized access", "danger")
        return redirect(url_for("login"))

    # Get current admin user
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    if current_admin.admin_level == 'super_admin':
        # Super admin: Release all problems globally
        problems_released = 0
        for problem_num in range(1, 12):
            visibility = ProblemVisibility.query.filter_by(
                problem_number=problem_num, admin_id=None).first()
            if not visibility:
                visibility = ProblemVisibility(
                    problem_number=problem_num, is_released=True, admin_id=None)
                db.session.add(visibility)
                problems_released += 1
            elif not visibility.is_released:
                visibility.is_released = True
                problems_released += 1

        db.session.commit()
        flash(
            f"✅ Released all {problems_released} problems for ALL users globally!", "success")

    else:
        # Sub admin: Release all problems for their assigned students only
        problems_released = 0
        for problem_num in range(1, 12):
            visibility = ProblemVisibility.query.filter_by(
                problem_number=problem_num, admin_id=current_admin.id).first()
            if not visibility:
                visibility = ProblemVisibility(
                    problem_number=problem_num, is_released=True, admin_id=current_admin.id)
                db.session.add(visibility)
                problems_released += 1
            elif not visibility.is_released:
                visibility.is_released = True
                problems_released += 1

        db.session.commit()

        # Get count of assigned students for display
        assigned_students = get_students_for_admin(current_admin)
        student_count = len(assigned_students)

        flash(
            f"✅ Released all {problems_released} problems for your {student_count} assigned students!", "success")

    return redirect(url_for("admin_dashboard"))


@app.route("/lock_all_problems_for_all_users", methods=["POST"])
def lock_all_problems_for_all_users():
    """Lock all problems for all users (or assigned students for sub-admins)"""
    if "loggedin" not in session or session.get("role") != "admin":
        flash("Unauthorized access", "danger")
        return redirect(url_for("login"))

    # Get current admin user
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()

    if current_admin.admin_level == 'super_admin':
        # Super admin: Lock all problems globally
        problems_locked = 0
        for problem_num in range(1, 12):
            visibility = ProblemVisibility.query.filter_by(
                problem_number=problem_num, admin_id=None).first()
            if not visibility:
                # If visibility record doesn't exist, create it as locked
                visibility = ProblemVisibility(
                    problem_number=problem_num, is_released=False, admin_id=None)
                db.session.add(visibility)
                problems_locked += 1
            elif visibility.is_released:
                visibility.is_released = False
                problems_locked += 1

        db.session.commit()
        flash(
            f"✅ Locked all {problems_locked} problems for ALL users globally!", "success")

    else:
        # Sub admin: Lock all problems for their assigned students only
        problems_locked = 0
        for problem_num in range(1, 12):
            visibility = ProblemVisibility.query.filter_by(
                problem_number=problem_num, admin_id=current_admin.id).first()
            if not visibility:
                # If visibility record doesn't exist, create it as locked
                visibility = ProblemVisibility(
                    problem_number=problem_num, is_released=False, admin_id=current_admin.id)
                db.session.add(visibility)
                problems_locked += 1
            elif visibility.is_released:
                visibility.is_released = False
                problems_locked += 1

        db.session.commit()

        # Get count of assigned students for display
        assigned_students = get_students_for_admin(current_admin)
        student_count = len(assigned_students)

        flash(
            f"✅ Locked all {problems_locked} problems for your {student_count} assigned students!", "success")

    return redirect(url_for("admin_dashboard"))


@app.route("/change_admin_password", methods=["GET", "POST"])
def change_admin_password():
    if "rollnumber" not in session:
        flash("Please login first.", "warning")
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Access denied.", "danger")
        return redirect(url_for("login"))

    if request.method == "POST":
        current_password = request.form.get("current_password")
        new_password = request.form.get("new_password")
        confirm_password = request.form.get("confirm_password")

        if user.password != current_password:
            flash("❌ Current password is incorrect.", "danger")
        elif new_password != confirm_password:
            flash("⚠️ New passwords do not match.", "warning")
        else:
            user.password = new_password
            db.session.commit()
            flash("✅ Password updated successfully!", "success")
            return redirect(url_for("admin_dashboard"))

    return render_template("change_admin_password.html")


# ======================== ADMIN MANAGEMENT ROUTES ========================

# Helper functions for roll management
def get_assigned_roll_numbers(admin_id):
    """Get all roll numbers assigned to a specific admin"""
    try:
        assignments = AdminRollAssignment.query.filter_by(
            admin_id=admin_id).all()
        assigned_rolls = set()

        for assignment in assignments:
            try:
                if assignment.assignment_type == 'range':
                    # Generate all rolls in range
                    start = assignment.roll_start
                    end = assignment.roll_end
                    assigned_rolls.update(generate_roll_range(start, end))
                elif assignment.assignment_type == 'individual':
                    # Add individual rolls (handle both single numbers and comma-separated lists)
                    if assignment.roll_numbers:
                        individual_rolls = [
                            roll.strip() for roll in assignment.roll_numbers.split(',') if roll.strip()]
                        assigned_rolls.update(individual_rolls)
            except Exception as assignment_error:
                print(
                    f"⚠️ Error processing assignment for admin {admin_id}: {assignment_error}")
                continue

        return list(assigned_rolls)
    except Exception as e:
        print(
            f"⚠️ Error getting assigned roll numbers for admin {admin_id}: {e}")
        return []


def generate_roll_range(start_roll, end_roll):
    """Generate all roll numbers between start and end (inclusive)"""
    try:
        # Try pure numeric conversion first
        start_num = int(start_roll)
        end_num = int(end_roll)
        return [str(i) for i in range(start_num, end_num + 1)]
    except (ValueError, TypeError):
        try:
            # Handle mixed alphanumeric roll numbers (e.g., CSE2021001)
            import re

            # Extract pattern and numbers
            start_match = re.match(r'([A-Za-z]*)(\d+)', str(start_roll))
            end_match = re.match(r'([A-Za-z]*)(\d+)', str(end_roll))

            if start_match and end_match:
                start_prefix = start_match.group(1)
                start_num = int(start_match.group(2))
                end_prefix = end_match.group(1)
                end_num = int(end_match.group(2))

                # Only generate range if prefixes match
                if start_prefix == end_prefix:
                    return [f"{start_prefix}{i}" for i in range(start_num, end_num + 1)]

            # If pattern matching fails, return individual values
            return [str(start_roll), str(end_roll)]
        except:
            # Last resort: return empty list
            return []


def count_rolls_in_assignment(assignment):
    """Count rolls for an assignment without expanding the entire range."""
    try:
        if assignment.assignment_type == 'individual':
            if not assignment.roll_numbers:
                return 0
            return len([
                roll.strip()
                for roll in assignment.roll_numbers.split(',')
                if roll.strip()
            ])

        if assignment.assignment_type != 'range':
            return 0

        start_roll = assignment.roll_start
        end_roll = assignment.roll_end

        try:
            start_num = int(start_roll)
            end_num = int(end_roll)
            return max(0, end_num - start_num + 1)
        except (ValueError, TypeError):
            start_match = re.match(r'([A-Za-z]*)(\d+)', str(start_roll))
            end_match = re.match(r'([A-Za-z]*)(\d+)', str(end_roll))

            if start_match and end_match:
                start_prefix = start_match.group(1)
                end_prefix = end_match.group(1)
                start_num = int(start_match.group(2))
                end_num = int(end_match.group(2))

                if start_prefix == end_prefix:
                    return max(0, end_num - start_num + 1)

            return 2 if start_roll and end_roll else 0
    except Exception as e:
        print(f"Error counting assignment {getattr(assignment, 'id', 'unknown')}: {e}")
        return 0


def get_students_for_admin(admin_user):
    """Get all students that an admin can manage"""
    if admin_user.admin_level == 'super_admin':
        # Super admin sees all students
        return User.query.filter_by(role='student').all()
    elif admin_user.admin_level == 'admin':
        # Regular admin sees only assigned students
        assigned_rolls = get_assigned_roll_numbers(admin_user.id)
        if assigned_rolls:
            return User.query.filter(User.rollnumber.in_(assigned_rolls), User.role == 'student').all()
        else:
            return []  # No assignments yet
    else:
        return []  # Not an admin


def find_admin_for_student(student_rollnumber):
    """Find which admin is responsible for a student"""
    print(
        f"🔍 find_admin_for_student: Looking for admin for student {student_rollnumber}")

    # First, check regular admins to see if student is specifically assigned
    regular_admins = User.query.filter_by(admin_level='admin').all()
    print(f"   📋 Found {len(regular_admins)} regular admins to check")

    for admin in regular_admins:
        assigned_rolls = get_assigned_roll_numbers(admin.id)
        print(
            f"   👤 Admin {admin.rollnumber}: has {len(assigned_rolls)} assigned rolls")
        print(
            f"      Assigned rolls: {assigned_rolls[:5]}{'...' if len(assigned_rolls) > 5 else ''}")

        if student_rollnumber in assigned_rolls:
            # Found specific admin assignment
            print(
                f"   ✅ Found! Student {student_rollnumber} is assigned to admin {admin.rollnumber}")
            return admin

    # If no specific assignment, fall back to super admin
    super_admin = User.query.filter_by(admin_level='super_admin').first()
    print(
        f"   ⚠️ No specific assignment found, falling back to super admin: {super_admin.rollnumber if super_admin else 'None'}")
    return super_admin


def count_students_for_admin(admin_id):
    """Count total roll numbers assigned to an admin (including ranges)"""
    admin = User.query.get(admin_id)
    if admin:
        if admin.admin_level == 'super_admin':
            # Super admin sees all students in database
            return User.query.filter_by(role='student').count()
        elif admin.admin_level == 'admin':
            assignments = AdminRollAssignment.query.filter_by(admin_id=admin_id).all()
            return sum(count_rolls_in_assignment(assignment) for assignment in assignments)
    return 0


# Register template function immediately after definition
app.jinja_env.globals['count_students_for_admin'] = count_students_for_admin


def debug_admin_assignments(admin_id):
    """Debug function to show assignment details"""
    assignments = AdminRollAssignment.query.filter_by(admin_id=admin_id).all()
    assigned_rolls = get_assigned_roll_numbers(admin_id)

    debug_info = {
        'admin_id': admin_id,
        'total_assignments': len(assignments),
        'total_rolls': len(assigned_rolls),
        'assignments_detail': [],
        # Show first 10 rolls
        'sample_rolls': assigned_rolls[:10] if assigned_rolls else []
    }

    for assignment in assignments:
        if assignment.assignment_type == 'range':
            range_rolls = generate_roll_range(
                assignment.roll_start, assignment.roll_end)
            debug_info['assignments_detail'].append({
                'type': 'range',
                'start': assignment.roll_start,
                'end': assignment.roll_end,
                'count': len(range_rolls),
                # Show first 5 from range
                'sample': range_rolls[:5] if range_rolls else []
            })
        elif assignment.assignment_type == 'individual':
            individual_rolls = [roll.strip() for roll in assignment.roll_numbers.split(
                ',') if roll.strip()] if assignment.roll_numbers else []
            debug_info['assignments_detail'].append({
                'type': 'individual',
                'rolls': assignment.roll_numbers,
                'count': len(individual_rolls),
                'parsed': individual_rolls
            })

    return debug_info


def get_problem_title(problem_number):
    """Get the title for a specific problem"""
    problem_titles = {
        1: "Watt governor",
        2: "Porter governor",
        3: "Gyroscopic effect on aeroplane",
        4: "Gyroscopic effect on ship",
        5: "Engine force analysis of slider crank mechanism",
        6: "Turning moment diagram",
        7: "Transverse vibration",
        8: "Longitudinal, Transverse and Torsional vibration",
        9: "Balancing of reciprocating masses",
        10: "Damped vibrations",
        11: "Balancing of rotating masses"
    }
    return problem_titles.get(problem_number, f"Problem {problem_number}")


# Register template function immediately after definition
app.jinja_env.globals['get_problem_title'] = get_problem_title

# Make early defined functions available in templates
app.jinja_env.globals['debug_admin_assignments'] = debug_admin_assignments


@app.route("/add_admin", methods=["GET", "POST"])
def add_admin():
    """Super admin can add new admins"""
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level != "super_admin":
        flash("Access denied. Super admin only.", "danger")
        return redirect(url_for("admin_dashboard"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        phone = request.form.get("phone", "").strip()

        # Validation
        if not username or not password:
            flash("Username and password are required.", "danger")
            return render_template("add_admin.html")

        # Check if username already exists
        existing_user = User.query.filter_by(rollnumber=username).first()
        if existing_user:
            flash("Username already exists.", "danger")
            return render_template("add_admin.html")

        # Create new admin
        new_admin = User(
            rollnumber=username,
            password=password,  # Store plain text as per your current system
            phonenumber=phone,
            role="admin",
            admin_level="admin",  # Sub admin by default
            created_by=user.id
        )

        db.session.add(new_admin)
        db.session.commit()

        flash(f"✅ Admin '{username}' created successfully!", "success")
        return redirect(url_for("admin_dashboard"))

    return render_template("add_admin.html")


@app.route("/edit_admin/<int:admin_id>", methods=["GET", "POST"])
def edit_admin(admin_id):
    """Super admin can edit admin details"""
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level != "super_admin":
        flash("Access denied. Super admin only.", "danger")
        return redirect(url_for("admin_dashboard"))

    admin_to_edit = User.query.get_or_404(admin_id)
    if admin_to_edit.admin_level not in ["admin", "super_admin"]:
        flash("This user is not an admin.", "danger")
        return redirect(url_for("admin_dashboard"))

    if request.method == "POST":
        new_username = request.form.get("username", "").strip()
        new_phone = request.form.get("phone", "").strip()

        if not new_username:
            flash("Username is required.", "danger")
            return render_template("edit_admin.html", admin=admin_to_edit)

        # Check if username is taken by another user
        existing_user = User.query.filter(
            User.rollnumber == new_username, User.id != admin_id).first()
        if existing_user:
            flash("Username already exists.", "danger")
            return render_template("edit_admin.html", admin=admin_to_edit)

        admin_to_edit.rollnumber = new_username
        admin_to_edit.phonenumber = new_phone
        db.session.commit()

        flash("✅ Admin details updated successfully!", "success")
        return redirect(url_for("admin_dashboard"))

    return render_template("edit_admin.html", admin=admin_to_edit)


@app.route("/reset_admin_password/<int:admin_id>", methods=["POST"])
def reset_admin_password(admin_id):
    """Super admin can reset admin password to 'admin123'"""
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level != "super_admin":
        flash("Access denied. Super admin only.", "danger")
        return redirect(url_for("admin_dashboard"))

    admin_to_reset = User.query.get_or_404(admin_id)
    if admin_to_reset.admin_level not in ["admin", "super_admin"]:
        flash("This user is not an admin.", "danger")
        return redirect(url_for("admin_dashboard"))

    admin_to_reset.password = "admin123"
    db.session.commit()

    flash(
        f"✅ Password for '{admin_to_reset.rollnumber}' reset to 'admin123'", "success")
    return redirect(url_for("admin_dashboard"))


@app.route("/delete_admin/<int:admin_id>", methods=["POST"])
def delete_admin(admin_id):
    """Super admin can delete admins"""
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level != "super_admin":
        flash("Access denied. Super admin only.", "danger")
        return redirect(url_for("admin_dashboard"))

    admin_to_delete = User.query.get_or_404(admin_id)
    if admin_to_delete.admin_level not in ["admin", "super_admin"]:
        flash("This user is not an admin.", "danger")
        return redirect(url_for("admin_dashboard"))

    if admin_to_delete.id == user.id:
        flash("You cannot delete yourself.", "danger")
        return redirect(url_for("admin_dashboard"))

    username = admin_to_delete.rollnumber
    db.session.delete(admin_to_delete)
    db.session.commit()

    flash(f"✅ Admin '{username}' deleted successfully!", "success")
    return redirect(url_for("admin_dashboard"))


@app.route("/assign_rolls/<int:admin_id>", methods=["GET", "POST"])
def assign_rolls(admin_id):
    """Super admin can assign roll numbers to admins"""
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level != "super_admin":
        flash("Access denied. Super admin only.", "danger")
        return redirect(url_for("admin_dashboard"))

    admin_user = User.query.get_or_404(admin_id)
    if admin_user.admin_level not in ["admin", "super_admin"]:
        flash("This user is not an admin.", "danger")
        return redirect(url_for("admin_dashboard"))

    # Prevent assignment for super admin - they see all students by default
    if admin_user.admin_level == "super_admin":
        flash("Super admins don't need student assignments - they see all students automatically.", "info")
        return redirect(url_for("admin_dashboard"))

    if request.method == "POST":
        assignment_type = request.form.get("assignment_type")

        # ---------------------- RANGE ASSIGNMENT ----------------------
        if assignment_type == "range":
            roll_start = request.form.get("roll_start", "").strip()
            roll_end = request.form.get("roll_end", "").strip()

            if not roll_start or not roll_end:
                flash("Both start and end roll numbers are required for range assignment.", "danger")
            else:
                import re
                match_start = re.match(r"^(.*?)(\d+)$", roll_start)
                match_end = re.match(r"^(.*?)(\d+)$", roll_end)

                if match_start and match_end and match_start.group(1) == match_end.group(1):
                    prefix = match_start.group(1)
                    start_num = int(match_start.group(2))
                    end_num = int(match_end.group(2))
                    width = len(match_start.group(2))

                    if start_num <= end_num:
                        created_count = 0
                        for num in range(start_num, end_num + 1):
                            roll_number = f"{prefix}{num:0{width}d}"
                            assignment = AdminRollAssignment(
                                admin_id=admin_id,
                                assignment_type="individual",
                                roll_numbers=roll_number
                            )
                            db.session.add(assignment)

                            # 🔹 NEW: update user's admin_id
                            student = User.query.filter_by(rollnumber=roll_number).first()
                            if student and user_model_has_admin_id():
                                student.admin_id = admin_id

                            created_count += 1
                        db.session.commit()
                        flash(
                            f"✅ Assigned {created_count} roll numbers ({roll_start}-{roll_end}) to {admin_user.rollnumber}",
                            "success"
                        )
                    else:
                        flash(f"Invalid range {roll_start}-{roll_end}: start must be <= end", "danger")
                else:
                    assignment = AdminRollAssignment(
                        admin_id=admin_id,
                        assignment_type="range",
                        roll_start=roll_start,
                        roll_end=roll_end
                    )
                    db.session.add(assignment)
                    db.session.commit()
                    flash(
                        f"✅ Range {roll_start}-{roll_end} assigned to {admin_user.rollnumber} (unexpanded pattern)",
                        "success"
                    )

        # ---------------------- INDIVIDUAL ASSIGNMENT ----------------------
        elif assignment_type == "individual":
            roll_numbers = request.form.get("roll_numbers", "").strip()

            if not roll_numbers:
                flash("Roll numbers list is required for individual assignment.", "danger")
            else:
                roll_list = [roll.strip() for roll in roll_numbers.split(",") if roll.strip()]
                if roll_list:
                    for roll in roll_list:
                        assignment = AdminRollAssignment(
                            admin_id=admin_id,
                            assignment_type="individual",
                            roll_numbers=roll
                        )
                        db.session.add(assignment)

                        # 🔹 NEW: update user's admin_id
                        student = User.query.filter_by(rollnumber=roll).first()
                        if student and user_model_has_admin_id():
                            student.admin_id = admin_id

                    db.session.commit()
                    flash(f"✅ Individual rolls assigned to {admin_user.rollnumber}", "success")
                else:
                    flash("No valid roll numbers found.", "danger")

        # ---------------------- BULK ASSIGNMENT ----------------------
        elif assignment_type == "bulk":
            bulk_assignment = request.form.get("bulk_assignment", "").strip()

            if not bulk_assignment:
                flash("Bulk assignment text is required.", "danger")
            else:
                try:
                    import re
                    items = [item.strip() for item in bulk_assignment.split(",") if item.strip()]
                    ranges_created = 0
                    individuals_created = 0

                    for item in items:
                        item = item.replace(" to ", "-")

                        if "-" in item and not item.startswith("-"):
                            parts = item.split("-", 1)
                            if len(parts) == 2:
                                roll_start = parts[0].strip()
                                roll_end = parts[1].strip()

                                match_start = re.match(r"^(.*?)(\d+)$", roll_start)
                                match_end = re.match(r"^(.*?)(\d+)$", roll_end)

                                if match_start and match_end and match_start.group(1) == match_end.group(1):
                                    prefix = match_start.group(1)
                                    start_num = int(match_start.group(2))
                                    end_num = int(match_end.group(2))
                                    width = len(match_start.group(2))

                                    if start_num <= end_num:
                                        for num in range(start_num, end_num + 1):
                                            roll_number = f"{prefix}{num:0{width}d}"
                                            assignment = AdminRollAssignment(
                                                admin_id=admin_id,
                                                assignment_type="individual",
                                                roll_numbers=roll_number
                                            )
                                            db.session.add(assignment)
                                            individuals_created += 1

                                            # 🔹 NEW: update user's admin_id
                                            student = User.query.filter_by(rollnumber=roll_number).first()
                                            if student and user_model_has_admin_id():
                                                student.admin_id = admin_id
                                    else:
                                        flash(f"Invalid range {roll_start}-{roll_end}: start must be <= end", "danger")
                                else:
                                    assignment = AdminRollAssignment(
                                        admin_id=admin_id,
                                        assignment_type="range",
                                        roll_start=roll_start,
                                        roll_end=roll_end
                                    )
                                    db.session.add(assignment)
                                    ranges_created += 1
                        else:
                            if item:
                                assignment = AdminRollAssignment(
                                    admin_id=admin_id,
                                    assignment_type="individual",
                                    roll_numbers=item
                                )
                                db.session.add(assignment)
                                individuals_created += 1

                                # 🔹 NEW: update user's admin_id
                                student = User.query.filter_by(rollnumber=item).first()
                                if student and user_model_has_admin_id():
                                    student.admin_id = admin_id

                    if ranges_created > 0 or individuals_created > 0:
                        db.session.commit()
                        flash(
                            f"✅ Bulk assignment completed: {ranges_created} ranges, {individuals_created} individual numbers assigned to {admin_user.rollnumber}",
                            "success"
                        )
                    else:
                        flash("No valid assignments found in bulk text.", "danger")

                except Exception as e:
                    flash(f"Error processing bulk assignment: {str(e)}", "danger")

        return redirect(url_for("assign_rolls", admin_id=admin_id))

    # ---------------------- DISPLAY ASSIGNMENTS ----------------------
    assignments = AdminRollAssignment.query.filter_by(admin_id=admin_id).all()

    return render_template(
        "assign_rolls.html",
        admin=admin_user,
        assignments=assignments,
        count_students_for_admin=count_students_for_admin
    )



@app.route("/delete_assignment/<int:assignment_id>", methods=["POST"])
def delete_assignment(assignment_id):
    """Delete a roll assignment"""
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level != "super_admin":
        flash("Access denied. Super admin only.", "danger")
        return redirect(url_for("admin_dashboard"))

    assignment = AdminRollAssignment.query.get_or_404(assignment_id)
    admin_id = assignment.admin_id

    db.session.delete(assignment)
    db.session.commit()

    flash("✅ Assignment deleted successfully!", "success")
    return redirect(url_for("assign_rolls", admin_id=admin_id))


@app.route("/edit_assignment/<int:assignment_id>", methods=["GET", "POST"])
def edit_assignment(assignment_id):
    """Edit a roll assignment"""
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level != "super_admin":
        flash("Access denied. Super admin only.", "danger")
        return redirect(url_for("admin_dashboard"))

    assignment = AdminRollAssignment.query.get_or_404(assignment_id)
    admin_user = User.query.get_or_404(assignment.admin_id)

    if request.method == "POST":
        assignment_type = request.form.get("assignment_type")

        if assignment_type == "range":
            roll_start = request.form.get("roll_start", "").strip()
            roll_end = request.form.get("roll_end", "").strip()

            if not roll_start or not roll_end:
                flash(
                    "Both start and end roll numbers are required for range assignment.", "danger")
            else:
                assignment.assignment_type = "range"
                assignment.roll_start = roll_start
                assignment.roll_end = roll_end
                assignment.roll_numbers = None  # Clear individual numbers
                db.session.commit()
                flash(
                    f"✅ Range assignment updated: {roll_start}-{roll_end}", "success")
                return redirect(url_for("assign_rolls", admin_id=assignment.admin_id))

        elif assignment_type == "individual":
            roll_numbers = request.form.get("roll_numbers", "").strip()

            if not roll_numbers:
                flash("Roll numbers are required for individual assignment.", "danger")
            else:
                assignment.assignment_type = "individual"
                assignment.roll_numbers = roll_numbers
                assignment.roll_start = None  # Clear range values
                assignment.roll_end = None
                db.session.commit()
                flash("✅ Individual assignment updated successfully!", "success")
                return redirect(url_for("assign_rolls", admin_id=assignment.admin_id))

    return render_template("edit_assignment.html", assignment=assignment, admin=admin_user)


@app.route("/impersonate_admin/<int:admin_id>", methods=["POST"])
def impersonate_admin(admin_id):
    """Super admin can impersonate other admins"""
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level != "super_admin":
        flash("Access denied. Super admin only.", "danger")
        return redirect(url_for("admin_dashboard"))

    target_admin = User.query.get_or_404(admin_id)
    if target_admin.admin_level not in ["admin", "super_admin"]:
        flash("This user is not an admin.", "danger")
        return redirect(url_for("admin_dashboard"))

    # Store original admin info for later restoration
    session["original_admin"] = {
        "rollnumber": session["rollnumber"],
        "role": session["role"]
    }

    # Switch to target admin
    session["rollnumber"] = target_admin.rollnumber
    session["role"] = target_admin.role
    session["impersonating"] = True
    session["impersonated_admin_name"] = target_admin.rollnumber

    flash(f"🎭 Now impersonating admin: {target_admin.rollnumber}", "warning")
    return redirect(url_for("admin_dashboard"))


@app.route("/stop_impersonation", methods=["POST"])
def stop_impersonation():
    """Stop impersonating and return to original admin"""
    if "rollnumber" not in session or "original_admin" not in session:
        return redirect(url_for("login"))

    # Restore original admin
    original_admin = session["original_admin"]
    session["rollnumber"] = original_admin["rollnumber"]
    session["role"] = original_admin["role"]

    # Clean up impersonation data
    session.pop("original_admin", None)
    session.pop("impersonating", None)
    session.pop("impersonated_admin_name", None)

    flash("✅ Stopped impersonation, returned to your account", "success")
    return redirect(url_for("admin_dashboard"))


# --- CO Management Routes ---
@app.route("/admin/co_management", methods=['GET', 'POST'])
def admin_co_management():
    if 'rollnumber' not in session:
        flash('Please log in to access this page.', 'error')
        return redirect(url_for('login'))

    user = User.query.filter_by(rollnumber=session['rollnumber']).first()
    if not user or user.admin_level != 'super_admin':
        flash('Access denied. Super admin only.', 'error')
        return redirect(url_for('admin_dashboard'))

    if request.method == 'POST':
        try:
            # Update CO mappings
            for i in range(1, 12):  # Questions 1-11
                co_assignment = request.form.get(f'q{i}_co')
                max_marks = float(request.form.get(f'q{i}_marks', 10.0))

                mapping = COMapping.query.filter_by(question_number=i).first()
                if mapping:
                    mapping.co_number = co_assignment
                    mapping.max_marks = max_marks
                else:
                    mapping = COMapping(
                        question_number=i,
                        co_number=co_assignment,
                        max_marks=max_marks
                    )
                    db.session.add(mapping)

            db.session.commit()

            # Update CO marks for all students
            students = User.query.filter_by(role='student').all()
            for student in students:
                update_co_marks(student.id)

            flash('CO mappings updated successfully!', 'success')

        except Exception as e:
            db.session.rollback()
            flash(f'Error updating CO mappings: {str(e)}', 'error')

    # Get current mappings
    mappings = COMapping.query.order_by(COMapping.question_number).all()
    mapping_dict = {m.question_number: m for m in mappings}

    return render_template('admin_co_management.html', mappings=mapping_dict)


@app.route("/admin/co_analysis")
def admin_co_analysis():
    # Clear any irrelevant flash messages from previous actions
    get_flashed_messages()  # This consumes and clears existing flash messages

    if 'rollnumber' not in session:
        flash('Please log in to access this page.', 'error')
        return redirect(url_for('login'))

    user = User.query.filter_by(rollnumber=session['rollnumber']).first()
    if not user or user.role not in ['admin', 'super_admin']:
        flash('Access denied. Admin access required.', 'error')
        return redirect(url_for('student_dashboard'))

    # Get students based on admin access
    try:
        if user.admin_level == 'super_admin':  # Super admin sees all
            students = User.query.filter_by(role='student').all()
        else:  # Other admins see their assigned students
            # Get assigned rolls for this admin
            assignments = AdminRollAssignment.query.filter_by(
                admin_id=user.id).all()

            if not assignments:
                flash(
                    'No student assignments found for your admin account. Please contact super admin.', 'warning')
                return render_template('admin_co_analysis.html',
                                       student_co_data=[],
                                       co_summary={co_num: {'max_marks': 0, 'questions': []} for co_num in [
                                           'CO1', 'CO2', 'CO3', 'CO4', 'CO5']},
                                       is_super_admin=False)

            assigned_rolls = []

            for assignment in assignments:
                try:
                    if assignment.assignment_type == 'individual':
                        if assignment.roll_numbers:
                            rolls = [roll.strip() for roll in assignment.roll_numbers.split(
                                ',') if roll.strip()]
                            assigned_rolls.extend(rolls)
                            app.logger.debug(
                                f"Admin {user.id} individual assignment rolls: {rolls}")
                    elif assignment.assignment_type == 'range':
                        # Parse range (assuming numeric rolls like 21001A0501-21001A0550)
                        start_roll = assignment.roll_start
                        end_roll = assignment.roll_end
                        if start_roll and end_roll:
                            # For now, get all students and filter - could be optimized
                            all_students = User.query.filter_by(
                                role='student').all()
                            for student in all_students:
                                if start_roll <= student.rollnumber <= end_roll:
                                    assigned_rolls.append(student.rollnumber)
                            app.logger.debug(
                                f"Admin {user.id} range assignment {start_roll}-{end_roll}, found rolls: {[s.rollnumber for s in all_students if start_roll <= s.rollnumber <= end_roll]}")
                except Exception as e:
                    app.logger.error(
                        f"Error processing assignment {assignment.id}: {str(e)}")
                    continue

            app.logger.debug(
                f"Admin {user.id} total assigned_rolls: {assigned_rolls}")

            if not assigned_rolls:
                flash(
                    'No valid student roll numbers found in your assignments. Please contact super admin.', 'warning')
                return render_template('admin_co_analysis.html',
                                       student_co_data=[],
                                       co_summary={co_num: {'max_marks': 0, 'questions': []} for co_num in [
                                           'CO1', 'CO2', 'CO3', 'CO4', 'CO5']},
                                       is_super_admin=False)

            # Get students using both string and converted rollnumbers for comparison
            students = User.query.filter(User.rollnumber.in_(
                assigned_rolls), User.role == 'student').all()

            # Debug: Check what students were found
            app.logger.debug(
                f"Found students for admin {user.id}: {[s.rollnumber for s in students]}")

            # Also try finding students with string conversion in case of data type mismatch
            if not students:
                app.logger.debug(
                    "No students found with direct match, trying string conversion")
                all_students = User.query.filter_by(role='student').all()
                app.logger.debug(
                    f"All students in database: {[(s.id, s.rollnumber, type(s.rollnumber)) for s in all_students]}")

                # Try to match by converting both to strings
                string_assigned_rolls = [str(roll) for roll in assigned_rolls]
                students = [s for s in all_students if str(
                    s.rollnumber) in string_assigned_rolls]
                app.logger.debug(
                    f"Students found with string conversion: {[s.rollnumber for s in students]}")

    except Exception as e:
        app.logger.error(f"Error getting student assignments: {str(e)}")
        flash('Error loading student assignments. Please try again or contact super admin.', 'error')
        return render_template('admin_co_analysis.html',
                               student_co_data=[],
                               co_summary={co_num: {'max_marks': 0, 'questions': []}
                                           for co_num in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']},
                               is_super_admin=False)

    # Get CO mappings and calculate summary
    try:
        co_mappings = COMapping.query.all()
        co_summary = {}

        for co_num in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
            total_max = sum(
                m.max_marks for m in co_mappings if m.co_number == co_num)
            co_summary[co_num] = {'max_marks': total_max, 'questions': []}

            for mapping in co_mappings:
                if mapping.co_number == co_num:
                    co_summary[co_num]['questions'].append({
                        'number': mapping.question_number,
                        'max_marks': mapping.max_marks
                    })
    except Exception as e:
        print(f"Error accessing CO mapping table: {str(e)}")
        flash('CO tables not found. Please ask super admin to create the CO tables first.', 'error')
        return render_template('admin_co_analysis.html',
                               student_co_data=[],
                               co_summary={co_num: {'max_marks': 0, 'questions': []}
                                           for co_num in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']},
                               is_super_admin=(user.admin_level == 'super_admin'))

    # Sort students by rollnumber ascending
    students_sorted = sorted(students, key=lambda s: s.rollnumber)

    # Get CO marks for students
    student_co_data = []
    for student in students_sorted:
        try:
            # Update CO marks for this student first
            update_co_marks(student.id)

            co_marks = COMarks.query.filter_by(user_id=student.id).all()
            co_marks_dict = {cm.co_number: cm for cm in co_marks}

            student_data = {
                'student': student,
                'co_marks': co_marks_dict,
                'total_obtained': sum(cm.marks_obtained for cm in co_marks),
                'total_max': sum(cm.max_marks for cm in co_marks),
                'overall_percentage': 0
            }

            if student_data['total_max'] > 0:
                student_data['overall_percentage'] = (
                    student_data['total_obtained'] / student_data['total_max']) * 100

            student_co_data.append(student_data)
        except Exception as e:
            print(
                f"Error processing CO marks for student {student.rollnumber}: {str(e)}")
            # Continue with other students even if one fails
            continue

    return render_template('admin_co_analysis.html',
                           student_co_data=student_co_data,
                           co_summary=co_summary,
                           is_super_admin=(user.admin_level == 'super_admin'))


@app.route("/admin/co_csv_download")
def admin_co_csv_download():
    """Download CO analysis data as CSV"""
    if 'rollnumber' not in session:
        flash('Please log in to access this page.', 'error')
        return redirect(url_for('login'))

    user = User.query.filter_by(rollnumber=session['rollnumber']).first()
    if not user or user.role not in ['admin', 'super_admin']:
        flash('Access denied. Admin access required.', 'error')
        return redirect(url_for('student_dashboard'))

    # Get students based on admin access (same logic as analysis)
    if user.admin_level == 'super_admin':  # Super admin sees all
        students = User.query.filter_by(role='student').all()
    else:  # Other admins see their assigned students
        # Get assigned rolls for this admin
        assignments = AdminRollAssignment.query.filter_by(
            admin_id=user.id).all()
        assigned_rolls = []

        for assignment in assignments:
            if assignment.assignment_type == 'individual':
                assigned_rolls.extend(
                    [roll.strip() for roll in assignment.roll_numbers.split(',')])
            elif assignment.assignment_type == 'range':
                # Parse range (assuming numeric rolls like 21001A0501-21001A0550)
                start_roll = assignment.roll_start
                end_roll = assignment.roll_end
                # For now, get all students and filter - could be optimized
                all_students = User.query.filter_by(role='student').all()
                for student in all_students:
                    if start_roll <= student.rollnumber <= end_roll:
                        assigned_rolls.append(student.rollnumber)

        students = User.query.filter(User.rollnumber.in_(
            assigned_rolls), User.role == 'student').all()

    # Get CO mappings and calculate max marks per CO
    co_mappings = COMapping.query.all()
    co_summary = {}

    for co_num in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
        total_max = sum(
            m.max_marks for m in co_mappings if m.co_number == co_num)
        co_summary[co_num] = total_max

    # Prepare CSV data in the exact format requested
    import csv
    from io import StringIO

    output = StringIO()
    writer = csv.writer(output)

    # Write header row with CO columns
    writer.writerow(['', 'CO1', 'CO2', 'CO3', 'CO4', 'CO5'])

    # Write CO MAX MARKS row
    max_marks_row = ['CO MAX MARKS']
    for co_num in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
        max_marks_row.append(int(co_summary[co_num]))
    writer.writerow(max_marks_row)

    # Write empty row header for student data
    writer.writerow(['Student Roll Number', '', '', '', '', ''])

    # Sort students by rollnumber ascending
    students_sorted = sorted(students, key=lambda s: s.rollnumber)
    # Write student data
    for student in students_sorted:
        # Update CO marks for this student first
        update_co_marks(student.id)

        co_marks = COMarks.query.filter_by(user_id=student.id).all()
        co_marks_dict = {cm.co_number: cm for cm in co_marks}

        row = [student.rollnumber]

        for co_num in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
            co_mark = co_marks_dict.get(co_num)
            if co_mark:
                row.append(int(co_mark.marks_obtained))
            else:
                row.append(0)

        writer.writerow(row)

    # Create response
    from flask import Response
    from datetime import datetime

    output.seek(0)

    # Generate filename with timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    admin_type = 'SuperAdmin' if user.admin_level == 'super_admin' else f'Admin_{user.rollnumber}'
    filename = f'CO_Summary_{admin_type}_{timestamp}.csv'

    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={
            'Content-Disposition': f'attachment; filename={filename}'
        }
    )


@app.route("/admin/co_detailed_csv_download")
def admin_co_detailed_csv_download():
    """Download detailed CO analysis data with individual problem scores as CSV"""
    if 'rollnumber' not in session:
        flash('Please log in to access this page.', 'error')
        return redirect(url_for('login'))

    user = User.query.filter_by(rollnumber=session['rollnumber']).first()
    if not user or user.role != 'admin':
        flash('Access denied. Admin access required.', 'error')
        return redirect(url_for('student_dashboard'))

    # Get students based on admin access (same logic as analysis)
    if user.admin_level == 'super_admin':  # Super admin sees all
        students = User.query.filter_by(role='student').all()
    else:  # Other admins see their assigned students
        # Get assigned rolls for this admin
        assignments = AdminRollAssignment.query.filter_by(
            admin_id=user.id).all()
        assigned_rolls = []

        for assignment in assignments:
            if assignment.assignment_type == 'individual':
                assigned_rolls.extend(
                    [roll.strip() for roll in assignment.roll_numbers.split(',')])
            elif assignment.assignment_type == 'range':
                start_roll = assignment.roll_start
                end_roll = assignment.roll_end
                all_students = User.query.filter_by(role='student').all()
                for student in all_students:
                    if start_roll <= student.rollnumber <= end_roll:
                        assigned_rolls.append(student.rollnumber)

        students = User.query.filter(User.rollnumber.in_(
            assigned_rolls), User.role == 'student').all()

    # Get CO mappings
    co_mappings = COMapping.query.order_by(COMapping.question_number).all()

    # Prepare CSV data
    import csv
    from io import StringIO

    output = StringIO()
    writer = csv.writer(output)

    # Write header - first the student info, then individual problems, then CO summaries
    header = ['Student Roll Number', 'Email', 'Phone', 'Total Marks']

    # Individual problem scores
    for i in range(1, 12):
        header.extend([f'P{i} Score', f'P{i} Attempts'])

    # CO mapping info
    header.append('CO Mappings')
    for mapping in co_mappings:
        header.append(f'Q{mapping.question_number}→{mapping.co_number}')

    # CO summaries
    for co_num in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
        header.extend([f'{co_num} Obtained', f'{co_num} Max', f'{co_num} %'])

    header.append('Overall %')
    writer.writerow(header)

    # Write student data
    for student in students:
        # Update CO marks for this student first
        update_co_marks(student.id)

        co_marks = COMarks.query.filter_by(user_id=student.id).all()
        co_marks_dict = {cm.co_number: cm for cm in co_marks}

        row = [
            student.rollnumber,
            student.email or 'N/A',
            student.phonenumber or 'N/A',
            student.marks or 0
        ]

        # Individual problem scores and attempts
        problem_scores = [
            student.p1_score or 0, student.p2_score or 0, student.p3_score or 0,
            student.p4_score or 0, student.p5_score or 0, student.p6_score or 0,
            student.p7_score or 0, student.p8_score or 0, student.p9_score or 0,
            student.p10_score or 0, student.p11_score or 0
        ]
        problem_attempts = [
            student.p1_attempts or 0, student.p2_attempts or 0, student.p3_attempts or 0,
            student.p4_attempts or 0, student.p5_attempts or 0, student.p6_attempts or 0,
            student.p7_attempts or 0, student.p8_attempts or 0, student.p9_attempts or 0,
            student.p10_attempts or 0, student.p11_attempts or 0
        ]

        for i in range(11):
            row.extend([problem_scores[i], problem_attempts[i]])

        # CO mapping indicator
        row.append('Questions→COs:')
        for mapping in co_mappings:
            row.append(f'Q{mapping.question_number}→{mapping.co_number}')

        # CO summaries
        total_obtained = 0
        total_max = 0

        for co_num in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
            co_mark = co_marks_dict.get(co_num)
            if co_mark:
                row.extend([
                    round(co_mark.marks_obtained, 1),
                    round(co_mark.max_marks, 1),
                    round(co_mark.percentage, 1)
                ])
                total_obtained += co_mark.marks_obtained
                total_max += co_mark.max_marks
            else:
                row.extend([0.0, 0.0, 0.0])

        # Overall percentage
        overall_percentage = (total_obtained / total_max *
                              100) if total_max > 0 else 0
        row.append(round(overall_percentage, 1))

        writer.writerow(row)

    # Add CO mapping summary at the end
    writer.writerow([])  # Empty row
    writer.writerow(['CO MAPPING SUMMARY'])
    writer.writerow(['Question', 'Course Outcome', 'Max Marks'])
    for mapping in co_mappings:
        writer.writerow(
            [f'Problem {mapping.question_number}', mapping.co_number, mapping.max_marks])

    # Create response
    from flask import Response
    from datetime import datetime

    output.seek(0)

    # Generate filename with timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    admin_type = 'SuperAdmin' if user.admin_level == 'super_admin' else f'Admin_{user.rollnumber}'
    filename = f'CO_Detailed_Analysis_{admin_type}_{timestamp}.csv'

    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={
            'Content-Disposition': f'attachment; filename={filename}'
        }
    )


@app.route("/change_my_password", methods=["GET", "POST"])
def change_my_password():
    """Sub admins can change their own password"""
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level not in ["admin", "super_admin"]:
        flash("Access denied.", "danger")
        return redirect(url_for("login"))

    if request.method == "POST":
        current_password = request.form.get("current_password", "").strip()
        new_password = request.form.get("new_password", "").strip()
        confirm_password = request.form.get("confirm_password", "").strip()

        # Validation
        if not current_password or not new_password or not confirm_password:
            flash("All fields are required.", "danger")
            return render_template("change_my_password.html")

        if current_password != user.password:
            flash("Current password is incorrect.", "danger")
            return render_template("change_my_password.html")

        if new_password != confirm_password:
            flash("New passwords do not match.", "danger")
            return render_template("change_my_password.html")

        if len(new_password) < 4:
            flash("Password must be at least 4 characters long.", "danger")
            return render_template("change_my_password.html")

        # Update password
        user.password = new_password
        db.session.commit()

        flash("✅ Password changed successfully!", "success")
        return redirect(url_for("admin_dashboard"))

    return render_template("change_my_password.html")


# ======================== END ADMIN MANAGEMENT ROUTES ========================


def get_max_marks_for_problem(problem_no):
    """Get maximum marks for a specific problem"""
    if problem_no in [1, 2, 3]:
        return 5
    elif problem_no == 8:
        return 15
    elif problem_no == 10:
        return 10  # Problem 10: 4 answers × 2.5 marks = 10 total, +2 for wrong attempts
    else:
        return 10  # Problems 4, 5, 6, 7, 9


# Register template function immediately after definition
app.jinja_env.globals['get_max_marks_for_problem'] = get_max_marks_for_problem


@app.route("/admin_quiz_co_csv_download")
def admin_quiz_co_csv_download():
    """Download Quiz CO analysis data as CSV"""
    if 'rollnumber' not in session:
        flash('Please log in to access this page.', 'error')
        return redirect(url_for('login'))

    user = User.query.filter_by(rollnumber=session['rollnumber']).first()
    if not user or user.role not in ['admin', 'super_admin']:
        flash('Access denied. Admin access required.', 'error')
        return redirect(url_for('student_dashboard'))

    try:
        # Get students based on admin access
        if user.admin_level == 'super_admin':
            students = User.query.filter_by(role='student').all()
        else:
            students = get_students_for_admin(user)

        # ✅ Sort students by rollnumber (ascending)
        students = sorted(students, key=lambda s: s.rollnumber)

        # Prepare CSV data
        import csv
        from io import StringIO

        output = StringIO()
        writer = csv.writer(output)

        writer.writerow([
            'Student Roll Number',
            'Completed Attempts',
            'Total Marks Obtained',
            'Total Possible Marks',
            'Overall Percentage',
            'CO1 Obtained', 'CO1 Max', 'CO1 Percentage',
            'CO2 Obtained', 'CO2 Max', 'CO2 Percentage',
            'CO3 Obtained', 'CO3 Max', 'CO3 Percentage',
            'CO4 Obtained', 'CO4 Max', 'CO4 Percentage',
            'CO5 Obtained', 'CO5 Max', 'CO5 Percentage',
        ])

        # Write student data
        for student in students:
            attempts = QuizAttempt.query.filter_by(user_id=student.id).filter(
                QuizAttempt.completed_at.isnot(None)).all()
            total_obtained = sum((attempt.score or 0) for attempt in attempts)
            total_possible = sum((attempt.total_points or 0)
                                 for attempt in attempts)
            overall_percentage = (
                total_obtained / total_possible * 100
                if total_possible > 0 else 0
            )
            co_performance = calculate_student_quiz_co_performance(student.id)

            row = [
                student.rollnumber,
                len(attempts),
                round(total_obtained, 2),
                round(total_possible, 2),
                round(overall_percentage, 2),
            ]

            for co_num in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
                co_data = co_performance.get(co_num, {})
                row.extend([
                    round(co_data.get('marks_obtained', 0), 2),
                    round(co_data.get('total_marks', 0), 2),
                    round(co_data.get('percentage', 0), 2),
                ])

            writer.writerow(row)

        # Create response
        from flask import Response
        from datetime import datetime

        output.seek(0)

        # Generate filename with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        admin_type = 'SuperAdmin' if user.admin_level == 'super_admin' else f'Admin_{user.rollnumber}'
        filename = f'Quiz_CO_Summary_{admin_type}_{timestamp}.csv'

        return Response(
            output.getvalue(),
            mimetype='text/csv',
            headers={
                'Content-Disposition': f'attachment; filename={filename}',
                'Content-Type': 'text/csv; charset=utf-8'
            }
        )

    except Exception as e:
        flash(f"Error generating quiz CO CSV: {str(e)}", "error")
        return redirect(url_for("admin_quiz_co_analysis"))


def calculate_total_max_marks():
    """Calculate total maximum marks across all problems"""
    return sum(get_max_marks_for_problem(i) for i in range(1, 12))  # Include Problem 11


@app.route("/student_dashboard")
def student_dashboard():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user:
        flash("User not found", "danger")
        return redirect(url_for("login"))

    total_max_marks = calculate_total_max_marks()

    # Get problem visibility settings for this student
    problem_visibility = {}
    for i in range(1, 12):  # Problems 1-11
        problem_visibility[i] = get_problem_visibility_for_user(
            i, user.rollnumber)

    # Get quiz information for this student (with error handling for missing tables)
    quiz_released = False
    quiz_attempts = 0
    quiz_best_score = 0
    quiz_max_score = 10  # Default quiz max score

    try:
        # For the student dashboard display, show the actual quiz release status
        # using the proper per-admin visibility check
        quiz_released = check_quiz_visibility_for_student(user)
        print(
            f"📊 Student dashboard for {user.rollnumber} ({user.role}): quiz_released = {quiz_released}"
        )

        # Get quiz attempts and best score
        quiz_attempts = QuizAttempt.query.filter_by(user_id=user.id).count()

        if quiz_attempts > 0:
            completed_attempts = (
                QuizAttempt.query.filter_by(user_id=user.id)
                .filter(QuizAttempt.completed_at.isnot(None))
                .all()
            )

            if completed_attempts:
                quiz_best_score = max(
                    attempt.score for attempt in completed_attempts)
                quiz_max_score = completed_attempts[0].total_points
    except Exception as e:
        # Quiz tables don't exist yet - set default values
        quiz_released = False
        quiz_attempts = 0
        quiz_best_score = 0
        quiz_max_score = 10

    setting = GlobalSettings.query.first()
    allow_student_mid_exam_questions = (
        setting.allow_student_mid_exam_questions if setting else False
    )
    allow_student_python_lab = (
        setting.allow_student_python_lab if setting else True
    )

    # ✅ --- KDM LAB VISIBILITY CHECK ---
    kdm_visible = False
    try:
        settings = KDMLabSettings.query.first()
        if settings and settings.student_visibility:
            kdm_visible = True
    except Exception as e:
        print("⚠️ Error checking KDM LAB visibility:", e)
        kdm_visible = False
    # ✅ --------------------------------

    # Calculate current completed problems count dynamically
    completed_problems = 0
    for i in range(1, 12):
        if getattr(user, f"p{i}_attempts", 0) > 0:
            completed_problems += 1

        # --- Apply dashboard visibility from admin_dashboard_visibility ---
    try:
        # 1️⃣ Find which admin manages this student
        from sqlalchemy import text
        admin_user = None

        # You likely already have a helper for this; if not, we can use:
        if hasattr(user, "admin_id") and user.admin_id:
            admin_user = User.query.get(user.admin_id)
        else:
            # fallback: use your helper if exists (used in other places)
            try:
                admin_user = find_admin_for_student(user.rollnumber)
            except Exception as e:
                print(f"⚠️ Error finding admin for student {user.rollnumber}: {e}")

        # 2️⃣ Default all dashboards visible if no admin found
        show_dom_dashboard = True
        show_kdm_dashboard = True
        show_python_dashboard = True

        # 3️⃣ Read from visibility table if admin found
        if admin_user:
            visibilities = AdminDashboardVisibility.query.filter_by(admin_id=admin_user.id).all()
            visibility_map = {v.dashboard_type: v.is_visible for v in visibilities}

            show_dom_dashboard = visibility_map.get("dom", True)
            show_kdm_dashboard = visibility_map.get("kdm", True)
            show_python_dashboard = visibility_map.get("python", True)

        print(f"🎓 Student {user.rollnumber}: DOM={show_dom_dashboard}, KDM={show_kdm_dashboard}, PYTHON={show_python_dashboard}")

    except Exception as e:
        print(f"⚠️ Error applying dashboard visibility for student {user.rollnumber}: {e}")
        show_dom_dashboard = show_kdm_dashboard = show_python_dashboard = True


    return render_template(
    "student_dashboard.html",
    user=user,
    completed_problems=completed_problems,
    total_max_marks=total_max_marks,
    problem_visibility=problem_visibility,
    quiz_released=quiz_released,
    quiz_attempts=quiz_attempts,
    quiz_best_score=quiz_best_score,
    quiz_max_score=quiz_max_score,
    allow_student_mid_exam_questions=allow_student_mid_exam_questions,
    allow_student_python_lab=allow_student_python_lab,
    kdm_visible=kdm_visible,
    show_dom_dashboard=show_dom_dashboard,
    show_kdm_dashboard=show_kdm_dashboard,
    show_python_dashboard=show_python_dashboard
)


@app.route("/dom_subject_student_dashboard")
def dom_subject_student_dashboard():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    # Reuse all variables from student_dashboard
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user:
        flash("User not found", "danger")
        return redirect(url_for("login"))

    total_max_marks = calculate_total_max_marks()

    problem_visibility = {}
    for i in range(1, 12):
        problem_visibility[i] = get_problem_visibility_for_user(i, user.rollnumber)

    quiz_released = check_quiz_visibility_for_student(user)
    quiz_attempts = QuizAttempt.query.filter_by(user_id=user.id).count()
    quiz_best_score = 0
    quiz_max_score = 10
    quiz_summaries = get_student_quiz_summaries(user)
    if quiz_attempts > 0:
        completed_attempts = (
            QuizAttempt.query.filter_by(user_id=user.id)
            .filter(QuizAttempt.completed_at.isnot(None))
            .all()
        )
        if completed_attempts:
            quiz_best_score = max(a.score for a in completed_attempts)
            quiz_max_score = completed_attempts[0].total_points

    setting = GlobalSettings.query.first()
    allow_student_mid_exam_questions = (
        setting.allow_student_mid_exam_questions if setting else False
    )

    completed_problems = sum(
        1 for i in range(1, 12) if getattr(user, f"p{i}_attempts", 0) > 0
    )

    active_conceptual_session = get_dom_conceptual_active_session_for_student(user)
    conceptual_response = (
        get_dom_conceptual_response(active_conceptual_session.id, user.id)
        if active_conceptual_session else None
    )

    return render_template(
        "dom_subject_student_dashboard.html",
        user=user,
        completed_problems=completed_problems,
        total_max_marks=total_max_marks,
        problem_visibility=problem_visibility,
        quiz_released=quiz_released,
        quiz_attempts=quiz_attempts,
        quiz_best_score=quiz_best_score,
        quiz_max_score=quiz_max_score,
        quiz_summaries=quiz_summaries,
        active_conceptual_session=active_conceptual_session,
        conceptual_response=conceptual_response,
        allow_student_mid_exam_questions=allow_student_mid_exam_questions,
    )

@app.route("/recalculate_completed_counts")
def recalculate_completed_counts():
    """Recalculate completed problem counts for all users"""
    if "loggedin" not in session or session.get("admin_level") != "super_admin":
        return "Access denied"

    try:
        users = User.query.filter_by(role='student').all()
        updated_count = 0

        for user in users:
            old_completed = user.completed
            # Recalculate completed count based on attempts > 0
            completed_count = 0
            for i in range(1, 12):
                if getattr(user, f"p{i}_attempts", 0) > 0:
                    completed_count += 1
            user.completed = completed_count

            if old_completed != completed_count:
                updated_count += 1

        db.session.commit()
        return f"Successfully recalculated completed counts for {updated_count} users out of {len(users)} total users."

    except Exception as e:
        return f"Error: {str(e)}"


@app.route("/debug_settings")
def debug_settings():
    """Debug route to check database settings"""
    if "loggedin" not in session:
        return "Please login first"

    # Check GlobalSettings
    global_setting = GlobalSettings.query.first()

    # Check PythonLabQuizSettings
    python_settings = PythonLabQuizSettings.query.all()

    debug_info = {
        "GlobalSettings exists": global_setting is not None,
        "allow_student_python_lab": global_setting.allow_student_python_lab if global_setting else "No GlobalSettings record",
        "allow_student_mid_exam_questions": global_setting.allow_student_mid_exam_questions if global_setting else "No GlobalSettings record",
        "PythonLabQuizSettings count": len(python_settings),
        "questions_visible_to_students": []
    }

    for setting in python_settings:
        debug_info["questions_visible_to_students"].append({
            "admin_id": setting.admin_id,
            "module_number": setting.module_number,
            "questions_visible_to_students": getattr(setting, 'questions_visible_to_students', 'Field not found')
        })

    return f"<pre>{debug_info}</pre>"


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        rollnumber = request.form["rollnumber"].strip()
        password = request.form["password"]
        email = request.form.get("email")
        phone = request.form.get("phonenumber")

        # --- Roll number restriction ---
        ENFORCE_ALLOWED_ROLLS = True  # Restrict registration to allowed roll numbers only

        # Check if current user is admin (admins can register anyone)
        is_admin_registering = False
        if "loggedin" in session and session.get("role") == "admin":
            is_admin_registering = True

        if ENFORCE_ALLOWED_ROLLS and not is_admin_registering and rollnumber not in ALLOWED_ROLLS:
            flash("This roll number is not authorized for registration.", "danger")
            return redirect(url_for("register"))

        existing_user = User.query.filter_by(rollnumber=rollnumber).first()
        if existing_user:
            flash(
                f"Roll number {rollnumber} is already registered. Please login instead.",
                "warning",
            )
            return redirect(url_for("login"))

        try:
            # --- Create new user record ---
            new_user = User(
                rollnumber=rollnumber,
                password=password,
                email=email,
                phonenumber=phone,
                marks=0,
                completed=0,
                role="student"  # All new registrations are students
            )

            # 🔹 STEP 1: Auto-assign admin_id from AdminRollAssignment
            import re
            linked_admin_id = None
            assignment = AdminRollAssignment.query.filter_by(roll_numbers=rollnumber).first()

            if assignment:
                linked_admin_id = assignment.admin_id
            else:
                # Check if roll falls under any range assignment
                all_ranges = AdminRollAssignment.query.filter_by(assignment_type="range").all()
                for a in all_ranges:
                    if a.roll_start and a.roll_end:
                        match_start = re.match(r"^(.*?)(\d+)$", a.roll_start)
                        match_end = re.match(r"^(.*?)(\d+)$", a.roll_end)
                        match_user = re.match(r"^(.*?)(\d+)$", rollnumber)
                        if match_start and match_end and match_user:
                            if match_start.group(1) == match_end.group(1) == match_user.group(1):
                                num_start = int(match_start.group(2))
                                num_end = int(match_end.group(2))
                                num_user = int(match_user.group(2))
                                if num_start <= num_user <= num_end:
                                    linked_admin_id = a.admin_id
                                    break

            if linked_admin_id and user_model_has_admin_id():
                new_user.admin_id = linked_admin_id

            db.session.add(new_user)
            db.session.commit()

            # --- Optional Debug / Info ---
            if linked_admin_id:
                flash(f"Automatically linked to admin ID {linked_admin_id}.", "info")

            # 🔹 STEP 2: Ensure entry exists in KDMLabStudent
            student_exists = KDMLabStudent.query.filter_by(rollnumber=new_user.rollnumber).first()
            if not student_exists:
                try:
                    new_student = KDMLabStudent(
                        rollnumber=new_user.rollnumber,
                        email=new_user.email,
                        phonenumber=new_user.phonenumber,
                        password=new_user.password
                    )
                    db.session.add(new_student)
                    db.session.commit()
                    print(f"✅ Added {new_user.rollnumber} to KDMLabStudent automatically.")
                except Exception as e:
                    db.session.rollback()
                    print(f"⚠️ Error creating KDMLabStudent record: {e}")

            # 🔹 STEP 4: Ensure entry exists in PythonLabStudent
            py_exists = PythonLabStudent.query.filter_by(rollnumber=new_user.rollnumber).first()
            if not py_exists:
                try:
                    new_py = PythonLabStudent(
                        rollnumber=new_user.rollnumber,
                        email=new_user.email,
                        phonenumber=new_user.phonenumber,
                        password=new_user.password
                    )
                    db.session.add(new_py)
                    db.session.commit()
                    print(f"✅ Added {new_user.rollnumber} to PythonLabStudent automatically.")
                except Exception as e:
                    db.session.rollback()
                    print(f"⚠️ Error creating PythonLabStudent record: {e}")

            flash("Registered successfully. Please login.", "success")
            return redirect(url_for("login"))



        except Exception as e:
            db.session.rollback()
            flash(f"Error during registration: {e}", "danger")

    return render_template("register.html")


# ---------------- Login ----------------


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        rollnumber = request.form.get("rollnumber")
        password = request.form.get("password")

        user = User.query.filter_by(rollnumber=rollnumber).first()

        # Debug: Check if user exists
        if not user:
            flash(f"No user found with roll number: {rollnumber}", "danger")
            return render_template("login.html")

        # Debug: Check password
        if user.password != password:
            flash(f"Password mismatch for user: {rollnumber}", "danger")
            return render_template("login.html")

        # Debug: Check role
        if not user.role:
            flash(f"User {rollnumber} has no role assigned", "danger")
            return render_template("login.html")

        # Login successful
        session["rollnumber"] = user.rollnumber
        session["role"] = user.role or ("student" if hasattr(user, "student_id") else None)
        session["loggedin"] = True

        if user.role in ["admin", "super_admin"]:
            flash(
                f"Admin login successful for {user.rollnumber} (role: {user.role}, admin_level: {getattr(user, 'admin_level', 'N/A')})", "success")
            return redirect(url_for("admin_dashboard"))
        else:
            session["role"] = "student"
            flash(f"Student login successful for {user.rollnumber}", "success")
            return redirect(url_for("student_dashboard"))

    return render_template("login.html")


# ---------------- Welcome ----------------
@app.route("/student_consolidated_marks")
def student_consolidated_marks():
    if "loggedin" not in session:
        return redirect(url_for("login"))

    student = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not student or student.role != "student":
        flash("Unauthorized access!", "danger")
        return redirect(url_for("login"))

    responsible_admin = find_admin_for_student(student.rollnumber)
    uploaded_lookup = {}
    settings = get_default_dom_consolidation_settings(responsible_admin)
    quiz_names = [DEFAULT_QUIZ_NAME]
    quiz_bank_map = {}

    if responsible_admin:
        settings = get_dom_consolidation_settings(responsible_admin)
        quiz_names = get_dom_admin_quiz_names(responsible_admin)
        quiz_bank_map = get_dom_admin_quiz_bank_map(responsible_admin)
        _, uploaded_lookup = load_uploaded_internal_marks(responsible_admin.id)

    row = build_dom_consolidated_row(
        student,
        current_admin=responsible_admin or student,
        uploaded_lookup=uploaded_lookup,
        settings=settings,
        quiz_names=quiz_names,
        quiz_bank_map=quiz_bank_map
    )

    data = [row]

    return render_template(
        "student_consolidated_marks.html",
        student=student,
        data=data,
        settings=settings,
        dom_consolidation_total=DOM_CONSOLIDATION_TOTAL
    )

@app.route("/welcome")
def welcome():
    if "loggedin" not in session:
        flash("Please login first.", "warning")
        return redirect(url_for("login"))

    # Fetch user from DB
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    if user:
        next_problem = (user.completed or 0) + 1

        if next_problem == 1:
            return redirect(url_for("problem1"))
        elif next_problem == 2:
            return redirect(url_for("problem2"))
        elif next_problem == 3:
            return redirect(url_for("problem3"))
        elif next_problem == 4:
            return redirect(url_for("problem4"))
        elif next_problem == 5:
            return redirect(url_for("problem5"))
        elif next_problem == 6:
            return redirect(url_for("problem6"))
        elif next_problem == 7:
            return redirect(url_for("problem7"))
        else:
            # ✅ Show congratulations inside welcome.html
            flash("🎉 Congratulations! You have completed all problems.", "success")
            return render_template("welcome.html", rollnumber=user.rollnumber, user=user)

    # fallback
    return render_template("welcome.html", rollnumber=session["rollnumber"])


# ---------------- Logout ----------------


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# ---------------- Problem 11 Helper Function ----------------


def build_steps_problem11(vars):
    m = vars['m']
    r = vars['r']
    angles = vars['angles']
    rb = vars['rb']

    steps = []

    # Step 1: Compute moments (m*r)
    M = [m[i] * r[i] for i in range(4)]
    steps.append("1. Compute each unbalance moment (M = m·r):")
    for i in range(4):
        steps.append(f"   M{i+1} = {m[i]} × {r[i]} = {M[i]:.3f} kg·m")

    # Step 2: Assign angular positions
    theta = [0] * 4
    theta[0] = 0
    for i in range(1, 4):
        theta[i] = theta[i-1] + angles[i-1]
    steps.append("2. Angular positions (cumulative): " +
                 ", ".join([f"{t}°" for t in theta]))

    # Step 3: Resolve into x,y components
    Sx, Sy = 0, 0
    steps.append("3. Resolve each moment into components:")
    for i in range(4):
        tx = M[i] * math.cos(math.radians(theta[i]))
        ty = M[i] * math.sin(math.radians(theta[i]))
        steps.append(
            f"   M{i+1}x = {M[i]:.3f}cos({theta[i]}°) = {tx:.3f},  M{i+1}y = {M[i]:.3f}sin({theta[i]}°) = {ty:.3f}")
        Sx += tx
        Sy += ty

    steps.append(f"   Resultant Sx = {Sx:.3f}, Sy = {Sy:.3f}")

    # Step 4: Resultant vector
    S = math.sqrt(Sx**2 + Sy**2)
    thetaS = math.degrees(math.atan2(Sy, Sx))
    steps.append(
        f"4. Resultant vector: |S| = √(Sx²+Sy²) = {S:.3f}, angle = atan2({Sy:.3f},{Sx:.3f}) = {thetaS:.3f}°")

    # Step 5: Balancing mass
    mb = S / rb
    theta_b = thetaS + 180
    if theta_b >= 360:
        theta_b -= 360
    steps.append(
        "5. Balance mass condition: mb·rb = |S|, placed opposite to resultant")
    steps.append(f"   mb = {S:.3f} / {rb} = {mb:.3f} kg")
    steps.append(f"   Position θb = {thetaS:.3f}° + 180° = {theta_b:.3f}°")

    return steps, M, theta, Sx, Sy, S, thetaS, mb, theta_b

# ---------------- Problem 1 ----------------


@app.route("/problem1", methods=["GET", "POST"])
def problem1():
    # Check problem access
    access_granted, user_or_redirect = check_problem_access(1)
    if not access_granted:
        return user_or_redirect

    user = user_or_redirect

    # Clear any lingering flash messages from previous admin actions when accessing problems
    # This prevents admin dashboard messages from appearing on problem pages
    if request.method == "GET" and request.referrer and ("admin" in request.referrer or "dashboard" in request.referrer):
        session.pop('_flashes', None)

    feedback, steps = None, None

    # Check if this is a new attempt (clear session for new problem)
    current_attempts = user.p1_attempts if user else 0
    if "p1_current_attempt" not in session or session["p1_current_attempt"] != current_attempts:
        # Clear problem session for new attempt
        for key in ["upper_arm", "inclination", "rise", "h1", "h2", "n1", "n2", "p1_answer"]:
            session.pop(key, None)
        session["p1_current_attempt"] = current_attempts

    # --- Handle Answer Submission (POST) ---
    if request.method == "POST":
        try:
            ans = float(request.form["answer"])
        except ValueError:
            flash("Invalid input", "danger")
            return redirect(url_for("problem1"))

        correct = session.get("p1_answer")
        if correct is None:
            flash("Session expired. Please refresh.", "warning")
            return redirect(url_for("problem1"))

        if abs(ans - correct) <= 0.05 * correct:
            score = 5
            feedback = f"✅ Correct! Answer ≈ {correct}%"
            # Check if this is first attempt with perfect score
            is_perfect_first_attempt = (user.p1_attempts == 0 and score == 5)
        else:
            score = 2
            feedback = f"❌ Incorrect. Correct ≈ {correct}%"
            is_perfect_first_attempt = False

        # Update score & attempts
        if not update_score(user, 1, score):
            flash("❌ Maximum attempts (2) exceeded for this problem!", "danger")
            # This should never happen for admin users due to unlimited attempts
            # Redirect based on user role
            if user.role == "admin":
                return redirect(url_for("admin_assignment_view"))
            else:
                return redirect(url_for("student_dashboard"))

        # Always build solution steps
        steps = build_steps_problem1(session)

        # Render template with feedback and steps (don't redirect)
        return render_template(
            "problem1.html",
            upper_arm=session["upper_arm"],
            inclination=session["inclination"],
            rise=session["rise"],
            feedback=feedback,
            steps=steps,
            user=user,
            show_celebration=is_perfect_first_attempt
        )

    # --- Generate new random values only on GET request (first time) ---
    # Only generate new problem if not already exists (preserve during POST)
    if "p1_answer" not in session:
        session["upper_arm"] = random.randint(350, 450)     # L (mm)
        session["inclination"] = random.randint(25, 35)     # θ (deg)
        session["rise"] = random.randint(15, 25)            # h (mm)

        cosi = math.cos(math.radians(session["inclination"]))
        h1 = session["upper_arm"] * 1e-3 * cosi
        h2 = h1 - session["rise"] * 1e-3
        n1 = math.sqrt(895 / h1)
        n2 = math.sqrt(895 / h2)
        perincspeed = round(((n2 - n1) / n1) * 100, 2)

        session["h1"], session["h2"] = h1, h2
        session["n1"], session["n2"] = n1, n2
        session["p1_answer"] = perincspeed

    return render_template(
        "problem1.html",
        upper_arm=session["upper_arm"],
        inclination=session["inclination"],
        rise=session["rise"],
        feedback=feedback,
        steps=steps,
        user=user,
        show_celebration=False
    )


def build_steps_problem1(vars):
    steps = []
    L = vars["upper_arm"]
    θ = vars["inclination"]
    h = vars["rise"]
    h1 = vars["h1"]
    h2 = vars["h2"]
    n1 = vars["n1"]
    n2 = vars["n2"]
    perinc = vars["p1_answer"]

    steps.append(
        f"Given: Upper arm length L = {L} mm, Inclination θ = {θ}°, Rise h = {h} mm")
    steps.append(f"h₁ = L·cosθ = {L}×cos({θ}°) = {h1:.4f} m")
    steps.append(f"h₂ = h₁ - (rise) = {h1:.4f} - {h*1e-3:.3f} = {h2:.4f} m")
    steps.append(f"n₁ = √(895/h₁) = √(895/{h1:.4f}) = {n1:.2f} rpm")
    steps.append(f"n₂ = √(895/h₂) = √(895/{h2:.4f}) = {n2:.2f} rpm")
    steps.append(f"% Increase in speed = ((n₂ - n₁)/n₁)×100 = {perinc}%")

    return steps


# ---------------- Problems 2–7 ----------------

# ---------------- Problem 2 ----------------


@app.route("/problem2", methods=["GET", "POST"])
def problem2():
    # Check problem access
    access_granted, user_or_redirect = check_problem_access(2)
    if not access_granted:
        return user_or_redirect

    user = user_or_redirect

    # Check if this is a new attempt (clear session for new problem)
    current_attempts = user.p2_attempts if user else 0
    if "p2_current_attempt" not in session or session["p2_current_attempt"] != current_attempts:
        # Clear problem session for new attempt
        for key in ["var21", "var22", "var23", "var24", "var25", "var26", "rangeofspeed"]:
            session.pop(key, None)
        session["p2_current_attempt"] = current_attempts

    # ---------------- Handle form submission ----------------
    if request.method == "POST" and session.get("rangeofspeed") is not None:
        try:
            answer = round(float(request.form["Answer"]), 2)
        except:
            flash("Invalid input", "danger")
            return redirect(url_for("problem2"))

        correct = session["rangeofspeed"]
        tol = 0.05 * correct
        if correct - tol <= answer <= correct + tol:
            feedback = f"✅ Correct! The answer is approx {correct} rpm"
            score = 5  # Changed from marks to score for consistency
            # Check if this is first attempt with perfect score
            is_perfect_first_attempt = (user.p2_attempts == 0 and score == 5)
        else:
            feedback = f"❌ Incorrect. Correct answer: {correct} rpm"
            score = 2  # Changed from marks to score for consistency
            is_perfect_first_attempt = False

        # Save best score
        if user:
            if not update_score(user, 2, score):  # Now using score consistently
                flash("❌ Maximum attempts (2) exceeded for this problem!", "danger")
                # Redirect based on user role
                if user.role == "admin":
                    return redirect(url_for("admin_assignment_view"))
                else:
                    return redirect(url_for("student_dashboard"))

        # Generate solution steps
        steps = build_steps_problem2(session)

        # Render template with feedback and steps (don't redirect)
        return render_template(
            "problem2.html",
            rollnumber=session["rollnumber"],
            central_mass=session["var21"],
            ball_mass=session["var22"],
            top_arms=session["var23"],
            bottom_arms=session["var24"],
            sleeve_friction=session["var25"],
            arms_angle=session["var26"],
            feedback=feedback,
            steps=steps,
            user=user,
            show_celebration=is_perfect_first_attempt
        )

    # ---------------- Generate new problem only if not exists ----------------
    # Only generate new variables if this is the first time loading this attempt
    if "rangeofspeed" not in session:
        session["var21"] = random.randint(14, 20)    # central load mass
        session["var22"] = random.randint(1, 5)      # ball mass
        session["var23"] = random.randint(200, 300)  # top arms (mm)
        session["var24"] = random.randint(250, 350)  # bottom arms (mm)
        session["var25"] = random.randint(10, 20)    # sleeve friction (N)
        session["var26"] = random.randint(40, 50)    # arm angle (deg)

    cosi2 = math.cos(session["var26"] * math.pi / 180)
    sini2 = math.sin(session["var26"] * math.pi / 180)

    h21 = session["var23"] * 1e-3 * cosi2
    h22 = session["var23"] * 1e-3 * sini2

    const = h22 / (session["var24"] * 1e-3)
    if const < -1 or const > 1:
        flash("Error: arcsin input out of range.", "danger")
        session["rangeofspeed"] = None
    else:
        b2 = math.degrees(math.asin(const))
        q2 = math.tan(math.radians(b2)) / \
            math.tan(session["var26"] * math.pi / 180)

        n21 = math.sqrt((895 / h21) * (
            (session["var22"] * 9.81 + ((session["var21"] * 9.81 - session["var25"]) / 2) * (1 + q2)) /
            (session["var22"] * 9.81)
        ))

        n22 = math.sqrt((895 / h21) * (
            (session["var22"] * 9.81 + ((session["var21"] * 9.81 + session["var25"]) / 2) * (1 + q2)) /
            (session["var22"] * 9.81)
        ))

        session["rangeofspeed"] = round(n22 - n21, 2)

    # ---------------- Render template ----------------
    return render_template(
        "problem2.html",
        rollnumber=session["rollnumber"],
        central_mass=session["var21"],
        ball_mass=session["var22"],
        top_arms=session["var23"],
        bottom_arms=session["var24"],
        sleeve_friction=session["var25"],
        arms_angle=session["var26"],
        feedback=None,
        steps=None,
        user=user,
        show_celebration=False
    )


# ---------------- Helper: steps for Problem 2 ----------------


def build_steps_problem2(vars):
    steps = []

    # Step 1: Extract data from session
    steps.append("Given Data:")
    steps.append(f"Central load mass (M) = {vars['var21']} kg")
    steps.append(f"Ball mass (m) = {vars['var22']} kg")
    steps.append(f"Length of upper arm (a) = {vars['var23']} mm")
    steps.append(f"Length of lower arm (b) = {vars['var24']} mm")
    steps.append(f"Sleeve friction (F) = {vars['var25']} N")
    steps.append(f"Arm angle (θ) = {vars['var26']}°")

    # Step 2: Compute trigonometric values
    cos_theta = math.cos(vars["var26"] * math.pi / 180)
    sin_theta = math.sin(vars["var26"] * math.pi / 180)
    steps.append("\nStep 1️⃣: Geometry of the governor arms")
    steps.append(
        f"cos(θ) = {round(cos_theta, 4)},  sin(θ) = {round(sin_theta, 4)}")

    # Step 3: Compute height (h) and radius (r)
    h21 = vars["var23"] * 1e-3 * cos_theta   # convert mm → m
    h22 = vars["var23"] * 1e-3 * sin_theta
    steps.append(
        f"Height of governor, h = a·cos(θ) = {vars['var23']}×10⁻³×{round(cos_theta, 4)} = {round(h21, 4)} m")
    steps.append(f"Radius of rotation, r = a·sin(θ) = {round(h22, 4)} m")

    # Step 4: Relation between upper & lower arms
    steps.append("\nStep 2️⃣: Relation between upper and lower arms")
    const = h22 / (vars["var24"] * 1e-3)
    steps.append(
        f"sin(β) = r/b = {round(h22, 4)}/{vars['var24']}×10⁻³ = {round(const, 4)}")

    if -1 <= const <= 1:
        b2 = math.degrees(math.asin(const))
        q2 = math.tan(math.radians(b2)) / \
            math.tan(vars["var26"] * math.pi / 180)
        steps.append(f"β = sin⁻¹({round(const, 4)}) = {round(b2, 4)}°")
        steps.append(f"Ratio, q = tan(β)/tan(θ) = {round(q2, 4)}")

        # Step 5: Compute governor speeds
        steps.append("\nStep 3️⃣: Calculation of governor speeds")
        steps.append("For a Porter governor:")
        steps.append(
            "n₁ = √[(895/h) × ((m·g + ((M·g − F)/2) × (1 + q)) / (m·g))]")
        steps.append(
            "n₂ = √[(895/h) × ((m·g + ((M·g + F)/2) × (1 + q)) / (m·g))]")

        n21 = math.sqrt((895 / h21) * (
            (vars["var22"] * 9.81 + ((vars["var21"] * 9.81 - vars["var25"]) / 2) * (1 + q2)) /
            (vars["var22"] * 9.81)
        ))

        n22 = math.sqrt((895 / h21) * (
            (vars["var22"] * 9.81 + ((vars["var21"] * 9.81 + vars["var25"]) / 2) * (1 + q2)) /
            (vars["var22"] * 9.81)
        ))

        steps.append(
            f"n₁ = {round(n21, 3)} rpm (minimum speed, friction acting downward)")
        steps.append(
            f"n₂ = {round(n22, 3)} rpm (maximum speed, friction acting upward)")

        # Step 6: Range of speed
        steps.append("\nStep 4️⃣: Range of speed")
        steps.append(
            f"Range = n₂ − n₁ = {round(n22, 3)} − {round(n21, 3)} = {round(n22 - n21, 2)} rpm")
        steps.append(
            f"✅ Therefore, Range of Speed = {round(n22 - n21, 2)} rpm")

    else:
        steps.append("⚠️ Error: arcsin input out of range — invalid geometry.")

    return steps

# ---------------- Problem 3 ----------------


@app.route("/problem3", methods=["GET", "POST"])
def problem3():
    # Check problem access
    access_granted, user_or_redirect = check_problem_access(3)
    if not access_granted:
        return user_or_redirect

    user = user_or_redirect

    # Check if this is a new attempt (clear session for new problem)
    current_attempts = user.p3_attempts if user else 0
    if "p3_current_attempt" not in session or session["p3_current_attempt"] != current_attempts:
        # Clear problem session for new attempt
        for key in ["mass", "gyration_radius", "speed_rpm", "velocity_kph", "curvature_radius", "turn_direction", "rotation_direction", "gyroscopic_couple", "effect"]:
            session.pop(key, None)
        session["p3_current_attempt"] = current_attempts

    # Generate new variables only if not already exists
    if "gyroscopic_couple" not in session:
        session["mass"] = random.randint(4000, 5000)
        session["gyration_radius"] = random.randint(10, 90) / 100
        session["speed_rpm"] = random.randint(2000, 3000)
        session["velocity_kph"] = random.randint(20, 80)
        session["curvature_radius"] = random.randint(20, 100)
        session["turn_direction"] = "left" if random.randint(0, 1) else "right"
        session["rotation_direction"] = (
            "clockwise" if random.randint(0, 1) else "anticlockwise"
        )

        mass = session["mass"]
        k = session["gyration_radius"]
        N = session["speed_rpm"]
        V_kph = session["velocity_kph"]
        R = session["curvature_radius"]

        V = V_kph * 5 / 18  # km/hr → m/s
        I = mass * (k**2)  # moment of inertia
        omega = 2 * math.pi * N / 60
        omega_p = V / R

        session["gyroscopic_couple"] = round(I * omega * omega_p, 2)

        if (
            (session["turn_direction"] ==
             "left" and session["rotation_direction"] == "clockwise")
            or (session["turn_direction"] == "right" and session["rotation_direction"] == "anticlockwise")
        ):
            session["effect"] = "raise the nose and dip the tail"
        else:
            session["effect"] = "dip the nose and raise the tail"

    # Fetch user progress
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    feedback = None
    steps = None
    if request.method == "POST":
        try:
            # don't round student's input here
            answer = float(request.form["Answer"])
            effect = request.form["Effect"]
        except:
            flash("Invalid input", "danger")
            return redirect(url_for("student_dashboard"))

        # --- Recompute correct couple from current session values (no rounding) ---
        m = session["mass"]
        k = session["gyration_radius"]
        N = session["speed_rpm"]
        V = session["velocity_kph"] * 5 / 18
        R = session["curvature_radius"]

        I = m * (k ** 2)
        omega = 2 * math.pi * N / 60
        omega_p = V / R

        correct_couple_exact = I * omega * omega_p          # precise physics value
        # only for display / feedback
        correct_couple_disp = round(correct_couple_exact, 2)

        # keep effect from session
        correct_effect = session["effect"]

        # 5% tolerance based on exact value
        tol = 0.05 * correct_couple_exact
        is_correct_couple = (correct_couple_exact -
                             tol) <= answer <= (correct_couple_exact + tol)
        is_correct_effect = (effect == correct_effect)

        # optional: keep session display value aligned (not required)
        session["gyroscopic_couple"] = correct_couple_disp

        if is_correct_couple and is_correct_effect:
            feedback = f"✅ Correct! Couple ≈ {correct_couple_disp} Nm, Effect = '{correct_effect}'"
            marks = 5
            is_perfect_first_attempt = (user.p3_attempts == 0 and marks == 5)
        elif is_correct_couple or is_correct_effect:
            feedback = f"⚠️ Partially correct. Correct answer: {correct_couple_disp} Nm, Effect = '{correct_effect}'"
            marks = 3
            is_perfect_first_attempt = False
        else:
            feedback = f"❌ Incorrect. Correct answer: {correct_couple_disp} Nm, Effect = '{correct_effect}'"
            marks = 2
            is_perfect_first_attempt = False

        if user:
            if not update_score(user, 3, marks):
                flash("❌ Maximum attempts (2) exceeded for this problem!", "danger")
                if user.role == "admin":
                    return redirect(url_for("admin_assignment_view"))
                else:
                    return redirect(url_for("student_dashboard"))

        steps = build_steps_problem3(session)

    return render_template(
        "problem3.html",
        rollnumber=session["rollnumber"],
        vars=session,
        feedback=feedback,
        steps=steps,
        user=user,
        show_celebration=feedback and is_perfect_first_attempt if 'is_perfect_first_attempt' in locals() else False
    )


# ---------------- Helper: steps for Problem 3 ----------------


def build_steps_problem3(vars):
    steps = []

    # Extract variables
    m = vars["mass"]
    k = vars["gyration_radius"]
    N = vars["speed_rpm"]
    V_kph = vars["velocity_kph"]
    R = vars["curvature_radius"]
    effect = vars["effect"]

    # Step calculations (no rounding)
    I = m * (k ** 2)                         # kg·m²
    omega = 2 * math.pi * N / 60             # rad/s
    V = V_kph * 5 / 18                       # m/s
    omega_p = V / R                          # rad/s
    couple = I * omega * omega_p             # Nm

    # Final couple rounded only once for display
    rounded_couple = round(couple, 2)

    # Steps (display with full precision where logical)
    steps.append(f"I = m·k² = {m}·{k}² = {I:.6f} kg·m²")
    steps.append(f"ω = 2πN/60 = {omega:.6f} rad/s")
    steps.append(f"V = {V_kph} km/h × 5/18 = {V:.6f} m/s")
    steps.append(f"ωp = V/R = {omega_p:.6f} rad/s")
    steps.append(f"C = I·ω·ωp = {rounded_couple} Nm")
    steps.append(f"Effect: {effect}")

    return steps


# ---------------- Problem 4 ----------------


@app.route("/problem4", methods=["GET", "POST"])
def problem4():
    # Check problem access
    access_granted, user_or_redirect = check_problem_access(4)
    if not access_granted:
        return user_or_redirect

    user = user_or_redirect

    # Check if this is a new attempt (clear session for new problem)
    current_attempts = user.p4_attempts if user else 0
    if "p4_current_attempt" not in session or session["p4_current_attempt"] != current_attempts:
        # Clear problem session for new attempt
        for key in ["mass", "gyration_radius", "speed_rpm", "velocity_kph", "curvature_radius", "turn_direction", "rotation_direction", "pitching_period", "pitching_angle", "pitching_direction", "gyroscopic_couple_steering", "effect_steering", "gyroscopic_couple_pitching", "effect_pitching"]:
            session.pop(key, None)
        session["p4_current_attempt"] = current_attempts

    # Generate new variables only if not already exists
    if "gyroscopic_couple_steering" not in session:
        session["mass"] = random.randint(3000, 10000)
        session["gyration_radius"] = random.randint(10, 90) / 100
        session["speed_rpm"] = random.randint(1500, 3000)
        session["velocity_kph"] = random.randint(20, 60)
        session["curvature_radius"] = random.randint(30, 100)
        session["turn_direction"] = "left" if random.randint(0, 1) else "right"
        session["rotation_direction"] = "clockwise" if random.randint(
            0, 1) else "anticlockwise"
        session["pitching_period"] = random.randint(20, 60)
        session["pitching_angle"] = random.randint(4, 10)
        session["pitching_direction"] = "bow rising" if random.randint(
            0, 1) else "bow descending"

        mass = session["mass"]
        k = session["gyration_radius"]
        N = session["speed_rpm"]
        V_kph = session["velocity_kph"]
        R = session["curvature_radius"]

        V = V_kph * 5 / 18
        I = mass * (k ** 2)
        omega = 2 * math.pi * N / 60
        omega_p_steering = V / R
        session["gyroscopic_couple_steering"] = round(
            I * omega * omega_p_steering, 2)

        if (session["turn_direction"] == "left" and session["rotation_direction"] == "clockwise") or \
           (session["turn_direction"] == "right" and session["rotation_direction"] == "anticlockwise"):
            session["effect_steering"] = "raise the bow and dip the stern"
        else:
            session["effect_steering"] = "dip the bow and raise the stern"

        pitching_angle_rad = session["pitching_angle"] * math.pi / 180
        omega_p_pitching = (
            2 * math.pi / session["pitching_period"]) * pitching_angle_rad
        session["gyroscopic_couple_pitching"] = round(
            I * omega * omega_p_pitching, 2)

        if (session["pitching_direction"] == "bow descending" and session["rotation_direction"] == "anticlockwise") or \
           (session["pitching_direction"] == "bow rising" and session["rotation_direction"] == "clockwise"):
            session["effect_pitching"] = "ship turning starboard side"
        else:
            session["effect_pitching"] = "ship turning port side"

    # Fetch user progress
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    feedback = None
    steps = None
    if request.method == "POST":
        try:
            # Do NOT round user inputs
            ans_steering = float(request.form["Answer_steering"])
            eff_steering = request.form["Effect_steering"]
            ans_pitching = float(request.form["Answer_pitching"])
            eff_pitching = request.form["Effect_pitching"]
        except:
            flash("Invalid input", "danger")
            return redirect(url_for("student_dashboard"))

        # Pull current params
        m = session["mass"]
        k = session["gyration_radius"]
        N = session["speed_rpm"]
        V_kph = session["velocity_kph"]
        R = session["curvature_radius"]
        T = session["pitching_period"]
        theta_deg = session["pitching_angle"]

        # Physics (NO rounding here)
        I = m * (k ** 2)                      # kg·m²
        omega = 2 * math.pi * N / 60          # rad/s
        V = V_kph * 5 / 18                    # m/s
        omega_p_steering = V / R              # rad/s

        theta = theta_deg * math.pi / 180.0   # rad
        # rad/s (peak small-angle rate)
        omega_p_pitching = (2 * math.pi / T) * theta

        # Exact couples for checking
        c_st_exact = I * omega * omega_p_steering
        c_pi_exact = I * omega * omega_p_pitching

        # Rounded only for display/feedback
        c_st_disp = round(c_st_exact, 2)
        c_pi_disp = round(c_pi_exact, 2)

        # Keep effects from session
        e_st = session["effect_steering"]
        e_pi = session["effect_pitching"]

        # 5% tolerance (on exact values)
        tol = 0.05
        st_c_ok = (c_st_exact * (1 - tol) <=
                   ans_steering <= c_st_exact * (1 + tol))
        st_e_ok = (eff_steering == e_st)
        pi_c_ok = (c_pi_exact * (1 - tol) <=
                   ans_pitching <= c_pi_exact * (1 + tol))
        pi_e_ok = (eff_pitching == e_pi)

        # Also update the session's display values so they match what you show
        session["gyroscopic_couple_steering"] = c_st_disp
        session["gyroscopic_couple_pitching"] = c_pi_disp

        # --- scoring logic (unchanged) ---
        if st_c_ok and st_e_ok and pi_c_ok and pi_e_ok:
            marks, feedback = 10, "All answers are correct!"
            is_perfect_first_attempt = (user.p4_attempts == 0 and marks == 10)
        elif (st_c_ok and st_e_ok and pi_c_ok and not pi_e_ok) or (st_c_ok and not st_e_ok and pi_c_ok and pi_e_ok):
            marks = 9
            feedback = f"Partially correct: mistake in one effect. Correct: Steering Couple = {c_st_disp} Nm, Effect = '{e_st}'; Pitching Couple = {c_pi_disp} Nm, Effect = '{e_pi}'"
            is_perfect_first_attempt = False
        elif (st_c_ok and st_e_ok and not pi_c_ok and pi_e_ok) or (not st_c_ok and st_e_ok and pi_c_ok and pi_e_ok):
            marks = 7
            feedback = f"Partially correct: mistake in one couple. Correct: Steering Couple = {c_st_disp} Nm, Effect = '{e_st}'; Pitching Couple = {c_pi_disp} Nm, Effect = '{e_pi}'"
            is_perfect_first_attempt = False
        elif (not st_c_ok and not st_e_ok and pi_c_ok and pi_e_ok) or (st_c_ok and st_e_ok and not pi_c_ok and not pi_e_ok):
            marks = 6
            feedback = f"Partially correct: mistake in one couple+effect. Correct: Steering Couple = {c_st_disp} Nm, Effect = '{e_st}'; Pitching Couple = {c_pi_disp} Nm, Effect = '{e_pi}'"
            is_perfect_first_attempt = False
        elif st_c_ok and not st_e_ok and pi_c_ok and not pi_e_ok:
            marks = 8
            feedback = f"Partially correct: both effects wrong. Correct: Steering Couple = {c_st_disp} Nm, Effect = '{e_st}'; Pitching Couple = {c_pi_disp} Nm, Effect = '{e_pi}'"
            is_perfect_first_attempt = False
        elif (st_c_ok and not st_e_ok and not pi_c_ok and not pi_e_ok) or (not st_c_ok and not st_e_ok and pi_c_ok and not pi_e_ok):
            marks = 5
            feedback = f"Partially correct: only one couple right. Correct: Steering Couple = {c_st_disp} Nm, Effect = '{e_st}'; Pitching Couple = {c_pi_disp} Nm, Effect = '{e_pi}'"
            is_perfect_first_attempt = False
        elif not st_c_ok and st_e_ok and not pi_c_ok and pi_e_ok:
            marks = 4
            feedback = f"Partially correct: both couples wrong. Correct: Steering Couple = {c_st_disp} Nm, Effect = '{e_st}'; Pitching Couple = {c_pi_disp} Nm, Effect = '{e_pi}'"
            is_perfect_first_attempt = False
        elif (not st_c_ok and st_e_ok and not pi_c_ok and not pi_e_ok) or (not st_c_ok and not st_e_ok and not pi_c_ok and pi_e_ok):
            marks = 3
            feedback = f"Partially correct: only one effect right. Correct: Steering Couple = {c_st_disp} Nm, Effect = '{e_st}'; Pitching Couple = {c_pi_disp} Nm, Effect = '{e_pi}'"
            is_perfect_first_attempt = False
        else:
            marks = 2
            feedback = f"All answers incorrect. Correct: Steering Couple = {c_st_disp} Nm, Effect = '{e_st}'; Pitching Couple = {c_pi_disp} Nm, Effect = '{e_pi}'"
            is_perfect_first_attempt = False

        if user:
            if not update_score(user, 4, marks):
                flash("❌ Maximum attempts (2) exceeded for this problem!", "danger")
                if user.role == "admin":
                    return redirect(url_for("admin_assignment_view"))
                else:
                    return redirect(url_for("student_dashboard"))

        steps = build_steps_problem4(session)

    return render_template(
        "problem4.html",
        rollnumber=session["rollnumber"],
        vars=session,
        feedback=feedback,
        steps=steps,
        user=user,
        show_celebration=is_perfect_first_attempt if 'is_perfect_first_attempt' in locals() else False
    )


# ---------------- Helper: steps for Problem 4 ----------------
def build_steps_problem4(vars):
    steps = []

    m = vars["mass"]
    k = vars["gyration_radius"]
    N = vars["speed_rpm"]
    V_kph = vars["velocity_kph"]
    R = vars["curvature_radius"]
    T = vars["pitching_period"]
    theta_deg = vars["pitching_angle"]

    I = m * (k ** 2)
    omega = 2 * math.pi * N / 60
    V = V_kph * 5 / 18
    omega_p_steering = V / R
    theta = theta_deg * math.pi / 180.0
    omega_p_pitching = (2 * math.pi / T) * theta

    # Couples rounded ONLY for display (values already stored in session by POST)
    c_st_disp = vars["gyroscopic_couple_steering"]
    c_pi_disp = vars["gyroscopic_couple_pitching"]

    steps.append(f"I = m·k² = {m}·{k}² = {I:.6f} kg·m²")
    steps.append(f"ω = 2πN/60 = {omega:.6f} rad/s")
    steps.append(f"V = {V_kph} km/h × 5/18 = {V:.6f} m/s")
    steps.append(f"ωp (steering) = V/R = {omega_p_steering:.6f} rad/s")
    steps.append(
        f"C (steering) = I·ω·ωp = {c_st_disp} Nm, Effect: {vars['effect_steering']}")
    steps.append(f"ωp (pitching) = (2π/T)·θ = {omega_p_pitching:.6f} rad/s")
    steps.append(
        f"C (pitching) = I·ω·ωp = {c_pi_disp} Nm, Effect: {vars['effect_pitching']}")
    return steps

# -------------------------------------------------------------------------------------------
# ---------------- Problem 5 ----------------


@app.route("/problem5", methods=["GET", "POST"])
def problem5():
    # Check problem access
    access_granted, user_or_redirect = check_problem_access(5)
    if not access_granted:
        return user_or_redirect

    user = user_or_redirect

    # Check if this is a new attempt (clear session for new problem)
    current_attempts = user.p5_attempts if user else 0
    if "p5_current_attempt" not in session or session["p5_current_attempt"] != current_attempts:
        # Clear problem session for new attempt
        for key in ["P", "D", "mass", "crank_radius", "engine_speed", "angle_theta", "connecting_rod_length", "FN", "FQ", "FT", "T"]:
            session.pop(key, None)
        session["p5_current_attempt"] = current_attempts

    # Generate new variables only if not already exists
    if "FN" not in session:
        while True:
            P = random.randint(2, 9) * 1e5
            D = random.randint(20, 90) / 100
            mass = random.randint(150, 500)
            r = random.randint(10, 60) / 100
            N = random.randint(200, 500)
            theta_deg = random.randint(30, 70)
            l = random.randint(110, 190) / 100

            omega = 2 * math.pi * N / 60
            n = l / r

            FL = P * (math.pi * D**2 / 4)
            FI = mass * (omega**2) * r * (math.cos(math.radians(theta_deg)
                                                   ) + math.cos(math.radians(2*theta_deg)) / n)
            FP = FL - FI

            if FP > 0:  # must be positive
                phi_rad = math.asin(math.sin(math.radians(theta_deg)) / n)
                FN = FP * math.tan(phi_rad)
                FQ = FP / math.cos(phi_rad)
                FT = FQ * math.sin(phi_rad + math.radians(theta_deg))
                T = FT * r

                session["P"] = P
                session["D"] = D
                session["mass"] = mass
                session["crank_radius"] = r
                session["engine_speed"] = N
                session["angle_theta"] = theta_deg
                session["connecting_rod_length"] = l
                session["FN"] = round(FN, 2)
                session["FQ"] = round(FQ, 2)
                session["FT"] = round(FT, 2)
                session["T"] = round(T, 2)
                break

    feedback = None
    steps = None
    if request.method == "POST":
        try:
            ans_FN = float(request.form["Answer_FN"])
            ans_FQ = float(request.form["Answer_FQ"])
            ans_FT = float(request.form["Answer_FT"])
            ans_T = float(request.form["Answer_T"])
        except:
            flash("Invalid input", "danger")
            return redirect(url_for("student_dashboard"))

        FNc, FQc, FTc, Tc = session["FN"], session["FQ"], session["FT"], session["T"]
        tol = 0.05
        FN_ok = FNc*(1-tol) <= ans_FN <= FNc*(1+tol)
        FQ_ok = FQc*(1-tol) <= ans_FQ <= FQc*(1+tol)
        FT_ok = FTc*(1-tol) <= ans_FT <= FTc*(1+tol)
        T_ok = Tc*(1-tol) <= ans_T <= Tc*(1+tol)

        if FN_ok and FQ_ok and FT_ok and T_ok:
            marks, feedback = 10, "All answers are correct!"
            # Check if this is first attempt with perfect score
            is_perfect_first_attempt = (user.p5_attempts == 0 and marks == 10)
        elif FN_ok and FQ_ok and FT_ok:
            marks, feedback = 8, "Partially correct: FN, FQ, FT correct."
            is_perfect_first_attempt = False
        elif FN_ok and FQ_ok:
            marks, feedback = 6, "Partially correct: FN and FQ correct."
            is_perfect_first_attempt = False
        elif FN_ok:
            marks, feedback = 4, "Partially correct: only FN correct."
            is_perfect_first_attempt = False
        else:
            marks, feedback = 2, "Incorrect answers."
            is_perfect_first_attempt = False

        if user:
            if not update_score(user, 5, marks):
                flash("❌ Maximum attempts (2) exceeded for this problem!", "danger")
                # Redirect based on user role
                if user.role == "admin":
                    return redirect(url_for("admin_assignment_view"))
                else:
                    return redirect(url_for("student_dashboard"))
        steps = build_steps_problem5(session, ans_FN, ans_FQ, ans_FT, ans_T)

    return render_template(
        "problem5.html",
        rollnumber=session["rollnumber"],
        vars=session,
        feedback=feedback,
        steps=steps,
        user=user,
        show_celebration=is_perfect_first_attempt if 'is_perfect_first_attempt' in locals() else False
    )


# ---------------- Helper: steps for Problem 5 ----------------
def build_steps_problem5(vars, ans_FN, ans_FQ, ans_FT, ans_T):
    steps = []
    steps.append(f"Your FN={round(ans_FN, 2)} vs Correct {vars['FN']}")
    steps.append(f"Your FQ={round(ans_FQ, 2)} vs Correct {vars['FQ']}")
    steps.append(f"Your FT={round(ans_FT, 2)} vs Correct {vars['FT']}")
    steps.append(f"Your T={round(ans_T, 2)} vs Correct {vars['T']}")
    steps.append("Steps:")
    steps.append("1. FL = P * (πD²/4)")
    steps.append("2. FI = mω²r(cosθ + cos2θ/n)")
    steps.append("3. FP = FL - FI")
    steps.append("4. FN = FP tanφ, φ=asin(sinθ/n)")
    steps.append("5. FQ = FP / cosφ")
    steps.append("6. FT = FQ sin(φ+θ)")
    steps.append("7. T = FT * r")
    return steps


# ---------------- Problem 6 ----------------
@app.route("/problem6", methods=["GET", "POST"])
def problem6():
    # Check problem access
    access_granted, user_or_redirect = check_problem_access(6)
    if not access_granted:
        return user_or_redirect

    user = user_or_redirect

    # Check if this is a new attempt (clear session for new problem)
    current_attempts = user.p6_attempts if user else 0
    if "p6_current_attempt" not in session or session["p6_current_attempt"] != current_attempts:
        # Clear problem session for new attempt
        for key in ["radius", "engine_speed", "fluctuation_percent", "vertical_scale", "horizontal_scale", "energy", "DELTA_E", "omega", "CS", "mass"]:
            session.pop(key, None)
        session["p6_current_attempt"] = current_attempts

    # Generate new variables only if not already exists
    if "mass" not in session:
        session["radius"] = random.randint(20, 80) / 100
        session["engine_speed"] = random.randint(300, 800)
        session["fluctuation_percent"] = random.randint(11, 21) / 10
        session["vertical_scale"] = random.randint(300, 800)
        session["horizontal_scale"] = random.randint(2, 6)

        areas = [
            random.randint(45, 55),
            random.randint(-130, -100),
            random.randint(80, 100),
            random.randint(-160, -130),
            random.randint(70, 100),
            random.randint(-85, -60),
            random.randint(90, 120)
        ]
        session["energy"] = areas

        energies = [areas[0]]
        for i in range(1, len(areas)):
            energies.append(energies[-1] + areas[i])

        maxE, minE = max(energies), min(energies)
        DELTA_E = (maxE - minE) * session["vertical_scale"] * \
            session["horizontal_scale"] * math.pi / 180
        session["DELTA_E"] = DELTA_E

        omega = 2 * math.pi * session["engine_speed"] / 60
        session["omega"] = omega
        CS = 2 * (session["fluctuation_percent"] / 100)
        session["CS"] = CS

        session["mass"] = DELTA_E / (session["radius"]**2 * omega**2 * CS)

    # Fetch user progress
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    feedback = None
    steps = None
    if request.method == "POST":
        try:
            ans_mass = float(request.form["Answer_mass"])
        except:
            flash("Invalid input", "danger")
            return redirect(url_for("student_dashboard"))

        correct_mass = round(session["mass"], 2)
        tol = 0.05
        min_m, max_m = correct_mass * (1 - tol), correct_mass * (1 + tol)
        if min_m <= ans_mass <= max_m:
            marks, feedback = 10, "✅ Correct! Mass within range."
            # Check if this is first attempt with perfect score
            is_perfect_first_attempt = (user.p6_attempts == 0 and marks == 10)
        else:
            marks, feedback = 2, f"❌ Incorrect. Correct mass ≈ {correct_mass} kg."
            is_perfect_first_attempt = False

        if user:
            if not update_score(user, 6, marks):
                flash("❌ Maximum attempts (2) exceeded for this problem!", "danger")
                # Redirect based on user role
                if user.role == "admin":
                    return redirect(url_for("admin_assignment_view"))
                else:
                    return redirect(url_for("student_dashboard"))
        steps = build_steps_problem6(session, ans_mass)

    return render_template(
        "problem6.html",
        rollnumber=session["rollnumber"],
        vars=session,
        feedback=feedback,
        steps=steps,
        user=user,
        show_celebration=is_perfect_first_attempt if 'is_perfect_first_attempt' in locals() else False
    )


# ---------------- Helper: steps for Problem 6 ----------------
def build_steps_problem6(vars, ans_mass):
    steps = []
    steps.append(
        f"Your Answer: {round(ans_mass, 2)} kg vs Correct ≈ {round(vars['mass'], 2)} kg")
    steps.append("Steps:")
    steps.append("1. Compute intercepted areas → cumulative energies.")
    steps.append(
        "2. Find ΔE = (max-min) * vertical_scale * horizontal_scale * π/180")
    steps.append(f"   ΔE = {round(vars['DELTA_E'], 2)} J")
    steps.append(f"3. ω = 2πN/60 = {round(vars['omega'], 2)} rad/s")
    steps.append(f"4. C_s = 2*(fluctuation%) = {round(vars['CS'], 4)}")
    steps.append("5. Mass = ΔE / (R² * ω² * C_s)")
    steps.append(f"   Mass ≈ {round(vars['mass'], 2)} kg")
    return steps

# (Add problem2(), problem3(), … using the conversions we already built)
# Due to space, I’ll not paste all here, but each problem follows the
# same template: generate vars → check DB → process POST → update marks → show steps.
# -----------------------------------------------------------------
# Example: Problem 7 (already converted)


@app.route("/problem7", methods=["GET", "POST"])
def problem7():
    # Check problem access
    access_granted, user_or_redirect = check_problem_access(7)
    if not access_granted:
        return user_or_redirect

    user = user_or_redirect

    # Clear any lingering flash messages from previous admin actions when accessing problems
    # This prevents admin dashboard messages from appearing on problem pages
    if request.method == "GET" and request.referrer and ("admin" in request.referrer or "dashboard" in request.referrer):
        session.pop('_flashes', None)

    # Check if this is a new attempt (clear session for new problem)
    current_attempts = user.p7_attempts if user else 0
    if "p7_current_attempt" not in session or session["p7_current_attempt"] != current_attempts:
        # Clear problem session for new attempt
        for key in ["shaft_diameter", "shaft_length", "load1", "load2", "load3", "distance1", "distance2", "distance3", "young_modulus", "frequency", "deflections", "I"]:
            session.pop(key, None)
        session["p7_current_attempt"] = current_attempts

    # Generate new variables only if not already exists
    if "frequency" not in session:
        session["shaft_diameter"] = random.randint(40, 60)
        session["shaft_length"] = random.randint(6, 10)
        session["load1"] = random.randint(500, 1500)
        session["load2"] = random.randint(1000, 2000)
        session["load3"] = random.randint(500, 1000)
        session["distance1"] = random.randint(1, 2)
        session["distance2"] = random.randint(3, 4)
        session["distance3"] = random.randint(5, 6)
        session["young_modulus"] = random.randint(150, 250)

        d_m = session["shaft_diameter"] / 1000
        L = session["shaft_length"]
        E = session["young_modulus"] * 1e9
        I = math.pi * d_m**4 / 64

        δ1 = (session["load1"] * session["distance1"]**2 *
              (L - session["distance1"])**2) / (3 * E * I * L)
        δ2 = (session["load2"] * session["distance2"]**2 *
              (L - session["distance2"])**2) / (3 * E * I * L)
        δ3 = (session["load3"] * session["distance3"]**2 *
              (L - session["distance3"])**2) / (3 * E * I * L)

        δ_total = δ1 + δ2 + δ3
        f_n = 0.4985 / math.sqrt(δ_total)

        session["frequency"] = round(f_n, 2)
        session["deflections"] = [δ1, δ2, δ3, δ_total]
        session["I"] = I

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    feedback, steps = None, None
    if request.method == "POST":
        try:
            # Extract numeric value from answer (handle cases like "19.7 Hz", "19.7", etc.)
            answer_str = request.form["answer"].strip()
            # Remove common units and extra text
            import re
            # Extract the first number from the string
            number_match = re.search(r'[-+]?(?:\d*\.\d+|\d+)', answer_str)
            if number_match:
                ans = float(number_match.group())
            else:
                raise ValueError("No numeric value found in answer")
        except (ValueError, TypeError) as e:
            flash(
                f"❌ Invalid answer format. Please enter a numeric value (units are optional). Error: {str(e)}", "error")
            return redirect(url_for("problem7"))

        correct = session["frequency"]
        if correct*0.95 <= ans <= correct*1.05:
            marks, feedback = 10, "✅ Correct"
            # Check if this is first attempt with perfect score
            is_perfect_first_attempt = (user.p7_attempts == 0 and marks == 10)
        else:
            marks, feedback = 2, f"❌ Wrong. Correct ≈ {correct} Hz"
            is_perfect_first_attempt = False

        if user:
            if not update_score(user, 7, marks):
                flash("❌ Maximum attempts (2) exceeded for this problem!", "danger")
                # Redirect based on user role
                if user.role == "admin":
                    return redirect(url_for("admin_assignment_view"))
                else:
                    return redirect(url_for("student_dashboard"))
         # ✅ Generate solution steps
        steps = build_steps_problem7(session)

    return render_template("problem7.html", vars=session, feedback=feedback, user=user, steps=steps,
                           show_celebration=is_perfect_first_attempt if 'is_perfect_first_attempt' in locals() else False)


def build_steps_problem7(vars):
    steps = []

    # Extract values
    D_mm = vars["shaft_diameter"]
    D_m = D_mm / 1000
    L = vars["shaft_length"]
    W1, W2, W3 = vars["load1"], vars["load2"], vars["load3"]
    a1, a2, a3 = vars["distance1"], vars["distance2"], vars["distance3"]
    E_gpa = vars["young_modulus"]
    E = E_gpa * 1e9

    # Step 1: Moment of inertia
    I = (math.pi * (D_m**4)) / 64
    steps.append(f"Moment of inertia (I) = π·D⁴/64 = {round(I, 6)} m⁴")

    # Step 2: Deflection for each load
    δ1 = (W1 * (a1**2) * ((L - a1)**2)) / (3 * E * I * L)
    δ2 = (W2 * (a2**2) * ((L - a2)**2)) / (3 * E * I * L)
    δ3 = (W3 * (a3**2) * ((L - a3)**2)) / (3 * E * I * L)
    steps.append(f"Deflection δ₁ = {round(δ1, 8)} m")
    steps.append(f"Deflection δ₂ = {round(δ2, 8)} m")
    steps.append(f"Deflection δ₃ = {round(δ3, 8)} m")

    # Step 3: Total deflection
    δ_total = δ1 + δ2 + δ3
    steps.append(f"Total deflection δ_total = {round(δ_total, 8)} m")

    # Step 4: Frequency calculation
    freq = 0.4985 / math.sqrt(δ_total)
    steps.append(
        f"Natural frequency fₙ = 0.4985 / √δ_total = {round(freq, 2)} Hz")

    return steps


@app.route("/problem8", methods=["GET", "POST"])
def problem8():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    # Clear any lingering flash messages from previous admin actions when accessing problems
    # This prevents admin dashboard messages from appearing on problem pages
    if request.method == "GET" and request.referrer and ("admin" in request.referrer or "dashboard" in request.referrer):
        session.pop('_flashes', None)

    # ---------------- Generate random problem values only on GET request ----------------
    if request.method == "GET":
        session["d"] = random.randint(
            80, 120) / 1000     # m (diameter 80–120 mm)
        session["L"] = random.randint(800, 1200) / 1000   # m (0.8–1.2 m)
        session["M"] = random.randint(400, 600)           # kg (disc mass)
        session["k"] = random.randint(
            400, 500) / 1000    # m (radius of gyration)
        session["E"] = random.randint(190, 210) * 1e9     # Pa
        session["G"] = random.randint(75, 85) * 1e9       # Pa

    # Use the stored values (either newly generated on GET or existing on POST)
    d, L, M, k, E, G = session["d"], session["L"], session["M"], session["k"], session["E"], session["G"]

    # Section properties
    A = math.pi * d**2 / 4
    I = math.pi * d**4 / 64
    J = math.pi * d**4 / 32

    # Calculate answers only on GET request (when new problem is generated)
    if request.method == "GET":
        # 1. Longitudinal
        kL = A * E / L
        fL = (1 / (2 * math.pi)) * math.sqrt(kL / M)

        # 2. Transverse
        W = M * 9.81
        delta = (W * L**3) / (3 * E * I)
        fT = 0.4985 / math.sqrt(delta)

        # 3. Torsional
        kT = G * J / L
        Ip = M * (k**2)
        fTor = (1 / (2 * math.pi)) * math.sqrt(kT / Ip)

        # Store answers
        session["fL"] = round(fL, 2)
        session["fT"] = round(fT, 2)
        session["fTor"] = round(fTor, 2)

    # ---------------- Fetch user ----------------
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    feedback, steps = None, None

    if request.method == "POST":
        try:
            ans_fL = float(request.form["Answer_fL"])
            ans_fT = float(request.form["Answer_fT"])
            ans_fTor = float(request.form["Answer_fTor"])
        except:
            flash("Invalid input", "danger")
            return redirect(url_for("student_dashboard"))

        correct_L = session["fL"]
        correct_T = session["fT"]
        correct_Tor = session["fTor"]

        tol = 0.05
        okL = correct_L * (1 - tol) <= ans_fL <= correct_L * (1 + tol)
        okT = correct_T * (1 - tol) <= ans_fT <= correct_T * (1 + tol)
        okTor = correct_Tor * (1 - tol) <= ans_fTor <= correct_Tor * (1 + tol)

        # Calculate marks: 5 marks per correct answer (total 15)
        marks = 0
        if okL:
            marks += 5
        if okT:
            marks += 5
        if okTor:
            marks += 5

        # If no attempt or all wrong, give 2 marks for attempt
        if marks == 0:
            marks = 2

        # Generate feedback based on marks
        if marks == 15:
            feedback = "✅ All answers correct!"
            # Check if this is first attempt with perfect score
            is_perfect_first_attempt = (user.p8_attempts == 0 and marks == 15)
        elif marks > 2:
            feedback = f"⚠️ Partially correct ({marks}/15 marks). Correct: fL={correct_L}, fT={correct_T}, fTor={correct_Tor}"
            is_perfect_first_attempt = False
        else:
            feedback = f"❌ Incorrect. Correct: fL={correct_L}, fT={correct_T}, fTor={correct_Tor}"
            is_perfect_first_attempt = False

        if user:
            if not update_score(user, 8, marks):
                flash("❌ Maximum attempts (2) exceeded for this problem!", "danger")
                # Redirect based on user role
                if user.role == "admin":
                    return redirect(url_for("admin_assignment_view"))
                else:
                    return redirect(url_for("student_dashboard"))

        # ✅ Generate solution steps
        steps = build_steps_problem8(session)

    return render_template("problem8.html", vars=session, feedback=feedback, steps=steps, user=user,
                           show_celebration=is_perfect_first_attempt if 'is_perfect_first_attempt' in locals() else False)


def build_steps_problem8(vars):
    steps = []
    d, L, M, k, E, G = vars["d"], vars["L"], vars["M"], vars["k"], vars["E"], vars["G"]

    # Step 1: Given values
    steps.append(
        f"Given: Shaft diameter d = {d*1000:.1f} mm, Length L = {L*1000:.1f} mm")
    steps.append(
        f"Disc mass M = {M} kg, Radius of gyration k = {k*1000:.1f} mm")
    steps.append(
        f"Young's Modulus E = {E/1e9:.0f} GN/m², Shear Modulus G = {G/1e9:.0f} GN/m²")

    # Step 2: Section properties
    A = math.pi * d**2 / 4
    I = math.pi * d**4 / 64
    J = math.pi * d**4 / 32
    steps.append(f"Cross-sectional area A = πd²/4 = {A:.6f} m²")
    steps.append(f"Second moment of area I = πd⁴/64 = {I:.8e} m⁴")
    steps.append(f"Polar moment of inertia J = πd⁴/32 = {J:.8e} m⁴")

    # Step 3: Longitudinal vibration
    kL = A * E / L
    fL = (1 / (2 * math.pi)) * math.sqrt(kL / M)
    steps.append(f"Longitudinal stiffness kL = AE/L = {kL:.2e} N/m")
    steps.append(
        f"Longitudinal frequency fL = (1/2π)√(kL/M) = {round(fL, 2)} Hz")

    # Step 4: Transverse vibration
    W = M * 9.81
    delta = (W * L**3) / (3 * E * I)
    fT = 0.4985 / math.sqrt(delta)
    steps.append(f"Weight W = Mg = {M} × 9.81 = {W:.2f} N")
    steps.append(f"Static deflection δ = WL³/(3EI) = {delta:.6e} m")
    steps.append(f"Transverse frequency fT = 0.4985/√δ = {round(fT, 2)} Hz")

    # Step 5: Torsional vibration
    kT = G * J / L
    Ip = M * k**2
    fTor = (1 / (2 * math.pi)) * math.sqrt(kT / Ip)
    steps.append(f"Torsional stiffness kT = GJ/L = {kT:.2e} Nm/rad")
    steps.append(
        f"Polar moment of inertia of disc Ip = Mk² = {M} × {k**2:.6f} = {Ip:.2f} kg·m²")
    steps.append(
        f"Torsional frequency fTor = (1/2π)√(kT/Ip) = {round(fTor, 2)} Hz")

    return steps


# from yourapp import app, User, update_score  # adjust imports to your project


def generate_problem9():
    """Generate new problem 9 values and store correct answers in session"""
    session["p9_speed_rpm"] = random.choice([240, random.randint(200, 400)])
    session["p9_m_rec"] = random.randint(30, 80)
    session["p9_m_rev"] = random.randint(20, 60)
    session["p9_rev_radius_mm"] = random.randint(100, 200)
    session["p9_stroke_mm"] = 2 * session["p9_rev_radius_mm"]
    session["p9_balance_radius_mm"] = random.randint(300, 500)

    # Convert & compute
    stroke_m = session["p9_stroke_mm"] / 1000.0
    crank_radius_m = stroke_m / 2.0
    rev_radius_m = session["p9_rev_radius_mm"] / 1000.0
    balance_radius_m = session["p9_balance_radius_mm"] / 1000.0
    m_rec = session["p9_m_rec"]
    m_rev = session["p9_m_rev"]
    N = session["p9_speed_rpm"]
    omega = 2 * math.pi * N / 60.0

    # 1) Balance mass
    numerator = m_rev * rev_radius_m + (2.0/3.0 * m_rec) * crank_radius_m
    m_b = numerator / balance_radius_m

    # 2) Residual force at θ = 60°
    c = 2.0 / 3.0
    theta_rad = math.radians(60)
    residual_force = (
        m_rec * crank_radius_m * omega**2 *
        math.sqrt(((1 - c)**2) * (math.cos(theta_rad)**2) +
                  (c**2) * (math.sin(theta_rad)**2))
    )

    # Store in session (rounded)
    session["p9_answer_mb"] = round(m_b, 3)
    session["p9_answer_residual"] = round(abs(residual_force), 2)


def build_steps_problem9(vars):
    """Build solution steps for Problem 9"""
    steps = []

    # Extract values from session/dict
    N = vars["p9_speed_rpm"]
    stroke_mm = vars["p9_stroke_mm"]
    m_rec = vars["p9_m_rec"]
    m_rev = vars["p9_m_rev"]
    rev_radius_mm = vars["p9_rev_radius_mm"]
    balance_radius_mm = vars["p9_balance_radius_mm"]

    # Conversions
    stroke_m = stroke_mm / 1000.0
    crank_radius_m = stroke_m / 2.0
    rev_radius_m = rev_radius_mm / 1000.0
    balance_radius_m = balance_radius_mm / 1000.0
    omega = 2 * math.pi * N / 60.0

    # Step 1: Given values and conversions
    steps.append(f"Given: Speed = {N} rpm, Stroke = {stroke_mm} mm")
    steps.append(
        f"Reciprocating mass = {m_rec} kg, Revolving mass = {m_rev} kg at {rev_radius_mm} mm"
    )
    steps.append(f"Balancing mass location = {balance_radius_mm} mm radius")
    steps.append(f"Angular velocity ω = 2π × {N}/60 = {omega:.3f} rad/s")
    steps.append(
        f"Crank radius r = stroke/2 = {stroke_mm}/2 = {crank_radius_m*1000:.1f} mm"
    )

    # Step 2: Balance mass calculation
    numerator = m_rev * rev_radius_m + (2.0 / 3.0 * m_rec) * crank_radius_m
    m_b = numerator / balance_radius_m
    steps.append("Balance mass calculation:")
    steps.append("m_b × R_b = m_rev × r_rev + (2/3 × m_rec) × r_crank")
    steps.append(
        f"m_b × {balance_radius_m:.3f} = {m_rev} × {rev_radius_m:.3f} + (2/3 × {m_rec}) × {crank_radius_m:.3f}"
    )
    steps.append(
        f"m_b = {numerator:.6f} / {balance_radius_m:.3f} = {m_b:.3f} kg")

    # Step 3: Residual force calculation
    c = 2.0 / 3.0  # fraction of reciprocating mass balanced
    theta_rad = math.radians(60)

    residual_force = (
        m_rec
        * omega**2
        * crank_radius_m
        * math.sqrt(
            ((1 - c) ** 2) * (math.cos(theta_rad) ** 2)
            + (c**2) * (math.sin(theta_rad) ** 2)
        )
    )

    steps.append("Residual unbalanced force at θ = 60°:")
    steps.append(
        "F_residual = m_rec × r_crank × ω² × sqrt((1-c)² cos²θ + c² sin²θ)")
    steps.append(
        f"F_residual = {m_rec} × {crank_radius_m:.3f} × {omega:.3f}² × "
        f"sqrt((1-{c:.2f})² cos²(60°) + {c:.2f}² sin²(60°))"
    )
    steps.append(f"F_residual = {residual_force:.2f} N")

    return steps


@app.route("/problem9", methods=["GET", "POST"])
def problem9():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    # Clear any lingering flash messages from previous admin actions when accessing problems
    # This prevents admin dashboard messages from appearing on problem pages
    if request.method == "GET" and request.referrer and ("admin" in request.referrer or "dashboard" in request.referrer):
        session.pop('_flashes', None)

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    feedback = None
    steps = None

    # Detect new attempt
    current_attempts = user.p9_attempts if user else 0
    if "p9_current_attempt" not in session or session["p9_current_attempt"] != current_attempts:
        # Clear old session values
        for key in [
            "p9_speed_rpm", "p9_stroke_mm", "p9_m_rec", "p9_m_rev",
            "p9_rev_radius_mm", "p9_balance_radius_mm",
            "p9_answer_mb", "p9_answer_residual"
        ]:
            session.pop(key, None)
        session["p9_current_attempt"] = current_attempts
        generate_problem9()  # ✅ immediately generate new problem

    # --- Handle Answer Submission ---
    if request.method == "POST":
        try:
            student_mb = float(request.form.get("balance_mass", "").strip())
            student_residual = float(
                request.form.get("residual_force", "").strip())
        except Exception:
            flash("Invalid numeric input. Please enter numbers.", "danger")
            return redirect(url_for("problem9"))

        correct_mb = session.get("p9_answer_mb")
        correct_res = session.get("p9_answer_residual")

        if correct_mb is None or correct_res is None:
            flash("Session expired. Please refresh.", "warning")
            return redirect(url_for("problem9"))

        tol = 0.05
        mb_ok = (correct_mb * (1 - tol) <=
                 student_mb <= correct_mb * (1 + tol))
        res_ok = (correct_res * (1 - tol) <=
                  student_residual <= correct_res * (1 + tol))

        marks = 0
        if mb_ok:
            marks += 5
        if res_ok:
            marks += 5
        if marks == 0:
            marks = 2

        if marks == 10:
            feedback = "✅ Both answers correct!"
            is_perfect_first_attempt = (user.p9_attempts == 0 and marks == 10)
        elif marks > 2:
            feedback = f"⚠️ Partially correct ({marks}/10 marks). Correct: Balance mass ≈ {correct_mb} kg, Residual ≈ {correct_res} N"
            is_perfect_first_attempt = False
        else:
            feedback = f"❌ Both incorrect. Correct: Balance mass ≈ {correct_mb} kg, Residual ≈ {correct_res} N"
            is_perfect_first_attempt = False

        if user:
            if not update_score(user, 9, marks):
                flash("❌ Maximum attempts (2) exceeded for this problem! No more marks will be awarded, but you can still view the solution.", "danger")
                marks = 0  # don’t award new marks

        steps = build_steps_problem9(session)
        return render_template(
            "problem9.html",
            vars=session,
            feedback=feedback,
            steps=steps,
            user=user,
            show_celebration=is_perfect_first_attempt
        )

    # --- First GET: generate problem if missing ---
    if "p9_answer_mb" not in session:
        generate_problem9()

    return render_template(
        "problem9.html",
        vars=session,
        feedback=feedback,
        steps=steps,
        user=user,
        show_celebration=False
    )


@app.route("/problem10", methods=["GET", "POST"])
def problem10():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    # Clear any lingering flash messages from previous admin actions when accessing problems
    # This prevents admin dashboard messages from appearing on problem pages
    if request.method == "GET" and request.referrer and ("admin" in request.referrer or "dashboard" in request.referrer):
        session.pop('_flashes', None)

    # Load user DB record
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()

    # Initialize variables
    feedback = None
    steps = None

    # Check if this is a new attempt (clear session for new problem)
    current_attempts = user.p10_attempts if user else 0
    if "p10_current_attempt" not in session or session["p10_current_attempt"] != current_attempts:
        # Clear problem session for new attempt
        for key in ["p10_m", "p10_k", "p10_c", "p10_c_cr", "p10_zeta", "p10_delta", "p10_ratio"]:
            session.pop(key, None)
        session["p10_current_attempt"] = current_attempts

    # --- Handle Answer Submission (POST) ---
    if request.method == "POST":
        # Parse student entries
        try:
            student_c_cr = float(request.form.get("c_cr", "").strip())
            student_zeta = float(request.form.get("zeta", "").strip())
            student_delta = float(request.form.get("delta", "").strip())
            student_ratio = float(request.form.get("ratio", "").strip())
        except Exception:
            flash("Invalid numeric input. Please enter numbers.", "danger")
            return redirect(url_for("problem10"))

        correct_c_cr = session.get("p10_c_cr")
        correct_zeta = session.get("p10_zeta")
        correct_delta = session.get("p10_delta")
        correct_ratio = session.get("p10_ratio")

        if any(x is None for x in [correct_c_cr, correct_zeta, correct_delta, correct_ratio]):
            flash("Session expired. Please refresh.", "warning")
            return redirect(url_for("problem10"))

        tol = 0.05  # 5% tolerance (consistent with other problems)

        c_cr_ok = abs(student_c_cr - correct_c_cr) <= tol * correct_c_cr
        zeta_ok = abs(student_zeta - correct_zeta) <= tol * correct_zeta
        delta_ok = abs(student_delta - correct_delta) <= tol * correct_delta
        ratio_ok = abs(student_ratio - correct_ratio) <= tol * correct_ratio

        # Calculate marks: 2.5 marks per correct answer (total 10)
        marks = 0
        if c_cr_ok:
            marks += 2.5
        if zeta_ok:
            marks += 2.5
        if delta_ok:
            marks += 2.5
        if ratio_ok:
            marks += 2.5

        # If no attempt or all wrong, give 2 marks for attempt
        if marks == 0:
            marks = 2

        # Generate detailed feedback based on marks
        if marks == 10:
            feedback = "✅ All answers correct!"
            # Check if this is first attempt with perfect score
            is_perfect_first_attempt = (user.p10_attempts == 0 and marks == 10)
        elif marks > 2:
            # Detailed partial feedback
            correct_answers = []
            incorrect_answers = []

            correct_c_cr = session.get("p10_c_cr", 0)
            correct_zeta = session.get("p10_zeta", 0)
            correct_delta = session.get("p10_delta", 0)
            correct_ratio = session.get("p10_ratio", 0)

            if c_cr_ok:
                correct_answers.append("c_cr")
            else:
                incorrect_answers.append(
                    f"c_cr (correct: {correct_c_cr} N·s/m)")

            if zeta_ok:
                correct_answers.append("ζ")
            else:
                incorrect_answers.append(f"ζ (correct: {correct_zeta})")

            if delta_ok:
                correct_answers.append("δ")
            else:
                incorrect_answers.append(f"δ (correct: {correct_delta})")

            if ratio_ok:
                correct_answers.append("Ratio")
            else:
                incorrect_answers.append(f"Ratio (correct: {correct_ratio})")

            feedback = f"⚠️ Partially correct ({marks}/10 marks). "
            if correct_answers:
                feedback += f"Correct: {', '.join(correct_answers)}. "
            if incorrect_answers:
                feedback += f"Incorrect: {', '.join(incorrect_answers)}."

            is_perfect_first_attempt = False
        else:
            # Show correct answers when all are wrong
            correct_c_cr = session.get("p10_c_cr", 0)
            correct_zeta = session.get("p10_zeta", 0)
            correct_delta = session.get("p10_delta", 0)
            correct_ratio = session.get("p10_ratio", 0)
            feedback = f"❌ All incorrect. Correct answers: c_cr = {correct_c_cr} N·s/m, ζ = {correct_zeta}, δ = {correct_delta}, Ratio = {correct_ratio}"
            is_perfect_first_attempt = False

        if user:
            if not update_score(user, 10, marks):
                flash("❌ Maximum attempts (2) exceeded for this problem!", "danger")
                # Redirect based on user role
                if user.role == "admin":
                    return redirect(url_for("admin_assignment_view"))
                else:
                    return redirect(url_for("student_dashboard"))

        # ✅ Generate solution steps
        vars = {"m": session["p10_m"],
                "k": session["p10_k"], "c": session["p10_c"]}
        steps, _, _, _, _ = build_steps_problem10(vars)

        # Render template with feedback and steps (don't redirect)
        return render_template("problem10.html", vars=session, feedback=feedback, steps=steps, user=user,
                               show_celebration=is_perfect_first_attempt)

    # --- Generate new random values only on GET request (first time) ---
    # Only generate new problem if not already exists (preserve during POST)
    if "p10_c_cr" not in session:
        # Generate random values
        session["p10_m"] = random.randint(5, 15)  # kg
        session["p10_k"] = random.randint(3000, 8000)  # N/m
        session["p10_c"] = random.randint(20, 80)  # N·s/m

        # Calculate correct answers
        m = session["p10_m"]
        k = session["p10_k"]
        c = session["p10_c"]

        # Critical damping coefficient
        c_cr = 2 * math.sqrt(m * k)
        # Damping factor
        zeta = c / c_cr
        # Logarithmic decrement
        delta = (2 * math.pi * zeta) / math.sqrt(1 - zeta**2)
        # Ratio of consecutive amplitudes
        ratio = math.exp(delta)

        # Store correct answers in session
        session["p10_c_cr"] = round(c_cr, 3)
        session["p10_zeta"] = round(zeta, 4)
        session["p10_delta"] = round(delta, 4)
        session["p10_ratio"] = round(ratio, 4)

    return render_template("problem10.html", vars=session, feedback=feedback, steps=steps, user=user,
                           show_celebration=False)


def build_steps_problem10(vars):
    m = vars['m']
    k = vars['k']
    c = vars['c']

    steps = []

    # Step 1: Critical damping coefficient
    steps.append(f"1. Critical damping coefficient formula: c_cr = 2√(m·k)")
    c_cr = 2 * math.sqrt(m * k)
    steps.append(f"   Substituting: c_cr = 2√({m} × {k}) = {c_cr:.3f} N·s/m")

    # Step 2: Damping factor
    steps.append("2. Damping factor formula: ζ = c / c_cr")
    zeta = c / c_cr
    steps.append(f"   Substituting: ζ = {c} / {c_cr:.3f} = {zeta:.4f}")

    # Step 3: Logarithmic decrement
    steps.append("3. Logarithmic decrement formula: δ = (2πζ)/√(1-ζ²)")
    delta = (2 * math.pi * zeta) / math.sqrt(1 - zeta**2)
    steps.append(
        f"   Substituting: δ = (2π×{zeta:.4f})/√(1-{zeta:.4f}²) = {delta:.4f}")

    # Step 4: Ratio of consecutive amplitudes
    steps.append("4. Ratio of consecutive amplitudes (A_n / A_(n+1)) = e^δ")
    ratio = math.exp(delta)
    steps.append(f"   Substituting: Ratio = e^{delta:.4f} = {ratio:.4f}")

    return steps, c_cr, zeta, delta, ratio


# ---------------- Problem 11 ----------------

@app.route("/problem11", methods=["GET", "POST"])
def problem11():
    has_access, result = check_problem_access(11)
    if not has_access:
        return result

    user = result

    # Clear flash messages if redirected from admin/dashboard
    if request.method == "GET" and request.referrer and ("admin" in request.referrer or "dashboard" in request.referrer):
        session.pop('_flashes', None)

    # -------------------- POST (Submission) --------------------
    if request.method == "POST":
        try:
            submitted_mb = round(float(request.form.get("mb", 0)), 3)
            submitted_theta_b = round(float(request.form.get("theta_b", 0)), 3)

            correct_mb = session.get("p11_mb", 0)
            correct_theta_b = session.get("p11_theta_b", 0)

            feedback = ""
            score = 0

            mb_correct = abs(submitted_mb - correct_mb) <= 0.05 * correct_mb
            theta_correct = abs(submitted_theta_b - correct_theta_b) <= 1.0

            if mb_correct and theta_correct:
                score = 10
                feedback = "✅ Excellent! Both balance mass and position are correct."
            elif mb_correct:
                score = 7
                feedback = f"⚠️ Balance mass is correct, but position is incorrect. Correct position ≈ {correct_theta_b:.1f}°"
            elif theta_correct:
                score = 7
                feedback = f"⚠️ Position is correct, but balance mass is incorrect. Correct balance mass ≈ {correct_mb:.3f} kg"
            else:
                score = 2
                feedback = f"❌ Both answers are incorrect. Correct: Balance mass ≈ {correct_mb:.3f} kg, Position ≈ {correct_theta_b:.1f}°"

            success = update_score(user, 11, score)
            if not success:
                flash("You have already attempted this problem 2 times!", "warning")
                return redirect(url_for("problem11"))

            # ✅ Save submitted data and flag
            session["p11_show_results"] = True
            session["p11_feedback"] = feedback
            session["p11_submitted_mb"] = submitted_mb
            session["p11_submitted_theta_b"] = submitted_theta_b
            # next refresh → generate new problem
            session["p11_generate_new"] = True

            return redirect(url_for("problem11"))

        except (ValueError, TypeError):
            flash("Please enter valid numerical values!", "danger")
            return redirect(url_for("problem11"))

    # -------------------- GET (Display) --------------------
    # Step 1: If results need to be shown (just after submission)
    if session.get("p11_show_results"):
        vars = {
            "m": session["p11_m"],
            "r": session["p11_r"],
            "angles": session["p11_angles"],
            "rb": session["p11_rb"]
        }
        M = session["p11_M"]
        theta = session["p11_theta"]
        Sx = session["p11_Sx"]
        Sy = session["p11_Sy"]
        S = session["p11_S"]
        thetaS = session["p11_thetaS"]
        mb = session["p11_mb"]
        theta_b = session["p11_theta_b"]

        steps, _, _, _, _, _, _, _, _ = build_steps_problem11(vars)

        feedback = session.pop("p11_feedback", "")
        submitted_mb = session.pop("p11_submitted_mb", None)
        submitted_theta_b = session.pop("p11_submitted_theta_b", None)
        session.pop("p11_show_results", None)  # clear flag

        # ✅ Keep show_inputs=False (hide form)
        return render_template(
            "problem11.html",
            vars=vars,
            M=M, theta=theta,
            Sx=Sx, Sy=Sy,
            S=S, thetaS=thetaS,
            mb=mb, theta_b=theta_b,
            steps=steps,
            user=user,
            feedback=feedback,
            submitted_mb=submitted_mb,
            submitted_theta_b=submitted_theta_b,
            show_inputs=False,
            show_celebration=False
        )

    # Step 2: If flagged for new problem
    if session.pop("p11_generate_new", False):
        # Generate new random problem now
        m = [random.randint(150, 300) for _ in range(4)]
        r = [round(random.uniform(0.15, 0.35), 2) for _ in range(4)]
        rb = round(random.uniform(0.15, 0.25), 2)
        angles = [random.choice([30, 45, 60, 75, 90, 120, 135])
                  for _ in range(3)]

        vars = {"m": m, "r": r, "angles": angles, "rb": rb}
        steps, M, theta, Sx, Sy, S, thetaS, mb, theta_b = build_steps_problem11(
            vars)

        # Store new problem
        session["p11_m"] = m
        session["p11_r"] = r
        session["p11_angles"] = angles
        session["p11_rb"] = rb
        session["p11_M"] = M
        session["p11_theta"] = theta
        session["p11_Sx"] = Sx
        session["p11_Sy"] = Sy
        session["p11_S"] = S
        session["p11_thetaS"] = thetaS
        session["p11_mb"] = mb
        session["p11_theta_b"] = theta_b

        return render_template(
            "problem11.html",
            vars=vars,
            M=M, theta=theta,
            Sx=Sx, Sy=Sy,
            S=S, thetaS=thetaS,
            mb=mb, theta_b=theta_b,
            steps=steps,
            user=user,
            feedback="",
            show_inputs=True,
            show_celebration=False
        )

    # Step 3: If neither results nor flag set (normal first load)
    # Use existing session if available, else generate new
    if all(k in session for k in ["p11_m", "p11_r", "p11_angles", "p11_rb"]):
        vars = {
            "m": session["p11_m"],
            "r": session["p11_r"],
            "angles": session["p11_angles"],
            "rb": session["p11_rb"]
        }
        M = session["p11_M"]
        theta = session["p11_theta"]
        Sx = session["p11_Sx"]
        Sy = session["p11_Sy"]
        S = session["p11_S"]
        thetaS = session["p11_thetaS"]
        mb = session["p11_mb"]
        theta_b = session["p11_theta_b"]
        steps, _, _, _, _, _, _, _, _ = build_steps_problem11(vars)
    else:
        # First visit — generate new problem
        m = [random.randint(150, 300) for _ in range(4)]
        r = [round(random.uniform(0.15, 0.35), 2) for _ in range(4)]
        rb = round(random.uniform(0.15, 0.25), 2)
        angles = [random.choice([30, 45, 60, 75, 90, 120, 135])
                  for _ in range(3)]
        vars = {"m": m, "r": r, "angles": angles, "rb": rb}
        steps, M, theta, Sx, Sy, S, thetaS, mb, theta_b = build_steps_problem11(
            vars)

        session["p11_m"] = m
        session["p11_r"] = r
        session["p11_angles"] = angles
        session["p11_rb"] = rb
        session["p11_M"] = M
        session["p11_theta"] = theta
        session["p11_Sx"] = Sx
        session["p11_Sy"] = Sy
        session["p11_S"] = S
        session["p11_thetaS"] = thetaS
        session["p11_mb"] = mb
        session["p11_theta_b"] = theta_b

    return render_template(
        "problem11.html",
        vars=vars,
        M=M, theta=theta,
        Sx=Sx, Sy=Sy,
        S=S, thetaS=thetaS,
        mb=mb, theta_b=theta_b,
        steps=steps,
        user=user,
        feedback="",
        show_inputs=True,
        show_celebration=False
    )


def initialize_problem_visibility():
    """Initialize problem visibility records for all problems 1-10"""
    try:
        with app.app_context():
            # Create tables if they don't exist
            db.create_all()

            # Initialize problem visibility for problems 1-10
            problems_initialized = 0
            for problem_num in range(1, 12):
                existing = ProblemVisibility.query.filter_by(
                    problem_number=problem_num, admin_id=None).first()

                if not existing:
                    # Create global visibility setting (default: released for problems 1-9, locked for 10)
                    # Problems 1-9 released by default, Problem 10 locked
                    is_released = problem_num <= 9
                    visibility = ProblemVisibility(
                        problem_number=problem_num,
                        is_released=is_released,
                        admin_id=None  # Global setting
                    )
                    db.session.add(visibility)
                    problems_initialized += 1

            if problems_initialized > 0:
                db.session.commit()

            # Initialize CO mappings
            initialize_co_mappings()

    except Exception as e:
        try:
            db.session.rollback()
        except:
            pass


# --- Quiz Management Routes ---

@app.route("/quiz_question_management")
def quiz_question_management():

    # Clear any old flash messages to avoid message pile-up
    from flask import get_flashed_messages
    get_flashed_messages()

    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Access denied. Only admins can manage quiz questions.", "error")
        return redirect(url_for("admin_dashboard"))

    allow_admin_add_question = get_config('allow_admin_add_question', '0') == '1'
    quiz_banks = get_admin_quiz_banks(user)
    selected_quiz_name = request.args.get("quiz_name") or (quiz_banks[0].quiz_name if quiz_banks else DEFAULT_QUIZ_NAME)
    selected_quiz = get_quiz_bank_by_name(selected_quiz_name)

    if selected_quiz and not can_admin_access_quiz_bank(user, selected_quiz):
        flash("You do not have access to that quiz.", "error")
        return redirect(url_for("quiz_question_management"))

    questions = get_quiz_questions(selected_quiz)
    question_counts = get_quiz_question_counts(selected_quiz)
    total_questions = sum(question_counts.values())
    quiz_release_status = get_quiz_release_status_for_admin(
        user, selected_quiz.quiz_name if selected_quiz else DEFAULT_QUIZ_NAME)
    all_admins = User.query.filter(User.admin_level.in_(['admin', 'super_admin'])).order_by(User.rollnumber.asc()).all()
    shared_admin_ids = {
        row.admin_id for row in QuizAdminAccess.query.filter_by(quiz_id=selected_quiz.id).all()
    } if selected_quiz else set()

    return render_template(
        "quiz_question_management.html",
        questions=questions,
        question_counts=question_counts,
        total_questions=total_questions,
        configured_questions_per_attempt=get_quiz_questions_per_attempt(
            selected_quiz, available_count=total_questions
        ) if selected_quiz else 0,
        user=user,
        allow_admin_add_question=allow_admin_add_question,
        quiz_banks=quiz_banks,
        selected_quiz=selected_quiz,
        quiz_release_status=quiz_release_status,
        all_admins=all_admins,
        shared_admin_ids=shared_admin_ids,
        can_edit_selected_quiz=can_admin_edit_quiz_bank(user, selected_quiz),
        quiz_co_options=QUIZ_CO_OPTIONS
    )

# Route for super admin to toggle add question for admins


@app.route("/toggle_admin_add_question")
def toggle_admin_add_question():
    if "rollnumber" not in session:
        return redirect(url_for("login"))
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.admin_level != 'super_admin':
        flash("Access denied. Only super admins can toggle this setting.", "error")
        return redirect(url_for("admin_dashboard"))
    current = get_config('allow_admin_add_question', '0')
    set_config('allow_admin_add_question', '0' if current == '1' else '1')
    return redirect(url_for("admin_dashboard"))


@app.route("/create_quiz_bank", methods=["POST"])
def create_quiz_bank():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Access denied.", "error")
        return redirect(url_for("admin_dashboard"))

    quiz_name = normalize_quiz_name(request.form.get("quiz_name"))
    quiz_title = (request.form.get("quiz_title") or "").strip()
    if not quiz_title:
        quiz_title = quiz_name.replace('_', ' ').title()

    if QuizBank.query.filter_by(quiz_name=quiz_name).first():
        flash("Quiz name already exists. Please use a different quiz name.", "error")
        return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

    quiz_bank = QuizBank(
        quiz_name=quiz_name,
        quiz_title=quiz_title,
        description=(request.form.get("description") or "").strip() or None,
        covered_cos=parse_quiz_cos(request.form.getlist("covered_cos")),
        questions_per_attempt=max(1, int(request.form.get("questions_per_attempt", 10) or 10)),
        is_shared_with_all_admins=bool(request.form.get("is_shared_with_all_admins")) if user.admin_level == 'super_admin' else False,
        created_by=user.id
    )
    db.session.add(quiz_bank)
    db.session.commit()

    flash(f"Quiz '{quiz_bank.quiz_title}' created successfully.", "success")
    return redirect(url_for("quiz_question_management", quiz_name=quiz_bank.quiz_name))


@app.route("/update_quiz_bank/<quiz_name>", methods=["POST"])
def update_quiz_bank(quiz_name):
    if "rollnumber" not in session:
        return redirect(url_for("login"))
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    quiz_bank = get_quiz_bank_by_name(quiz_name)
    if not user or not can_admin_edit_quiz_bank(user, quiz_bank):
        flash("Access denied.", "error")
        return redirect(url_for("quiz_question_management"))

    quiz_bank.quiz_title = (request.form.get("quiz_title") or quiz_bank.quiz_title).strip()
    quiz_bank.description = (request.form.get("description") or "").strip() or None
    quiz_bank.covered_cos = parse_quiz_cos(request.form.getlist("covered_cos"))
    quiz_bank.questions_per_attempt = max(1, int(request.form.get("questions_per_attempt", quiz_bank.questions_per_attempt or 10) or 10))
    if user.admin_level == 'super_admin':
        quiz_bank.is_shared_with_all_admins = bool(request.form.get("is_shared_with_all_admins"))
    db.session.commit()
    flash("Quiz details updated successfully.", "success")
    return redirect(url_for("quiz_question_management", quiz_name=quiz_bank.quiz_name))


@app.route("/delete_quiz_bank/<quiz_name>", methods=["POST"])
def delete_quiz_bank(quiz_name):
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    quiz_bank = get_quiz_bank_by_name(quiz_name)
    if not user or not can_admin_edit_quiz_bank(user, quiz_bank):
        flash("Access denied. You cannot delete this quiz.", "error")
        return redirect(url_for("quiz_question_management"))

    try:
        assigned_question_ids = [
            row.question_id for row in QuizQuestionAssignment.query.filter_by(quiz_id=quiz_bank.id).all()
        ]

        attempt_ids = [
            row.id for row in QuizAttempt.query.filter_by(quiz_name=quiz_bank.quiz_name).all()
        ]
        if attempt_ids:
            QuizResponse.query.filter(QuizResponse.attempt_id.in_(attempt_ids)).delete(
                synchronize_session=False
            )
            QuizAttempt.query.filter(QuizAttempt.id.in_(attempt_ids)).delete(
                synchronize_session=False
            )

        QuizVisibility.query.filter_by(quiz_name=quiz_bank.quiz_name).delete()
        QuizAdminAccess.query.filter_by(quiz_id=quiz_bank.id).delete()
        QuizQuestionAssignment.query.filter_by(quiz_id=quiz_bank.id).delete()

        for question_id in assigned_question_ids:
            if not QuizQuestionAssignment.query.filter_by(question_id=question_id).first():
                QuizQuestion.query.filter_by(id=question_id).delete()

        quiz_title = quiz_bank.quiz_title
        db.session.delete(quiz_bank)
        db.session.commit()
        flash(f"Quiz '{quiz_title}' deleted successfully.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting quiz: {str(e)}", "error")

    return redirect(url_for("quiz_question_management"))


@app.route("/update_quiz_admin_access/<quiz_name>", methods=["POST"])
def update_quiz_admin_access(quiz_name):
    if "rollnumber" not in session:
        return redirect(url_for("login"))
    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    quiz_bank = get_quiz_bank_by_name(quiz_name)
    if not user or user.admin_level != 'super_admin' or not quiz_bank:
        flash("Only the main admin can change quiz sharing.", "error")
        return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

    selected_admin_ids = {
        int(admin_id) for admin_id in request.form.getlist("allowed_admin_ids") if str(admin_id).isdigit()
    }
    selected_admin_ids.discard(quiz_bank.created_by)

    QuizAdminAccess.query.filter_by(quiz_id=quiz_bank.id).delete()
    for admin_id in selected_admin_ids:
        db.session.add(QuizAdminAccess(
            quiz_id=quiz_bank.id,
            admin_id=admin_id,
            granted_by=user.id
        ))
    db.session.commit()

    flash("Quiz access updated successfully.", "success")
    return redirect(url_for("quiz_question_management", quiz_name=quiz_bank.quiz_name))


@app.route("/download_quiz_bulk_template")
def download_quiz_bulk_template():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    sample_quiz_name = request.args.get("quiz_name") or DEFAULT_QUIZ_NAME
    output = StringIO()
    fieldnames = [
        "quiz_name", "quiz_title", "question_text", "co_number", "points",
        "answer_type", "option_a", "option_b", "option_c", "option_d",
        "correct_answers", "explanation"
    ]
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    writer.writerow({
        "quiz_name": sample_quiz_name,
        "quiz_title": "Sample Quiz",
        "question_text": "Which law explains the relationship between stress and strain in the elastic region?",
        "co_number": "CO1",
        "points": 2,
        "answer_type": "single",
        "option_a": "Hooke's Law",
        "option_b": "Bernoulli's Principle",
        "option_c": "Newton's Second Law",
        "option_d": "Pascal's Law",
        "correct_answers": "A",
        "explanation": "Hooke's Law applies in the elastic region."
    })
    writer.writerow({
        "quiz_name": sample_quiz_name,
        "quiz_title": "Sample Quiz",
        "question_text": "Select the valid design considerations for shafts.",
        "co_number": "CO2",
        "points": 3,
        "answer_type": "multiple",
        "option_a": "Strength",
        "option_b": "Rigidity",
        "option_c": "Resonance check",
        "option_d": "Aesthetics only",
        "correct_answers": "A,B,C",
        "explanation": "Strength, rigidity and vibration checks are relevant."
    })

    response = make_response(output.getvalue())
    response.headers["Content-Type"] = "text/csv"
    response.headers["Content-Disposition"] = "attachment; filename=dom_quiz_bulk_template.csv"
    return response


@app.route("/download_quiz_bulk_export")
def download_quiz_bulk_export():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    quiz_name = request.args.get("quiz_name") or DEFAULT_QUIZ_NAME
    quiz_bank = get_quiz_bank_by_name(quiz_name)
    if not user or not can_admin_access_quiz_bank(user, quiz_bank):
        flash("Access denied.", "error")
        return redirect(url_for("quiz_question_management"))

    output = StringIO()
    fieldnames = [
        "quiz_name", "quiz_title", "question_text", "co_number", "points",
        "answer_type", "option_a", "option_b", "option_c", "option_d",
        "correct_answers", "explanation"
    ]
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()

    for question in get_quiz_questions(quiz_bank):
        choices = question.choices or []
        choice_map = {}
        for index, choice in enumerate(choices):
            letter = ['A', 'B', 'C', 'D'][index]
            choice_map[letter] = choice.get("text") if isinstance(choice, dict) else choice

        writer.writerow({
            "quiz_name": quiz_bank.quiz_name,
            "quiz_title": quiz_bank.quiz_title,
            "question_text": question.question_text,
            "co_number": question.co_number,
            "points": question.points,
            "answer_type": question.answer_type,
            "option_a": choice_map.get("A", ""),
            "option_b": choice_map.get("B", ""),
            "option_c": choice_map.get("C", ""),
            "option_d": choice_map.get("D", ""),
            "correct_answers": ",".join(question.correct_answers or []),
            "explanation": question.explanation or ""
        })

    response = make_response(output.getvalue())
    response.headers["Content-Type"] = "text/csv"
    response.headers["Content-Disposition"] = f"attachment; filename={quiz_bank.quiz_name}_questions.csv"
    return response


@app.route("/upload_quiz_bulk_csv", methods=["POST"])
def upload_quiz_bulk_csv():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    allow_admin_add_question = get_config('allow_admin_add_question', '0') == '1'
    if not user or (user.admin_level != 'super_admin' and not (user.admin_level == 'admin' and allow_admin_add_question)):
        flash("Access denied.", "error")
        return redirect(url_for("admin_dashboard"))

    file = request.files.get("csv_file")
    selected_quiz_name = request.form.get("selected_quiz_name")
    if not file or not file.filename.lower().endswith(".csv"):
        flash("Please upload a valid CSV file.", "error")
        return redirect(url_for("quiz_question_management", quiz_name=selected_quiz_name or DEFAULT_QUIZ_NAME))

    csv_reader = csv.DictReader(StringIO(file.read().decode("utf-8-sig")))
    required_columns = [
        "question_text", "co_number", "points", "answer_type",
        "option_a", "option_b", "option_c", "option_d", "correct_answers"
    ]
    if not csv_reader.fieldnames or not all(column in csv_reader.fieldnames for column in required_columns):
        flash("CSV format is invalid. Please use the downloaded template.", "error")
        return redirect(url_for("quiz_question_management", quiz_name=selected_quiz_name or DEFAULT_QUIZ_NAME))

    rows = list(csv_reader)
    if not rows:
        flash("CSV file is empty.", "error")
        return redirect(url_for("quiz_question_management", quiz_name=selected_quiz_name or DEFAULT_QUIZ_NAME))

    first_row = rows[0]
    quiz_name = normalize_quiz_name(selected_quiz_name or first_row.get("quiz_name") or DEFAULT_QUIZ_NAME)
    quiz_bank = QuizBank.query.filter_by(quiz_name=quiz_name).first()
    if not quiz_bank:
        quiz_bank = QuizBank(
            quiz_name=quiz_name,
            quiz_title=(first_row.get("quiz_title") or quiz_name.replace("_", " ").title()).strip(),
            description=None,
            covered_cos=[],
            questions_per_attempt=10,
            is_shared_with_all_admins=False,
            created_by=user.id
        )
        db.session.add(quiz_bank)
        db.session.flush()

    if not can_admin_edit_quiz_bank(user, quiz_bank):
        flash("You do not have permission to upload questions into that quiz.", "error")
        db.session.rollback()
        return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

    questions_added = 0
    covered_cos = set(quiz_bank.covered_cos or [])
    try:
        for row in rows:
            co_number = (row.get("co_number") or "").strip().upper()
            answer_type = (row.get("answer_type") or "single").strip().lower()
            correct_answers = [answer.strip().upper() for answer in (row.get("correct_answers") or "").split(",") if answer.strip()]
            if co_number not in QUIZ_CO_OPTIONS:
                continue
            if answer_type not in ["single", "multiple"]:
                continue
            if answer_type == "single" and len(correct_answers) != 1:
                continue

            choices = []
            for index, key in enumerate(["option_a", "option_b", "option_c", "option_d"]):
                choices.append({
                    "id": ['A', 'B', 'C', 'D'][index],
                    "text": (row.get(key) or "").strip()
                })

            if any(not choice["text"] for choice in choices):
                continue

            question = QuizQuestion(
                question_text=(row.get("question_text") or "").strip(),
                co_number=co_number,
                points=int(row.get("points") or 1),
                answer_type=answer_type,
                choices=choices,
                correct_answers=correct_answers,
                explanation=(row.get("explanation") or "").strip() or None,
                created_by=user.id
            )
            db.session.add(question)
            db.session.flush()
            assign_question_to_quiz(quiz_bank, question)
            covered_cos.add(co_number)
            questions_added += 1

        quiz_bank.covered_cos = sorted(covered_cos)
        db.session.commit()
        flash(f"Successfully uploaded {questions_added} questions into '{quiz_bank.quiz_title}'.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error uploading CSV: {str(e)}", "error")

    return redirect(url_for("quiz_question_management", quiz_name=quiz_bank.quiz_name))


@app.route("/add_quiz_question", methods=["POST"])
def add_quiz_question():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    allow_admin_add_question = get_config(
        'allow_admin_add_question', '0') == '1'
    # Only allow if super_admin, or admin and allowed by config
    if not user or (user.admin_level != 'super_admin' and not (user.admin_level == 'admin' and allow_admin_add_question)):
        flash("Access denied. Only super admins or allowed admins can create quiz questions.", "error")
        return redirect(url_for("admin_dashboard"))

    quiz_name = request.form.get("quiz_name") or DEFAULT_QUIZ_NAME
    quiz_bank = get_quiz_bank_by_name(quiz_name)
    if not quiz_bank or not can_admin_edit_quiz_bank(user, quiz_bank):
        flash("You do not have permission to add questions to that quiz.", "error")
        return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

    try:
        question_text = request.form.get("question_text")
        co_number = request.form.get("co_number")
        # Default to 1 if not provided
        points = int(request.form.get("points", 1))
        answer_type = request.form.get("answer_type")

        # Get choices and format them as JSON objects
        choices = []
        choice_ids = ['A', 'B', 'C', 'D']
        for i in range(4):
            choice = request.form.get(f"choice_{i}")
            if choice:
                choices.append({
                    'id': choice_ids[i],
                    'text': choice.strip()
                })

        if len(choices) != 4:
            flash("All 4 choices are required.", "error")
            return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

        # Get correct answers - convert to choice IDs
        correct_answer_indices = []
        for value in request.form.getlist("correct_answers"):
            correct_answer_indices.append(int(value))

        if not correct_answer_indices:
            flash("At least one correct answer must be selected.", "error")
            return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

        # Convert indices to choice IDs (0->A, 1->B, 2->C, 3->D)
        correct_answers = []
        for idx in correct_answer_indices:
            if 0 <= idx < 4:
                correct_answers.append(choice_ids[idx])

        # Validate answer type
        if answer_type == "single" and len(correct_answers) > 1:
            flash("Single answer type can only have one correct answer.", "error")
            return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

        # Create new question
        new_question = QuizQuestion(
            question_text=question_text,
            co_number=co_number,
            points=points,
            answer_type=answer_type,
            choices=choices,
            correct_answers=correct_answers,
            explanation=(request.form.get("explanation") or "").strip() or None,
            created_by=user.id
        )

        db.session.add(new_question)
        db.session.flush()
        assign_question_to_quiz(quiz_bank, new_question)
        covered_cos = set(quiz_bank.covered_cos or [])
        covered_cos.add(co_number)
        quiz_bank.covered_cos = sorted(covered_cos)
        db.session.commit()

        flash(
            f"Quiz question added successfully! (Q{new_question.id})", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error adding question: {str(e)}", "error")

    return redirect(url_for("quiz_question_management", quiz_name=quiz_name))


@app.route("/edit_quiz_question/<int:question_id>", methods=["POST"])
def edit_quiz_question(question_id):
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    question = QuizQuestion.query.get_or_404(question_id)
    quiz_name = request.form.get("quiz_name") or DEFAULT_QUIZ_NAME
    quiz_bank = get_quiz_bank_by_name(quiz_name)
    if not user or not can_admin_edit_quiz_bank(user, quiz_bank):
        flash("Access denied. You cannot edit questions for this quiz.", "error")
        return redirect(url_for("admin_dashboard"))

    try:
        question.question_text = request.form.get("question_text")
        question.co_number = request.form.get("co_number")
        question.points = int(request.form.get("points"))
        question.answer_type = request.form.get("answer_type")
        question.explanation = (request.form.get("explanation") or "").strip() or None

        # Get choices and format them as JSON objects
        choices = []
        choice_ids = ['A', 'B', 'C', 'D']
        for i in range(4):
            choice = request.form.get(f"choice_{i}")
            if choice:
                choices.append({
                    'id': choice_ids[i],
                    'text': choice.strip()
                })

        if len(choices) != 4:
            flash("All 4 choices are required.", "error")
            return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

        question.choices = choices

        # Get correct answers - convert to choice IDs
        correct_answer_indices = []
        for value in request.form.getlist("correct_answers"):
            correct_answer_indices.append(int(value))

        if not correct_answer_indices:
            flash("At least one correct answer must be selected.", "error")
            return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

        # Convert indices to choice IDs (0->A, 1->B, 2->C, 3->D)
        correct_answers = []
        for idx in correct_answer_indices:
            if 0 <= idx < 4:
                correct_answers.append(choice_ids[idx])

        # Validate answer type
        if question.answer_type == "single" and len(correct_answers) > 1:
            flash("Single answer type can only have one correct answer.", "error")
            return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

        question.correct_answers = correct_answers
        covered_cos = {q.co_number for q in get_quiz_questions(quiz_bank) if q.co_number}
        covered_cos.add(question.co_number)
        quiz_bank.covered_cos = sorted(covered_cos)

        db.session.commit()
        flash(f"Question Q{question_id} updated successfully!", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error updating question: {str(e)}", "error")

    return redirect(url_for("quiz_question_management", quiz_name=quiz_name))


@app.route("/delete_quiz_question/<int:question_id>", methods=["POST"])
def delete_quiz_question(question_id):
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    question = QuizQuestion.query.get_or_404(question_id)
    quiz_name = request.form.get("quiz_name") or DEFAULT_QUIZ_NAME
    quiz_bank = get_quiz_bank_by_name(quiz_name)
    if not user or not can_admin_edit_quiz_bank(user, quiz_bank):
        flash("Access denied. You cannot delete questions from this quiz.", "error")
        return redirect(url_for("admin_dashboard"))

    try:
        assignment_count = QuizQuestionAssignment.query.filter_by(question_id=question_id).count()
        if assignment_count > 1:
            QuizQuestionAssignment.query.filter_by(
                quiz_id=quiz_bank.id, question_id=question_id).delete()
            quiz_bank.covered_cos = sorted({q.co_number for q in get_quiz_questions(quiz_bank) if q.co_number})
            db.session.commit()
            flash(f"Question Q{question_id} removed from '{quiz_bank.quiz_title}'.", "success")
            return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

        # Find all responses for this question
        responses = QuizResponse.query.filter_by(question_id=question_id).all()
        # If any response is from a student, block deletion
        student_response = False
        admin_response_ids = []
        for resp in responses:
            attempt = QuizAttempt.query.get(resp.attempt_id)
            if attempt:
                user = User.query.get(attempt.user_id)
                if user and user.role == 'student':
                    student_response = True
                    break
                elif user and user.role in ['admin', 'super_admin']:
                    admin_response_ids.append(resp.id)
        if student_response:
            flash(
                f"Cannot delete question Q{question_id}. It has been answered by at least one student.", "error")
            return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

        # Delete all admin/super_admin responses for this question
        if admin_response_ids:
            QuizResponse.query.filter(QuizResponse.id.in_(
                admin_response_ids)).delete(synchronize_session=False)
            db.session.commit()

        # Now safe to delete the question
        db.session.delete(question)
        db.session.commit()
        flash(f"Question Q{question_id} deleted successfully!", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting question: {str(e)}", "error")

    return redirect(url_for("quiz_question_management", quiz_name=quiz_name))


@app.route("/bulk_assign_points", methods=["POST"])
def bulk_assign_points():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    quiz_name = request.form.get("quiz_name") or DEFAULT_QUIZ_NAME
    quiz_bank = get_quiz_bank_by_name(quiz_name)
    if not user or not can_admin_edit_quiz_bank(user, quiz_bank):
        flash("Access denied. You cannot bulk update this quiz.", "error")
        return redirect(url_for("admin_dashboard"))

    try:
        bulk_points = int(request.form.get("bulk_points"))
        bulk_co = request.form.get("bulk_co")

        query = QuizQuestion.query.join(
            QuizQuestionAssignment, QuizQuestionAssignment.question_id == QuizQuestion.id
        ).filter(QuizQuestionAssignment.quiz_id == quiz_bank.id)
        if bulk_co:
            query = query.filter_by(co_number=bulk_co)

        questions = query.all()

        if not questions:
            flash("No questions found to update.", "warning")
            return redirect(url_for("quiz_question_management", quiz_name=quiz_name))

        # Update points
        for question in questions:
            question.points = bulk_points

        db.session.commit()

        co_filter = f" for {bulk_co}" if bulk_co else ""
        flash(
            f"Successfully updated {len(questions)} questions{co_filter} to {bulk_points} points each.", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error updating points: {str(e)}", "error")

    return redirect(url_for("quiz_question_management", quiz_name=quiz_name))


@app.route("/toggle_quiz_release", methods=["POST"])
def toggle_quiz_release():
    print("🔄 Toggle quiz release route called")

    if "rollnumber" not in session:
        print("❌ No rollnumber in session")
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        print(
            f"❌ Access denied. User: {user}, Role: {user.role if user else 'None'}")
        flash("Access denied. Only admins can control quiz release.", "error")
        return redirect(url_for("login"))

    print(f"✅ Admin user found: {user.rollnumber}, Level: {user.admin_level}")

    try:
        # For this database structure, we use a single default quiz
        quiz_name = 'default_quiz'
        print(f"🔍 Looking for quiz visibility with quiz_name: {quiz_name}")

        # Determine current visibility status based on admin level
        current_status = False
        if user.admin_level == 'super_admin':
            # Super admin sees global quiz visibility status
            global_visibility = QuizVisibility.query.filter_by(
                quiz_name=quiz_name, admin_id=None).first()
            current_status = global_visibility.is_released if global_visibility else False
            print(f"📊 Super admin current quiz status: {current_status}")
        else:
            # Sub admin sees BOTH their own AND global quiz visibility status
            # First check for global release (admin_id=None)
            global_visibility = QuizVisibility.query.filter_by(
                quiz_name=quiz_name, admin_id=None).first()
            # Then check for their own release
            own_visibility = QuizVisibility.query.filter_by(
                quiz_name=quiz_name, admin_id=user.id).first()

            # Quiz is released if EITHER global OR own release exists and is released
            is_globally_released = global_visibility.is_released if global_visibility else False
            is_own_released = own_visibility.is_released if own_visibility else False

            current_status = is_globally_released or is_own_released
            print(
                f"📊 Sub admin {user.id} current quiz status - Global: {is_globally_released}, Own: {is_own_released}, Combined: {current_status}")

        # Determine new status
        new_status = not current_status
        print(f"🔄 Toggling from {current_status} to {new_status}")

        # Update visibility
        result = set_quiz_visibility(user.id, new_status)
        print(f"💾 Set quiz visibility result: {result}")

        action = "released" if new_status else "locked"
        flash(f"Quiz {action} successfully!", "success")
        print(f"✅ Success: Quiz {action}")

    except Exception as e:
        print(f"💥 Error in toggle_quiz_release: {str(e)}")
        import traceback
        traceback.print_exc()
        flash(f"Error toggling quiz release: {str(e)}", "error")

    return redirect(url_for("admin_dashboard"))


@app.route("/toggle_selected_quiz_release", methods=["POST"])
def toggle_selected_quiz_release():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    quiz_name = request.form.get("quiz_name") or DEFAULT_QUIZ_NAME
    quiz_bank = get_quiz_bank_by_name(quiz_name)
    if not user or not quiz_bank or not can_admin_access_quiz_bank(user, quiz_bank):
        flash("You do not have access to release that quiz.", "error")
        return redirect(url_for("quiz_question_management"))

    try:
        current_status = get_quiz_release_status_for_admin(user, quiz_name)
        new_status = not current_status
        set_quiz_visibility(user.id, new_status, quiz_name=quiz_name)
        flash(
            f"Quiz '{quiz_bank.quiz_title}' {'released' if new_status else 'locked'} successfully.",
            "success"
        )
    except Exception as e:
        flash(f"Error toggling quiz release: {str(e)}", "error")

    return redirect(url_for("quiz_question_management", quiz_name=quiz_name))


# --- Quiz Taking Routes ---

@app.route("/take_quiz", methods=["GET", "POST"])
def take_quiz():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user:
        return redirect(url_for("login"))
    return handle_take_quiz_v2(user)

    print(
        f"🎯 take_quiz: User {user.rollnumber} (role: {user.role}) attempting to take quiz")
    if session.get("impersonating"):
        print(
            f"🎭 take_quiz: Currently impersonating {session.get('impersonated_admin_name')}")

    # Check if user can take quiz
    can_take, message = can_user_take_quiz(user.rollnumber)
    if not can_take:
        print(
            f"🚫 User {user.rollnumber} ({user.role}) cannot take quiz: {message}")
        flash(f"Cannot access quiz: {message}", "warning")
        # Redirect based on user role
        if user.role == "admin":
            return redirect(url_for("admin_dashboard"))
        else:
            return redirect(url_for("student_dashboard"))

    if request.method == "POST" and request.form.get("action") == "start":
        # Start a new quiz attempt
        print("🚀 Starting new quiz attempt...")
        return start_new_quiz_attempt(user)

    # Check if user has an active attempt
    active_attempt = QuizAttempt.query.filter_by(
        user_id=user.id, completed_at=None
    ).order_by(QuizAttempt.id.desc()).first()

    if active_attempt:
        # Resume existing attempt
        print(f"📋 Found active attempt {active_attempt.id}, resuming...")
        return render_quiz_attempt(active_attempt)

    # Show quiz start page
    print("📄 Showing quiz start page...")
    return render_quiz_start_page(user)


def handle_take_quiz_v2(user):
    available_quizzes = get_available_quiz_banks_for_user(user)
    if not available_quizzes:
        flash("Cannot access quiz: Quiz has not been released yet", "warning")
        return redirect(url_for("admin_dashboard" if user.role == "admin" else "student_dashboard"))

    selected_quiz_name = request.values.get("quiz_name") or available_quizzes[0].quiz_name
    selected_quiz = next((quiz for quiz in available_quizzes if quiz.quiz_name == selected_quiz_name), available_quizzes[0])

    if request.method == "POST" and request.form.get("action") == "start":
        can_take, message = can_user_take_quiz(user.rollnumber, quiz_name=selected_quiz.quiz_name)
        if not can_take:
            flash(f"Cannot access quiz: {message}", "warning")
            return redirect(url_for("take_quiz", quiz_name=selected_quiz.quiz_name))
        return start_new_quiz_attempt_v2(user, selected_quiz.quiz_name)

    active_attempt = QuizAttempt.query.filter_by(
        user_id=user.id, completed_at=None
    ).order_by(QuizAttempt.id.desc()).first()

    if active_attempt:
        return render_quiz_attempt(active_attempt)

    return render_quiz_start_page_v2(user, selected_quiz.quiz_name, available_quizzes)


def render_quiz_start_page_v2(user, quiz_name, available_quizzes=None):
    available_quizzes = available_quizzes or get_available_quiz_banks_for_user(user)
    selected_quiz = get_quiz_bank_by_name(quiz_name)
    quiz_questions = get_quiz_questions(selected_quiz)
    previous_attempts = QuizAttempt.query.filter_by(
        user_id=user.id, quiz_name=quiz_name
    ).filter(QuizAttempt.completed_at.isnot(None)).order_by(QuizAttempt.attempt_number).all()

    next_attempt_number = len(previous_attempts) + 1
    max_attempts = None if user.role == "admin" else 2
    total_questions = get_quiz_questions_per_attempt(selected_quiz, available_count=len(quiz_questions))
    total_marks = sum(
        sorted(((q.points or 1) for q in quiz_questions), reverse=True)[:total_questions]
    )

    return render_template(
        "quiz.html",
        quiz_mode="start",
        user=user,
        previous_attempts=previous_attempts,
        next_attempt_number=next_attempt_number,
        max_attempts=max_attempts,
        total_questions=total_questions,
        total_marks=int(total_marks),
        available_quizzes=available_quizzes,
        selected_quiz=selected_quiz
    )


def start_new_quiz_attempt_v2(user, quiz_name):
    selected_quiz = get_quiz_bank_by_name(quiz_name)
    selected_questions = select_random_questions_for_quiz(selected_quiz)
    choice_orders = {}

    if not selected_questions:
        flash("No questions available for this quiz. Please contact admin.", "error")
        return redirect(url_for("take_quiz", quiz_name=quiz_name))

    for question in selected_questions:
        shuffled_indices = list(range(len(question.choices)))
        random.shuffle(shuffled_indices)
        choice_orders[question.id] = shuffled_indices

    max_possible_score = sum(q.points for q in selected_questions)
    attempt_number = QuizAttempt.query.filter_by(
        user_id=user.id, quiz_name=quiz_name).count() + 1
    new_attempt = QuizAttempt(
        user_id=user.id,
        quiz_name=quiz_name,
        attempt_number=attempt_number,
        score=0,
        total_points=max_possible_score
    )

    db.session.add(new_attempt)
    db.session.commit()

    session[f'quiz_attempt_{new_attempt.id}_questions'] = [q.id for q in selected_questions]
    session[f'quiz_attempt_{new_attempt.id}_choice_orders'] = choice_orders
    return render_quiz_attempt(new_attempt)


def render_quiz_start_page(user):
    """Render the quiz start page with attempt info"""
    previous_attempts = QuizAttempt.query.filter_by(
        user_id=user.id
    ).filter(QuizAttempt.completed_at.isnot(None)).order_by(QuizAttempt.attempt_number).all()

    next_attempt_number = len(previous_attempts) + 1
    max_attempts = None if user.role == "admin" else 2

    # Count total questions and marks
    total_questions = 0
    total_marks = 0
    for co in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
        co_questions = QuizQuestion.query.filter_by(co_number=co).count()
        if co_questions >= 2:
            # We'll select 2 questions from this CO
            total_questions += 2
            # Get average points for this CO (we'll use this for estimate)
            avg_points = db.session.query(db.func.avg(
                QuizQuestion.points)).filter_by(co_number=co).scalar()
            total_marks += (avg_points or 1) * 2

    return render_template("quiz.html",
                           quiz_mode="start",
                           user=user,
                           previous_attempts=previous_attempts,
                           next_attempt_number=next_attempt_number,
                           max_attempts=max_attempts,
                           total_questions=total_questions,
                           total_marks=int(total_marks))


def start_new_quiz_attempt(user):
    """Start a new quiz attempt with randomly selected questions"""
    try:
        print(
            f"🚀 start_new_quiz_attempt: Starting quiz for user {user.rollnumber}")

        # Get 2 random questions from each CO that has questions
        selected_questions = []
        choice_orders = {}

        for co in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
            co_questions = QuizQuestion.query.filter_by(co_number=co).all()
            print(f"📊 Found {len(co_questions)} questions for {co}")
            if len(co_questions) >= 2:
                # Randomly select 2 questions from this CO
                selected_co_questions = random.sample(co_questions, 2)
                selected_questions.extend(selected_co_questions)

        print(f"📝 Total selected questions: {len(selected_questions)}")

        if len(selected_questions) == 0:
            print("❌ No questions available for quiz!")
            flash("No questions available for quiz. Please contact admin.", "error")
            # Redirect based on user role
            if user.role == "admin":
                return redirect(url_for("admin_dashboard"))
            else:
                return redirect(url_for("student_dashboard"))

        # Shuffle choices for each question and store the order
        for question in selected_questions:
            shuffled_indices = list(range(len(question.choices)))
            random.shuffle(shuffled_indices)
            choice_orders[question.id] = shuffled_indices

        # Calculate total possible score
        max_possible_score = sum(q.points for q in selected_questions)

        # Create new attempt
        attempt_number = QuizAttempt.query.filter_by(
            user_id=user.id).count() + 1

        new_attempt = QuizAttempt(
            user_id=user.id,
            attempt_number=attempt_number,
            score=0,
            total_points=max_possible_score
        )

        db.session.add(new_attempt)
        db.session.commit()

        # Store the questions and choice orders in session for this attempt
        session[f'quiz_attempt_{new_attempt.id}_questions'] = [
            q.id for q in selected_questions]
        session[f'quiz_attempt_{new_attempt.id}_choice_orders'] = choice_orders

        return render_quiz_attempt(new_attempt)

    except Exception as e:
        print(f"💥 Error in start_new_quiz_attempt: {str(e)}")
        import traceback
        traceback.print_exc()
        db.session.rollback()
        flash(f"Error starting quiz: {str(e)}", "error")
        # Redirect based on user role
        if user.role == "admin":
            return redirect(url_for("admin_dashboard"))
        else:
            return redirect(url_for("student_dashboard"))


def render_quiz_attempt(attempt):
    """Render the quiz taking interface for an active attempt"""
    user = User.query.get(attempt.user_id)
    questions = []

    print(f"🎯 render_quiz_attempt: Rendering quiz for attempt {attempt.id}")

    # Get questions for this attempt from session
    questions_given = session.get(f'quiz_attempt_{attempt.id}_questions', [])
    choice_orders = session.get(f'quiz_attempt_{attempt.id}_choice_orders', {})

    print(f"📋 Questions from session: {questions_given}")
    print(f"🔀 Choice orders from session: {choice_orders}")

    # If no session data found for this attempt, it means session was lost
    if not questions_given:
        print(
            f"❌ No session data found for attempt {attempt.id}. Session may have expired.")

        # Delete the incomplete attempt and force restart
        db.session.delete(attempt)
        db.session.commit()

        flash("Quiz session expired. Starting a new quiz attempt.", "warning")
        return redirect(url_for("take_quiz"))

    # Get questions for this attempt
    for question_id in questions_given:
        question = QuizQuestion.query.get(question_id)
        if question:
            print(
                f"✅ Found question {question_id}: {question.question_text[:50]}...")
            # Apply the shuffled choice order for this attempt
            shuffled_choices = []
            choice_order = choice_orders.get(
                str(question_id), list(range(len(question.choices))))

            for original_index in choice_order:
                if original_index < len(question.choices):
                    choice = question.choices[original_index]
                    # Extract text from choice object if it's a dict, otherwise use as-is
                    if isinstance(choice, dict) and 'text' in choice:
                        choice_text = choice['text']
                    else:
                        choice_text = str(choice)

                    shuffled_choices.append((original_index, choice_text))

            # Add shuffled choices to question object
            question.shuffled_choices = shuffled_choices
            questions.append(question)
        else:
            print(f"❌ Question {question_id} not found in database")

    print(f"📝 Final questions to render: {len(questions)}")

    return render_template("quiz.html",
                           quiz_mode="take",
                           user=user,
                           questions=questions,
                           attempt_id=attempt.id,
                           attempt_number=attempt.attempt_number)


@app.route("/submit_quiz", methods=["POST"])
def submit_quiz():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user:
        return redirect(url_for("login"))

    attempt_id = request.form.get("attempt_id")
    if not attempt_id:
        flash("Invalid quiz submission", "error")
        return redirect(url_for("take_quiz"))

    attempt = QuizAttempt.query.get(attempt_id)
    if not attempt or attempt.user_id != user.id:
        flash("Invalid quiz attempt", "error")
        return redirect(url_for("take_quiz"))

    if attempt.completed_at:
        flash("Quiz already submitted", "warning")
        return redirect(url_for("quiz_results", attempt_id=attempt_id))

    try:
        total_score = 0
        auto_submitted = request.form.get("auto_submit") == "1"

        # DEBUG: Print all form data received
        print("🔍 QUIZ SUBMISSION DEBUG:")
        print("Form data received:")
        for key, value in request.form.items():
            print(f"  {key}: {value}")
        print("Form lists:")
        for key in request.form.keys():
            if key.startswith('question_'):
                values = request.form.getlist(key)
                print(f"  {key}: {values}")

        # Get questions for this attempt from session
        questions_given = session.get(
            f'quiz_attempt_{attempt.id}_questions', [])
        print(f"Questions given from session: {questions_given}")

        if not questions_given:
            flash("Quiz session expired. Please retake the quiz.", "error")
            if user.role == "admin":
                return redirect(url_for("admin_dashboard"))
            else:
                return redirect(url_for("student_dashboard"))

        # Process each question's response
        for i, question_id in enumerate(questions_given):
            question = QuizQuestion.query.get(question_id)
            if not question:
                continue

            # Get user's selected answers
            selected_answers = request.form.getlist(f"question_{i}")
            selected_indices = [int(ans) for ans in selected_answers]

            # Convert selected indices to letter IDs to match database format
            choice_ids = ['A', 'B', 'C', 'D']
            selected_letter_ids = []
            for idx in selected_indices:
                if 0 <= idx < 4:
                    selected_letter_ids.append(choice_ids[idx])

            print(f"🔍 Question {i} (ID: {question_id}):")
            print(f"  Question text: {question.question_text[:50]}...")
            print(f"  Correct answers (letters): {question.correct_answers}")
            print(f"  Selected answers raw: {selected_answers}")
            print(f"  Selected indices: {selected_indices}")
            print(f"  Selected letter IDs: {selected_letter_ids}")

            # Calculate score for this question using letter IDs
            question_score = calculate_question_score(
                question, selected_letter_ids)

            # For multiple choice, consider partial credit as correct
            # For single choice, only full points count as correct
            if question.answer_type == 'multiple':
                is_correct = question_score > 0  # Any points earned = partially correct
            else:
                is_correct = question_score == question.points  # Full points only

            print(f"  Question score: {question_score}/{question.points}")
            print(f"  Is correct: {is_correct}")

            # Create response record (store the original indices for consistency)
            response = QuizResponse(
                attempt_id=attempt.id,
                question_id=question.id,
                selected_answer=selected_indices,  # Keep original indices for display
                points_earned=question_score,
                is_correct=is_correct
            )

            db.session.add(response)
            total_score += question_score

            print(
                f"  Added response: points_earned={question_score}, is_correct={is_correct}")
            print(f"  Running total_score: {total_score}")

        print(f"🏁 FINAL SCORING:")
        print(f"  Total score calculated: {total_score}")

        # Update attempt as completed
        attempt.score = total_score
        attempt.total_points = sum(question.points for question in QuizQuestion.query.filter(
            QuizQuestion.id.in_(questions_given)).all())
        attempt.completed_at = datetime.now()
        attempt.auto_submitted = auto_submitted

        print(f"  Attempt score set to: {attempt.score}")
        print(f"  Attempt total_points set to: {attempt.total_points}")
        print(f"  Attempt auto_submitted set to: {attempt.auto_submitted}")

        db.session.commit()
        print(f"  Database committed successfully")

        if auto_submitted:
            flash(
                f"Quiz auto-submitted due to anti-cheating protection. Score: {total_score}/{attempt.total_points}",
                "warning"
            )
        else:
            flash(
                f"Quiz submitted successfully! Score: {total_score}/{attempt.total_points}", "success")

        # Clean up session data
        session.pop(f'quiz_attempt_{attempt.id}_questions', None)
        session.pop(f'quiz_attempt_{attempt.id}_choice_orders', None)

        return redirect(url_for("quiz_results", attempt_id=attempt.id))

    except Exception as e:
        db.session.rollback()
        print(f"💥 Error submitting quiz: {str(e)}")
        import traceback
        traceback.print_exc()
        flash(f"Error submitting quiz: {str(e)}", "error")
        # Redirect based on user role
        if user.role == "admin":
            return redirect(url_for("admin_dashboard"))
        else:
            return redirect(url_for("student_dashboard"))


def calculate_question_score(question, selected_letter_ids):
    """Calculate score for a question based on selected answers (letter IDs)"""
    correct_answers = set(
        question.correct_answers)  # Set of letter IDs like ['A', 'B']
    # Set of letter IDs like ['A', 'B']
    selected_answers = set(selected_letter_ids)

    print(f"🔍 SCORING DEBUG:")
    print(f"  Question ID: {question.id}")
    print(f"  Answer type: {question.answer_type}")
    print(f"  Correct answers: {correct_answers}")
    print(f"  Selected answers: {selected_answers}")
    print(f"  Max points: {question.points}")

    if question.answer_type == 'single':
        # For single answer questions: all or nothing
        if selected_answers == correct_answers:
            print(f"  Single answer: CORRECT - Full points")
            return question.points
        else:
            print(f"  Single answer: WRONG - No points")
            return 0
    else:
        # For multiple answer questions: points divided equally among correct answers
        if not correct_answers:
            print(f"  Multiple answer: No correct answers defined - No points")
            return 0

        total_correct = len(correct_answers)
        points_per_correct = question.points / total_correct

        # Calculate score based on individual correct selections
        correct_selected = len(selected_answers.intersection(correct_answers))
        incorrect_selected = len(selected_answers - correct_answers)

        # Award points for each correct selection
        earned_points = correct_selected * points_per_correct

        # Apply penalty for incorrect selections (50% penalty per wrong selection)
        penalty = incorrect_selected * points_per_correct * 0.5

        final_score = max(0, earned_points - penalty)

        print(f"  Multiple answer calculation:")
        print(f"    Total correct answers: {total_correct}")
        print(f"    Points per correct: {points_per_correct:.3f}")
        print(f"    Correct selected: {correct_selected}")
        print(f"    Incorrect selected: {incorrect_selected}")
        print(f"    Earned points: {earned_points:.3f}")
        print(f"    Penalty: {penalty:.3f}")
        print(f"    Final score before rounding: {final_score:.3f}")

        # Round to 2 decimal places for better precision
        rounded_score = round(final_score, 2)
        print(f"    Final score after rounding: {rounded_score}")

        return rounded_score


@app.route("/quiz_results/<int:attempt_id>")
def quiz_results(attempt_id):
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    initialize_quiz_bank_system()

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user:
        return redirect(url_for("login"))

    attempt = QuizAttempt.query.get_or_404(attempt_id)

    # Check access permissions
    if attempt.user_id != user.id and user.role != "admin":
        flash("Access denied", "error")
        return redirect(url_for("student_dashboard"))

    # For admins viewing other students' results, we should get the actual student user
    attempt_user = User.query.get(
        attempt.user_id) if user.role == "admin" and attempt.user_id != user.id else user

    if not attempt.completed_at:
        flash("Quiz not yet completed", "warning")
        return redirect(url_for("take_quiz"))

    # Get all responses for this attempt
    responses = QuizResponse.query.filter_by(attempt_id=attempt.id).all()

    # Calculate CO-wise performance
    co_performance = {}
    for co in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
        co_responses = [r for r in responses if r.question.co_number == co]
        if co_responses:
            total_scored = sum(r.points_earned for r in co_responses)
            total_possible = sum(r.question.points for r in co_responses)
            percentage = (total_scored / total_possible) * \
                100 if total_possible > 0 else 0

            co_performance[co] = {
                'scored': total_scored,
                'total': total_possible,
                'percentage': round(percentage, 1)
            }

    # Check if user can retake
    quiz_attempts = QuizAttempt.query.filter_by(
        user_id=attempt.user_id, quiz_name=attempt.quiz_name
    ).count()
    can_retake = user.role in ["admin", "super_admin"] or quiz_attempts < 2

    correct_answers = sum(1 for r in responses if r.is_correct)
    total_questions = len(responses)
    hide_correct_answers = bool(attempt.auto_submitted)

    return render_template("quiz.html",
                           quiz_mode="results",
                           user=user,
                           attempt_user=attempt_user,
                           attempt=attempt,
                           responses=responses,
                           co_performance=co_performance,
                           correct_answers=correct_answers,
                           total_questions=total_questions,
                           can_retake=can_retake,
                           retake_quiz_name=attempt.quiz_name,
                           hide_correct_answers=hide_correct_answers)


@app.route("/reset_quiz_attempts", methods=["POST"])
def reset_quiz_attempts():
    """Reset all quiz attempts for the current admin (for testing purposes)"""
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash(
            "Access denied. Only admins and super admins can reset quiz attempts.", "error")
        return redirect(url_for("take_quiz"))

    try:
        # Get all attempt IDs for this user (completed and active)
        attempt_ids = [a.id for a in QuizAttempt.query.filter_by(
            user_id=user.id).all()]

        # Delete all quiz responses for these attempts
        deleted_responses = 0
        if attempt_ids:
            deleted_responses = db.session.query(QuizResponse).filter(
                QuizResponse.attempt_id.in_(attempt_ids)
            ).delete(synchronize_session=False)

        # Delete all quiz attempts for this user
        deleted_attempts = QuizAttempt.query.filter_by(
            user_id=user.id).delete()

        db.session.commit()

        flash(
            f"Successfully reset quiz attempts! Deleted {deleted_attempts} attempts and {deleted_responses} responses.", "success")
        print(
            f"🔄 Admin {user.rollnumber} reset quiz attempts: {deleted_attempts} attempts, {deleted_responses} responses")

    except Exception as e:
        db.session.rollback()
        print(
            f"💥 Error resetting quiz attempts for {user.rollnumber}: {str(e)}")
        flash(f"Error resetting quiz attempts: {str(e)}", "error")

    return redirect(url_for("take_quiz"))


@app.route("/admin_quiz_co_analysis")
def admin_quiz_co_analysis():
    if "rollnumber" not in session:
        return redirect(url_for("login"))

    user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
    if not user or user.role not in ["admin", "super_admin"]:
        flash("Access denied. Only admins can view quiz CO analysis.", "error")
        return redirect(url_for("student_dashboard"))

    # Clear flash messages to avoid showing on CO analysis page
    get_flashed_messages()

    try:
        # Debug: Check what quiz questions we have
        all_quiz_questions = QuizQuestion.query.all()
        co_debug = {}
        for co in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
            co_questions = [q for q in all_quiz_questions if q.co_number == co]
            co_debug[co] = {
                'count': len(co_questions),
                'total_marks': sum(q.points for q in co_questions)
            }
        print(f"🔍 DEBUG Quiz CO Distribution: {co_debug}")

        # Get students based on admin level
        if user.admin_level == 'super_admin':
            students = User.query.filter_by(role='student').all()
        else:
            students = get_students_for_admin(user)

        # Calculate summary statistics
        total_students = len(students)
        students_attempted = 0
        students_passed = 0
        total_attempts = 0
        total_scores = []

        # Get all quiz attempts for these students
        student_ids = [s.id for s in students]
        all_attempts = QuizAttempt.query.filter(
            QuizAttempt.user_id.in_(student_ids),
            QuizAttempt.completed_at.isnot(None)
        ).all()

        total_attempts = len(all_attempts)

        # Calculate student-level statistics
        student_performance = []
        attempted_students = set()

        for student in students:
            student_attempts = [
                a for a in all_attempts if a.user_id == student.id]
            attempt_count = len(student_attempts)

            if attempt_count > 0:
                attempted_students.add(student.id)
                best_attempts = list(get_best_attempts_by_quiz(student_attempts).values())
                total_score = sum((attempt.score or 0)
                                  for attempt in best_attempts)
                total_points = sum((attempt.total_points or 0)
                                   for attempt in best_attempts)
                percentage = (total_score / total_points) * \
                    100 if total_points > 0 else 0

                total_scores.append(percentage)
                if percentage >= 60:
                    students_passed += 1

                # Calculate CO-wise performance for this student
                co_performance = calculate_student_quiz_co_performance(
                    student.id)

                student_performance.append({
                    'rollnumber': student.rollnumber,
                    'quiz_count': len(best_attempts),
                    'total_score': total_score,
                    'total_points': total_points,
                    'percentage': percentage,
                    'co_performance': co_performance
                })
            else:
                student_performance.append({
                    'rollnumber': student.rollnumber,
                    'quiz_count': 0,
                    'total_score': 0,
                    'total_points': 0,
                    'percentage': 0,
                    'co_performance': {}
                })

        students_attempted = len(attempted_students)
        avg_score = sum(total_scores) / \
            len(total_scores) if total_scores else 0
        pass_percentage = (students_passed / students_attempted) * \
            100 if students_attempted > 0 else 0

        # Calculate CO analysis
        co_analysis = calculate_quiz_co_analysis(student_ids)

        # Calculate performance distribution
        # excellent, good, average, poor, not_attempted
        performance_distribution = [0, 0, 0, 0, 0]
        for student in student_performance:
            if student['quiz_count'] == 0:
                performance_distribution[4] += 1
            elif student['percentage'] >= 80:
                performance_distribution[0] += 1
            elif student['percentage'] >= 60:
                performance_distribution[1] += 1
            elif student['percentage'] >= 40:
                performance_distribution[2] += 1
            else:
                performance_distribution[3] += 1

        # Calculate attempt distribution
        attempt_distribution = [0, 0, 0, 0]  # 0, 1, 2, 3+ attempts
        for student in student_performance:
            if student['quiz_count'] == 0:
                attempt_distribution[0] += 1
            elif student['quiz_count'] == 1:
                attempt_distribution[1] += 1
            elif student['quiz_count'] == 2:
                attempt_distribution[2] += 1
            else:
                attempt_distribution[3] += 1

        # Get question bank statistics
        question_bank_stats = {}
        for co in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
            question_bank_stats[co] = QuizQuestion.query.filter_by(
                co_number=co).count()

        quiz_bank_titles = {
            quiz.quiz_name: quiz.quiz_title
            for quiz in QuizBank.query.all()
        }

        quiz_summary = {}
        for attempt in all_attempts:
            quiz_name = attempt.quiz_name or DEFAULT_QUIZ_NAME
            summary = quiz_summary.setdefault(quiz_name, {
                'quiz_name': quiz_name,
                'quiz_title': quiz_bank_titles.get(quiz_name, quiz_name),
                'students': {},
                'highest_score': 0,
                'highest_total_points': 0,
                'top_rollnumbers': [],
                'total_percentage': 0,
                'average_percentage': 0
            })
            current_best = summary['students'].get(attempt.user_id)
            current_score = attempt.score or 0
            current_total = attempt.total_points or 0
            current_percentage = (current_score / current_total * 100) if current_total > 0 else 0
            if (
                current_best is None
                or current_score > current_best['score']
                or (current_score == current_best['score'] and current_percentage > current_best['percentage'])
            ):
                summary['students'][attempt.user_id] = {
                    'rollnumber': attempt.user.rollnumber if attempt.user else '',
                    'score': current_score,
                    'total_points': current_total,
                    'percentage': current_percentage
                }

        for summary in quiz_summary.values():
            summary['students_count'] = len(summary['students'])
            best_attempts = list(summary['students'].values())
            if best_attempts:
                summary['highest_score'] = max(item['score'] for item in best_attempts)
                highest_attempt = max(
                    best_attempts,
                    key=lambda item: (item['score'], item['total_points'], item['percentage'])
                )
                summary['highest_total_points'] = highest_attempt['total_points']
                summary['top_rollnumbers'] = sorted(
                    item['rollnumber']
                    for item in best_attempts
                    if item['score'] == summary['highest_score']
                )
                summary['average_percentage'] = (
                    sum(item['percentage'] for item in best_attempts) / len(best_attempts)
                )
            else:
                summary['highest_score'] = 0
                summary['highest_total_points'] = 0
                summary['top_rollnumbers'] = []
                summary['average_percentage'] = 0
            del summary['students']

        quiz_summary = sorted(
            quiz_summary.values(),
            key=lambda item: item['quiz_title'].lower()
        )

        recent_attempts = []
        for attempt in sorted(
            all_attempts,
            key=lambda item: item.completed_at or item.started_at or datetime.min,
            reverse=True
        ):
            quiz_name = attempt.quiz_name or DEFAULT_QUIZ_NAME
            recent_attempts.append({
                'user': attempt.user,
                'quiz_name': quiz_name,
                'quiz_title': quiz_bank_titles.get(quiz_name, quiz_name),
                'attempt_number': attempt.attempt_number,
                'score': attempt.score or 0,
                'total_points': attempt.total_points or 0,
                'percentage': ((attempt.score or 0) / attempt.total_points) * 100 if attempt.total_points > 0 else 0,
                'completed_at': attempt.completed_at,
                'started_at': attempt.started_at
            })

        # Prepare data for charts
        co_labels = list(co_analysis.keys())
        co_percentages = [co_analysis[co]['average_percentage']
                          for co in co_labels]

        return render_template("admin_quiz_co_analysis.html",
                               total_students=total_students,
                               students_attempted=students_attempted,
                               students_passed=students_passed,
                               total_attempts=total_attempts,
                               avg_score=avg_score,
                               pass_percentage=pass_percentage,
                               co_analysis=co_analysis,
                               student_performance=student_performance,
                               performance_distribution=performance_distribution,
                               attempt_distribution=attempt_distribution,
                               question_bank_stats=question_bank_stats,
                               quiz_summary=quiz_summary,
                               recent_attempts=recent_attempts,
                               co_labels=co_labels,
                               co_percentages=co_percentages)

    except Exception as e:
        print(f"🔴 ERROR in admin_quiz_co_analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return f"<h1>Error in Quiz CO Analysis</h1><p>Error: {str(e)}</p><pre>{traceback.format_exc()}</pre>"


def calculate_student_quiz_co_performance(user_id):
    """Calculate CO-wise performance across all completed quiz attempts."""
    co_totals = {
        co: {"marks_obtained": 0.0, "total_marks": 0.0}
        for co in QUIZ_CO_OPTIONS
    }

    # Get all completed attempts for this user
    attempts = QuizAttempt.query.filter_by(
        user_id=user_id
    ).filter(QuizAttempt.completed_at.isnot(None)).all()

    if not attempts:
        return {}

    best_attempts = list(get_best_attempts_by_quiz(attempts).values())
    attempt_ids = [attempt.id for attempt in best_attempts]
    attempt_quiz_name = {
        attempt.id: (attempt.quiz_name or DEFAULT_QUIZ_NAME)
        for attempt in best_attempts
    }
    quiz_names = sorted(set(attempt_quiz_name.values()))
    quiz_bank_map = {
        quiz.quiz_name: quiz
        for quiz in QuizBank.query.filter(QuizBank.quiz_name.in_(quiz_names)).all()
    }
    responses = QuizResponse.query.filter(
        QuizResponse.attempt_id.in_(attempt_ids)
    ).all()
    responses_by_attempt = {}
    for response in responses:
        responses_by_attempt.setdefault(response.attempt_id, []).append(response)

    for attempt_id, attempt_responses in responses_by_attempt.items():
        quiz_name = attempt_quiz_name.get(attempt_id, DEFAULT_QUIZ_NAME)
        quiz_cos = get_dom_quiz_cos(quiz_name, quiz_bank_map)
        if not quiz_cos:
            quiz_cos = sorted({
                response.question.co_number
                for response in attempt_responses
                if response.question and response.question.co_number in QUIZ_CO_OPTIONS
            })
        quiz_cos = [co for co in quiz_cos if co in QUIZ_CO_OPTIONS]
        if not quiz_cos:
            continue

        co_shares = get_quiz_co_question_share(quiz_bank_map.get(quiz_name))
        if not co_shares:
            co_shares = {co: 1.0 for co in quiz_cos}
        total_share = sum(co_shares.get(co, 0) for co in quiz_cos) or len(quiz_cos)
        attempt_scored = sum(response.points_earned for response in attempt_responses)
        attempt_possible = sum(response.question.points for response in attempt_responses)
        for co in quiz_cos:
            weight = co_shares.get(co, 0) / total_share
            co_totals[co]["marks_obtained"] += attempt_scored * weight
            co_totals[co]["total_marks"] += attempt_possible * weight

    # Calculate CO-wise scores
    co_performance = {}
    for co in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
        total_scored = co_totals[co]["marks_obtained"]
        total_possible = co_totals[co]["total_marks"]
        if total_possible > 0:
            percentage = (total_scored / total_possible) * \
                100
            co_performance[co] = {
                'marks_obtained': total_scored,
                'total_marks': total_possible,
                'percentage': percentage
            }
        else:
            co_performance[co] = {
                'marks_obtained': 0,
                'total_marks': 0,
                'percentage': 0
            }

    return co_performance


def calculate_quiz_co_analysis(student_ids):
    """Calculate overall CO analysis across every completed quiz attempt."""
    co_analysis = {}

    attempt_rows = db.session.query(
        QuizAttempt.id,
        QuizAttempt.user_id,
        QuizAttempt.quiz_name
    ).filter(
            QuizAttempt.user_id.in_(student_ids),
            QuizAttempt.completed_at.isnot(None)
    ).all()
    attempt_ids = [row.id for row in attempt_rows]
    attempt_owner = {row.id: row.user_id for row in attempt_rows}
    attempt_quiz_name = {row.id: (row.quiz_name or DEFAULT_QUIZ_NAME) for row in attempt_rows}
    quiz_names = sorted(set(attempt_quiz_name.values()))
    quiz_bank_map = {
        quiz.quiz_name: quiz
        for quiz in QuizBank.query.filter(QuizBank.quiz_name.in_(quiz_names)).all()
    }
    # Get all QuizResponses for every completed attempt
    if attempt_ids:
        asked_responses = QuizResponse.query.filter(
            QuizResponse.attempt_id.in_(attempt_ids)).all()
    else:
        asked_responses = []

    responses_by_attempt = {}
    for response in asked_responses:
        responses_by_attempt.setdefault(response.attempt_id, []).append(response)

    responses_by_student = {}
    for response in asked_responses:
        student_id = attempt_owner.get(response.attempt_id)
        if student_id is not None:
            responses_by_student.setdefault(student_id, []).append(response)

    quiz_posted_counts = {
        co: {"questions": 0.0, "marks": 0.0}
        for co in QUIZ_CO_OPTIONS
    }

    for quiz_name in quiz_names:
        quiz_cos = get_dom_quiz_cos(quiz_name, quiz_bank_map)
        quiz_cos = [co for co in quiz_cos if co in QUIZ_CO_OPTIONS]
        if not quiz_cos:
            continue

        quiz_bank = quiz_bank_map.get(quiz_name)
        co_shares = get_quiz_co_question_share(quiz_bank)
        if not co_shares:
            base_quota = 10 // len(quiz_cos)
            remainder = 10 % len(quiz_cos)
            co_shares = {co: float(base_quota) for co in quiz_cos}
            for co in quiz_cos:
                if remainder <= 0:
                    break
                co_shares[co] += 1.0
                remainder -= 1

        for co, co_share_count in co_shares.items():
            quiz_posted_counts[co]["questions"] += co_share_count
            quiz_posted_counts[co]["marks"] += co_share_count

    for co in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
        total_questions = round(quiz_posted_counts[co]["questions"], 2)
        max_marks = round(quiz_posted_counts[co]["marks"], 2)

        # Debug print for admin: show which questions are being counted for this CO
        print(
            f"[CO-DEBUG] {co}: posted_questions={total_questions}, posted_max_marks={max_marks}")
        print(
            f"[CO-DEBUG] {co}: total_questions={total_questions}, max_marks={max_marks}")

        # Get all responses for this CO from student attempts
        co_percentages = []

        for student_id in student_ids:
            co_scored = 0.0
            co_possible = 0.0
            for attempt_id, responses in responses_by_attempt.items():
                if attempt_owner.get(attempt_id) != student_id:
                    continue
                quiz_name = attempt_quiz_name.get(attempt_id, DEFAULT_QUIZ_NAME)
                quiz_cos = get_dom_quiz_cos(quiz_name, quiz_bank_map)
                if not quiz_cos:
                    quiz_cos = sorted({
                        response.question.co_number
                        for response in responses
                        if response.question and response.question.co_number in QUIZ_CO_OPTIONS
                    })
                quiz_cos = [quiz_co for quiz_co in quiz_cos if quiz_co in QUIZ_CO_OPTIONS]
                if co not in quiz_cos:
                    continue

                co_shares = get_quiz_co_question_share(quiz_bank_map.get(quiz_name))
                if not co_shares:
                    co_shares = {quiz_co: 1.0 for quiz_co in quiz_cos}
                total_share = sum(co_shares.get(quiz_co, 0) for quiz_co in quiz_cos) or len(quiz_cos)
                weight = co_shares.get(co, 0) / total_share
                attempt_scored = sum(response.points_earned for response in responses)
                attempt_possible = sum(response.question.points for response in responses)
                co_scored += attempt_scored * weight
                co_possible += attempt_possible * weight

            if co_possible > 0:
                percentage = (co_scored / co_possible) * 100
                co_percentages.append(percentage)

        # Calculate statistics
        if co_percentages:
            average_percentage = sum(co_percentages) / len(co_percentages)
            excellent_count = sum(1 for p in co_percentages if p >= 80)
            good_count = sum(1 for p in co_percentages if 60 <= p < 80)
            poor_count = sum(1 for p in co_percentages if p < 60)
        else:
            average_percentage = 0
            excellent_count = 0
            good_count = 0
            poor_count = 0

        co_analysis[co] = {
            'total_questions': total_questions,
            'max_marks': max_marks,
            'average_percentage': average_percentage,
            'excellent_count': excellent_count,
            'good_count': good_count,
            'poor_count': poor_count
        }

    return co_analysis


def ensure_database_columns():
    """Simple check for database columns - non-intrusive"""
    print("🔍 Database columns check: p11_attempts and p11_score should exist")
    # Note: Columns should be added manually via SQL if they don't exist


@app.route("/python_lab_reset_individual_release/<int:experiment_id>/<int:student_id>", methods=["POST"])
def python_lab_reset_individual_release(experiment_id, student_id):
    """Reset release for a single student for a given experiment"""
    if "loggedin" not in session:
        return redirect(url_for("login"))
    release = PythonLabStudentExperimentRelease.query.filter_by(
        experiment_id=experiment_id, student_id=student_id
    ).first()
    if release:
        release.is_released = False
        # Delete all responses for this student and experiment
        PythonLabResponse.query.filter_by(
            user_id=student_id, experiment_id=experiment_id
        ).delete()
        # Delete all manual marks for this student and experiment
        PythonLabManualMarks.query.filter_by(
            student_id=student_id, experiment_id=experiment_id
        ).delete()
        # Delete attempt record for this student and experiment
        PythonLabAttempt.query.filter_by(
            student_id=student_id, experiment_id=experiment_id
        ).delete()
        db.session.commit()
        flash(f"Release and marks reset for student {student_id}.", "success")
    else:
        flash("No release record found for this student.", "warning")
    return redirect(url_for("python_lab_individual_releases", experiment_id=experiment_id))


@app.route("/python_lab_bulk_reset_releases/<int:experiment_id>", methods=["POST"])
def python_lab_bulk_reset_releases(experiment_id):
    """Bulk reset releases for all students for a given experiment"""
    if "loggedin" not in session:
        return redirect(url_for("login"))
    releases = PythonLabStudentExperimentRelease.query.filter_by(
        experiment_id=experiment_id
    ).all()
    count = 0
    student_ids = []
    for release in releases:
        if release.is_released:
            release.is_released = False
            count += 1
        student_ids.append(release.student_id)
    # Delete all responses, manual marks, and attempts for all students in this experiment
    for student_id in student_ids:
        PythonLabResponse.query.filter_by(
            user_id=student_id, experiment_id=experiment_id
        ).delete()
        PythonLabManualMarks.query.filter_by(
            student_id=student_id, experiment_id=experiment_id
        ).delete()
        PythonLabAttempt.query.filter_by(
            student_id=student_id, experiment_id=experiment_id
        ).delete()
    db.session.commit()
    flash(
        f"Bulk reset completed for {count} students. All marks and attempts reset.", "success")
    return redirect(url_for("python_lab_individual_releases", experiment_id=experiment_id))


@app.route("/kdm_lab_individual_releases/<int:experiment_id>")
def kdm_lab_individual_releases(experiment_id):
    """Manage individual student releases for a specific KDM Lab experiment"""
    if "loggedin" not in session:
        return redirect(url_for("login"))

    # Ensure only admin or super_admin can access
    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        flash("Admin access required.", "danger")
        return redirect(url_for("login"))

    # Get the experiment
    experiment = KDMLabExperiment.query.get_or_404(experiment_id)

    # Get students based on admin level
    if current_admin.admin_level == "super_admin":
        students = User.query.filter_by(role="student").all()
    else:
        students = get_students_for_admin(current_admin)

    # Sort students by roll number
    students.sort(key=lambda s: roll_sort_key(s.rollnumber))

    # Get existing individual releases for this experiment
    individual_releases = {}
    releases = KDMLabStudentExperimentRelease.query.filter_by(
        admin_id=current_admin.id,
        experiment_id=experiment_id
    ).all()

    for release in releases:
        individual_releases[release.student_id] = release.is_released

    # Check if experiment is globally/admin released (bulk release)
    bulk_released = False
    if current_admin.admin_level == "super_admin":
        # Check if any admin has released this experiment
        bulk_released = KDMlabAdminRelease.query.filter_by(
            experiment_id=experiment_id, is_released=True).first() is not None
    else:
        # Check if this admin has released the experiment in bulk
        admin_release = KDMlabAdminRelease.query.filter_by(
            admin_id=current_admin.id,
            experiment_id=experiment_id,
            is_released=True
        ).first()
        bulk_released = admin_release is not None

    return render_template(
        "kdm_lab_individual_releases.html",
        students=students,
        experiment=experiment,
        individual_releases=individual_releases,
        bulk_released=bulk_released,
        current_admin=current_admin
    )


@app.route("/kdm_lab_toggle_individual_release", methods=["POST"])
def kdm_lab_toggle_individual_release():
    """Toggle individual experiment release for a specific student"""
    if "loggedin" not in session:
        return jsonify({"error": "Not logged in"}), 401

    current_admin = User.query.filter_by(
        rollnumber=session["rollnumber"]).first()
    if not current_admin or current_admin.role not in ["admin", "super_admin"]:
        return jsonify({"error": "Admin access required"}), 403

    try:
        data = request.get_json()
        student_id = data.get('student_id')
        experiment_id = data.get('experiment_id')
        is_released = data.get('is_released')

        # Find existing release record
        release_record = KDMLabStudentExperimentRelease.query.filter_by(
            admin_id=current_admin.id,
            student_id=student_id,
            experiment_id=experiment_id
        ).first()

        if release_record:
            release_record.is_released = is_released
        else:
            # Create new release record
            release_record = KDMLabStudentExperimentRelease(
                admin_id=current_admin.id,
                student_id=student_id,
                experiment_id=experiment_id,
                is_released=is_released
            )
            db.session.add(release_record)

        db.session.commit()
        return jsonify({"success": True, "is_released": is_released})

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


# ---------------- Run ----------------
if __name__ == "__main__":
    # Simple database column check
    ensure_database_columns()

    # Initialize problem visibility on startup (but don't crash if it fails)
    # Debug route to check KDM quiz attempts
    @app.route("/debug_kdm_quiz_attempts")
    def debug_kdm_quiz_attempts():
        if "loggedin" not in session:
            return "Please login first"

        user = User.query.filter_by(rollnumber=session["rollnumber"]).first()
        if not user or user.role not in ["admin", "super_admin"]:
            return "Admin access required"

        attempts = KDMLabQuizAttempt.query.all()
        result = f"<h3>KDM Lab Quiz Attempts ({len(attempts)} total)</h3><br>"

        for attempt in attempts:
            student = User.query.get(attempt.student_id)
            admin = User.query.get(attempt.admin_id)
            result += f"Student: {student.rollnumber if student else 'Unknown'} | Admin: {admin.rollnumber if admin else 'Unknown'} | Score: {attempt.score} | Completed: {attempt.is_completed}<br>"

        return result

    try:
        initialize_problem_visibility()
    except Exception as e:
        print(f"Warning: Could not initialize problem visibility: {e}")

    app.run(debug=True)
